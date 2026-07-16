"""
formatting_stage2.py

Stage 2 formatting engine for EUR-Lex / Official Journal XHTML.

This version uses only open-source Python tooling:
    - beautifulsoup4 / lxml for source parsing
    - python-docx for document construction
    - zipfile + lxml direct OOXML patching for native Word footnotes

It does NOT use Aspose.Words.

Key features:
    - Rebuilds a clean Word document from the EUR-Lex XHTML source.
    - Uses the EUR-Lex ID schema: cit_*, rct_*, cpt_*, art_*, anx_*.
    - Adds the missing adoption formula: "HAVE ADOPTED THIS REGULATION:".
    - Adds Chapter and Article level-1 headings, with line breaks between number and title.
    - Converts XHTML footnote references into native Word footnotes using OOXML.
    - Converts Annex VI into a two-column Word table from the underlying HTML table.

Run:
    python formatting_stage2.py L_2017168EN.01001201.xml.html Prospectus_Regulation_Stage2.docx
"""

from __future__ import annotations

from pathlib import Path
import argparse
import os
import re
import shutil
import tempfile
import zipfile

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


# =============================================================================
# General helpers
# =============================================================================

NBSP = "\u00A0"
FN_PLACEHOLDER_RE = re.compile(r"\[\[FN:(\d+)\]\]")


def clean_text(text: str | None) -> str:
    """Normalise whitespace without changing legal substance."""
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    return " ".join(text.split())


def load_html(path: str | Path) -> BeautifulSoup:
    return BeautifulSoup(Path(path).read_text(encoding="utf-8", errors="replace"), "lxml")


def set_run_font(run, size: int = 11, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_base_format(paragraph, *, before=6, after=6, line_spacing=1.16) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def apply_non_breaking_spaces(text: str) -> str:
    """Apply NBSP rules for thousands, currency and percentages."""
    text = re.sub(r"(?<=\d) (?=\d{3}\b)", NBSP, text)
    text = re.sub(r"\b(EUR)\s+(?=\d)", r"\1" + NBSP, text)
    text = re.sub(r"€\s+(?=\d)", "€" + NBSP, text)
    text = re.sub(r"(?<=\d)\s+%", NBSP + "%", text)
    return text


def new_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


# =============================================================================
# Native Word footnote placeholders during document creation
# =============================================================================

def extract_footnotes_map(soup: BeautifulSoup) -> dict[str, str]:
    """Extract actual footnotes from p.oj-note only.

    This deliberately ignores hr.oj-note separators.
    Returns a mapping {"1": "footnote body", ...}.
    """
    footnotes: dict[str, str] = {}
    for fallback_idx, note in enumerate(soup.select("p.oj-note"), start=1):
        text = clean_text(note.get_text(" ", strip=True))
        match = re.match(r"^\(?\s*(\d+)\s*\)?\s*(.*)$", text)
        if match:
            num, body = match.groups()
        else:
            num, body = str(fallback_idx), text
        footnotes[num] = body.strip()
    return footnotes


def add_footnote_placeholder(paragraph, number: str) -> None:
    """Insert a placeholder run that will later become a native Word footnote."""
    run = paragraph.add_run(f"[[FN:{number}]]")
    set_run_font(run, size=11)
    run.font.superscript = True


def is_footnote_reference_tag(node: Tag) -> str | None:
    """Return footnote number if an XHTML node is a footnote reference."""
    if not isinstance(node, Tag):
        return None
    classes = node.get("class", [])
    if node.name == "sup" or "oj-super" in classes:
        text = clean_text(node.get_text(" ", strip=True))
        match = re.search(r"\(?\s*(\d+)\s*\)?", text)
        if match:
            return match.group(1)
    return None


def add_inline_runs(paragraph, element: Tag, *, default_size=11) -> None:
    """Add XHTML inline content to Word, preserving basic bold/italic/superscript.

    Footnote reference tags (usually span.oj-super) are emitted as placeholders,
    which are converted to native Word footnote references after the .docx is saved.
    """

    def walk(node, bold=False, italic=False, superscript=False):
        if isinstance(node, str):
            if node:
                run = paragraph.add_run(apply_non_breaking_spaces(node))
                set_run_font(run, size=default_size, bold=bold, italic=italic)
                run.font.superscript = superscript
            return

        if not isinstance(node, Tag):
            return

        footnote_number = is_footnote_reference_tag(node)
        if footnote_number:
            add_footnote_placeholder(paragraph, footnote_number)
            return

        node_classes = node.get("class", [])
        next_bold = bold or node.name in {"b", "strong"} or "bold" in node_classes
        next_italic = italic or node.name in {"i", "em"} or "italic" in node_classes
        next_sup = superscript or node.name == "sup" or "oj-super" in node_classes

        for child in node.children:
            walk(child, next_bold, next_italic, next_sup)

    for child in element.children:
        walk(child)


# =============================================================================
# Paragraph / heading builders
# =============================================================================

def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_base_format(p)
    run = p.add_run(clean_text(text))
    set_run_font(run, bold=True)
    return p


def add_centered_heading(doc: Document, first_line: str, second_line: str | None = None):
    """Add centred level-1 style heading using Shift+Return between lines."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_base_format(p)
    run = p.add_run(clean_text(first_line))
    set_run_font(run, bold=True)
    if second_line:
        run.add_break(WD_BREAK.LINE)  # line break, not paragraph break
        run2 = p.add_run(clean_text(second_line))
        set_run_font(run2, bold=True)
    return p


def add_justified_paragraph(doc: Document, text: str, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    run = p.add_run(apply_non_breaking_spaces(clean_text(text)))
    set_run_font(run)
    return p


def add_justified_element_paragraph(doc: Document, element: Tag, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    add_inline_runs(p, element)
    return p


def add_recital(doc: Document, recital_div: Tag):
    rid = recital_div.get("id", "")
    number = rid.replace("rct_", "") if rid.startswith("rct_") else ""
    text = clean_text(recital_div.get_text(" ", strip=True))
    text = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", text).strip() if number else text

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)

    run_num = p.add_run(f"({number})\t" if number else "")
    set_run_font(run_num)
    run_text = p.add_run(apply_non_breaking_spaces(text))
    set_run_font(run_text)
    return p


def paragraph_level(text: str) -> int:
    txt = text.strip()
    if re.match(r"^\d+\.\s+", txt):
        return 1
    if re.match(r"^\([a-z]\)\s+", txt):
        return 2
    if re.match(r"^\([ivxlcdm]+\)\s+", txt, flags=re.I):
        return 3
    return 0


def add_article_text_paragraph(doc: Document, element: Tag):
    text = clean_text(element.get_text(" ", strip=True))
    if not text:
        return None
    level = paragraph_level(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)

    if level == 1:
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
    elif level == 2:
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.first_line_indent = Cm(-1)
    elif level == 3:
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.first_line_indent = Cm(-1)
    else:
        p.paragraph_format.left_indent = Cm(1)

    add_inline_runs(p, element)
    return p


# =============================================================================
# Schema extraction helpers: adoption formula / chapters / articles
# =============================================================================

def add_citations(doc: Document, soup: BeautifulSoup) -> None:
    for citation in soup.select('div.eli-subdivision[id^="cit_"]'):
        add_justified_element_paragraph(doc, citation, left_cm=0)


def add_recitals(doc: Document, soup: BeautifulSoup) -> None:
    for recital in soup.select('div.eli-subdivision[id^="rct_"]'):
        add_recital(doc, recital)


def add_adoption_formula(doc: Document, soup: BeautifulSoup) -> None:
    """Add text between the last recital and Article 1, including adoption formula.

    This captures the missing line such as:
        HAVE ADOPTED THIS REGULATION:
    """
    recitals = soup.select('div.eli-subdivision[id^="rct_"]')
    articles = soup.select('div.eli-subdivision[id^="art_"]')
    if not recitals or not articles:
        return

    last_recital = recitals[-1]
    first_article = articles[0]
    seen: set[int] = set()

    for el in last_recital.next_elements:
        if el is first_article:
            break
        if not isinstance(el, Tag):
            continue
        if id(el) in seen:
            continue
        seen.add(id(el))
        classes = el.get("class", [])
        if "oj-normal" in classes or "oj-final" in classes:
            text = clean_text(el.get_text(" ", strip=True))
            if text:
                # Avoid accidentally repeating the final recital text.
                if text == clean_text(last_recital.get_text(" ", strip=True)):
                    continue
                add_justified_element_paragraph(doc, el, left_cm=0)


def chapter_heading_text(chapter: Tag) -> tuple[str, str | None]:
    """Return (chapter number, chapter title) from a cpt_* subdivision."""
    # Common EUR-Lex pattern: first title-like node contains CHAPTER I, second contains title.
    title_nodes = []
    for selector in [".oj-ti-section-1", ".oj-ti-section-2", ".eli-title", ".oj-doc-ti"]:
        for node in chapter.select(selector):
            txt = clean_text(node.get_text(" ", strip=True))
            if txt and txt not in title_nodes:
                title_nodes.append(txt)
    if not title_nodes:
        cid = chapter.get("id", "cpt")
        return cid.replace("cpt_", "CHAPTER "), None
    if len(title_nodes) == 1:
        return title_nodes[0], None
    return title_nodes[0], title_nodes[1]


def article_heading_parts(article: Tag, soup: BeautifulSoup) -> tuple[str, str | None]:
    art_id = article.get("id", "")
    number = art_id.replace("art_", "") if art_id.startswith("art_") else art_id
    article_number_line = f"Article {number}" if number else "Article"

    title_tag = article.select_one(".oj-sti-art") or soup.find(id=f"{art_id}.tit_1")
    title = clean_text(title_tag.get_text(" ", strip=True)) if title_tag else None

    # If the title tag itself says "Article 1", do not duplicate.
    if title and re.match(r"^Article\s+", title, flags=re.I):
        return title, None
    return article_number_line, title


def add_single_article(doc: Document, article: Tag, soup: BeautifulSoup) -> None:
    first, second = article_heading_parts(article, soup)
    add_centered_heading(doc, first, second)

    for p in article.find_all(["p", "div"], recursive=True):
        classes = p.get("class", [])
        if any(cls in classes for cls in ["oj-ti-art", "oj-sti-art", "eli-title", "oj-ti-section-1", "oj-ti-section-2"]):
            continue
        if "oj-normal" in classes or "oj-enumeration-spacing" in classes:
            add_article_text_paragraph(doc, p)


def add_operatives(doc: Document, soup: BeautifulSoup) -> None:
    """Add chapters and articles in structural order."""
    processed_articles: set[str] = set()
    chapters = soup.select('div.eli-subdivision[id^="cpt_"]')

    for chapter in chapters:
        first, second = chapter_heading_text(chapter)
        add_centered_heading(doc, first, second)
        for article in chapter.select('div.eli-subdivision[id^="art_"]'):
            art_id = article.get("id", "")
            if art_id in processed_articles:
                continue
            add_single_article(doc, article, soup)
            processed_articles.add(art_id)

    # Fallback for any articles not contained in detected chapter containers.
    for article in soup.select('div.eli-subdivision[id^="art_"]'):
        art_id = article.get("id", "")
        if art_id in processed_articles:
            continue
        add_single_article(doc, article, soup)
        processed_articles.add(art_id)


# =============================================================================
# Annex VI table builder
# =============================================================================

def set_cell_font(cell, bold: bool = False, size: int = 11) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def cell_text_from_html(cell: Tag) -> str:
    parts: list[str] = []
    children = cell.find_all(["p", "div"], recursive=False)
    if children:
        for child in children:
            text = clean_text(child.get_text(" ", strip=True))
            if text:
                parts.append(apply_non_breaking_spaces(text))
    else:
        text = clean_text(cell.get_text(" ", strip=True))
        if text:
            parts.append(apply_non_breaking_spaces(text))
    return "\n".join(parts)


def find_largest_html_table(annex: Tag):
    tables = annex.select("table.oj-table") or annex.find_all("table")
    if not tables:
        return None
    return max(tables, key=lambda t: len(t.find_all("tr")) * 10 + len(t.find_all(["td", "th"])))


def build_annex_vi_table(doc: Document, annex: Tag) -> None:
    html_table = find_largest_html_table(annex)
    if html_table is None:
        add_justified_paragraph(doc, annex.get_text(" ", strip=True))
        return

    word_table = doc.add_table(rows=0, cols=2)
    word_table.style = "Table Grid"
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.autofit = True

    for tr in html_table.find_all("tr"):
        html_cells = tr.find_all(["th", "td"], recursive=False)
        if not html_cells:
            continue
        extracted = [cell_text_from_html(cell) for cell in html_cells]
        extracted = [text for text in extracted if text]
        if not extracted:
            continue
        left_text = extracted[0]
        right_text = "\n".join(extracted[1:]) if len(extracted) > 1 else ""

        row = word_table.add_row()
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].text = left_text
        row.cells[1].text = right_text

        is_header_row = any(cell.name == "th" or "oj-tbl-hdr" in cell.get("class", []) for cell in html_cells)
        set_cell_font(row.cells[0], bold=is_header_row)
        set_cell_font(row.cells[1], bold=is_header_row)


def add_annexes(doc: Document, soup: BeautifulSoup) -> None:
    for annex in soup.select('div.eli-container[id^="anx_"]'):
        heading = annex.select_one("p.oj-doc-ti")
        heading_text = clean_text(heading.get_text(" ", strip=True)) if heading else annex.get("id", "Annex")
        add_centered_heading(doc, heading_text)

        if annex.get("id") == "anx_VI":
            build_annex_vi_table(doc, annex)
        else:
            body_text = clean_text(annex.get_text(" ", strip=True))
            if heading_text and body_text.startswith(heading_text):
                body_text = body_text[len(heading_text):].strip()
            add_justified_paragraph(doc, body_text, left_cm=0)


# =============================================================================
# OOXML native footnote patcher
# =============================================================================

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

NSMAP_W = {"w": W_NS}


def w_tag(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def rel_tag(local: str) -> str:
    return f"{{{REL_NS}}}{local}"


def ct_tag(local: str) -> str:
    return f"{{{CT_NS}}}{local}"


def make_text_run(text: str) -> etree._Element:
    r = etree.Element(w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_footnote_reference_run(fid: int) -> etree._Element:
    r = etree.Element(w_tag("r"))
    rpr = etree.SubElement(r, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    ref = etree.SubElement(r, w_tag("footnoteReference"))
    ref.set(w_tag("id"), str(fid))
    return r


def replace_placeholders_in_document_xml(document_xml: bytes, footnote_texts: dict[str, str]) -> tuple[bytes, dict[int, str]]:
    """Replace [[FN:n]] placeholder runs with w:footnoteReference runs."""
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(document_xml, parser)
    next_id = 1
    id_to_text: dict[int, str] = {}
    number_to_id: dict[str, int] = {}

    for t in list(root.xpath(".//w:t", namespaces=NSMAP_W)):
        if not t.text or "[[FN:" not in t.text:
            continue
        run = t.getparent()
        if run is None or run.tag != w_tag("r"):
            continue
        paragraph = run.getparent()
        if paragraph is None:
            continue
        original = t.text
        parts = []
        pos = 0
        for match in FN_PLACEHOLDER_RE.finditer(original):
            if match.start() > pos:
                parts.append(("text", original[pos:match.start()]))
            number = match.group(1)
            if number not in number_to_id:
                number_to_id[number] = next_id
                id_to_text[next_id] = footnote_texts.get(number, "")
                next_id += 1
            parts.append(("fn", number_to_id[number]))
            pos = match.end()
        if pos < len(original):
            parts.append(("text", original[pos:]))

        insert_index = paragraph.index(run)
        paragraph.remove(run)
        for offset, (kind, value) in enumerate(parts):
            new_run = make_text_run(value) if kind == "text" else make_footnote_reference_run(int(value))
            paragraph.insert(insert_index + offset, new_run)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"), id_to_text


def create_footnotes_xml(id_to_text: dict[int, str]) -> bytes:
    root = etree.Element(w_tag("footnotes"), nsmap={"w": W_NS})

    # Required separator footnotes.
    for fid, ftype in [(-1, "separator"), (0, "continuationSeparator")]:
        footnote = etree.SubElement(root, w_tag("footnote"))
        footnote.set(w_tag("id"), str(fid))
        footnote.set(w_tag("type"), ftype)
        p = etree.SubElement(footnote, w_tag("p"))
        r = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r, w_tag("separator" if fid == -1 else "continuationSeparator"))

    for fid, text in sorted(id_to_text.items()):
        footnote = etree.SubElement(root, w_tag("footnote"))
        footnote.set(w_tag("id"), str(fid))
        p = etree.SubElement(footnote, w_tag("p"))
        ppr = etree.SubElement(p, w_tag("pPr"))
        etree.SubElement(ppr, w_tag("pStyle")).set(w_tag("val"), "FootnoteText")
        ind = etree.SubElement(ppr, w_tag("ind"))
        ind.set(w_tag("left"), "567")     # approx 1 cm
        ind.set(w_tag("hanging"), "567")  # approx 1 cm hanging indent
        spacing = etree.SubElement(ppr, w_tag("spacing"))
        spacing.set(w_tag("before"), "0")
        spacing.set(w_tag("after"), "0")
        spacing.set(w_tag("line"), "240")
        spacing.set(w_tag("lineRule"), "auto")

        r_ref = etree.SubElement(p, w_tag("r"))
        rpr = etree.SubElement(r_ref, w_tag("rPr"))
        etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
        etree.SubElement(r_ref, w_tag("footnoteRef"))

        # Tab after footnote number.
        r_tab = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r_tab, w_tag("tab"))

        r_text = etree.SubElement(p, w_tag("r"))
        rpr_text = etree.SubElement(r_text, w_tag("rPr"))
        rfonts = etree.SubElement(rpr_text, w_tag("rFonts"))
        rfonts.set(w_tag("ascii"), "Arial")
        rfonts.set(w_tag("hAnsi"), "Arial")
        sz = etree.SubElement(rpr_text, w_tag("sz"))
        sz.set(w_tag("val"), "18")  # 9 pt
        t = etree.SubElement(r_text, w_tag("t"))
        t.text = apply_non_breaking_spaces(text)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_relationship(rels_xml: bytes) -> bytes:
    root = etree.fromstring(rels_xml)
    existing = root.xpath("./rel:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']", namespaces={"rel": REL_NS})
    if existing:
        return rels_xml
    used_ids = {rel.get("Id") for rel in root.findall(rel_tag("Relationship"))}
    rid = "rIdFootnotes"
    counter = 1
    while rid in used_ids:
        counter += 1
        rid = f"rIdFootnotes{counter}"
    rel = etree.SubElement(root, rel_tag("Relationship"))
    rel.set("Id", rid)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    rel.set("Target", "footnotes.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = etree.fromstring(content_types_xml)
    existing = root.xpath("./ct:Override[@PartName='/word/footnotes.xml']", namespaces={"ct": CT_NS})
    if existing:
        return content_types_xml
    override = etree.SubElement(root, ct_tag("Override"))
    override.set("PartName", "/word/footnotes.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_with_native_footnotes(docx_path: str | Path, footnote_texts: dict[str, str]) -> None:
    """Patch a python-docx-generated file to contain native Word footnotes."""
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)

        document_xml_path = tmp / "word" / "document.xml"
        rels_path = tmp / "word" / "_rels" / "document.xml.rels"
        content_types_path = tmp / "[Content_Types].xml"
        footnotes_path = tmp / "word" / "footnotes.xml"

        new_document_xml, id_to_text = replace_placeholders_in_document_xml(document_xml_path.read_bytes(), footnote_texts)
        document_xml_path.write_bytes(new_document_xml)

        if id_to_text:
            footnotes_path.write_bytes(create_footnotes_xml(id_to_text))
            rels_path.write_bytes(ensure_footnotes_relationship(rels_path.read_bytes()))
            content_types_path.write_bytes(ensure_footnotes_content_type(content_types_path.read_bytes()))

        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


# =============================================================================
# Build document
# =============================================================================

def build_document(html_path: str | Path, output_path: str | Path) -> Path:
    soup = load_html(html_path)
    footnotes = extract_footnotes_map(soup)
    doc = new_document()

    title = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if title:
        add_title(doc, title.get_text(" ", strip=True))

    add_citations(doc, soup)
    add_recitals(doc, soup)
    add_adoption_formula(doc, soup)
    add_operatives(doc, soup)
    add_annexes(doc, soup)

    output_path = Path(output_path)
    doc.save(output_path)
    patch_docx_with_native_footnotes(output_path, footnotes)
    return output_path


# =============================================================================
# CLI
# =============================================================================

def main() -> int:
    parser = argparse.ArgumentParser(description="Stage 2 EUR-Lex XHTML to formatted Word document with native OOXML footnotes.")
    parser.add_argument("source", help="Path to EUR-Lex XHTML/HTML input file")
    parser.add_argument("output", help="Path to output .docx file")
    args = parser.parse_args()

    output = build_document(args.source, args.output)
    print(f"Stage 2 document saved: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
