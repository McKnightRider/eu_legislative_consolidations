"""
formatting_stage2.py

Stage 2 formatting engine for EUR-Lex / Official Journal XHTML.

Corrected version addressing:
    1. Excess Shift+Returns caused by overly broad chapter/article heading extraction.
    2. Recital 89 being inserted twice by over-broad adoption-formula traversal.
    3. Missing adoption formula: "HAVE ADOPTED THIS REGULATION:".
    4. Missing Chapter / Article level-1 numbers and titles.
    5. Native Word footnotes without Aspose.Words, using python-docx plus direct OOXML patching.
    6. Annex VI conversion into a two-column Word table from the source XHTML table.

Open-source dependencies only:
    beautifulsoup4
    lxml
    python-docx

Run:
    python formatting_stage2.py L_2017168EN.01001201.xml.html Prospectus_Regulation_Stage2.docx
"""

from __future__ import annotations

from pathlib import Path
import argparse
import copy
import re
import shutil
import tempfile
import zipfile

from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


# =============================================================================
# Constants
# =============================================================================

NBSP = "\u00A0"
FN_PLACEHOLDER_RE = re.compile(r"\[\[FN:(\d+)\]\]")
STRUCTURAL_ID_RE = re.compile(r"^(cit|rct|cpt|art|anx)_")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
NSMAP_W = {"w": W_NS}


# =============================================================================
# General helpers
# =============================================================================

def clean_text(text: str | None) -> str:
    """Normalise whitespace without changing legal substance."""
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    return " ".join(text.split())


def load_html(path: str | Path) -> BeautifulSoup:
    return BeautifulSoup(Path(path).read_text(encoding="utf-8", errors="replace"), "lxml")


def apply_non_breaking_spaces(text: str) -> str:
    """Apply non-breaking-space rules for thousands separators, currency and percentages."""
    if not text:
        return ""
    # Thousands separator: 200 000 -> 200 NBSP 000; works repeatedly in long numbers.
    text = re.sub(r"(?<=\d) (?=\d{3}\b)", NBSP, text)
    # Currency spacing.
    text = re.sub(r"\b(EUR)\s+(?=\d)", r"\1" + NBSP, text)
    text = re.sub(r"€\s+(?=\d)", "€" + NBSP, text)
    # Percentages.
    text = re.sub(r"(?<=\d)\s+%", NBSP + "%", text)
    return text


def set_run_font(run, size: int = 11, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_base_format(paragraph, *, before=6, after=6, line_spacing=1.16) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def new_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def tag_classes(tag: Tag | None) -> list[str]:
    if not isinstance(tag, Tag):
        return []
    classes = tag.get("class", [])
    if isinstance(classes, str):
        return classes.split()
    return list(classes)


def has_class(tag: Tag | None, class_name: str) -> bool:
    return class_name in tag_classes(tag)


def element_without_child_classes(tag: Tag, classes_to_remove: set[str]) -> str:
    """Return text from a copied node after removing descendants with specified classes."""
    soup_copy = BeautifulSoup(str(tag), "lxml")
    for removable in soup_copy.find_all(lambda t: isinstance(t, Tag) and set(tag_classes(t)).intersection(classes_to_remove)):
        removable.decompose()
    return clean_text(soup_copy.get_text(" ", strip=True))


# =============================================================================
# Structural scoping helpers
# =============================================================================

def nearest_structural_ancestor_id(tag: Tag | None) -> str | None:
    """Return nearest ancestor-or-self ID matching a EUR-Lex structural container.

    This prevents chapter extraction from accidentally collecting Article headings
    nested inside a chapter, and prevents article text extraction from walking into
    unrelated nested containers.
    """
    current = tag
    while isinstance(current, Tag):
        cid = current.get("id")
        if cid and STRUCTURAL_ID_RE.match(cid):
            return cid
        current = current.parent
    return None


def belongs_to_container(tag: Tag, container: Tag) -> bool:
    return nearest_structural_ancestor_id(tag) == container.get("id")


def structural_id_sort_key(tag: Tag) -> tuple[int, str]:
    """Sort helper preserving document order where BeautifulSoup select already does.

    Kept for clarity if later sorting is needed; currently not used for reordering.
    """
    return (0, tag.get("id", ""))


# =============================================================================
# Footnotes: extraction and placeholders
# =============================================================================

def extract_footnotes_map(soup: BeautifulSoup) -> dict[str, str]:
    """Extract actual footnotes from p.oj-note only.

    The diagnostic files showed an hr.oj-note separator plus actual p.oj-note
    footnotes.  Using p.oj-note avoids the earlier off-by-one error.
    """
    footnotes: dict[str, str] = {}
    for fallback_idx, note in enumerate(soup.select("p.oj-note"), start=1):
        marker_tag = note.select_one(".oj-note-tag")
        if marker_tag:
            marker_text = clean_text(marker_tag.get_text(" ", strip=True))
            marker_match = re.search(r"\(?\s*(\d+)\s*\)?", marker_text)
            number = marker_match.group(1) if marker_match else str(fallback_idx)
            body = element_without_child_classes(note, {"oj-note-tag"})
        else:
            full_text = clean_text(note.get_text(" ", strip=True))
            match = re.match(r"^\(?\s*(\d+)\s*\)?\s*(.*)$", full_text)
            if match:
                number, body = match.groups()
            else:
                number, body = str(fallback_idx), full_text
        footnotes[number] = body.strip()
    return footnotes


def is_footnote_reference_tag(node: Tag) -> str | None:
    """Return the footnote number if this node is a footnote reference."""
    if not isinstance(node, Tag):
        return None
    classes = tag_classes(node)
    if node.name == "sup" or "oj-super" in classes:
        text = clean_text(node.get_text(" ", strip=True))
        match = re.search(r"\(?\s*(\d+)\s*\)?", text)
        if match:
            return match.group(1)
    return None


def add_footnote_placeholder(paragraph, number: str) -> None:
    """Insert a placeholder that will later be replaced by native OOXML footnote reference."""
    run = paragraph.add_run(f"[[FN:{number}]]")
    set_run_font(run, size=11)
    run.font.superscript = True


# =============================================================================
# Inline formatting
# =============================================================================

def add_inline_runs(paragraph, element: Tag, *, default_size=11) -> None:
    """Add XHTML inline content to a Word paragraph.

    Preserves basic bold, italic and superscript.  Footnote reference tags become
    placeholders and are converted to native Word footnotes after the document is
    saved.
    """

    def walk(node, bold=False, italic=False, superscript=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(apply_non_breaking_spaces(text))
                set_run_font(run, size=default_size, bold=bold, italic=italic)
                run.font.superscript = superscript
            return

        if not isinstance(node, Tag):
            return

        footnote_number = is_footnote_reference_tag(node)
        if footnote_number:
            add_footnote_placeholder(paragraph, footnote_number)
            return

        classes = tag_classes(node)
        next_bold = bold or node.name in {"b", "strong"} or "bold" in classes
        next_italic = italic or node.name in {"i", "em"} or "italic" in classes
        next_sup = superscript or node.name == "sup" or "oj-super" in classes

        for child in node.children:
            walk(child, next_bold, next_italic, next_sup)

    for child in element.children:
        walk(child)


# =============================================================================
# Paragraph and heading builders
# =============================================================================

def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_paragraph_base_format(p)
    run = p.add_run(clean_text(text))
    set_run_font(run, bold=True)
    return p


def add_centered_heading(doc: Document, first_line: str, second_line: str | None = None):
    """Add a centred heading.

    A Shift+Return is inserted only where we explicitly pass a second line, i.e.
    for Chapter/Article number + title.  This avoids the previous broad-title bug
    that inserted line breaks in many unexpected places.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_paragraph_base_format(p)

    run1 = p.add_run(clean_text(first_line))
    set_run_font(run1, bold=True)

    if second_line:
        run1.add_break(WD_BREAK.LINE)
        run2 = p.add_run(clean_text(second_line))
        set_run_font(run2, bold=True)
    return p


def add_justified_paragraph(doc: Document, text: str, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    set_paragraph_base_format(p)
    run = p.add_run(apply_non_breaking_spaces(clean_text(text)))
    set_run_font(run)
    return p


def add_justified_element_paragraph(doc: Document, element: Tag, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    set_paragraph_base_format(p)
    add_inline_runs(p, element)
    return p


def add_recital(doc: Document, recital_div: Tag):
    rid = recital_div.get("id", "")
    number = rid.replace("rct_", "") if rid.startswith("rct_") else ""

    # Recital text is often duplicated within nested table/p elements. Prefer the
    # most specific p.oj-normal belonging to this recital; fall back to container text.
    normal = None
    for candidate in recital_div.select("p.oj-normal"):
        if belongs_to_container(candidate, recital_div):
            normal = candidate
            break

    raw_text = clean_text(normal.get_text(" ", strip=True)) if normal else clean_text(recital_div.get_text(" ", strip=True))
    raw_text = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", raw_text).strip() if number else raw_text

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    set_paragraph_base_format(p)

    run_num = p.add_run(f"({number})\t" if number else "")
    set_run_font(run_num)

    if normal is not None:
        add_inline_runs(p, normal)
    else:
        run_text = p.add_run(apply_non_breaking_spaces(raw_text))
        set_run_font(run_text)
    return p


def paragraph_level(text: str) -> int:
    txt = text.strip()
    if re.match(r"^\d+\.\s+", txt):
        return 1
    if re.match(r"^\([a-z]\)\s+", txt):
        return 2
    if re.match(r"^\([ivxlcdm]+\)\s+", txt, flags=re.I):
        return 3
    return 0


def add_article_text_paragraph(doc: Document, element: Tag):
    text = clean_text(element.get_text(" ", strip=True))
    if not text:
        return None

    level = paragraph_level(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)

    if level == 1:
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
    elif level == 2:
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.first_line_indent = Cm(-1)
    elif level == 3:
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.first_line_indent = Cm(-1)
    else:
        p.paragraph_format.left_indent = Cm(1)

    add_inline_runs(p, element)
    return p


# =============================================================================
# EUR-Lex schema extraction
# =============================================================================

def add_citations(doc: Document, soup: BeautifulSoup) -> None:
    for citation in soup.select('div.eli-subdivision[id^="cit_"]'):
        # Prefer specific p.oj-normal within the citation to avoid container duplication.
        specific = citation.select_one("p.oj-normal")
        add_justified_element_paragraph(doc, specific or citation, left_cm=0)


def add_recitals(doc: Document, soup: BeautifulSoup) -> None:
    for recital in soup.select('div.eli-subdivision[id^="rct_"]'):
        add_recital(doc, recital)


def add_adoption_formula(doc: Document, soup: BeautifulSoup) -> None:
    """Add the adoption formula only once.

    The previous implementation walked from the last recital through next_elements,
    which traversed descendants of recital 89 and caused recital 89 to be duplicated.
    This version directly searches for the explicit formula paragraph and inserts
    only that paragraph.
    """
    formula_re = re.compile(r"\bHAVE\s+ADOPTED\s+THIS\s+REGULATION\s*:", re.I)
    for tag in soup.find_all(["p", "div"]):
        text = clean_text(tag.get_text(" ", strip=True))
        if not text or not formula_re.search(text):
            continue
        # Prefer a leaf-ish paragraph. Avoid enormous containers whose text merely
        # contains the formula together with the whole act.
        if len(text) > 200:
            continue
        add_justified_element_paragraph(doc, tag, left_cm=0)
        return


def title_nodes_belonging_to(container: Tag, selectors: list[str]) -> list[Tag]:
    """Find heading/title nodes scoped to a structural container only.

    This avoids selecting Article headings while extracting Chapter headings because
    Article nodes are nested below Chapter containers.
    """
    nodes: list[Tag] = []
    for selector in selectors:
        for node in container.select(selector):
            if belongs_to_container(node, container):
                text = clean_text(node.get_text(" ", strip=True))
                if text and node not in nodes:
                    nodes.append(node)
    return nodes


def chapter_heading_text(chapter: Tag) -> tuple[str, str | None]:
    cid = chapter.get("id", "")

    # Prefer explicit title IDs if present, e.g. cpt_I.tit_1, cpt_I.tit_2.
    explicit_titles = []
    for i in range(1, 5):
        node = chapter.find(id=f"{cid}.tit_{i}")
        if isinstance(node, Tag):
            txt = clean_text(node.get_text(" ", strip=True))
            if txt:
                explicit_titles.append(txt)
    if explicit_titles:
        return explicit_titles[0], explicit_titles[1] if len(explicit_titles) > 1 else None

    nodes = title_nodes_belonging_to(chapter, [".oj-ti-section-1", ".oj-ti-section-2", ".eli-title", ".oj-doc-ti"])
    texts = []
    for node in nodes:
        txt = clean_text(node.get_text(" ", strip=True))
        # Exclude Article headings defensively.
        if re.match(r"^Article\s+\d+", txt, flags=re.I):
            continue
        if txt and txt not in texts:
            texts.append(txt)
    if not texts:
        label = cid.replace("cpt_", "CHAPTER ") if cid else "CHAPTER"
        return label, None
    return texts[0], texts[1] if len(texts) > 1 else None


def article_heading_parts(article: Tag) -> tuple[str, str | None]:
    art_id = article.get("id", "")
    number = art_id.replace("art_", "") if art_id.startswith("art_") else art_id

    article_number = None
    article_title = None

    # Prefer the actual article heading/subheading nodes scoped to this article.
    ti_node = None
    sti_node = None
    for node in article.select(".oj-ti-art"):
        if belongs_to_container(node, article):
            ti_node = node
            break
    for node in article.select(".oj-sti-art"):
        if belongs_to_container(node, article):
            sti_node = node
            break

    if ti_node:
        article_number = clean_text(ti_node.get_text(" ", strip=True))
    else:
        article_number = f"Article {number}" if number else "Article"

    if sti_node:
        article_title = clean_text(sti_node.get_text(" ", strip=True))
    else:
        # Fallback to explicit title ID, but only if it is not the Article number.
        title_node = article.find(id=f"{art_id}.tit_1")
        if isinstance(title_node, Tag):
            candidate = clean_text(title_node.get_text(" ", strip=True))
            if candidate and not re.match(r"^Article\s+\d+", candidate, flags=re.I):
                article_title = candidate

    # Avoid duplicated headings if article_title repeats Article number.
    if article_title and article_title == article_number:
        article_title = None
    return article_number, article_title


def article_content_paragraphs(article: Tag) -> list[Tag]:
    """Return non-heading article content paragraphs, scoped to the article.

    The earlier file used find_all(["p", "div"], recursive=True), which picked up
    containers and caused duplicate/odd output.  This version collects only p tags
    carrying content classes and belonging to the current article container.
    """
    out: list[Tag] = []
    for p in article.find_all("p"):
        if not belongs_to_container(p, article):
            continue
        classes = tag_classes(p)
        if any(cls in classes for cls in ["oj-ti-art", "oj-sti-art", "eli-title", "oj-ti-section-1", "oj-ti-section-2"]):
            continue
        if "oj-note" in classes:
            continue
        if "oj-normal" in classes or "oj-enumeration-spacing" in classes:
            text = clean_text(p.get_text(" ", strip=True))
            if text:
                out.append(p)
    return out


def add_single_article(doc: Document, article: Tag) -> None:
    first, second = article_heading_parts(article)
    add_centered_heading(doc, first, second)
    for p in article_content_paragraphs(article):
        add_article_text_paragraph(doc, p)


def add_operatives(doc: Document, soup: BeautifulSoup) -> None:
    """Add chapters and articles in structural order."""
    processed_articles: set[str] = set()

    for chapter in soup.select('div.eli-subdivision[id^="cpt_"]'):
        first, second = chapter_heading_text(chapter)
        add_centered_heading(doc, first, second)

        for article in chapter.select('div.eli-subdivision[id^="art_"]'):
            art_id = article.get("id", "")
            if art_id in processed_articles:
                continue
            add_single_article(doc, article)
            processed_articles.add(art_id)

    # Fallback: if any article is outside detected chapter containers, include it once.
    for article in soup.select('div.eli-subdivision[id^="art_"]'):
        art_id = article.get("id", "")
        if art_id in processed_articles:
            continue
        add_single_article(doc, article)
        processed_articles.add(art_id)


# =============================================================================
# Annex VI table builder
# =============================================================================

def set_cell_font(cell, bold: bool = False, size: int = 11) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def cell_text_from_html(cell: Tag) -> str:
    parts: list[str] = []
    children = cell.find_all(["p", "div"], recursive=False)
    if children:
        for child in children:
            text = clean_text(child.get_text(" ", strip=True))
            if text:
                parts.append(apply_non_breaking_spaces(text))
    else:
        text = clean_text(cell.get_text(" ", strip=True))
        if text:
            parts.append(apply_non_breaking_spaces(text))
    return "\n".join(parts)


def find_largest_html_table(annex: Tag):
    tables = annex.select("table.oj-table") or annex.find_all("table")
    if not tables:
        return None
    return max(tables, key=lambda t: len(t.find_all("tr")) * 10 + len(t.find_all(["td", "th"])))


def build_annex_vi_table(doc: Document, annex: Tag) -> None:
    html_table = find_largest_html_table(annex)
    if html_table is None:
        add_justified_paragraph(doc, annex.get_text(" ", strip=True))
        return

    word_table = doc.add_table(rows=0, cols=2)
    word_table.style = "Table Grid"
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.autofit = True

    for tr in html_table.find_all("tr"):
        html_cells = tr.find_all(["th", "td"], recursive=False)
        if not html_cells:
            continue
        extracted = [cell_text_from_html(cell) for cell in html_cells]
        extracted = [text for text in extracted if text]
        if not extracted:
            continue

        left_text = extracted[0]
        right_text = "\n".join(extracted[1:]) if len(extracted) > 1 else ""

        row = word_table.add_row()
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].text = left_text
        row.cells[1].text = right_text

        is_header_row = any(cell.name == "th" or "oj-tbl-hdr" in tag_classes(cell) for cell in html_cells)
        set_cell_font(row.cells[0], bold=is_header_row)
        set_cell_font(row.cells[1], bold=is_header_row)


def add_annexes(doc: Document, soup: BeautifulSoup) -> None:
    for annex in soup.select('div.eli-container[id^="anx_"]'):
        heading = annex.select_one("p.oj-doc-ti")
        heading_text = clean_text(heading.get_text(" ", strip=True)) if heading else annex.get("id", "Annex")
        add_centered_heading(doc, heading_text)

        if annex.get("id") == "anx_VI":
            build_annex_vi_table(doc, annex)
        else:
            body_text = clean_text(annex.get_text(" ", strip=True))
            if heading_text and body_text.startswith(heading_text):
                body_text = body_text[len(heading_text):].strip()
            add_justified_paragraph(doc, body_text, left_cm=0)


# =============================================================================
# OOXML native footnote patcher
# =============================================================================

def w_tag(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def rel_tag(local: str) -> str:
    return f"{{{REL_NS}}}{local}"


def ct_tag(local: str) -> str:
    return f"{{{CT_NS}}}{local}"


def make_text_run(text: str) -> etree._Element:
    r = etree.Element(w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_footnote_reference_run(fid: int) -> etree._Element:
    r = etree.Element(w_tag("r"))
    rpr = etree.SubElement(r, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    ref = etree.SubElement(r, w_tag("footnoteReference"))
    ref.set(w_tag("id"), str(fid))
    return r


def replace_placeholders_in_document_xml(document_xml: bytes, footnote_texts: dict[str, str]) -> tuple[bytes, dict[int, str]]:
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(document_xml, parser)
    next_id = 1
    id_to_text: dict[int, str] = {}
    number_to_id: dict[str, int] = {}

    for t in list(root.xpath(".//w:t", namespaces={"w": W_NS})):
        if not t.text or "[[FN:" not in t.text:
            continue
        run = t.getparent()
        if run is None or run.tag != w_tag("r"):
            continue
        paragraph = run.getparent()
        if paragraph is None:
            continue

        original = t.text
        parts: list[tuple[str, str | int]] = []
        pos = 0
        for match in FN_PLACEHOLDER_RE.finditer(original):
            if match.start() > pos:
                parts.append(("text", original[pos:match.start()]))
            number = match.group(1)
            if number not in number_to_id:
                number_to_id[number] = next_id
                id_to_text[next_id] = footnote_texts.get(number, "")
                next_id += 1
            parts.append(("fn", number_to_id[number]))
            pos = match.end()
        if pos < len(original):
            parts.append(("text", original[pos:]))

        insert_index = paragraph.index(run)
        paragraph.remove(run)
        for offset, (kind, value) in enumerate(parts):
            if kind == "text":
                new_run = make_text_run(str(value))
            else:
                new_run = make_footnote_reference_run(int(value))
            paragraph.insert(insert_index + offset, new_run)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"), id_to_text


def create_footnotes_xml(id_to_text: dict[int, str]) -> bytes:
    root = etree.Element(w_tag("footnotes"), nsmap={"w": W_NS})

    # Required separator footnotes.
    for fid, ftype, marker in [(-1, "separator", "separator"), (0, "continuationSeparator", "continuationSeparator")]:
        footnote = etree.SubElement(root, w_tag("footnote"))
        footnote.set(w_tag("id"), str(fid))
        footnote.set(w_tag("type"), ftype)
        p = etree.SubElement(footnote, w_tag("p"))
        r = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r, w_tag(marker))

    for fid, text in sorted(id_to_text.items()):
        footnote = etree.SubElement(root, w_tag("footnote"))
        footnote.set(w_tag("id"), str(fid))
        p = etree.SubElement(footnote, w_tag("p"))
        ppr = etree.SubElement(p, w_tag("pPr"))
        etree.SubElement(ppr, w_tag("pStyle")).set(w_tag("val"), "FootnoteText")
        ind = etree.SubElement(ppr, w_tag("ind"))
        ind.set(w_tag("left"), "567")
        ind.set(w_tag("hanging"), "567")
        spacing = etree.SubElement(ppr, w_tag("spacing"))
        spacing.set(w_tag("before"), "0")
        spacing.set(w_tag("after"), "0")
        spacing.set(w_tag("line"), "240")
        spacing.set(w_tag("lineRule"), "auto")

        r_ref = etree.SubElement(p, w_tag("r"))
        rpr = etree.SubElement(r_ref, w_tag("rPr"))
        etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
        etree.SubElement(r_ref, w_tag("footnoteRef"))

        r_tab = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r_tab, w_tag("tab"))

        r_text = etree.SubElement(p, w_tag("r"))
        rpr_text = etree.SubElement(r_text, w_tag("rPr"))
        rfonts = etree.SubElement(rpr_text, w_tag("rFonts"))
        rfonts.set(w_tag("ascii"), "Arial")
        rfonts.set(w_tag("hAnsi"), "Arial")
        sz = etree.SubElement(rpr_text, w_tag("sz"))
        sz.set(w_tag("val"), "18")  # 9pt
        t = etree.SubElement(r_text, w_tag("t"))
        t.text = apply_non_breaking_spaces(text)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_relationship(rels_xml: bytes) -> bytes:
    root = etree.fromstring(rels_xml)
    existing = root.xpath(
        "./rel:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']",
        namespaces={"rel": REL_NS},
    )
    if existing:
        return rels_xml
    used_ids = {rel.get("Id") for rel in root.findall(rel_tag("Relationship"))}
    rid = "rIdFootnotes"
    counter = 1
    while rid in used_ids:
        counter += 1
        rid = f"rIdFootnotes{counter}"
    rel = etree.SubElement(root, rel_tag("Relationship"))
    rel.set("Id", rid)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    rel.set("Target", "footnotes.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = etree.fromstring(content_types_xml)
    existing = root.xpath("./ct:Override[@PartName='/word/footnotes.xml']", namespaces={"ct": CT_NS})
    if existing:
        return content_types_xml
    override = etree.SubElement(root, ct_tag("Override"))
    override.set("PartName", "/word/footnotes.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_with_native_footnotes(docx_path: str | Path, footnote_texts: dict[str, str]) -> None:
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)

        document_xml_path = tmp / "word" / "document.xml"
        rels_path = tmp / "word" / "_rels" / "document.xml.rels"
        content_types_path = tmp / "[Content_Types].xml"
        footnotes_path = tmp / "word" / "footnotes.xml"

        new_document_xml, id_to_text = replace_placeholders_in_document_xml(document_xml_path.read_bytes(), footnote_texts)
        document_xml_path.write_bytes(new_document_xml)

        if id_to_text:
            footnotes_path.write_bytes(create_footnotes_xml(id_to_text))
            rels_path.write_bytes(ensure_footnotes_relationship(rels_path.read_bytes()))
            content_types_path.write_bytes(ensure_footnotes_content_type(content_types_path.read_bytes()))

        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


# =============================================================================
# Build document
# =============================================================================

def build_document(html_path: str | Path, output_path: str | Path) -> Path:
    soup = load_html(html_path)
    footnotes = extract_footnotes_map(soup)
    doc = new_document()

    title = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if title:
        add_title(doc, title.get_text(" ", strip=True))

    add_citations(doc, soup)
    add_recitals(doc, soup)
    add_adoption_formula(doc, soup)
    add_operatives(doc, soup)
    add_annexes(doc, soup)

    output_path = Path(output_path)
    doc.save(output_path)
    patch_docx_with_native_footnotes(output_path, footnotes)
    return output_path


# =============================================================================
# CLI
# =============================================================================

def main() -> int:
    parser = argparse.ArgumentParser(description="Stage 2 EUR-Lex XHTML to formatted Word document with native OOXML footnotes.")
    parser.add_argument("source", help="Path to EUR-Lex XHTML/HTML input file")
    parser.add_argument("output", help="Path to output .docx file")
    args = parser.parse_args()

    output = build_document(args.source, args.output)
    print(f"Stage 2 document saved: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
