
"""
formatting_stage2_v2.py

EUR-Lex / Official Journal XHTML -> formatted Word document.

Updated after review of Prospectus_Regulation_Stage2.docx.

Fixes in this version:
    - Footnote references are native Word footnotes, with the reference run set to
      Arial 9 pt superscript; footnote numbers in the footnote text are also 9 pt
      superscript.
    - Source bracketed references such as "(1)" are converted to bare superscript
      Word footnote references, with no space before the reference and no space
      between the reference and following punctuation.
    - Preamble citations and the adoption formula are no longer italicised.
    - Article containers and chapter containers are matched by exact IDs only, so
      metadata nodes such as art_31.tit_1 / cpt_I.tit_1 are not emitted as content.
    - Article 2-style introductory text is not indented.
    - Annex headings have page_break_before applied.
    - Lettered Annex items such as A. / B. are compacted to A.Offer statistics,
      i.e. no space after the letter/full stop.

Dependencies:
    beautifulsoup4
    lxml
    python-docx

Run:
    python formatting_stage2_v2.py L_2017168EN.01001201.xml.html Prospectus_Regulation_Stage2.docx
"""

from __future__ import annotations

from pathlib import Path
import argparse
import re
import shutil
import tempfile
import zipfile

from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


NBSP = "\u00A0"
FN_TOKEN_RE = re.compile(r"\[\[FN:(\d+)\]\]")
STRUCTURAL_ID_RE = re.compile(
    r"^(cit_\d+|rct_\d+|cpt_[IVXLCDM]+|art_\d+[A-Za-z]?|anx_[IVXLCDM]+)$"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def clean_text(text: str | None) -> str:
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    return " ".join(text.split()).strip()


def load_html(path: str | Path) -> BeautifulSoup:
    return BeautifulSoup(Path(path).read_text(encoding="utf-8", errors="replace"), "lxml")


def tag_classes(tag: Tag | None) -> list[str]:
    if not isinstance(tag, Tag):
        return []
    classes = tag.get("class", [])
    if isinstance(classes, str):
        return classes.split()
    return list(classes)


def has_class(tag: Tag | None, class_name: str) -> bool:
    return class_name in tag_classes(tag)


def apply_nbsp_rules(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"(?<=\d) (?=\d{3}\b)", NBSP, text)
    text = re.sub(r"\b(EUR)\s+(?=\d)", r"\1" + NBSP, text)
    text = re.sub(r"€\s+(?=\d)", "€" + NBSP, text)
    text = re.sub(r"(?<=\d)\s+%", NBSP + "%", text)
    return text


def set_run_font(run, *, size: int = 11, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, *, before=6, after=6, line_spacing=1.16) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def new_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def nearest_structural_id(tag: Tag | None) -> str | None:
    cur = tag
    while isinstance(cur, Tag):
        sid = str(cur.get("id", ""))
        if sid and STRUCTURAL_ID_RE.match(sid):
            return sid
        cur = cur.parent
    return None


def is_within_structural_container(tag: Tag, container_id: str) -> bool:
    return nearest_structural_id(tag) == container_id


def has_ancestor_with_id_prefix(tag: Tag, prefix: str) -> bool:
    cur = tag.parent
    while isinstance(cur, Tag):
        if str(cur.get("id", "")).startswith(prefix):
            return True
        cur = cur.parent
    return False


# =============================================================================
# Footnote tokenisation
# =============================================================================

def footnote_number_from_reference(node: Tag) -> str | None:
    if not isinstance(node, Tag):
        return None
    classes = tag_classes(node)
    if node.name == "sup" or "oj-super" in classes:
        txt = clean_text(node.get_text(" ", strip=True))
        m = re.search(r"\(?\s*(\d+)\s*\)?", txt)
        if m:
            return m.group(1)
    return None


def text_with_footnote_tokens(element: Tag) -> str:
    """Extract text and convert footnote refs to [[FN:n]] tokens.

    The source often has a visual reference such as "(1)".  This function removes
    the brackets and stores only a token, so that the final Word output shows only
    the native superscript footnote reference.
    """
    parts: list[str] = []

    def walk(node):
        if isinstance(node, NavigableString):
            txt = clean_text(str(node))
            if txt:
                parts.append(txt)
            return
        if not isinstance(node, Tag):
            return
        if "oj-note" in tag_classes(node):
            return
        fn = footnote_number_from_reference(node)
        if fn:
            parts.append(f"[[FN:{fn}]]")
            return
        for child in node.children:
            walk(child)

    walk(element)
    text = clean_text(" ".join(parts))

    # Normalise bracketed source references.
    text = re.sub(r"\s*\(\s*\[\[FN:(\d+)\]\]\s*\)", r"[[FN:\1]]", text)
    text = re.sub(r"\s*\(\s*(\[\[FN:\d+\]\])\s*\)", r"\1", text)

    # No space before a footnote reference.
    text = re.sub(r"\s+(\[\[FN:\d+\]\])", r"\1", text)

    # No space between a footnote reference and following punctuation, but preserve
    # ordinary word spacing after the reference.
    text = re.sub(r"(\[\[FN:\d+\]\])\s+([,.;:])", r"\1\2", text)

    return apply_nbsp_rules(text)


def add_text_with_footnotes(paragraph, text: str, *, size=11, bold=False, italic=False) -> None:
    if not text:
        return
    pos = 0
    for m in FN_TOKEN_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_run_font(run, size=size, bold=bold, italic=italic)
        run = paragraph.add_run(f"[[FN:{m.group(1)}]]")
        set_run_font(run, size=9)
        run.font.superscript = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=bold, italic=italic)


# =============================================================================
# Word paragraph helpers
# =============================================================================

def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_para_spacing(p)
    add_text_with_footnotes(p, clean_text(text), bold=True)
    return p


def add_center_heading(doc: Document, first_line: str, second_line: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_para_spacing(p)
    r1 = p.add_run(clean_text(first_line))
    set_run_font(r1, bold=True)
    if second_line:
        # Deliberate Word line break: only for Chapter/Article heading number + title.
        r1.add_break(WD_BREAK.LINE)
        r2 = p.add_run(clean_text(second_line))
        set_run_font(r2, bold=True)
    return p


def add_para(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    left_cm=0,
    first_line_cm=None,
    size=11,
    bold=False,
    italic=False,
    before=6,
    after=6,
    line_spacing=1.16,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    set_para_spacing(p, before=before, after=after, line_spacing=line_spacing)
    add_text_with_footnotes(p, text, size=size, bold=bold, italic=italic)
    return p


# =============================================================================
# Footnote bodies
# =============================================================================

def element_text_without_classes(tag: Tag, classes_to_remove: set[str]) -> str:
    tmp = BeautifulSoup(str(tag), "lxml")
    for removable in tmp.find_all(
        lambda t: isinstance(t, Tag) and set(tag_classes(t)).intersection(classes_to_remove)
    ):
        removable.decompose()
    return clean_text(tmp.get_text(" ", strip=True))


def extract_footnotes_map(soup: BeautifulSoup) -> dict[str, str]:
    notes: dict[str, str] = {}
    for fallback, note in enumerate(soup.select("p.oj-note"), start=1):
        marker = note.select_one(".oj-note-tag")
        if marker:
            marker_text = clean_text(marker.get_text(" ", strip=True))
            m = re.search(r"\(?\s*(\d+)\s*\)?", marker_text)
            number = m.group(1) if m else str(fallback)
            body = element_text_without_classes(note, {"oj-note-tag"})
        else:
            whole = clean_text(note.get_text(" ", strip=True))
            m = re.match(r"^\(?\s*(\d+)\s*\)?\s*(.*)$", whole)
            if m:
                number, body = m.groups()
            else:
                number, body = str(fallback), whole
        notes[number] = apply_nbsp_rules(body.strip())
    return notes


# =============================================================================
# Source paragraph selection and numbering
# =============================================================================

def provision_paragraphs(container: Tag, classes: tuple[str, ...] = ("oj-normal", "oj-enumeration-spacing")) -> list[Tag]:
    cid = str(container.get("id", ""))
    out: list[Tag] = []
    for p in container.find_all("p"):
        pid = str(p.get("id", ""))
        if ".tit_" in pid:
            continue
        if cid and not is_within_structural_container(p, cid):
            continue
        p_classes = tag_classes(p)
        if "oj-note" in p_classes:
            continue
        if any(c in p_classes for c in classes):
            txt = text_with_footnote_tokens(p)
            if txt:
                out.append(p)
    return out


def is_standalone_marker(text: str) -> bool:
    t = clean_text(text)
    return bool(re.match(r"^(\([a-z]\)|\([ivxlcdm]+\)|[A-Z]\.\s*|[IVXLCDM]+\.|—|–)$", t, flags=re.I))


def paragraph_level(text: str) -> int:
    t = clean_text(text)
    if re.match(r"^\d+\.\s+", t) or re.match(r"^\d+\.\t", t):
        return 1
    if re.match(r"^\([a-z]\)\s+", t) or re.match(r"^\([a-z]\)\t", t):
        return 2
    if re.match(r"^\([ivxlcdm]+\)\s+", t, flags=re.I) or re.match(r"^\([ivxlcdm]+\)\t", t, flags=re.I):
        return 3
    if t.startswith("—") or t.startswith("–"):
        return 3
    if re.match(r"^[A-Z]\.\s+", t) or re.match(r"^[A-Z]\.\t", t) or re.match(r"^[IVXLCDM]+\.\s+", t) or re.match(r"^[IVXLCDM]+\.\t", t):
        return 1
    return 0


def normalise_numbered_text(text: str, *, compact_lettered: bool = False) -> str:
    letter_repl = r"\1" if compact_lettered else r"\1\t"
    patterns = [
        (r"^(\d+\.)\s+", r"\1\t"),
        (r"^(\([a-z]\))\s+", r"\1\t"),
        (r"^(\([ivxlcdm]+\))\s+", r"\1\t"),
        (r"^([A-Z]\.)\s+", letter_repl),
        (r"^([IVXLCDM]+\.)\s+", letter_repl),
    ]
    out = text
    for pat, repl in patterns:
        out = re.sub(pat, repl, out, flags=re.I)
    return out


def emit_numbered_paragraph_sequence(doc: Document, paragraphs: list[Tag], *, compact_lettered: bool = False) -> None:
    """Emit paragraphs, merging standalone markers with the following body."""
    i = 0
    while i < len(paragraphs):
        txt = text_with_footnote_tokens(paragraphs[i])
        if not txt:
            i += 1
            continue

        if is_standalone_marker(txt) and i + 1 < len(paragraphs):
            next_txt = text_with_footnote_tokens(paragraphs[i + 1])
            if compact_lettered and re.match(r"^([A-Z]\.|[IVXLCDM]+\.)$", txt, flags=re.I):
                combined = f"{txt}{next_txt}" if next_txt else txt
            else:
                combined = f"{txt}\t{next_txt}" if next_txt else txt
            i += 2
        else:
            combined = normalise_numbered_text(txt, compact_lettered=compact_lettered)
            i += 1

        level = paragraph_level(combined)
        if level == 1:
            add_para(doc, combined, left_cm=1, first_line_cm=-1)
        elif level == 2:
            add_para(doc, combined, left_cm=2, first_line_cm=-1)
        elif level == 3:
            add_para(doc, combined, left_cm=3, first_line_cm=-1)
        else:
            # Introductory text such as Article 2's chapeau should not be indented.
            add_para(doc, combined, left_cm=0)


# =============================================================================
# Main sections
# =============================================================================

def add_citations(doc: Document, soup: BeautifulSoup) -> None:
    for cit in soup.find_all(id=re.compile(r"^cit_\d+$")):
        paras = provision_paragraphs(cit, classes=("oj-normal",))
        if paras:
            for p in paras:
                add_para(doc, text_with_footnote_tokens(p), left_cm=0)
        else:
            add_para(doc, text_with_footnote_tokens(cit), left_cm=0)


def add_recitals(doc: Document, soup: BeautifulSoup) -> None:
    for rec in soup.find_all(id=re.compile(r"^rct_\d+$")):
        rid = str(rec.get("id", ""))
        number = rid.replace("rct_", "")
        paras = provision_paragraphs(rec, classes=("oj-normal",))
        body_texts: list[str] = []
        for p in paras:
            txt = text_with_footnote_tokens(p)
            if txt in {number, f"({number})"}:
                continue
            txt = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", txt).strip()
            if txt:
                body_texts.append(txt)
        body = " ".join(body_texts) if body_texts else re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", text_with_footnote_tokens(rec)).strip()
        add_para(doc, f"({number})\t{body}", left_cm=1, first_line_cm=-1)


def add_adoption_formula(doc: Document, soup: BeautifulSoup) -> None:
    for p in soup.find_all("p"):
        txt = text_with_footnote_tokens(p)
        if re.match(r"^HAVE\s+ADOPTED\s+THIS\s+REGULATION\s*:$", txt, flags=re.I):
            add_para(doc, txt, left_cm=0)
            return
    for tag in soup.find_all(["p", "div"]):
        txt = text_with_footnote_tokens(tag)
        if "HAVE ADOPTED THIS REGULATION" in txt and len(txt) <= 120:
            add_para(doc, txt, left_cm=0)
            return


def chapter_heading_parts(chapter: Tag) -> tuple[str, str | None]:
    cid = str(chapter.get("id", ""))
    roman = cid.replace("cpt_", "") if cid.startswith("cpt_") else ""
    first = f"CHAPTER {roman}" if roman else "CHAPTER"

    title_candidates: list[str] = []

    for node in chapter.find_all(id=re.compile(rf"^{re.escape(cid)}\.tit_\d+$")):
        txt = clean_text(node.get_text(" ", strip=True))
        if txt and not re.match(r"^CHAPTER\s+", txt, re.I):
            title_candidates.append(txt)

    if not title_candidates:
        for selector in [".oj-ti-section-1", ".oj-ti-section-2", ".eli-title", ".oj-doc-ti"]:
            for node in chapter.select(selector):
                if has_ancestor_with_id_prefix(node, "art_"):
                    continue
                txt = clean_text(node.get_text(" ", strip=True))
                if not txt:
                    continue
                if re.match(r"^CHAPTER\s+", txt, re.I) or re.match(r"^Article\s+\d+", txt, re.I):
                    continue
                if ".tit_" in txt:
                    continue
                if len(txt) <= 120 and txt not in title_candidates:
                    title_candidates.append(txt)

    return first, title_candidates[0] if title_candidates else None


def article_heading_parts(article: Tag) -> tuple[str, str | None]:
    aid = str(article.get("id", ""))
    num = aid.replace("art_", "") if aid.startswith("art_") else ""
    first = f"Article {num}" if num else "Article"

    title = None
    for node in article.select(".oj-sti-art"):
        if is_within_structural_container(node, aid):
            candidate = clean_text(node.get_text(" ", strip=True))
            if candidate and not re.match(r"^Article\s+\d+", candidate, re.I):
                title = candidate
                break

    if not title:
        for node in article.find_all(id=re.compile(rf"^{re.escape(aid)}\.tit_\d+$")):
            candidate = clean_text(node.get_text(" ", strip=True))
            if not candidate:
                continue
            m = re.match(r"^Article\s+\d+\s*(.+)$", candidate, flags=re.I)
            if m and m.group(1).strip():
                title = m.group(1).strip()
                break
            if not re.match(r"^Article\s+\d+$", candidate, flags=re.I):
                title = candidate
                break

    return first, title


def add_operatives(doc: Document, soup: BeautifulSoup) -> None:
    processed: set[str] = set()
    chapters = soup.find_all(id=re.compile(r"^cpt_[IVXLCDM]+$"))

    for cpt in chapters:
        first, second = chapter_heading_parts(cpt)
        add_center_heading(doc, first, second)
        for art in cpt.find_all(id=re.compile(r"^art_\d+[A-Za-z]?$")):
            aid = str(art.get("id", ""))
            if aid in processed:
                continue
            a1, a2 = article_heading_parts(art)
            add_center_heading(doc, a1, a2)
            emit_numbered_paragraph_sequence(doc, provision_paragraphs(art))
            processed.add(aid)

    for art in soup.find_all(id=re.compile(r"^art_\d+[A-Za-z]?$")):
        aid = str(art.get("id", ""))
        if aid in processed:
            continue
        a1, a2 = article_heading_parts(art)
        add_center_heading(doc, a1, a2)
        emit_numbered_paragraph_sequence(doc, provision_paragraphs(art))
        processed.add(aid)


# =============================================================================
# Annexes
# =============================================================================

def set_cell_font(cell, *, bold=False, size=11) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            set_run_font(r, size=size, bold=bold)


def cell_text_from_html(cell: Tag) -> str:
    parts: list[str] = []
    children = cell.find_all(["p", "div"], recursive=False)
    if children:
        for child in children:
            txt = text_with_footnote_tokens(child)
            if txt:
                parts.append(txt)
    else:
        txt = text_with_footnote_tokens(cell)
        if txt:
            parts.append(txt)
    return "\n".join(parts)


def find_largest_html_table(annex: Tag):
    tables = annex.select("table.oj-table") or annex.find_all("table")
    if not tables:
        return None
    return max(tables, key=lambda t: len(t.find_all("tr")) * 10 + len(t.find_all(["td", "th"])))


def build_annex_vi_table(doc: Document, annex: Tag) -> None:
    html_table = find_largest_html_table(annex)
    if html_table is None:
        add_para(doc, text_with_footnote_tokens(annex), left_cm=0)
        return

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for tr in html_table.find_all("tr"):
        src_cells = tr.find_all(["th", "td"], recursive=False)
        if not src_cells:
            continue
        extracted = [cell_text_from_html(c) for c in src_cells]
        extracted = [e for e in extracted if e]
        if not extracted:
            continue
        left = extracted[0]
        right = "\n".join(extracted[1:]) if len(extracted) > 1 else ""
        row = table.add_row()
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].text = left
        row.cells[1].text = right
        is_header = any(c.name == "th" or "oj-tbl-hdr" in tag_classes(c) for c in src_cells)
        set_cell_font(row.cells[0], bold=is_header)
        set_cell_font(row.cells[1], bold=is_header)


def add_annexes(doc: Document, soup: BeautifulSoup) -> None:
    for annex in soup.find_all(id=re.compile(r"^anx_[IVXLCDM]+$")):
        heading = annex.select_one("p.oj-doc-ti")
        heading_text = clean_text(heading.get_text(" ", strip=True)) if heading else str(annex.get("id", "Annex"))
        heading_p = add_center_heading(doc, heading_text)
        heading_p.paragraph_format.page_break_before = True

        if annex.get("id") == "anx_VI":
            build_annex_vi_table(doc, annex)
            continue

        paras: list[Tag] = []
        for p in annex.find_all("p"):
            if has_class(p, "oj-note"):
                continue
            pid = str(p.get("id", ""))
            if ".tit_" in pid:
                continue
            txt = text_with_footnote_tokens(p)
            if not txt or txt == heading_text:
                continue
            classes = tag_classes(p)
            if any(c in classes for c in ["oj-normal", "oj-enumeration-spacing", "oj-doc-ti", "oj-ti-section-1", "oj-ti-section-2"]):
                paras.append(p)

        for p in paras:
            txt = text_with_footnote_tokens(p)
            if any(c in tag_classes(p) for c in ["oj-doc-ti", "oj-ti-section-1", "oj-ti-section-2"]):
                add_center_heading(doc, txt)
            else:
                emit_numbered_paragraph_sequence(doc, [p], compact_lettered=True)


# =============================================================================
# OOXML footnote patcher
# =============================================================================

def w_tag(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def rel_tag(local: str) -> str:
    return f"{{{REL_NS}}}{local}"


def ct_tag(local: str) -> str:
    return f"{{{CT_NS}}}{local}"


def make_text_run(text: str) -> etree._Element:
    r = etree.Element(w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_footnote_reference_run(fid: int) -> etree._Element:
    r = etree.Element(w_tag("r"))
    rpr = etree.SubElement(r, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    rfonts = etree.SubElement(rpr, w_tag("rFonts"))
    rfonts.set(w_tag("ascii"), "Arial")
    rfonts.set(w_tag("hAnsi"), "Arial")
    sz = etree.SubElement(rpr, w_tag("sz"))
    sz.set(w_tag("val"), "18")
    szcs = etree.SubElement(rpr, w_tag("szCs"))
    szcs.set(w_tag("val"), "18")
    vert = etree.SubElement(rpr, w_tag("vertAlign"))
    vert.set(w_tag("val"), "superscript")
    ref = etree.SubElement(r, w_tag("footnoteReference"))
    ref.set(w_tag("id"), str(fid))
    return r


def replace_placeholders_in_document_xml(document_xml: bytes, footnote_texts: dict[str, str]) -> tuple[bytes, dict[int, str]]:
    root = etree.fromstring(document_xml, etree.XMLParser(remove_blank_text=False))
    next_id = 1
    number_to_id: dict[str, int] = {}
    id_to_text: dict[int, str] = {}

    for t in list(root.xpath(".//w:t", namespaces={"w": W_NS})):
        if not t.text or "[[FN:" not in t.text:
            continue
        run = t.getparent()
        para = run.getparent() if run is not None else None
        if run is None or para is None:
            continue
        original = t.text
        parts: list[tuple[str, str | int]] = []
        pos = 0
        for m in FN_TOKEN_RE.finditer(original):
            if m.start() > pos:
                parts.append(("text", original[pos:m.start()]))
            num = m.group(1)
            if num not in number_to_id:
                number_to_id[num] = next_id
                id_to_text[next_id] = footnote_texts.get(num, "")
                next_id += 1
            parts.append(("fn", number_to_id[num]))
            pos = m.end()
        if pos < len(original):
            parts.append(("text", original[pos:]))

        idx = para.index(run)
        para.remove(run)
        for offset, (kind, value) in enumerate(parts):
            new_run = make_text_run(str(value)) if kind == "text" else make_footnote_reference_run(int(value))
            para.insert(idx + offset, new_run)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"), id_to_text


def create_footnotes_xml(id_to_text: dict[int, str]) -> bytes:
    root = etree.Element(w_tag("footnotes"), nsmap={"w": W_NS})

    for fid, ftype, marker in [(-1, "separator", "separator"), (0, "continuationSeparator", "continuationSeparator")]:
        fn = etree.SubElement(root, w_tag("footnote"))
        fn.set(w_tag("id"), str(fid))
        fn.set(w_tag("type"), ftype)
        p = etree.SubElement(fn, w_tag("p"))
        r = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r, w_tag(marker))

    for fid, text in sorted(id_to_text.items()):
        fn = etree.SubElement(root, w_tag("footnote"))
        fn.set(w_tag("id"), str(fid))
        p = etree.SubElement(fn, w_tag("p"))
        ppr = etree.SubElement(p, w_tag("pPr"))
        etree.SubElement(ppr, w_tag("pStyle")).set(w_tag("val"), "FootnoteText")
        ind = etree.SubElement(ppr, w_tag("ind"))
        ind.set(w_tag("left"), "567")
        ind.set(w_tag("hanging"), "567")
        spacing = etree.SubElement(ppr, w_tag("spacing"))
        spacing.set(w_tag("before"), "0")
        spacing.set(w_tag("after"), "0")
        spacing.set(w_tag("line"), "240")
        spacing.set(w_tag("lineRule"), "auto")

        r_ref = etree.SubElement(p, w_tag("r"))
        rpr = etree.SubElement(r_ref, w_tag("rPr"))
        etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
        rfonts_ref = etree.SubElement(rpr, w_tag("rFonts"))
        rfonts_ref.set(w_tag("ascii"), "Arial")
        rfonts_ref.set(w_tag("hAnsi"), "Arial")
        sz_ref = etree.SubElement(rpr, w_tag("sz"))
        sz_ref.set(w_tag("val"), "18")
        szcs_ref = etree.SubElement(rpr, w_tag("szCs"))
        szcs_ref.set(w_tag("val"), "18")
        vert_ref = etree.SubElement(rpr, w_tag("vertAlign"))
        vert_ref.set(w_tag("val"), "superscript")
        etree.SubElement(r_ref, w_tag("footnoteRef"))

        r_tab = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r_tab, w_tag("tab"))

        r_text = etree.SubElement(p, w_tag("r"))
        rpr_text = etree.SubElement(r_text, w_tag("rPr"))
        rfonts = etree.SubElement(rpr_text, w_tag("rFonts"))
        rfonts.set(w_tag("ascii"), "Arial")
        rfonts.set(w_tag("hAnsi"), "Arial")
        sz = etree.SubElement(rpr_text, w_tag("sz"))
        sz.set(w_tag("val"), "18")
        szcs = etree.SubElement(rpr_text, w_tag("szCs"))
        szcs.set(w_tag("val"), "18")
        t = etree.SubElement(r_text, w_tag("t"))
        t.text = text

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_relationship(rels_xml: bytes) -> bytes:
    root = etree.fromstring(rels_xml)
    existing = root.xpath(
        "./rel:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']",
        namespaces={"rel": REL_NS},
    )
    if existing:
        return rels_xml
    used = {rel.get("Id") for rel in root.findall(rel_tag("Relationship"))}
    rid = "rIdFootnotes"
    i = 1
    while rid in used:
        i += 1
        rid = f"rIdFootnotes{i}"
    rel = etree.SubElement(root, rel_tag("Relationship"))
    rel.set("Id", rid)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    rel.set("Target", "footnotes.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = etree.fromstring(content_types_xml)
    existing = root.xpath("./ct:Override[@PartName='/word/footnotes.xml']", namespaces={"ct": CT_NS})
    if existing:
        return content_types_xml
    override = etree.SubElement(root, ct_tag("Override"))
    override.set("PartName", "/word/footnotes.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_with_native_footnotes(docx_path: str | Path, footnotes: dict[str, str]) -> None:
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)

        document_xml = tmp / "word" / "document.xml"
        rels_xml = tmp / "word" / "_rels" / "document.xml.rels"
        content_types = tmp / "[Content_Types].xml"
        footnotes_xml = tmp / "word" / "footnotes.xml"

        new_doc_xml, id_to_text = replace_placeholders_in_document_xml(document_xml.read_bytes(), footnotes)
        document_xml.write_bytes(new_doc_xml)

        if id_to_text:
            footnotes_xml.write_bytes(create_footnotes_xml(id_to_text))
            rels_xml.write_bytes(ensure_footnotes_relationship(rels_xml.read_bytes()))
            content_types.write_bytes(ensure_footnotes_content_type(content_types.read_bytes()))

        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


# =============================================================================
# Build document
# =============================================================================

def build_document(html_path: str | Path, output_path: str | Path) -> Path:
    soup = load_html(html_path)
    footnotes = extract_footnotes_map(soup)
    doc = new_document()

    title = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if title:
        add_title(doc, clean_text(title.get_text(" ", strip=True)))

    add_citations(doc, soup)
    add_recitals(doc, soup)
    add_adoption_formula(doc, soup)
    add_operatives(doc, soup)
    add_annexes(doc, soup)

    output_path = Path(output_path)
    doc.save(output_path)
    patch_docx_with_native_footnotes(output_path, footnotes)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="EUR-Lex XHTML to formatted Word with native footnotes.")
    parser.add_argument("source", help="Path to EUR-Lex XHTML/HTML input file")
    parser.add_argument("output", help="Path to output .docx file")
    args = parser.parse_args()
    output = build_document(args.source, args.output)
    print(f"Stage 2 document saved: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
