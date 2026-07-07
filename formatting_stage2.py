"""
formatting_stage2.py

Stage 2 formatting engine for EUR-Lex / Official Journal XHTML.

This version integrates the revised Annex VI table builder, which reads the
underlying XHTML table structure rather than trying to split plain text rows.

Input:
    EUR-Lex XHTML / HTML file, e.g. L_2017168EN.01001201.xml.html

Output:
    Rebuilt Word document.

Run:
    python formatting_stage2.py L_2017168EN.01001201.xml.html Prospectus_Regulation_Stage2.docx

Notes:
    - Annex VI is converted from the source HTML table into a two-column Word table.
    - Other annexes are currently written as text blocks.
    - This file is intended to replace the earlier formatting_stage2.py draft.
"""

from __future__ import annotations

from pathlib import Path
import argparse
import re

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# =============================================================================
# General helpers
# =============================================================================

NBSP = "\u00A0"


def clean_text(text: str | None) -> str:
    """Normalise whitespace without changing legal substance."""
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    return " ".join(text.split())


def load_html(path: str | Path) -> BeautifulSoup:
    """Load EUR-Lex XHTML/HTML."""
    return BeautifulSoup(Path(path).read_text(encoding="utf-8", errors="replace"), "lxml")


def set_run_font(run, size: int = 11, bold: bool | None = None, italic: bool | None = None) -> None:
    """Apply Arial font settings to a run."""
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_base_format(paragraph, *, before=6, after=6, line_spacing=1.16) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def apply_non_breaking_spaces(text: str) -> str:
    """Apply NBSP rules for thousands, currency, and percentages."""
    # Thousands separator: 200 000 -> 200 NBSP 000. Repeated to handle 1 000 000.
    text = re.sub(r"(?<=\d) (?=\d{3}\b)", NBSP, text)
    # Currency spacing: EUR 200 -> EUR NBSP 200; € 200 -> € NBSP 200.
    text = re.sub(r"\b(EUR)\s+(?=\d)", r"\1" + NBSP, text)
    text = re.sub(r"€\s+(?=\d)", "€" + NBSP, text)
    # Percentages: 20 % -> 20 NBSP %.
    text = re.sub(r"(?<=\d)\s+%", NBSP + "%", text)
    return text


def new_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


# =============================================================================
# Inline XHTML formatting
# =============================================================================

def add_inline_runs(paragraph, element: Tag, *, default_size=11) -> None:
    """Add text from an XHTML element to a Word paragraph, preserving basic b/i/sup.

    This is intentionally conservative: it preserves text order and applies bold,
    italic and superscript for common XHTML tags/classes.
    """

    def walk(node, bold=False, italic=False, superscript=False):
        if isinstance(node, str):
            if node:
                run = paragraph.add_run(apply_non_breaking_spaces(node))
                set_run_font(run, size=default_size, bold=bold, italic=italic)
                run.font.superscript = superscript
            return

        if not isinstance(node, Tag):
            return

        node_classes = node.get("class", [])
        next_bold = bold or node.name in {"b", "strong"} or "bold" in node_classes
        next_italic = italic or node.name in {"i", "em"} or "italic" in node_classes
        next_sup = superscript or node.name == "sup" or "oj-super" in node_classes

        for child in node.children:
            walk(child, next_bold, next_italic, next_sup)

    for child in element.children:
        walk(child)


# =============================================================================
# Headings and paragraph formatting
# =============================================================================

def add_centered_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_base_format(p)
    run = p.add_run(clean_text(text))
    set_run_font(run, bold=True)
    return p


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_base_format(p)
    run = p.add_run(clean_text(text))
    set_run_font(run, size=11, bold=True)
    return p


def add_justified_paragraph(doc: Document, text: str, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    run = p.add_run(apply_non_breaking_spaces(clean_text(text)))
    set_run_font(run)
    return p


def add_recital(doc: Document, recital_div: Tag):
    """Add a recital using rct_* ID and hanging-indent style."""
    rid = recital_div.get("id", "")
    number = rid.replace("rct_", "") if rid.startswith("rct_") else ""
    text = clean_text(recital_div.get_text(" ", strip=True))

    # If the source text already begins with (n), avoid doubling the number.
    text = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", text).strip() if number else text

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)

    run_num = p.add_run(f"({number})\t" if number else "")
    set_run_font(run_num)
    run_text = p.add_run(apply_non_breaking_spaces(text))
    set_run_font(run_text)
    return p


def paragraph_level(text: str) -> int:
    """Detect Article paragraph/list level from literal numbering. No auto-numbering."""
    txt = text.strip()
    if re.match(r"^\d+\.\s+", txt):
        return 1
    if re.match(r"^\([a-z]\)\s+", txt):
        return 2
    if re.match(r"^\([ivxlcdm]+\)\s+", txt, flags=re.I):
        return 3
    return 0


def add_article_text_paragraph(doc: Document, element: Tag):
    """Add an article paragraph with legal indentation levels."""
    text = clean_text(element.get_text(" ", strip=True))
    if not text:
        return None

    level = paragraph_level(text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_base_format(p)

    if level == 1:
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
    elif level == 2:
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.first_line_indent = Cm(-1)
    elif level == 3:
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.first_line_indent = Cm(-1)
    else:
        p.paragraph_format.left_indent = Cm(1)

    add_inline_runs(p, element)
    return p


# =============================================================================
# Annex VI table builder — integrated replacement
# =============================================================================

def set_cell_font(cell, bold: bool = False, size: int = 11) -> None:
    """Apply Arial formatting to all runs in a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def cell_text_from_html(cell: Tag) -> str:
    """Extract text from an HTML table cell, preserving useful internal line breaks."""
    parts: list[str] = []

    # Prefer immediate paragraph/div children, because EUR-Lex often uses multiple
    # paragraphs inside a single table cell.
    children = cell.find_all(["p", "div"], recursive=False)
    if children:
        for child in children:
            text = clean_text(child.get_text(" ", strip=True))
            if text:
                parts.append(apply_non_breaking_spaces(text))
    else:
        text = clean_text(cell.get_text(" ", strip=True))
        if text:
            parts.append(apply_non_breaking_spaces(text))

    return "\n".join(parts)


def find_largest_html_table(annex: Tag):
    """Find the most substantial table in Annex VI."""
    tables = annex.select("table.oj-table")
    if not tables:
        tables = annex.find_all("table")
    if not tables:
        return None

    def score(table: Tag) -> int:
        return len(table.find_all("tr")) * 10 + len(table.find_all(["td", "th"]))

    return max(tables, key=score)


def build_annex_vi_table(doc: Document, annex: Tag) -> None:
    """Build Annex VI as a proper 2-column Word table from the XHTML table.

    Behaviour:
        - one source cell  -> left column, blank right column;
        - two source cells -> direct mapping;
        - >2 source cells  -> first cell left, remaining cells combined in right.
    """
    html_table = find_largest_html_table(annex)

    if html_table is None:
        add_justified_paragraph(doc, annex.get_text(" ", strip=True))
        return

    word_table = doc.add_table(rows=0, cols=2)
    word_table.style = "Table Grid"
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.autofit = True

    for tr in html_table.find_all("tr"):
        html_cells = tr.find_all(["th", "td"], recursive=False)
        if not html_cells:
            continue

        extracted = [cell_text_from_html(cell) for cell in html_cells]
        extracted = [text for text in extracted if text]
        if not extracted:
            continue

        if len(extracted) == 1:
            left_text = extracted[0]
            right_text = ""
        else:
            left_text = extracted[0]
            right_text = "\n".join(extracted[1:])

        row = word_table.add_row()
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].text = left_text
        row.cells[1].text = right_text

        is_header_row = any(
            cell.name == "th" or "oj-tbl-hdr" in cell.get("class", [])
            for cell in html_cells
        )
        set_cell_font(row.cells[0], bold=is_header_row)
        set_cell_font(row.cells[1], bold=is_header_row)


# =============================================================================
# Document builders
# =============================================================================

def add_citations(doc: Document, soup: BeautifulSoup) -> None:
    citations = soup.select('div.eli-subdivision[id^="cit_"]')
    for citation in citations:
        add_justified_paragraph(doc, citation.get_text(" ", strip=True), left_cm=0)


def add_recitals(doc: Document, soup: BeautifulSoup) -> None:
    for recital in soup.select('div.eli-subdivision[id^="rct_"]'):
        add_recital(doc, recital)


def add_articles(doc: Document, soup: BeautifulSoup) -> None:
    for art in soup.select('div.eli-subdivision[id^="art_"]'):
        art_id = art.get("id", "")
        title_tag = soup.find(id=f"{art_id}.tit_1")
        heading_text = clean_text(title_tag.get_text(" ", strip=True)) if title_tag else art_id

        # Use Shift+Return between Article number and title if both appear in separate title lines.
        heading_text = heading_text.replace(" Article ", "\nArticle ") if heading_text.startswith("CHAPTER") else heading_text
        add_centered_heading(doc, heading_text, level=2)

        # Add only immediate oj-normal / enumeration paragraphs inside the article to avoid duplicate container text.
        for p in art.find_all(["p", "div"], recursive=True):
            classes = p.get("class", [])
            if "oj-ti-art" in classes or "oj-sti-art" in classes or "eli-title" in classes:
                continue
            if "oj-normal" in classes or "oj-enumeration-spacing" in classes:
                add_article_text_paragraph(doc, p)


def add_annexes(doc: Document, soup: BeautifulSoup) -> None:
    for annex in soup.select('div.eli-container[id^="anx_"]'):
        heading = annex.select_one("p.oj-doc-ti")
        heading_text = clean_text(heading.get_text(" ", strip=True)) if heading else annex.get("id", "Annex")
        add_centered_heading(doc, heading_text, level=2)

        if annex.get("id") == "anx_VI":
            build_annex_vi_table(doc, annex)
        else:
            # For non-Annex VI, preserve text block for now.
            # Later enhancement can convert all annex tables structurally.
            body_text = clean_text(annex.get_text(" ", strip=True))
            # Avoid repeating the heading at the start if present.
            if heading_text and body_text.startswith(heading_text):
                body_text = body_text[len(heading_text):].strip()
            add_justified_paragraph(doc, body_text, left_cm=0)


def add_footnotes_section(doc: Document, soup: BeautifulSoup) -> None:
    """Add footnotes as a separate section for inspection.

    This keeps p.oj-note only and ignores the hr.oj-note separator.
    Native Word footnote conversion should be handled by the dedicated footnote
    OOXML engine when that module is added; this section is retained as a QA aid.
    """
    notes = soup.select("p.oj-note")
    if not notes:
        return

    add_centered_heading(doc, "Footnotes", level=1)
    for note in notes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        run = p.add_run(clean_text(note.get_text(" ", strip=True)))
        set_run_font(run, size=9)


def build_document(html_path: str | Path, output_path: str | Path) -> Path:
    soup = load_html(html_path)
    doc = new_document()

    title = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if title:
        add_title(doc, title.get_text(" ", strip=True))

    add_citations(doc, soup)
    add_recitals(doc, soup)
    add_articles(doc, soup)
    add_annexes(doc, soup)
    add_footnotes_section(doc, soup)

    output_path = Path(output_path)
    doc.save(output_path)
    return output_path


# =============================================================================
# CLI
# =============================================================================

def main() -> int:
    parser = argparse.ArgumentParser(description="Stage 2 EUR-Lex XHTML to formatted Word document.")
    parser.add_argument("source", help="Path to EUR-Lex XHTML/HTML input file")
    parser.add_argument("output", help="Path to output .docx file")
    args = parser.parse_args()

    output = build_document(args.source, args.output)
    print(f"Stage 2 document saved: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
