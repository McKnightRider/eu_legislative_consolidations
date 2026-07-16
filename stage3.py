"""Stage 3 consolidation draft: apply one amending regulation onto a Stage 1/3 DOCX.

Inputs:
1) base DOCX (previous consolidated version)
2) amending regulation HTML
3) display color for this amendment layer

Behavior implemented in this draft:
- Inserts a clearly visible amending regulation section near the top that includes
    the full amending title and full amending recitals.
- Applies heuristic paragraph replacements for amendment blocks against the
    current document text.
- Performs paragraph-level matching and word-level diff analysis for transparency.
- Writes an updated DOCX plus a JSON analysis report.

Note:
This is a draft analysis-first consolidation stage. It does not yet perform full
legal instruction execution (replace/insert/delete at exact legal anchors).
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from dataclasses import asdict
from datetime import datetime, timezone
import json
from pathlib import Path
import re
import shutil
import tempfile
import zipfile
from difflib import SequenceMatcher

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from lxml import etree

from eurlex_parser import parse_eurlex_document
from stage1 import (
    article_heading_parts,
    clean_text,
    ensure_footnotes_content_type,
    ensure_footnotes_relationship,
    extract_footnotes_map,
    load_html,
    normalise_marker_spacing,
    provision_paragraphs,
    split_main_title,
    text_with_footnote_tokens,
    w_tag,
)

DEFAULT_OUTPUT_DIR = Path("outputs")
DEFAULT_ANALYSIS_DIR = DEFAULT_OUTPUT_DIR / "qa"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

RECITAL_RE = re.compile(r"^\(\d+\)")
TOKEN_RE = re.compile(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?")
FN_TOKEN_RE = re.compile(r"\[\[FN:(\d+)\]\]")
HIGH_CONFIDENCE_MATCH = 0.72
ENABLE_HEURISTIC_DOC_MUTATIONS = True
FORCE_ANALYSIS_ON_DOUBT = True
NON_LEGISLATIVE_PAYLOAD_RE = re.compile(
    r'"(?:missing_unique_tokens|missing_tokens|coverage_ratio|comparison_type|report_type|source_docx|source_html|token_coverage)"\s*:',
    flags=re.I,
)

COLOR_MAP: dict[str, tuple[int, int, int]] = {
    "red": (192, 0, 0),
    "blue": (0, 102, 204),
    "green": (0, 128, 0),
    "orange": (230, 120, 0),
    "purple": (128, 0, 128),
    "teal": (0, 128, 128),
    "brown": (128, 64, 0),
    "black": (0, 0, 0),
}
REVISION_COLOR = COLOR_MAP["orange"]


def revision_color_hex() -> str:
    return f"{REVISION_COLOR[0]:02X}{REVISION_COLOR[1]:02X}{REVISION_COLOR[2]:02X}"


def set_revision_color(rgb: tuple[int, int, int]) -> None:
    global REVISION_COLOR
    REVISION_COLOR = rgb


def normalize_text(text: str) -> str:
    return clean_text(text)


def strip_src_artifacts(text: str) -> str:
    txt = text or ""
    txt = re.sub(r"\s*\[\[SRC:[^\]]+\]\]", "", txt, flags=re.I)
    txt = re.sub(r"\bSRC:[A-Za-z0-9_.-]+\b", "", txt, flags=re.I)
    txt = re.sub(r"\bsrc\s+art\s+\d+[A-Za-z]?(?:\s+\d+[A-Za-z]?)?\b", "", txt, flags=re.I)
    return normalize_text(txt)


def strip_leading_roman_marker(text: str) -> str:
    txt = normalize_text(text)
    return re.sub(r"^\(([ivxlcdm]+)\)\s+", "", txt, count=1, flags=re.I)


def is_bare_amendment_marker(text: str) -> bool:
    txt = normalize_text(text)
    return bool(re.fullmatch(r"\([ivxlcdm]+\)|\([a-z]+\)|\([a-z]{2,}\)|\d+[A-Za-z]?\.", txt, flags=re.I))


def ensure_leading_marker_tab(text: str) -> str:
    """Normalize leading legal markers to use a tab separator.

    Examples:
    - "1. The ..." -> "1.\tThe ..."
    - "(b) text" -> "(b)\ttext"
    Handles optional opening quote characters that may wrap inserted blocks.
    """
    txt = (text or "").replace("\xa0", " ").strip()
    if not txt:
        return ""

    quote = r"[‘'\"“”]?"
    txt = re.sub(rf"^({quote}\(\d+\))\s+", r"\1\t", txt)
    txt = re.sub(rf"^({quote}\d+[A-Za-z]?\.)\s+", r"\1\t", txt)
    txt = re.sub(rf"^({quote}\([a-z]+\))\s+", r"\1\t", txt, flags=re.I)
    txt = re.sub(rf"^({quote}\([ivxlcdm]+\))\s+", r"\1\t", txt, flags=re.I)
    return txt


def ensure_marker_tabs_multiline(text: str) -> str:
    """Apply marker-tab normalization to each logical line."""
    raw = text or ""
    if not raw:
        return ""
    lines = raw.split("\n")
    normalized = [ensure_leading_marker_tab(line) if line.strip() else "" for line in lines]
    return "\n".join(normalized)


def format_amendment_item_text(text: str) -> str:
    txt = strip_src_artifacts(text)
    if not txt:
        return ""

    txt = sanitize_inserted_marker_text(txt)
    txt = normalize_quote_wrapping(txt)

    # Drop outer roman markers when they are only amendment scaffolding for a nested marker or instruction.
    if re.match(r"^\([ivxlcdm]+\)\s+(?=\([a-z]+\)|in\s+the\b|the\s+following\b|paragraph\b|point\b|Article\b)", txt, flags=re.I):
        txt = re.sub(r"^\([ivxlcdm]+\)\s+", "", txt, count=1, flags=re.I)

    if is_instructional_amendment_line(txt) or is_instructional_amendment_line(strip_leading_roman_marker(txt)):
        return ""

    if is_bare_amendment_marker(txt):
        return ""

    # Drop marker-only punctuation artifacts produced by split list rows,
    # e.g. "(iii);" or ";".
    if re.match(r"^\(([a-z]+|[ivxlcdm]+)\)\s*;\s*$", txt, flags=re.I):
        return ""
    if re.match(r"^;\s*$", txt):
        return ""

    return ensure_leading_marker_tab(txt)


def normalize_amendment_item(item: dict[str, str]) -> dict[str, str]:
    normalized = dict(item)
    if not normalized.get("amendment_kind"):
        normalized["amendment_kind"] = classify_amendment_kind(
            normalized.get("source_instruction"),
            normalized.get("text"),
        )
    if normalized.get("amendment_kind") == "deletion":
        normalized["text"] = normalize_text(normalized.get("text", "") or normalized.get("source_instruction", ""))
    else:
        normalized["text"] = format_amendment_item_text(normalized.get("text", ""))
    point_marker = normalize_text(str(normalized.get("target_point_marker") or ""))
    source_instruction = normalize_text(str(normalized.get("source_instruction") or ""))
    if (
        normalized["text"]
        and re.fullmatch(r"[IVXLC]+\.", normalized["text"], flags=re.I)
        and re.search(r"\bin\s+Annex\b", source_instruction, flags=re.I)
        and re.search(r"\breplaced by the following\b", source_instruction, flags=re.I)
    ):
        # Standalone Annex heading markers (e.g. "II.") are scaffolding lines,
        # not substantive inserted paragraphs.
        normalized["text"] = ""
    if point_marker and normalized["text"] and normalized.get("amendment_kind") != "deletion":
        existing_point_marker = first_top_level_point_marker(normalized["text"])
        explicit_point_reference = bool(
            re.search(rf"\bpoint\s*\({re.escape(point_marker)}\)", source_instruction, flags=re.I)
        )
        generic_point_addition_instruction = bool(
            re.search(
                r"\bthe following point(?:s)?\s+(?:is|are)\s+(?:added|inserted)\b",
                source_instruction,
                flags=re.I,
            )
        )
        if (
            (explicit_point_reference or generic_point_addition_instruction)
            and existing_point_marker is None
            and not re.match(
            rf"^\({re.escape(point_marker)}\)(?=\s|\t|$)",
            normalized["text"],
            flags=re.I,
            )
        ):
            normalized["text"] = f"({point_marker})\t{normalized['text']}"
    return normalized


def normalize_amendment_items(items: list[dict[str, str]]) -> list[dict[str, str]]:
    normalized: list[dict[str, str]] = []
    for item in items:
        cleaned = normalize_amendment_item(item)
        if not cleaned.get("text") and cleaned.get("amendment_kind") != "deletion":
            continue
        normalized.append(cleaned)
    return normalized


def normalize_replacement_target_points(items: list[dict[str, str]]) -> list[dict[str, str]]:
    for item in items:
        instruction = normalize_text(item.get("source_instruction", ""))
        if not instruction or not re.search(r"\breplaced by the following\b", instruction, flags=re.I):
            continue
        target_point = extract_target_point_from_instruction(instruction)
        if target_point:
            item["target_point_marker"] = target_point
        target_annex = extract_target_annex_from_instruction(instruction)
        if target_annex:
            item["target_annex_number"] = target_annex
            item["target_article_number"] = None
            item["target_paragraph_number"] = None
            item["target_point_marker"] = None
        target_annex_point = extract_target_annex_point_from_instruction(instruction)
        if target_annex_point:
            item["target_annex_point_marker"] = target_annex_point
    return items


def backfill_targets_from_instruction(items: list[dict[str, str]]) -> list[dict[str, str]]:
    last_target_article: str | None = None
    for item in items:
        instruction = normalize_text(item.get("source_instruction", ""))
        if not instruction:
            if item.get("target_article_number"):
                last_target_article = str(item.get("target_article_number"))
            continue
        # Prefer explicit target extraction from the instruction itself.
        explicit_art = extract_target_article_from_instruction(instruction)
        if explicit_art:
            item["target_article_number"] = explicit_art
        if not item.get("target_article_number"):
            art = extract_target_article_from_instruction(instruction)
            if art:
                item["target_article_number"] = art
        if not item.get("target_article_number") and last_target_article and re.search(
            r"\b(?:paragraph|subparagraph|point|points|introductory wording)\b",
            instruction,
            flags=re.I,
        ):
            # Relative instructions often omit the article and rely on the
            # previous instruction's article context.
            item["target_article_number"] = last_target_article
        if not item.get("target_paragraph_number"):
            par = extract_target_paragraph_from_instruction(instruction)
            if par:
                item["target_paragraph_number"] = par
        if not item.get("target_point_marker"):
            pt = extract_target_point_from_instruction(instruction)
            if pt:
                item["target_point_marker"] = pt
        if not item.get("target_annex_number"):
            annex = extract_target_annex_from_instruction(instruction)
            if annex:
                item["target_annex_number"] = annex
                item["target_article_number"] = None
                item["target_paragraph_number"] = None
                item["target_point_marker"] = None
        if not item.get("target_annex_point_marker"):
            annex_pt = extract_target_annex_point_from_instruction(instruction)
            if annex_pt:
                item["target_annex_point_marker"] = annex_pt
        if item.get("target_article_number"):
            last_target_article = str(item.get("target_article_number"))
    return items


def is_non_legislative_payload_text(text: str) -> bool:
    """Detect JSON/QA payload fragments that must never enter consolidated DOCX text."""
    txt = normalize_text(text)
    if not txt:
        return False
    if NON_LEGISLATIVE_PAYLOAD_RE.search(txt):
        return True
    if txt.lower().startswith('"missing_unique_tokens"'):
        return True
    if re.match(r"^[\[{]", txt) and re.search(r'"[^"\\n]+"\s*:\s*', txt):
        return True
    if ('{' in txt or '[' in txt) and txt.count('":') >= 2:
        return True
    return False


def extract_target_article_from_instruction(text: str) -> str | None:
    """Extract target base-legal-act article reference from an instruction line."""
    txt = normalize_text(text)
    if not txt:
        return None

    patterns = [
        r"\bin\s+Article\s+(\d+[A-Za-z]?)\b",
        r"\bArticle\s+(\d+[A-Za-z]?)\s*\(\d+[A-Za-z]?\)",
        r"\bArticle\s+(\d+[A-Za-z]?)\s*,\s*the following",
        r"\bArticle\s+(\d+[A-Za-z]?)\s+is\s+amended\b",
        r"\bArticle\s+(\d+[A-Za-z]?)\s+(?:is|are)\s+replaced\b",
        r"\bArticle\s+(\d+[A-Za-z]?)\s+(?:is|are)\s+inserted\b",
        r"\bArticle\s+(\d+[A-Za-z]?)\s+(?:is|are)\s+deleted\b",
    ]
    for pat in patterns:
        m = re.search(pat, txt, flags=re.I)
        if m:
            return m.group(1)
    return None


def extract_target_paragraph_from_instruction(text: str) -> str | None:
    """Extract target paragraph number from instruction text when present."""
    txt = normalize_text(text)
    if not txt:
        return None

    m = re.search(r"\bparagraph\s+(\d+[A-Za-z]?)\b", txt, flags=re.I)
    if m:
        return m.group(1)

    # Also support references encoded as Article X(Y), e.g. Article 15(1).
    m = re.search(r"\bArticle\s+\d+[A-Za-z]?\s*\((\d+[A-Za-z]?)\)", txt, flags=re.I)
    if m:
        return m.group(1)
    return None


def extract_target_point_from_instruction(text: str) -> str | None:
    """Extract point marker from instruction text when present (e.g. point (f))."""
    txt = normalize_text(text)
    if not txt:
        return None

    m = re.search(r"\bpoint\s*\(([a-z]+)\)", txt, flags=re.I)
    if m:
        return m.group(1).lower()
    return None


def extract_all_target_points_from_instruction(text: str) -> list[str]:
    """Extract all point markers referenced in instruction text."""
    txt = normalize_text(text)
    if not txt:
        return []
    hits = [m.lower() for m in re.findall(r"\(([a-z]+)\)", txt, flags=re.I)]
    return list(dict.fromkeys(hits))


def extract_target_annex_from_instruction(text: str) -> str | None:
    txt = normalize_text(text)
    if not txt:
        return None
    m = re.search(r"\bin\s+Annex\s+([IVXLC]+)\b", txt, flags=re.I)
    if m:
        return m.group(1).upper()
    return None


def extract_target_annex_point_from_instruction(text: str) -> str | None:
    txt = normalize_text(text)
    if not txt:
        return None
    m = re.search(r"\bpoint\s+([IVXLC]+)\b", txt, flags=re.I)
    if m:
        return m.group(1).lower()
    return None


def extract_inserted_article_marker_from_instruction_paragraph(p: Tag) -> str | None:
    """Extract inserted article marker (e.g. 14a) from nearby article div id."""
    txt = normalize_text(p.get_text(" ", strip=True)) if p else ""
    if txt:
        inline = re.search(
            r"\bthe following article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\s*:\s*[‘'\"\(\[]*\s*article\s+(\d+[A-Za-z])\b",
            txt,
            flags=re.I,
        )
        if inline:
            return inline.group(1).lower()

    parent = p.parent if isinstance(p.parent, Tag) else None
    if parent is None:
        return None

    # In EUR-Lex amend tables, the inserted article body usually follows in a
    # sibling <div id="014A"> ... </div> under the same table cell.
    candidate_div = p.find_next_sibling("div")
    if candidate_div is None:
        candidate_div = parent.find("div")
    if candidate_div is not None:
        div_id = normalize_text(str(candidate_div.get("id", "")))
        m = re.fullmatch(r"0*(\d+)([A-Za-z])", div_id)
        if m:
            return f"{int(m.group(1))}{m.group(2).lower()}"

    # Some EUR-Lex layouts keep the inserted article heading as nearby text
    # instead of a sibling div id. Probe the next local elements for a heading
    # that starts with "Article <num><letter>".
    lookahead_tags: list[Tag] = []
    for sib in p.next_siblings:
        if isinstance(sib, Tag):
            lookahead_tags.append(sib)
            if len(lookahead_tags) >= 10:
                break
    if not lookahead_tags:
        lookahead_tags = [n for n in p.find_all_next(["p", "div", "td", "li"], limit=10) if isinstance(n, Tag)]

    for tag in lookahead_tags:
        candidate_text = normalize_text(tag.get_text(" ", strip=True))
        mh = re.match(r"^[‘'\"\[]?Article\s+(\d+[A-Za-z])\b", candidate_text, flags=re.I)
        if mh:
            return mh.group(1).lower()

    return None


def is_instructional_amendment_line(text: str) -> bool:
    """Identify amendment instruction scaffolding, not substantive inserted text."""
    txt = strip_leading_roman_marker(text)
    if not txt:
        return False

    # Bare list item markers in amending acts, e.g. "(1)", "(2)".
    if re.fullmatch(r"\(\d+\)", txt):
        return True

    lowered = txt.lower()
    patterns = [
        r"\bis amended as follows\b",
        r"\bthe following paragraph(?:s)? (?:is|are) inserted\b",
        r"\bthe following point(?:s)? (?:is|are) inserted\b",
        r"\bthe following article(?:s)? (?:is|are) inserted\b",
        r"\bthe following paragraph(?:s)? (?:is|are) added\b",
        r"\bthe following point(?:s)? (?:is|are) added\b",
        r"\bthe following article(?:s)? (?:is|are) added\b",
        r"\bthe following subparagraph(?:s)? (?:is|are) added\b",
        r"\bin the first subparagraph, the following sentence is added\b",
        r"\bin the second subparagraph, the following sentence is added\b",
        r"\bin the second subparagraph, the following subparagraph(?:s)? are added\b",
        r"\bthe following subparagraph(?:s)? (?:is|are) inserted\b",
        r"\bis inserted as annex\s+[a-z0-9]+\b",
        r"\bare inserted as annex\s+[a-z0-9]+\b",
        r"\bannexes?\s+[ivxlcdm0-9\s,toand]+\s+are\s+amended\s+in\s+accordance\s+with\s+the\s+annex\s+to\s+this\s+regulation\b",
        r"\bis replaced by the following\b",
        r"\bare replaced by the following\b",
        r"\bis deleted\b",
        r"\bare deleted\b",
    ]
    return any(re.search(p, lowered) for p in patterns)


def classify_amendment_kind(source_instruction: str | None, text: str | None = None) -> str:
    """Classify an amendment as insertion, replacement, deletion, or structural scaffolding.

    The classification is intentionally conservative: replacement language wins
    whenever the instruction explicitly says a provision is "replaced".
    "Insertion" covers both newly inserted and added text, including inserted
    annexes and point lists. Any line that is only a marker or otherwise lacks a
    usable instruction defaults to insertion so it can still be placed by the
    structural insertion logic.
    """
    instruction = normalize_text(source_instruction or "")
    lowered = instruction.lower()
    payload = normalize_text(text or "")

    if not instruction:
        return "insertion"

    if re.search(r"\breplaced\s+by\s+the\s+following\b", lowered) or re.search(r"\bis\s+replaced\b", lowered) or re.search(r"\bare\s+replaced\b", lowered):
        return "replacement"

    if re.search(r"\b(?:is|are)\s+deleted\b", lowered) or re.search(r"\bdeleted\b", lowered):
        return "deletion"

    if re.search(r"\bis\s+inserted\s+as\s+annex\b", lowered) or re.search(r"\bare\s+inserted\s+as\s+annex\b", lowered):
        return "insertion"

    if re.search(r"\bthe following (?:paragraph(?:s)?|point(?:s)?|article(?:s)?|subparagraph(?:s)?)\s+(?:is|are)\s+(?:inserted|added)\b", lowered):
        return "insertion"

    if re.search(r"\bthe following sentence is added\b", lowered):
        return "insertion"

    if payload and is_bare_amendment_marker(payload):
        return "structural"

    return "insertion"


def iter_document_paragraphs(doc: Document):
    """Yield paragraphs from document body and all table cells (including nested tables)."""

    def iter_table_paragraphs(table):
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_cells:
                    continue
                seen_cells.add(tc_id)
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    yield from iter_table_paragraphs(nested)

    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from iter_table_paragraphs(table)


def parse_color(value: str) -> tuple[int, int, int]:
    key = value.strip().lower()
    if key in COLOR_MAP:
        return COLOR_MAP[key]

    if key.startswith("#"):
        key = key[1:]
    if len(key) == 6 and all(ch in "0123456789abcdef" for ch in key):
        return (int(key[0:2], 16), int(key[2:4], 16), int(key[4:6], 16))

    raise ValueError(
        "Unsupported color. Use a name like red/blue/green/orange/purple/teal or hex #RRGGBB."
    )


def apply_run_color(run, rgb: tuple[int, int, int]) -> None:
    run.font.color.rgb = RGBColor(*rgb)
    # Also write explicit OOXML color attributes to avoid theme-based overrides
    # in some office viewers.
    r_pr = run._element.get_or_add_rPr()
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    color.set(qn("w:val"), hex_color)
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        if qn(attr) in color.attrib:
            del color.attrib[qn(attr)]


def style_inserted_run(run) -> None:
    apply_run_color(run, REVISION_COLOR)
    run.font.strike = False
    run.font.underline = WD_UNDERLINE.DOUBLE
    run.bold = False
    # Mirror underline style/color in OOXML for better renderer compatibility.
    r_pr = run._element.get_or_add_rPr()
    u = r_pr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        r_pr.append(u)
    u.set(qn("w:val"), "double")
    u.set(qn("w:color"), revision_color_hex())


def style_inserted_run_no_bold_change(run) -> None:
    """Apply insertion mark styling without forcing bold on/off."""
    apply_run_color(run, REVISION_COLOR)
    run.font.strike = False
    run.font.underline = WD_UNDERLINE.DOUBLE

    r_pr = run._element.get_or_add_rPr()
    u = r_pr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        r_pr.append(u)
    u.set(qn("w:val"), "double")
    u.set(qn("w:color"), revision_color_hex())


def style_inserted_title_run(run) -> None:
    """Apply revision styling while preserving the bold title formatting."""
    apply_run_color(run, REVISION_COLOR)
    run.font.underline = WD_UNDERLINE.DOUBLE
    run.bold = True

    r_pr = run._element.get_or_add_rPr()
    u = r_pr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        r_pr.append(u)
    u.set(qn("w:val"), "double")
    u.set(qn("w:color"), revision_color_hex())


def force_paragraph_justification(paragraph) -> None:
    """Set justification in both python-docx and raw OOXML form.

    Some Office renderers respect the underlying XML justification more
    consistently than the high-level paragraph alignment property alone.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "both")


def paragraph_is_nested_roman_point(paragraph) -> bool:
    marker = first_top_level_point_marker(getattr(paragraph, "text", ""))
    return bool(marker and is_roman_point_marker(marker))


def should_leave_nested_roman_plain(paragraph, source_instruction: str | None) -> bool:
    if not paragraph_is_nested_roman_point(paragraph):
        return False
    instruction = normalize_text(source_instruction or "")
    # Only Roman items inside replacement instructions should remain plain.
    return bool(re.search(r"\breplaced by the following\b", instruction, flags=re.I))


def enforce_inserted_paragraph_styles(inserted_paragraphs: list[tuple]) -> None:
    """Guarantee inserted amendment paragraphs are visibly marked.

    Applied after insertion so formatting inheritance occurs first, then revision
    styling is layered on top. All inserted additions are styled the same way,
    regardless of whether they are alphabetic or Roman markers.
    """
    for p, source_instruction in inserted_paragraphs:
        if should_leave_nested_roman_plain(p, source_instruction):
            # Only Roman items that belong to a replacement instruction stay
            # plain; ordinary inserted additions are always styled.
            continue
        for run in p.runs:
            style_inserted_run(run)


def paragraph_has_inserted_style(paragraph) -> bool:
    for run in paragraph.runs:
        if run.font.underline == WD_UNDERLINE.DOUBLE and not run.font.strike:
            return True
    return False


def retrofit_styles_for_matching_amendment_paragraphs(
    doc: Document,
    amending_items: list[dict[str, str]],
    allowed_paragraph_ids: set[int] | None = None,
) -> int:
    """Force style on already-applied amendment paragraphs that are still plain.

    This is a compatibility safeguard for documents created before styling fixes.
    It applies the same visible revision styling to any inserted addition that
    still lacks it.
    """
    target_texts = {normalize_text(item.get("text", "")) for item in amending_items if normalize_text(item.get("text", ""))}
    changed = 0
    for p in doc.paragraphs:
        if allowed_paragraph_ids is not None and id(p._element) not in allowed_paragraph_ids:
            continue
        txt = normalize_text(p.text)
        if not txt or txt not in target_texts:
            continue
        matching_instructions = [
            item.get("source_instruction")
            for item in amending_items
            if normalize_text(item.get("text", "")) == txt
        ]
        if any(should_leave_nested_roman_plain(p, source_instruction) for source_instruction in matching_instructions):
            # Only replacement-only Roman paragraphs are exempt from retrofit
            # styling.
            continue
        if paragraph_has_inserted_style(p):
            continue
        if not p.runs:
            p.add_run(p.text)
        for run in p.runs:
            if run.text and run.text.strip():
                style_inserted_run(run)
                changed += 1
    return changed


def split_amending_title(title: str) -> list[str]:
    """Split amending title like Stage 1, but with generic date matching.

    Inserts line breaks around "of <date>" and before "(Text with EEA relevance)".
    """
    title = normalize_text(title)
    title = re.sub(
        r"\s+(of\s+\d{1,2}\s+[A-Za-z]+\s+\d{4})\s+",
        r"\n\1\n",
        title,
        flags=re.I,
    )
    title = re.sub(r"\s+(\(Text\s+with\s+EEA\s+relevance\))", r"\n\1", title, flags=re.I)
    return [line.strip() for line in title.split("\n") if line.strip()]


def next_amendment_number_from_base(base_docx: Path) -> int:
    stem = base_docx.stem
    if stem == "stage1":
        return 1
    match = re.fullmatch(r"stage3_(\d+)", stem)
    if match:
        return int(match.group(1))
    return 1


def default_output_path(base_docx: Path) -> Path:
    amendment_number = next_amendment_number_from_base(base_docx)
    return DEFAULT_OUTPUT_DIR / f"stage3_{amendment_number}.docx"


def default_analysis_path(output_docx: Path) -> Path:
    return DEFAULT_ANALYSIS_DIR / f"{output_docx.stem}_amendment_analysis.json"


def default_analysis_docx_path(output_docx: Path) -> Path:
    return DEFAULT_ANALYSIS_DIR / f"{output_docx.stem}_amendment_analysis.docx"


def resolve_base_docx(base_docx: Path) -> tuple[Path | None, list[str]]:
    """Resolve a usable base DOCX, with fallbacks for deleted stage3 files."""
    notes: list[str] = []
    if base_docx.exists():
        return base_docx, notes

    parent = base_docx.parent
    stem = base_docx.stem

    m = re.fullmatch(r"stage3_(\d+)", stem)
    if m:
        n = int(m.group(1))
        for prior in range(n - 1, 0, -1):
            candidate = parent / f"stage3_{prior}.docx"
            if candidate.exists():
                notes.append(f"Requested base missing; using previous Stage 3 base: {candidate}")
                return candidate, notes

        stage1_candidate = parent / "stage1.docx"
        if stage1_candidate.exists():
            notes.append(f"Requested base missing; using Stage 1 base: {stage1_candidate}")
            return stage1_candidate, notes

    if stem == "stage1":
        local_stage1 = parent / "stage1.docx"
        if local_stage1.exists():
            notes.append(f"Requested base missing; using available Stage 1 base: {local_stage1}")
            return local_stage1, notes

    return None, notes


def fallback_extract_recitals_from_html(amending_html: Path) -> list[str]:
    html = amending_html.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    rows: list[str] = []
    for tag in soup.select('div.eli-subdivision[id^="rct_"] p.oj-normal, div.eli-subdivision[id^="rct_"]'):
        text = normalize_text(tag.get_text(" ", strip=True))
        if text and RECITAL_RE.match(text) and not is_non_legislative_payload_text(text):
            rows.append(text)
    # Preserve order and uniqueness.
    unique: list[str] = []
    seen = set()
    for row in rows:
        if row not in seen:
            unique.append(row)
            seen.add(row)

    if unique:
        return unique

    # Fallback: numbered oj-normal paragraphs before first Article container.
    first_article = soup.select_one('div.eli-subdivision[id^="art_"]')
    for tag in soup.find_all(True):
        if first_article is not None and tag is first_article:
            break
        classes = tag.get("class", [])
        if tag.name == "p" and "oj-normal" in classes:
            text = normalize_text(tag.get_text(" ", strip=True))
            if text and RECITAL_RE.match(text) and not is_non_legislative_payload_text(text):
                unique.append(text)

    return unique


def extract_title_from_html(amending_html: Path) -> str:
    html = amending_html.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    node = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if node:
        text = normalize_text(node.get_text(" ", strip=True))
        if text:
            return text
    node = soup.select_one("title")
    if node:
        text = normalize_text(node.get_text(" ", strip=True))
        if text:
            return text
    return ""


def fallback_extract_amending_items_from_html(amending_html: Path) -> list[dict[str, str]]:
    html = amending_html.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    items: list[dict[str, str]] = []

    for art in soup.select('div.eli-subdivision[id^="art_"]'):
        art_id = str(art.get("id", ""))
        m = re.search(r"art_([A-Za-z0-9]+)", art_id)
        article_number = m.group(1) if m else art_id or "?"
        heading_tag = art.select_one(".oj-ti-art")
        heading = normalize_text(heading_tag.get_text(" ", strip=True)) if heading_tag else f"Article {article_number}"

        seen_block_texts: set[str] = set()
        for p in art.select("p.oj-normal"):
            text = normalize_text(p.get_text(" ", strip=True))
            if not text:
                continue
            if is_non_legislative_payload_text(text):
                continue
            if text in seen_block_texts:
                continue
            seen_block_texts.add(text)
            items.append(
                {
                    "article_number": article_number,
                    "article_heading": heading,
                    "text": text,
                    "source_instruction": None,
                }
            )

    return items


def sanitize_inserted_marker_text(text: str) -> str:
    """Remove leading quote chars before top-level markers (e.g. '6a.')."""
    txt = normalize_text(text)
    return re.sub(r"^[‘'\"](?=(\d+[A-Za-z]?\.|\([a-z]+\)))", "", txt)

def is_standalone_list_marker_text(text: str) -> bool:
    txt = normalize_text(text)
    return bool(re.fullmatch(r"\(([a-z]+|[ivxlcdm]+)\)", txt, flags=re.I))

def coalesce_marker_body_items(items: list[dict[str, str]]) -> list[dict[str, str]]:
    """Merge standalone marker rows with the immediately following body row.

    EUR-Lex often splits list points as:
      (a)
      body text...
    We coalesce to:
      (a)\tbody text...
    """
    merged: list[dict[str, str]] = []
    i = 0
    while i < len(items):
        cur = dict(items[i])
        txt = normalize_text(cur.get("text", ""))
        if is_standalone_list_marker_text(txt) and i + 1 < len(items):
            nxt = items[i + 1]
            nxt_txt = normalize_text(nxt.get("text", ""))
            if nxt_txt and not is_standalone_list_marker_text(nxt_txt):
                cur["text"] = f"{txt}\t{nxt_txt}"
                marker_match = re.fullmatch(r"\(([a-z]+|[ivxlcdm]+)\)", txt, flags=re.I)
                if marker_match:
                    cur["target_point_marker"] = marker_match.group(1).lower()
                # Prefer more specific context from body line when available.
                for key in ("target_article_number", "target_paragraph_number", "target_point_marker", "source_instruction"):
                    if not cur.get(key) and nxt.get(key):
                        cur[key] = nxt.get(key)
                if nxt.get("source_instruction"):
                    cur["source_instruction"] = nxt.get("source_instruction")
                merged.append(cur)
                i += 2
                continue
        merged.append(cur)
        i += 1
    return merged


def infer_target_regulation_from_base_docx(base_docx: Path) -> str | None:
    """Infer target regulation id like 2017/1129 from base consolidated DOCX."""
    doc = Document(str(base_docx))
    for p in doc.paragraphs[:30]:
        txt = normalize_text(p.text)
        if not txt:
            continue
        m = re.search(r"\b(\d{4}/\d{1,4})\b", txt)
        if m:
            return m.group(1)
    return None


def is_item_relevant_to_target_regulation(item: dict[str, str], target_regulation: str | None) -> bool:
    """Keep only amending items that target the base regulation.

    Uses amending-article heading/instruction context to stay generic.
    """
    if not target_regulation:
        return True

    heading = normalize_text(item.get("article_heading", ""))
    instruction = normalize_text(item.get("source_instruction", ""))
    needle = target_regulation.lower()
    return needle in heading.lower() or needle in instruction.lower()


def identify_relevant_amendments(
    base_docx: Path,
    amending_html: Path,
    *,
    target_regulation: str | None = None,
) -> dict:
    """Phase 1: identify amendment provisions relevant to the base regulation."""
    title, recitals, items = collect_amending_blocks(amending_html)
    resolved_target = target_regulation or infer_target_regulation_from_base_docx(base_docx)
    relevant_items = [it for it in items if is_item_relevant_to_target_regulation(it, resolved_target)]
    relevant_items = coalesce_marker_body_items(relevant_items)
    relevant_items = normalize_amendment_items(relevant_items)
    relevant_items = backfill_targets_from_instruction(relevant_items)
    relevant_items = normalize_replacement_target_points(relevant_items)
    entry_into_force_block = extract_entry_into_force_signoff_block(amending_html)

    return {
        "created_at_utc": datetime.now(timezone.utc).isoformat(),
        "base_docx": str(base_docx),
        "amending_html": str(amending_html),
        "target_regulation": resolved_target,
        "amending_title": title,
        "recitals": recitals,
        "entry_into_force_block": entry_into_force_block,
        "detected_items": len(items),
        "relevant_items": len(relevant_items),
        "items": relevant_items,
    }


def extract_entry_into_force_signoff_block(amending_html: Path) -> dict:
    """Extract amending entry-into-force article and final signatory details."""
    soup = load_html(amending_html)

    article_heading_line = ""
    article_subheading_line = ""
    article_body: list[str] = []

    candidate_articles = soup.find_all(id=re.compile(r"^art_[A-Za-z0-9]+$"))
    chosen_article = None
    for art in candidate_articles:
        a1, a2 = article_heading_parts(art)
        heading_line = normalize_text(a1 or "")
        subtitle = normalize_text(a2 or "")

        # Only accept explicit entry-into-force article labels. Using broad
        # body-text heuristics can misclassify substantive amendment articles
        # (e.g. Article 1 "Amendments to ...") as entry-into-force blocks.
        if (
            re.search(r"\bentry\s+into\s+force\b", subtitle, flags=re.I)
            or re.search(r"\bentry\s+into\s+force\b", heading_line, flags=re.I)
        ):
            chosen_article = art
            break

    if chosen_article is not None:
        a1, a2 = article_heading_parts(chosen_article)
        article_heading_line = normalize_text(a1)
        article_subheading_line = normalize_text(a2 or "")
        article_body = [
            format_amendment_item_text(normalize_text(text_with_footnote_tokens(p)))
            for p in provision_paragraphs(chosen_article)
        ]
        article_body = [line for line in article_body if line]

    final_lines: list[str] = []
    signatories: list[list[dict[str, object]]] = []
    final_block = soup.select_one("div.oj-final")
    if final_block is not None:
        for child in final_block.find_all(recursive=False):
            if not isinstance(child, Tag):
                continue

            classes = child.get("class", [])
            if isinstance(classes, str):
                classes = classes.split()

            if child.name == "p" and "oj-normal" in classes:
                txt = normalize_text(text_with_footnote_tokens(child))
                if txt and not is_instructional_amendment_line(txt):
                    final_lines.append(txt)
                continue

            if child.name == "div" and "oj-signatory" in classes:
                block: list[dict[str, object]] = []
                for p in child.find_all("p", class_="oj-signatory"):
                    txt = normalize_text(text_with_footnote_tokens(p))
                    if not txt:
                        continue
                    block.append({"text": txt, "italic": bool(p.select_one(".oj-italic"))})
                if block:
                    signatories.append(block)

    return {
        "article_heading_line": article_heading_line,
        "article_subheading_line": article_subheading_line,
        "article_body": article_body,
        "final_lines": final_lines,
        "signatories": signatories,
    }


def extract_inserted_annex_lines(amending_html: Path, annex_label: str = "VA") -> list[str]:
    """Extract inserted Annex content lines from the amending regulation HTML."""
    soup = load_html(amending_html)
    target = annex_label.upper()

    container = soup.find(id=re.compile(r"^anx_[A-Za-z0-9]+$"))
    if not isinstance(container, Tag):
        return []

    started = False
    raw_lines: list[str] = []
    heading_re = re.compile(rf"^[‘'\"]?ANNEX\s+{re.escape(target)}\b", flags=re.I)

    for p in container.find_all("p"):
        txt = normalize_text(text_with_footnote_tokens(p))
        if not txt:
            continue
        if not started:
            if not heading_re.match(txt):
                continue
            started = True
            raw_lines.append(f"ANNEX {target}")
            continue
        raw_lines.append(txt)

    if not raw_lines:
        return []

    # Coalesce split table rows such as "(a)" + "name;" -> "(a)\tname;".
    merged: list[str] = []
    i = 0
    while i < len(raw_lines):
        cur = normalize_text(raw_lines[i])
        if is_standalone_list_marker_text(cur) and i + 1 < len(raw_lines):
            nxt = normalize_text(raw_lines[i + 1])
            if nxt and not is_standalone_list_marker_text(nxt):
                merged.append(f"{cur}\t{nxt}")
                i += 2
                continue
        merged.append(cur)
        i += 1

    # Strip opening/closing quotation marks that wrap inserted annex blocks.
    cleaned: list[str] = []
    for idx, line in enumerate(merged):
        txt = line
        if idx == 0:
            txt = re.sub(r"^[‘'\"]+", "", txt)
        txt = re.sub(r"[’'\"]$", "", txt)
        txt = normalize_text(txt)
        if txt:
            cleaned.append(txt)

    return cleaned


def annex_heading_label(text: str) -> str | None:
    probe = normalize_quote_wrapping(normalize_text(text))
    m = re.match(r"^ANNEX\s+([IVXLC]+[A-Za-z]?)\b", probe, flags=re.I)
    return m.group(1).upper() if m else None


def clean_annex_block_lines(raw_lines: list[str]) -> list[str]:
    if not raw_lines:
        return []

    merged: list[str] = []
    i = 0
    while i < len(raw_lines):
        cur = normalize_text(raw_lines[i])
        standalone_annex_marker = bool(
            re.fullmatch(
                r"(?:\(([a-z]+|[ivxlcdm]+)\)|[IVXLC]+\.|[A-Z]\.|\d+[A-Za-z]?[\.)])",
                cur,
                flags=re.I,
            )
        )
        if standalone_annex_marker and i + 1 < len(raw_lines):
            nxt = normalize_text(raw_lines[i + 1])
            nxt_standalone_marker = bool(
                re.fullmatch(
                    r"(?:\(([a-z]+|[ivxlcdm]+)\)|[IVXLC]+\.|[A-Z]\.|\d+[A-Za-z]?[\.)])",
                    nxt,
                    flags=re.I,
                )
            )
            if nxt and not nxt_standalone_marker:
                merged.append(f"{cur}\t{nxt}")
                i += 2
                continue
        merged.append(cur)
        i += 1

    cleaned: list[str] = []
    for idx, line in enumerate(merged):
        txt = line
        if idx == 0:
            txt = re.sub(r"^[‘'\"“”]+", "", txt)
        txt = re.sub(r"[’'\"“”]+$", "", txt)
        txt = normalize_text(txt)
        if txt:
            cleaned.append(txt)
    return cleaned


def format_annex_line_text(line: str) -> str:
    txt = normalize_text(line)
    if not txt:
        return ""
    txt = normalise_marker_spacing(txt, annex_mode=True)
    txt = re.sub(r"^\(([a-z]+)\)\s+", r"(\1)\t", txt, flags=re.I)
    txt = re.sub(r"^([IVXLC]+\.)\s+", r"\1\t", txt, flags=re.I)
    txt = re.sub(r"^([A-Z]\.)\s+", r"\1\t", txt)
    txt = re.sub(r"^(\d+[A-Za-z]?\))\s+", r"\1\t", txt)
    txt = re.sub(r"^(\d+[A-Za-z]?\.)\s+", r"\1\t", txt, flags=re.I)
    return txt


def normalize_annex_lines_for_comparison(lines: list[str]) -> list[str]:
    """Normalize parsed annex lines so comparison remains paragraph-aligned.

    Some source annexes encode heading + title as one paragraph with a line
    break, while amendment HTML can split them into two lines.
    """
    prepared = [format_annex_line_text(x) for x in lines if normalize_text(x)]
    if len(prepared) < 2:
        return prepared

    first = prepared[0]
    second = prepared[1]
    if re.match(r"^ANNEX\s+[IVXLC]+[A-Za-z]?\b", first, flags=re.I):
        second_is_structural = bool(
            re.match(r"^(?:[IVXLC]+\.|[A-Z]\.|\d+[A-Za-z]?[\.)]|\([a-zivxlcdm]+\))(?=\s|\t|$)", second, flags=re.I)
        )
        if not second_is_structural:
            return [f"{first}\n{second}"] + prepared[2:]

    return prepared


def annex_alignment_key(line: str) -> tuple[str, str]:
    """Return a structural alignment key for annex lines.

    The key is marker-first (ANNEX heading, roman heading, numeric heading,
    point marker) so minor wording shifts do not desynchronise the whole block.
    """
    txt = normalize_text(strip_src_artifacts(line or ""))
    if not txt:
        return ("empty", "")

    heading = annex_heading_label(txt)
    if heading:
        return ("annex_heading", heading)

    m_roman = re.match(r"^([IVXLC]+)\.(?=\s|\t|$)", txt, flags=re.I)
    if m_roman:
        return ("roman_heading", m_roman.group(1).upper())

    m_num = re.match(r"^(\d+[A-Za-z]?)\.(?=\s|\t|$)", txt, flags=re.I)
    if m_num:
        return ("numeric_heading", m_num.group(1).lower())

    m_num_paren = re.match(r"^(\d+[A-Za-z]?)\)(?=\s|\t|$)", txt, flags=re.I)
    if m_num_paren:
        return ("numeric_heading", m_num_paren.group(1).lower())

    m_alpha = re.match(r"^([A-Z])\.(?=\s|\t|$)", txt)
    if m_alpha:
        return ("alpha_heading", m_alpha.group(1).upper())

    m_point = re.match(r"^\(([a-z]+|[ivxlcdm]+)\)(?=\s|\t|$)", txt, flags=re.I)
    if m_point:
        return ("point", m_point.group(1).lower())

    # Fallback body key: compact alnum-only text to absorb punctuation noise.
    body_key = re.sub(r"[^A-Za-z0-9]+", " ", txt).strip().lower()
    return ("body", body_key)


def find_annex_alignment_match(
    old_keys: list[tuple[str, str]],
    new_keys: list[tuple[str, str]],
    old_idx: int,
    new_idx: int,
    window: int = 10,
) -> tuple[int | None, int | None]:
    """Find nearest forward sync point in old/new key streams."""
    old_hit: int | None = None
    new_hit: int | None = None

    max_old = min(len(old_keys), old_idx + window)
    max_new = min(len(new_keys), new_idx + window)

    for oi in range(old_idx + 1, max_old):
        if old_keys[oi] == new_keys[new_idx]:
            old_hit = oi
            break

    for nj in range(new_idx + 1, max_new):
        if new_keys[nj] == old_keys[old_idx]:
            new_hit = nj
            break

    return old_hit, new_hit


def build_annex_contextual_keys(lines: list[str]) -> list[tuple[str, str]]:
    """Build alignment keys with local structural scope.

    Repeated markers like A./B./E. appear in multiple annex sections. We scope
    those keys to the latest roman/numeric heading so matching stays local.
    """
    keys: list[tuple[str, str]] = []
    current_roman: str = ""
    current_numeric: str = ""

    for line in lines:
        kind, value = annex_alignment_key(line)

        if kind == "roman_heading":
            current_roman = value
            current_numeric = ""
            keys.append((kind, value))
            continue

        if kind == "numeric_heading":
            current_numeric = value
            scoped_value = f"{current_roman}>{value}" if current_roman else value
            keys.append((kind, scoped_value))
            continue

        if kind in {"alpha_heading", "point", "body"}:
            scope_parts = [part for part in (current_roman, current_numeric) if part]
            scope = ">".join(scope_parts)
            scoped_value = f"{scope}>{value}" if scope else value
            keys.append((kind, scoped_value))
            continue

        keys.append((kind, value))

    return keys


def build_annex_alignment_units(lines: list[str]) -> list[str]:
    """Build stable comparison units for annex line alignment."""
    keys = build_annex_contextual_keys(lines)
    units: list[str] = []
    for (kind, value), line in zip(keys, lines):
        txt = normalize_text(strip_src_artifacts(line or "")).lower()
        txt = re.sub(r"\s+", " ", txt)
        txt_sig = re.sub(r"[^a-z0-9 ]+", "", txt)[:120]
        units.append(f"{kind}|{value}|{txt_sig}")
    return units


def extract_annex_blocks_from_amending_html(amending_html: Path) -> list[dict[str, object]]:
    """Extract full Annex replacement/addition blocks from amending HTML annex container."""
    soup = load_html(amending_html)
    container = soup.find(id=re.compile(r"^anx_[A-Za-z0-9]+$"))
    if not isinstance(container, Tag):
        return []

    blocks: list[dict[str, object]] = []
    mode: str | None = None
    current_label: str | None = None
    current_lines: list[str] = []

    def flush_current() -> None:
        nonlocal current_label, current_lines
        if not current_label:
            return
        cleaned = clean_annex_block_lines(current_lines)
        if cleaned:
            blocks.append(
                {
                    "mode": mode,
                    "annex_label": current_label,
                    "lines": cleaned,
                }
            )
        current_label = None
        current_lines = []

    for p in container.find_all("p"):
        txt = normalize_text(text_with_footnote_tokens(p))
        if not txt:
            continue

        if re.search(r"\bannexes?\s+[ivxlcdm0-9\s,toand]+\s+are\s+replaced\s+by\s+the\s+following\b", txt, flags=re.I):
            flush_current()
            mode = "replace"
            continue

        if re.search(r"\bthe\s+following\s+annex(?:es)?\s+are\s+added\b", txt, flags=re.I):
            flush_current()
            mode = "add"
            continue

        heading = annex_heading_label(txt)
        if heading:
            flush_current()
            if mode in {"replace", "add"}:
                current_label = heading
                current_lines = [f"ANNEX {heading}"]
            continue

        if current_label:
            current_lines.append(txt)

    flush_current()
    return blocks


def find_annex_section_bounds_by_label(doc: Document, annex_label: str) -> tuple[int, int] | None:
    start_idx = annex_heading_index(doc, annex_label)
    if start_idx is None:
        return None

    end_idx = len(doc.paragraphs)
    for idx in range(start_idx + 1, len(doc.paragraphs)):
        if annex_heading_label(doc.paragraphs[idx].text) is not None:
            end_idx = idx
            break
    return start_idx, end_idx


def find_last_annex_section_end_index(doc: Document) -> int:
    last_heading_idx: int | None = None
    for idx, para in enumerate(doc.paragraphs):
        if annex_heading_label(para.text) is not None:
            last_heading_idx = idx
    if last_heading_idx is None:
        return len(doc.paragraphs)

    bounds = find_annex_section_bounds_by_label(doc, annex_heading_label(doc.paragraphs[last_heading_idx].text) or "")
    if bounds is None:
        return len(doc.paragraphs)
    _, end_idx = bounds
    return end_idx


def insert_annex_block_lines(doc: Document, insertion_index: int, annex_lines: list[str]) -> int:
    """Insert a full annex block with Stage 1-like heading/body formatting."""
    if not annex_lines:
        return 0

    inserted_count = 0

    heading_txt = normalize_text(annex_lines[0])
    heading_label = annex_heading_label(heading_txt)
    heading_para = insert_plain_paragraph_before_index(doc, insertion_index)
    insertion_index += 1
    inserted_count += 1
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_para.paragraph_format.left_indent = Cm(0)
    heading_para.paragraph_format.first_line_indent = None
    heading_para.paragraph_format.page_break_before = True

    heading_run = heading_para.add_run(f"ANNEX {heading_label}" if heading_label else heading_txt)
    heading_run.bold = False
    heading_run.italic = True
    style_inserted_run_no_bold_change(heading_run)
    heading_run.italic = True

    line_idx = 1
    subtitle = normalize_text(annex_lines[1]) if len(annex_lines) > 1 else ""
    if subtitle and subtitle.isupper() and len(subtitle.split()) >= 3:
        heading_run.add_break(WD_BREAK.LINE)
        subtitle_run = heading_para.add_run(subtitle)
        subtitle_run.bold = True
        subtitle_run.italic = False
        style_inserted_run_no_bold_change(subtitle_run)
        subtitle_run.bold = True
        subtitle_run.italic = False
        line_idx = 2

    current_level = 1
    last_body_level = 1
    last_point_level: int | None = None
    previous_was_point = False

    for line in annex_lines[line_idx:]:
        p = insert_plain_paragraph_before_index(doc, insertion_index)
        insertion_index += 1
        inserted_count += 1

        txt = format_annex_line_text(line)
        marker = first_top_level_paragraph_marker(txt)
        point_marker = first_top_level_point_marker(txt)
        annex_heading_marker = bool(re.match(r"^[IVXLC]+\.(?=\s|\t|$)", txt, flags=re.I))
        digit_marker = marker is not None and not annex_heading_marker

        if annex_heading_marker:
            level = 1
            last_point_level = None
        elif digit_marker:
            level = 2
            last_point_level = None
        elif point_marker is not None:
            if last_point_level is None:
                level = 3 if current_level >= 2 else 2
            else:
                level = last_point_level
        else:
            level = current_level

        reset_to_body_after_point_series = bool(
            point_marker is None
            and previous_was_point
            and re.match(r"^[A-Z]", txt)
            and not re.match(r"^(and|or)\b", txt, flags=re.I)
        )

        if marker is not None or annex_heading_marker:
            force_paragraph_justification(p)
            p.paragraph_format.left_indent = Cm(level)
            p.paragraph_format.first_line_indent = Cm(-1)
            current_level = level
            last_body_level = level
            previous_was_point = False
        elif point_marker is not None:
            force_paragraph_justification(p)
            p.paragraph_format.left_indent = Cm(level)
            p.paragraph_format.first_line_indent = Cm(-1)
            current_level = level
            last_point_level = level
            previous_was_point = True
        else:
            if reset_to_body_after_point_series:
                level = max(1, last_body_level)
            force_paragraph_justification(p)
            p.paragraph_format.left_indent = Cm(level)
            p.paragraph_format.first_line_indent = None
            current_level = level
            last_body_level = level
            previous_was_point = False

        run = p.add_run(txt)
        style_inserted_run_no_bold_change(run)

    return inserted_count


def apply_annex_blocks_from_amending_html(doc: Document, amending_html: Path) -> dict[str, int]:
    """Apply annex replacement/addition blocks parsed from amending HTML annex container."""
    blocks = extract_annex_blocks_from_amending_html(amending_html)
    replaced = 0
    added = 0

    for block in blocks:
        mode = normalize_text(str(block.get("mode") or "")).lower()
        label = normalize_text(str(block.get("annex_label") or "")).upper()
        lines = [normalize_text(str(x)) for x in block.get("lines", []) if normalize_text(str(x))]
        if not label or not lines:
            continue

        if mode == "replace":
            bounds = find_annex_section_bounds_by_label(doc, label)
            if bounds is not None:
                start_idx, end_idx = bounds
                existing_paras = list(doc.paragraphs[start_idx:end_idx])
                new_lines = normalize_annex_lines_for_comparison(lines)

                old_lines = [normalize_text(strip_src_artifacts(p.text)) for p in existing_paras]
                old_keys = build_annex_contextual_keys(old_lines)
                new_keys = build_annex_contextual_keys(new_lines)
                def insert_new_line_at_current_anchor(new_text: str, old_pos: int) -> None:
                    anchor_para = existing_paras[old_pos] if old_pos < len(existing_paras) else None
                    if anchor_para is not None:
                        insertion_index = find_current_paragraph_index(doc, anchor_para)
                        if insertion_index is None:
                            insertion_index = find_annex_section_bounds_by_label(doc, label)[1]
                    else:
                        insertion_index = find_annex_section_bounds_by_label(doc, label)[1]

                    p = insert_plain_paragraph_before_index(doc, insertion_index)

                    if re.match(r"^ANNEX\s+[IVXLC]+[A-Za-z]?\b", new_text, flags=re.I):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.left_indent = Cm(0)
                        p.paragraph_format.first_line_indent = None
                    else:
                        marker = first_top_level_paragraph_marker(new_text)
                        point_marker = first_top_level_point_marker(new_text)
                        annex_heading_marker = bool(re.match(r"^[IVXLC]+\.(?=\s|\t|$)", new_text, flags=re.I))
                        if marker is not None or annex_heading_marker or point_marker is not None:
                            p.paragraph_format.left_indent = Cm(1 if annex_heading_marker else 2)
                            p.paragraph_format.first_line_indent = Cm(-1)
                        else:
                            p.paragraph_format.left_indent = Cm(1)
                            p.paragraph_format.first_line_indent = None

                    run = p.add_run(new_text)
                    style_inserted_run_no_bold_change(run)

                sm_keys = SequenceMatcher(None, old_keys, new_keys)
                for op, i1, i2, j1, j2 in sm_keys.get_opcodes():
                    if op == "equal":
                        for old_idx, new_idx in zip(range(i1, i2), range(j1, j2)):
                            para = existing_paras[old_idx]
                            old_text = strip_src_artifacts(para.text)
                            new_text = new_lines[new_idx]
                            if normalize_text(old_text) != normalize_text(new_text):
                                replace_paragraph_with_revision_marks(para, old_text, new_text)
                        continue

                    if op == "delete":
                        for old_idx in range(i1, i2):
                            para = existing_paras[old_idx]
                            old_text = strip_src_artifacts(para.text)
                            if normalize_text(old_text):
                                replace_paragraph_with_revision_marks(para, old_text, "")
                        continue

                    if op == "insert":
                        for new_idx in range(j1, j2):
                            insert_new_line_at_current_anchor(new_lines[new_idx], i1)
                        continue

                    # Replace: pair only by local index where possible,
                    # then treat residual tails as explicit delete/insert.
                    old_slice_keys = old_keys[i1:i2]
                    new_slice_keys = new_keys[j1:j2]
                    overlap = len(set(old_slice_keys) & set(new_slice_keys))
                    baseline = min(len(old_slice_keys), len(new_slice_keys))
                    overlap_ratio = (overlap / baseline) if baseline else 0.0

                    if overlap_ratio < 0.4:
                        for old_idx in range(i1, i2):
                            para = existing_paras[old_idx]
                            old_text = strip_src_artifacts(para.text)
                            if normalize_text(old_text):
                                replace_paragraph_with_revision_marks(para, old_text, "")
                        for new_idx in range(j1, j2):
                            insert_new_line_at_current_anchor(new_lines[new_idx], i1)
                        continue

                    paired = min(i2 - i1, j2 - j1)
                    for delta in range(paired):
                        old_idx = i1 + delta
                        new_idx = j1 + delta
                        if old_keys[old_idx] == new_keys[new_idx]:
                            para = existing_paras[old_idx]
                            old_text = strip_src_artifacts(para.text)
                            new_text = new_lines[new_idx]
                            replace_paragraph_with_revision_marks(para, old_text, new_text)
                        else:
                            para = existing_paras[old_idx]
                            old_text = strip_src_artifacts(para.text)
                            if normalize_text(old_text):
                                replace_paragraph_with_revision_marks(para, old_text, "")
                            insert_new_line_at_current_anchor(new_lines[new_idx], old_idx)

                    if (i2 - i1) > paired:
                        for old_idx in range(i1 + paired, i2):
                            para = existing_paras[old_idx]
                            old_text = strip_src_artifacts(para.text)
                            if normalize_text(old_text):
                                replace_paragraph_with_revision_marks(para, old_text, "")

                    if (j2 - j1) > paired:
                        for new_idx in range(j1 + paired, j2):
                            insert_new_line_at_current_anchor(new_lines[new_idx], i2)
                replaced += 1
                continue

            insertion_index = find_last_annex_section_end_index(doc)
            insert_annex_block_lines(doc, insertion_index, lines)
            replaced += 1
            continue

        if mode == "add":
            if annex_heading_exists(doc, label):
                continue
            insertion_index = find_last_annex_section_end_index(doc)
            insert_annex_block_lines(doc, insertion_index, lines)
            added += 1

    return {"replaced": replaced, "added": added}


def find_annex_va_insertion_index(doc: Document) -> int:
    """Return insertion index for Annex Va (before Annex VI when present)."""
    for idx, para in enumerate(doc.paragraphs):
        if annex_heading_number(para.text) == "VI":
            return idx

    bounds_v = find_annex_section_bounds(doc, "V")
    if bounds_v is not None:
        _, end_idx = bounds_v
        return end_idx

    return len(doc.paragraphs)


def insert_annex_va_from_amending_html(doc: Document, amending_html: Path) -> int:
    """Insert Annex Va block from amendment HTML when instructed by amendment."""
    if any(re.match(r"^ANNEX\s+VA\b", normalize_text(p.text), flags=re.I) for p in doc.paragraphs):
        return 0

    annex_lines = extract_inserted_annex_lines(amending_html, annex_label="VA")
    if not annex_lines:
        return 0

    insertion_index = find_annex_va_insertion_index(doc)
    inserted_count = 0

    # Stage 1-like Annex heading block: centered, page-break-before, with a
    # manual line break between annex number and title line.
    heading_para = insert_plain_paragraph_before_index(doc, insertion_index)
    insertion_index += 1
    inserted_count += 1
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_para.paragraph_format.left_indent = Cm(0)
    heading_para.paragraph_format.first_line_indent = None
    heading_para.paragraph_format.page_break_before = True

    heading_run = heading_para.add_run("ANNEX Va")
    heading_run.bold = False
    heading_run.italic = True
    style_inserted_run_no_bold_change(heading_run)
    heading_run.italic = True

    line_idx = 1
    subtitle = normalize_text(annex_lines[1]) if len(annex_lines) > 1 else ""
    if subtitle and subtitle.isupper() and len(subtitle.split()) >= 3:
        heading_run.add_break(WD_BREAK.LINE)
        subtitle_run = heading_para.add_run(subtitle)
        subtitle_run.bold = True
        subtitle_run.italic = False
        style_inserted_run_no_bold_change(subtitle_run)
        subtitle_run.bold = True
        subtitle_run.italic = False
        line_idx = 2

    current_level = 1
    last_body_level = 1
    last_point_level: int | None = None
    previous_was_point = False

    for line in annex_lines[line_idx:]:
        p = insert_plain_paragraph_before_index(doc, insertion_index)
        insertion_index += 1
        inserted_count += 1

        txt = format_annex_line_text(line)
        marker = first_top_level_paragraph_marker(txt)
        point_marker = first_top_level_point_marker(txt)
        annex_heading_marker = bool(re.match(r"^[IVXLC]+\.(?=\s|\t|$)", txt, flags=re.I))
        digit_marker = marker is not None and not annex_heading_marker

        if annex_heading_marker:
            level = 1
            last_point_level = None
        elif digit_marker:
            level = 2
            last_point_level = None
        elif point_marker is not None:
            if last_point_level is None:
                level = 3 if current_level >= 2 else 2
            else:
                level = last_point_level
        else:
            level = current_level

        reset_to_body_after_point_series = bool(
            point_marker is None
            and previous_was_point
            and re.match(r"^[A-Z]", txt)
            and not re.match(r"^(and|or)\b", txt, flags=re.I)
        )

        if marker is not None or annex_heading_marker:
            # Annex section headers like "I. Summary" follow stage1 hanging style.
            force_paragraph_justification(p)
            p.paragraph_format.left_indent = Cm(level)
            p.paragraph_format.first_line_indent = Cm(-1)
            current_level = level
            last_body_level = level
            previous_was_point = False
        elif point_marker is not None:
            force_paragraph_justification(p)
            p.paragraph_format.left_indent = Cm(level)
            p.paragraph_format.first_line_indent = Cm(-1)
            current_level = level
            last_point_level = level
            previous_was_point = True
        else:
            if reset_to_body_after_point_series:
                level = max(1, last_body_level)
            # Annex body paragraphs remain indented within the annex section.
            force_paragraph_justification(p)
            p.paragraph_format.left_indent = Cm(level)
            p.paragraph_format.first_line_indent = None
            current_level = level
            last_body_level = level
            previous_was_point = False

        run = p.add_run(txt)
        style_inserted_run_no_bold_change(run)

    normalize_annex_va_point_levels(doc)

    return inserted_count


def normalize_annex_va_point_levels(doc: Document) -> None:
    """Keep consecutive Annex Va point-marker siblings at the same indent level."""
    annex_idx = annex_heading_index(doc, "VA")
    if annex_idx is None:
        return

    next_annex_idx = len(doc.paragraphs)
    for idx in range(annex_idx + 1, len(doc.paragraphs)):
        if idx > annex_idx and annex_heading_number(doc.paragraphs[idx].text) is not None:
            next_annex_idx = idx
            break

    active_point_indent = None
    for idx in range(annex_idx + 1, next_annex_idx):
        paragraph = doc.paragraphs[idx]
        text = normalize_text(paragraph.text)
        if not text:
            continue

        if re.match(r"^[IVXLC]+\.(?=\s|\t|$)", text, flags=re.I) or first_top_level_paragraph_marker(text) is not None:
            active_point_indent = None
            continue

        point_marker = first_top_level_point_marker(text)
        if point_marker is None:
            continue

        if active_point_indent is None:
            active_point_indent = paragraph.paragraph_format.left_indent
            continue

        paragraph.paragraph_format.left_indent = active_point_indent
        paragraph.paragraph_format.first_line_indent = Cm(-1)


def annex_heading_exists(doc: Document, annex_label: str) -> bool:
    label = normalize_text(annex_label).upper()
    if not label:
        return False
    return any(
        re.match(rf"^ANNEX\s+{re.escape(label)}\b", normalize_text(p.text), flags=re.I)
        for p in doc.paragraphs
    )


def expected_inserted_annex_labels(amending_items: list[dict[str, str]]) -> set[str]:
    labels: set[str] = set()
    pat = re.compile(r"\b(?:is|are)\s+inserted\s+as\s+annex\s+([A-Za-z0-9]+)\b", flags=re.I)
    for item in amending_items:
        instruction = normalize_text(str(item.get("source_instruction") or ""))
        if not instruction:
            continue
        m = pat.search(instruction)
        if m:
            labels.add(m.group(1).upper())
    return labels


def assert_inserted_annexes_present(doc: Document, amending_items: list[dict[str, str]]) -> None:
    expected = expected_inserted_annex_labels(amending_items)
    if not expected:
        return

    missing = sorted(label for label in expected if not annex_heading_exists(doc, label))
    if missing:
        missing_text = ", ".join(f"ANNEX {label}" for label in missing)
        raise RuntimeError(
            f"QA check failed: amendment instructs insertion of {missing_text}, but heading not found in output DOCX."
        )


def annex_heading_index(doc: Document, annex_label: str) -> int | None:
    label = normalize_text(annex_label).upper()
    if not label:
        return None
    for idx, p in enumerate(doc.paragraphs):
        if re.match(rf"^ANNEX\s+{re.escape(label)}\b", normalize_text(p.text), flags=re.I):
            return idx
    return None


def roman_to_int(roman: str) -> int:
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100}
    total = 0
    prev = 0
    for ch in reversed(roman.upper()):
        val = values.get(ch, 0)
        if val < prev:
            total -= val
        else:
            total += val
            prev = val
    return total


def int_to_roman(value: int) -> str:
    table = [
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    n = value
    out: list[str] = []
    for arabic, numeral in table:
        while n >= arabic:
            out.append(numeral)
            n -= arabic
    return "".join(out)


def next_roman_label(label: str) -> str | None:
    txt = normalize_text(label).upper()
    if not txt or not re.fullmatch(r"[IVXLC]+", txt):
        return None
    n = roman_to_int(txt)
    if n <= 0:
        return None
    return int_to_roman(n + 1)


def assert_inserted_annex_order(doc: Document, amending_items: list[dict[str, str]]) -> None:
    """Assert inserted lettered annexes are placed in legal sequence.

    Example: if Annex VA is expected, it must be after Annex V and before Annex VI
    when those neighboring headings exist in the document.
    """
    expected = expected_inserted_annex_labels(amending_items)
    if not expected:
        return

    violations: list[str] = []
    for label in sorted(expected):
        m = re.fullmatch(r"([IVXLC]+)([A-Z]+)", label)
        if not m:
            continue
        base_roman = m.group(1)
        suffix = m.group(2)
        if suffix != "A":
            continue

        label_idx = annex_heading_index(doc, label)
        if label_idx is None:
            continue

        base_idx = annex_heading_index(doc, base_roman)
        if base_idx is not None and label_idx <= base_idx:
            violations.append(f"ANNEX {label} must appear after ANNEX {base_roman}")

        next_roman = next_roman_label(base_roman)
        if next_roman:
            next_idx = annex_heading_index(doc, next_roman)
            if next_idx is not None and label_idx >= next_idx:
                violations.append(f"ANNEX {label} must appear before ANNEX {next_roman}")

    if violations:
        raise RuntimeError("QA check failed: " + "; ".join(violations) + ".")


def contains_contiguous_token_sequence(haystack: list[str], needle: list[str]) -> bool:
    if not needle:
        return False
    n = len(needle)
    if n > len(haystack):
        return False
    first = needle[0]
    max_start = len(haystack) - n
    for i in range(max_start + 1):
        if haystack[i] != first:
            continue
        if haystack[i:i + n] == needle:
            return True
    return False


def has_probe_match_with_backoff(
    haystack_tokens: list[str],
    probe_tokens: list[str],
    *,
    fallback_lengths: list[int],
) -> tuple[bool, int]:
    if not probe_tokens:
        return False, 0

    lengths: list[int] = []
    seen: set[int] = set()
    for length in [len(probe_tokens)] + fallback_lengths:
        trimmed = min(length, len(probe_tokens))
        if trimmed <= 0 or trimmed in seen:
            continue
        seen.add(trimmed)
        lengths.append(trimmed)

    for length in lengths:
        if contains_contiguous_token_sequence(haystack_tokens, probe_tokens[:length]):
            return True, length
    return False, lengths[-1] if lengths else 0


def contains_ordered_token_subsequence(haystack: list[str], needle: list[str]) -> bool:
    if not needle:
        return False
    h_index = 0
    h_len = len(haystack)
    for token in needle:
        while h_index < h_len and haystack[h_index] != token:
            h_index += 1
        if h_index >= h_len:
            return False
        h_index += 1
    return True


def summarize_one_go_application_check(analysis_items: list[dict], doc: Document | None = None) -> dict[str, object]:
    """Summarize whether all detected substantive amendment items were applied in one run.

    Accepted applied modes:
    - inserted
    - replaced
    - already_applied

    Items left as analysis_only are treated as fully applied when their amendment
    text is already present in output, and checks are evaluated per instruction
    group (target + source instruction) rather than raw row-level.
    """
    substantive: list[dict] = []
    for item in analysis_items:
        amend_text = normalize_text(item.get("amending_text", ""))
        if not amend_text:
            continue
        # Skip formatting scaffolding rows (e.g. isolated dash separators).
        if len(tokenize(amend_text)) == 0:
            continue
        substantive.append(item)
    docx_tokens: list[str] = []
    if doc is not None:
        docx_visible = normalize_text("\n".join(normalize_text(p.text) for p in doc.paragraphs))
        docx_tokens = tokenize(docx_visible)

    def group_key(item: dict) -> tuple[str, str, str, str, str, str, str]:
        return (
            normalize_text(str(item.get("article_number", ""))),
            normalize_text(str(item.get("target_article_number", ""))),
            normalize_text(str(item.get("target_paragraph_number", ""))),
            normalize_text(str(item.get("target_point_marker", ""))),
            normalize_text(str(item.get("target_annex_number", ""))),
            normalize_text(str(item.get("target_annex_point_marker", ""))),
            normalize_text(str(item.get("source_instruction", ""))),
        )

    grouped_items: dict[tuple[str, str, str, str, str, str, str], list[dict]] = {}
    for item in substantive:
        grouped_items.setdefault(group_key(item), []).append(item)

    unresolved_groups: list[list[dict]] = []
    unapplied: list[dict] = []
    for group in grouped_items.values():
        group_has_effective_application = False
        unresolved_in_group: list[dict] = []
        for item in group:
            if normalize_text(item.get("applied_mode", "")) != "analysis_only":
                group_has_effective_application = True
                continue
            amend_text = normalize_text(item.get("amending_text", ""))
            probe = tokenize(amend_text)[:24]
            if docx_tokens and probe:
                found, _ = has_probe_match_with_backoff(
                    docx_tokens,
                    probe,
                    fallback_lengths=[18, 12, 8],
                )
                if (not found) and len(probe) >= 8:
                    for length in [12, 8]:
                        trimmed = min(length, len(probe))
                        if trimmed <= 0:
                            continue
                        if contains_ordered_token_subsequence(docx_tokens, probe[:trimmed]):
                            found = True
                            break
                if found:
                    group_has_effective_application = True
                    continue
            unresolved_in_group.append(item)
        if not group_has_effective_application:
            unresolved_groups.append(group)
            unapplied.extend(unresolved_in_group if unresolved_in_group else group)

    examples: list[dict[str, str]] = []
    for item in unapplied[:25]:
        examples.append(
            {
                "article_number": str(item.get("article_number", "")),
                "target_article_number": str(item.get("target_article_number", "")),
                "source_instruction": clean_text(str(item.get("source_instruction", "")))[:220],
                "amending_text": clean_text(str(item.get("amending_text", "")))[:220],
            }
        )

    return {
        "passed": len(unresolved_groups) == 0,
        "substantive_items_total": len(substantive),
        "substantive_items_not_fully_applied": len(unapplied),
        "not_fully_applied_examples": examples,
    }


def infer_operation_type(item: dict) -> str:
    source_instruction = normalize_text(str(item.get("source_instruction", "")))
    amendment_kind = normalize_text(str(item.get("amendment_kind", "")))

    if amendment_kind == "replacement" or re.search(r"\breplaced\b", source_instruction, flags=re.I):
        return "replace"
    if amendment_kind == "deletion" or re.search(r"\bdeleted\b", source_instruction, flags=re.I):
        return "delete"
    if amendment_kind == "insertion" or re.search(r"\b(?:added|inserted)\b", source_instruction, flags=re.I):
        return "insert"
    return "unknown"


def build_operation_target_path(item: dict) -> str:
    target_annex = normalize_text(str(item.get("target_annex_number", "")))
    target_annex_point = normalize_text(str(item.get("target_annex_point_marker", "")))
    target_article = normalize_text(str(item.get("target_article_number", "")))
    target_paragraph = normalize_text(str(item.get("target_paragraph_number", "")))
    target_point = normalize_text(str(item.get("target_point_marker", "")))

    if target_annex:
        path = f"annex/{target_annex}"
        if target_annex_point:
            path += f"/point/{target_annex_point}"
        return path

    if target_article:
        path = f"article/{target_article}"
        if target_paragraph:
            path += f"/paragraph/{target_paragraph}"
        if target_point:
            path += f"/point/{target_point}"
        return path

    return "unresolved"


def build_operation_proof_report(analysis_items: list[dict]) -> dict[str, object]:
    operations: list[dict[str, object]] = []
    precondition_failures = 0
    unresolved_operations = 0
    precondition_failure_examples: list[dict[str, str]] = []

    for idx, item in enumerate(analysis_items):
        operation_type = infer_operation_type(item)
        applied_mode = normalize_text(str(item.get("applied_mode", "")))
        source_instruction = normalize_text(str(item.get("source_instruction", "")))
        amending_text = normalize_text(str(item.get("amending_text", "")))

        target_article = normalize_text(str(item.get("target_article_number", "")))
        target_paragraph = normalize_text(str(item.get("target_paragraph_number", "")))
        target_point = normalize_text(str(item.get("target_point_marker", "")))
        target_annex = normalize_text(str(item.get("target_annex_number", "")))
        target_annex_point = normalize_text(str(item.get("target_annex_point_marker", "")))

        preconditions: dict[str, bool] = {
            "has_amending_text": bool(amending_text),
            "has_scope_target": True,
            "has_required_paragraph_scope": True,
        }

        if operation_type in {"replace", "delete"}:
            preconditions["has_scope_target"] = bool(target_article or target_annex)

        if operation_type in {"replace", "delete"} and re.search(
            r"\b(?:first|second|third|fourth|fifth)\s+(?:sub)?paragraph\b",
            source_instruction,
            flags=re.I,
        ):
            preconditions["has_required_paragraph_scope"] = bool(target_paragraph)

        preconditions_passed = all(preconditions.values())
        mutation_proven = applied_mode in {"replaced", "inserted", "already_applied"}
        status = "proven" if (preconditions_passed and mutation_proven) else "unresolved"

        if not preconditions_passed:
            precondition_failures += 1
            if len(precondition_failure_examples) < 20:
                missing = [k for k, v in preconditions.items() if not v]
                precondition_failure_examples.append(
                    {
                        "operation_id": f"op_{idx + 1:04d}",
                        "operation_type": operation_type,
                        "target_path": build_operation_target_path(item),
                        "missing_preconditions": ", ".join(missing),
                        "source_instruction": clean_text(source_instruction)[:220],
                        "amending_text": clean_text(amending_text)[:220],
                    }
                )

        if status != "proven":
            unresolved_operations += 1

        operations.append(
            {
                "operation_id": f"op_{idx + 1:04d}",
                "operation_type": operation_type,
                "target_path": build_operation_target_path(item),
                "target": {
                    "article": target_article or None,
                    "paragraph": target_paragraph or None,
                    "point": target_point or None,
                    "annex": target_annex or None,
                    "annex_point": target_annex_point or None,
                },
                "source_instruction": clean_text(source_instruction),
                "amending_text": clean_text(amending_text),
                "best_match_score": item.get("best_match_score"),
                "applied_mode": applied_mode,
                "preconditions": preconditions,
                "preconditions_passed": preconditions_passed,
                "mutation_proven": mutation_proven,
                "status": status,
            }
        )

    return {
        "summary": {
            "operations_total": len(operations),
            "precondition_failures": precondition_failures,
            "unresolved_operations": unresolved_operations,
            "all_operations_proven": unresolved_operations == 0,
        },
        "precondition_failure_examples": precondition_failure_examples,
        "operations": operations,
    }


def find_binding_clause_index(doc: Document) -> int | None:
    for idx, para in enumerate(doc.paragraphs):
        txt = normalize_text(para.text)
        if re.match(r"^This\s+Regulation\s+shall\s+be\s+binding\s+in\s+its\s+entirety\b", txt, flags=re.I):
            return idx
    return None


def insert_table_before_index(doc: Document, insertion_index: int, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    if insertion_index < len(doc.paragraphs):
        anchor = doc.paragraphs[insertion_index]._element
        anchor.addprevious(table._element)
    return table


def insert_amending_entry_into_force_block(doc: Document, insertion_index: int, block: dict) -> None:
    heading = normalize_text(str(block.get("article_heading_line", "")))
    subheading = normalize_text(str(block.get("article_subheading_line", "")))
    body_lines = [
        ensure_leading_marker_tab(str(x))
        for x in block.get("article_body", [])
        if ensure_leading_marker_tab(str(x))
    ]

    if not any([heading, subheading, body_lines]):
        return

    # Idempotency guard for repeated Stage 3 runs on an already amended output.
    heading_blob = normalize_text(f"{heading}{subheading}") if (heading or subheading) else ""
    if heading_blob and any(normalize_text(p.text) == heading_blob for p in doc.paragraphs):
        return

    if heading:
        p = insert_plain_paragraph_before_index(doc, insertion_index)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Cm(0)
        r1 = p.add_run(heading)
        r1.bold = False
        r1.italic = True
        style_inserted_run_no_bold_change(r1)
        if subheading:
            r1.add_break()
            r2 = p.add_run(subheading)
            r2.bold = True
            r2.italic = False
            style_inserted_run_no_bold_change(r2)
        insertion_index += 1

    for line in body_lines:
        p = insert_plain_paragraph_before_index(doc, insertion_index)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        marker = first_top_level_paragraph_marker(line)
        point_marker = first_top_level_point_marker(line)
        if marker is not None:
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-1)
        elif point_marker is not None:
            p.paragraph_format.left_indent = Cm(2)
            p.paragraph_format.first_line_indent = Cm(-1)
        else:
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = None
        run = p.add_run(line)
        style_inserted_run_no_bold_change(run)
        insertion_index += 1

def insert_amending_final_signoff_block(doc: Document, insertion_index: int, block: dict) -> None:
    final_lines = [normalize_text(str(x)) for x in block.get("final_lines", []) if normalize_text(str(x))]
    signatories = block.get("signatories", []) if isinstance(block.get("signatories", []), list) else []

    if not any([final_lines, signatories]):
        return

    # Idempotency guard for repeated Stage 3 runs.
    marker_candidates: list[str] = []
    if len(final_lines) > 1:
        marker_candidates.extend(final_lines[1:])
    for block_lines in signatories:
        for line in block_lines:
            txt = normalize_text(str(line.get("text", "")))
            if txt:
                marker_candidates.append(txt)
    existing_texts = {normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)}
    if marker_candidates and any(m in existing_texts for m in marker_candidates):
        return

    spacer = insert_plain_paragraph_before_index(doc, insertion_index)
    spacer.paragraph_format.left_indent = Cm(0)
    insertion_index += 1

    for line in final_lines:
        p = insert_plain_paragraph_before_index(doc, insertion_index)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(line)
        style_inserted_run_no_bold_change(run)
        insertion_index += 1

    if signatories:
        table = insert_table_before_index(doc, insertion_index, rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        tbl = table._tbl
        tbl_pr = tbl.tblPr
        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"), str(int(Cm(2).emu / 635)))
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_ind)

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            borders.append(border)
        tbl_pr.append(borders)

        col_width = Cm(6.5)
        for col in table.columns:
            col.width = col_width
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        for col_idx in range(2):
            cell = table.cell(0, col_idx)
            cell.text = ""
            block_lines = signatories[col_idx] if col_idx < len(signatories) else []
            for line_idx, line in enumerate(block_lines):
                txt = normalize_text(str(line.get("text", "")))
                if not txt:
                    continue
                is_italic = bool(line.get("italic", False))
                p = cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(txt)
                run.italic = is_italic
                style_inserted_run_no_bold_change(run)


def find_title_index(doc: Document) -> int:
    for idx, para in enumerate(doc.paragraphs):
        if normalize_text(para.text):
            return idx
    return 0


def find_title_block_end_index(doc: Document) -> int:
    """Return the index after the current opening title block.

    The title block is treated as the initial consecutive run of title-like
    paragraphs at the top of the document. This lets later Stage 3 runs append
    a new amendment title below previous titles instead of inserting it above
    them.
    """
    title_like_pattern = re.compile(r"^(REGULATION|DIRECTIVE|DECISION)\b", flags=re.I)
    seen_title = False
    last_title_idx = -1
    for idx, para in enumerate(doc.paragraphs):
        txt = normalize_text(para.text)
        if not txt:
            if seen_title:
                continue
            continue
        if title_like_pattern.match(txt):
            seen_title = True
            last_title_idx = idx
            continue
        if seen_title:
            break
        return idx
    return last_title_idx + 1 if last_title_idx >= 0 else 0


def find_last_recital_index(doc: Document) -> int | None:
    last_idx: int | None = None
    for idx, para in enumerate(doc.paragraphs):
        if RECITAL_RE.match(normalize_text(para.text)):
            last_idx = idx
    return last_idx


def find_adoption_formula_index(doc: Document) -> int | None:
    for idx, para in enumerate(doc.paragraphs):
        txt = normalize_text(para.text)
        if re.match(r"^HAVE\s+ADOPTED\s+THIS\s+REGULATION\s*:?$", txt, flags=re.I):
            return idx
    return None


def insert_plain_paragraph_before_index(doc: Document, insertion_index: int, style_name: str | None = None):
    if insertion_index < len(doc.paragraphs):
        anchor = doc.paragraphs[insertion_index]
        return anchor.insert_paragraph_before(style=style_name)
    if style_name:
        return doc.add_paragraph(style=style_name)
    return doc.add_paragraph()


def find_previous_top_level_paragraph_index(doc: Document, section_start: int, insertion_index: int) -> int | None:
    for idx in range(insertion_index - 1, section_start - 1, -1):
        if first_top_level_paragraph_marker(doc.paragraphs[idx].text) is not None:
            return idx
    return None


def clone_paragraph_layout(dst, src) -> None:
    dst.style = src.style
    dst.alignment = src.alignment
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.right_indent = src.paragraph_format.right_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def normalize_quote_wrapping(text: str) -> str:
    text = text.strip()
    text = text.lstrip("‘\"\'")
    text = re.sub(r"[’\"\']\s*;?$", "", text)
    text = re.sub(r"[’\"\']\.$", ".", text)
    text = re.sub(r"[’\"\']+\s*;\s*$", "", text)
    text = re.sub(r"(?<!\.)\.\.(?!\.)$", ".", text)
    return text.strip()


def normalize_amendment_quoted_text(text: str) -> str:
    text = normalize_quote_wrapping(text)
    text = re.sub(r"\[\[FN:\d+\]\]", "", text)
    text = clean_text(text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    return text


def strip_src_markers_from_docx(docx_path: Path) -> None:
    """Remove internal [[SRC:...]] anchors from final user-facing DOCX."""
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)

        document_xml = tmp / "word" / "document.xml"
        root = etree.fromstring(document_xml.read_bytes(), etree.XMLParser(remove_blank_text=False))
        changed = False
        for t in root.xpath(".//w:t", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            if t.text and "[[SRC:" in t.text:
                new_text = re.sub(r"\s*\[\[SRC:[^\]]+\]\]", "", t.text)
                if new_text != t.text:
                    t.text = new_text
                    changed = True
        if not changed:
            return

        document_xml.write_bytes(etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


def normalize_footnote_ids_by_appearance(docx_path: Path) -> None:
    """Renumber footnotes by first appearance order in document.xml.

    This keeps visual numbering consistent across Word and LibreOffice for
    amendment-marked references.
    """
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)

        document_xml = tmp / "word" / "document.xml"
        footnotes_xml = tmp / "word" / "footnotes.xml"
        if not document_xml.exists() or not footnotes_xml.exists():
            return

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        doc_root = etree.fromstring(document_xml.read_bytes(), etree.XMLParser(remove_blank_text=False))
        fn_root = etree.fromstring(footnotes_xml.read_bytes(), etree.XMLParser(remove_blank_text=False))

        refs = doc_root.xpath(".//w:footnoteReference", namespaces=ns)
        ordered_ids: list[int] = []
        seen: set[int] = set()
        for ref in refs:
            rid_raw = ref.get(w_tag("id"))
            if rid_raw is None or not str(rid_raw).lstrip("-").isdigit():
                continue
            rid = int(rid_raw)
            if rid < 1 or rid in seen:
                continue
            seen.add(rid)
            ordered_ids.append(rid)

        existing_ids = [
            int(fid)
            for fid in (n.get(w_tag("id")) for n in fn_root.findall(f".//{w_tag('footnote')}"))
            if fid is not None and str(fid).lstrip("-").isdigit() and int(fid) >= 1
        ]
        for fid in sorted(existing_ids):
            if fid not in seen:
                ordered_ids.append(fid)

        if not ordered_ids:
            return

        old_to_new = {old: idx + 1 for idx, old in enumerate(ordered_ids)}
        if all(old == new for old, new in old_to_new.items()):
            return

        for ref in refs:
            rid_raw = ref.get(w_tag("id"))
            if rid_raw is None or not str(rid_raw).lstrip("-").isdigit():
                continue
            rid = int(rid_raw)
            if rid >= 1 and rid in old_to_new:
                ref.set(w_tag("id"), str(old_to_new[rid]))

        footnote_nodes = fn_root.findall(f".//{w_tag('footnote')}")
        for fn in footnote_nodes:
            fid_raw = fn.get(w_tag("id"))
            if fid_raw is None or not str(fid_raw).lstrip("-").isdigit():
                continue
            fid = int(fid_raw)
            if fid >= 1 and fid in old_to_new:
                fn.set(w_tag("id"), str(old_to_new[fid]))

        # Keep special separators first, then numbered notes in ascending order.
        root_children = list(fn_root)
        special = [n for n in root_children if n.tag == w_tag("footnote") and int(n.get(w_tag("id"), "0")) < 1]
        numbered = [n for n in root_children if n.tag == w_tag("footnote") and int(n.get(w_tag("id"), "0")) >= 1]
        others = [n for n in root_children if n.tag != w_tag("footnote")]
        numbered.sort(key=lambda n: int(n.get(w_tag("id"), "0")))

        for child in root_children:
            fn_root.remove(child)
        for child in special + numbered + others:
            fn_root.append(child)

        document_xml.write_bytes(etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        footnotes_xml.write_bytes(etree.tostring(fn_root, xml_declaration=True, encoding="UTF-8", standalone="yes"))

        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


def normalize_footnote_body_text(text: str) -> str:
    txt = normalize_text(text)
    txt = re.sub(r"(?<=\d)\s+([\)\]\.,;:!?])", r"\1", txt)
    txt = re.sub(r"\(\s+(?=OJ\b)", "(", txt)
    return txt


def footnote_plain_text_from_node(node: etree._Element) -> str:
    chunks: list[str] = []
    for t in node.xpath(".//w:t", namespaces={"w": W_NS}):
        if t.text:
            chunks.append(t.text)
    return normalize_footnote_body_text(" ".join(chunks))


def make_standard_footnote_reference_run(fid: int, *, styled: bool = False) -> etree._Element:
    r = etree.Element(w_tag("r"))
    rpr = etree.SubElement(r, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    if styled:
        add_revision_run_style(rpr)
    sz = etree.SubElement(rpr, w_tag("sz"))
    sz.set(w_tag("val"), "22")
    szcs = etree.SubElement(rpr, w_tag("szCs"))
    szcs.set(w_tag("val"), "22")
    vert = etree.SubElement(rpr, w_tag("vertAlign"))
    vert.set(w_tag("val"), "superscript")
    ref = etree.SubElement(r, w_tag("footnoteReference"))
    ref.set(w_tag("id"), str(fid))
    return r


def make_text_run_with_rpr(rpr_template: etree._Element | None, text: str) -> etree._Element:
    r = etree.Element(w_tag("r"))
    if rpr_template is not None:
        r.append(deepcopy(rpr_template))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_tab_run(rpr_template: etree._Element | None = None) -> etree._Element:
    r = etree.Element(w_tag("r"))
    if rpr_template is not None:
        r.append(deepcopy(rpr_template))
    etree.SubElement(r, w_tag("tab"))
    return r


def add_revision_run_style(rpr: etree._Element) -> None:
    rfonts = rpr.find(w_tag("rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, w_tag("rFonts"))
    rfonts.set(w_tag("ascii"), "Arial")
    rfonts.set(w_tag("hAnsi"), "Arial")

    color = rpr.find(w_tag("color"))
    if color is None:
        color = etree.SubElement(rpr, w_tag("color"))
    color.set(w_tag("val"), revision_color_hex())

    u = rpr.find(w_tag("u"))
    if u is None:
        u = etree.SubElement(rpr, w_tag("u"))
    u.set(w_tag("val"), "double")
    u.set(w_tag("color"), revision_color_hex())


def run_has_revision_style(run: etree._Element) -> bool:
    rpr = run.find(w_tag("rPr"))
    if rpr is None:
        return False
    color = rpr.find(w_tag("color"))
    underline = rpr.find(w_tag("u"))
    return bool(
        (color is not None and color.get(w_tag("val")) == revision_color_hex())
        or (underline is not None and underline.get(w_tag("val")) == "double")
    )


def paragraph_has_revision_style_xml(paragraph: etree._Element) -> bool:
    return any(run_has_revision_style(r) for r in paragraph.xpath("./w:r", namespaces={"w": W_NS}))


def append_standard_footnote_node(root: etree._Element, fid: int, text: str, *, styled: bool = False) -> None:
    fn = etree.SubElement(root, w_tag("footnote"))
    fn.set(w_tag("id"), str(fid))

    p = etree.SubElement(fn, w_tag("p"))
    ppr = etree.SubElement(p, w_tag("pPr"))
    tabs = etree.SubElement(ppr, w_tag("tabs"))
    tab_stop = etree.SubElement(tabs, w_tag("tab"))
    tab_stop.set(w_tag("val"), "left")
    tab_stop.set(w_tag("pos"), "284")
    spacing = etree.SubElement(ppr, w_tag("spacing"))
    spacing.set(w_tag("before"), "0")
    spacing.set(w_tag("after"), "0")
    spacing.set(w_tag("line"), "240")
    spacing.set(w_tag("lineRule"), "auto")
    ind = etree.SubElement(ppr, w_tag("ind"))
    ind.set(w_tag("left"), "284")
    ind.set(w_tag("hanging"), "284")
    jc = etree.SubElement(ppr, w_tag("jc"))
    jc.set(w_tag("val"), "both")

    r_ref = etree.SubElement(p, w_tag("r"))
    rpr = etree.SubElement(r_ref, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    if styled:
        add_revision_run_style(rpr)
    sz_ref = etree.SubElement(rpr, w_tag("sz"))
    sz_ref.set(w_tag("val"), "18")
    szcs_ref = etree.SubElement(rpr, w_tag("szCs"))
    szcs_ref.set(w_tag("val"), "18")
    vert_ref = etree.SubElement(rpr, w_tag("vertAlign"))
    vert_ref.set(w_tag("val"), "superscript")
    etree.SubElement(r_ref, w_tag("footnoteRef"))

    r_tab = etree.SubElement(p, w_tag("r"))
    etree.SubElement(r_tab, w_tag("tab"))

    r_text = etree.SubElement(p, w_tag("r"))
    rpr_text = etree.SubElement(r_text, w_tag("rPr"))
    if styled:
        add_revision_run_style(rpr_text)
    sz = etree.SubElement(rpr_text, w_tag("sz"))
    sz.set(w_tag("val"), "18")
    szcs = etree.SubElement(rpr_text, w_tag("szCs"))
    szcs.set(w_tag("val"), "18")
    t = etree.SubElement(r_text, w_tag("t"))
    t.text = text


def style_footnote_node_revision(fn: etree._Element) -> None:
    for run in fn.xpath(".//w:r", namespaces={"w": W_NS}):
        rpr = run.find(w_tag("rPr"))
        if rpr is None:
            rpr = etree.SubElement(run, w_tag("rPr"))
        add_revision_run_style(rpr)


def ensure_minimal_footnotes_root(footnotes_xml_path: Path) -> etree._Element:
    if footnotes_xml_path.exists():
        return etree.fromstring(footnotes_xml_path.read_bytes(), etree.XMLParser(remove_blank_text=False))

    root = etree.Element(w_tag("footnotes"), nsmap={"w": W_NS})
    for fid, ftype, marker in [(-1, "separator", "separator"), (0, "continuationSeparator", "continuationSeparator")]:
        fn = etree.SubElement(root, w_tag("footnote"))
        fn.set(w_tag("id"), str(fid))
        fn.set(w_tag("type"), ftype)
        p = etree.SubElement(fn, w_tag("p"))
        r = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r, w_tag(marker))
    return root


def patch_docx_with_native_footnotes_merge(docx_path: Path, amending_footnotes: dict[str, str]) -> int:
    """Convert [[FN:n]] markers to native OOXML footnotes while preserving existing notes."""
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)

        document_xml = tmp / "word" / "document.xml"
        footnotes_xml = tmp / "word" / "footnotes.xml"
        rels_xml = tmp / "word" / "_rels" / "document.xml.rels"
        content_types = tmp / "[Content_Types].xml"

        doc_root = etree.fromstring(document_xml.read_bytes(), etree.XMLParser(remove_blank_text=False))

        token_total = 0
        for t in doc_root.xpath(".//w:t", namespaces={"w": W_NS}):
            if t.text:
                token_total += len(FN_TOKEN_RE.findall(t.text))
        if token_total == 0:
            return 0

        fn_root = ensure_minimal_footnotes_root(footnotes_xml)

        existing_ids: set[int] = set()
        existing_id_to_text: dict[int, str] = {}
        existing_text_to_id: dict[str, int] = {}
        for fn in fn_root.findall(f".//{w_tag('footnote')}"):
            fid_raw = fn.get(w_tag("id"))
            if fid_raw is None or not str(fid_raw).lstrip("-").isdigit():
                continue
            fid = int(fid_raw)
            if fid < 1:
                continue
            existing_ids.add(fid)
            body = footnote_plain_text_from_node(fn)
            existing_id_to_text[fid] = body
            if body and body not in existing_text_to_id:
                existing_text_to_id[body] = fid

        next_id = (max(existing_ids) + 1) if existing_ids else 1
        # Unstyled tokens (if any) may reuse IDs by body; styled amendment
        # tokens intentionally do not collapse and always get fresh IDs.
        number_to_id: dict[str, int | None] = {}

        footnote_runs = []
        for run in doc_root.xpath(".//w:r", namespaces={"w": W_NS}):
            original_parts: list[str] = []
            for child in run:
                if child.tag == w_tag("t") and child.text:
                    original_parts.append(child.text)
                elif child.tag == w_tag("tab"):
                    original_parts.append("\t")
                elif child.tag == w_tag("br"):
                    original_parts.append("\n")
            original = "".join(original_parts)
            if "[[FN:" in original:
                # Any [[FN:n]] token originates from amendment payload text,
                # so its native reference/body must be visibly revision-styled.
                footnote_runs.append((run, original, True))

        styled_ids: set[int] = set()

        for run, original, is_styled in footnote_runs:
            para = run.getparent() if run is not None else None
            if run is None or para is None:
                continue

            rpr_template = run.find(w_tag("rPr"))
            parts: list[tuple[str, str | int]] = []
            pos = 0
            for m in FN_TOKEN_RE.finditer(original):
                if m.start() > pos:
                    parts.append(("text", original[pos:m.start()]))

                number = m.group(1)
                if is_styled:
                    fallback_existing = ""
                    if number.isdigit():
                        fallback_existing = existing_id_to_text.get(int(number), "")
                    body = normalize_footnote_body_text(amending_footnotes.get(number, "") or fallback_existing)
                    fid = next_id
                    next_id += 1
                    existing_ids.add(fid)
                    append_standard_footnote_node(fn_root, fid, body, styled=True)
                else:
                    if number not in number_to_id:
                        body = normalize_footnote_body_text(amending_footnotes.get(number, ""))
                        if body:
                            fid = existing_text_to_id.get(body)
                            if fid is None:
                                fid = next_id
                                next_id += 1
                                existing_ids.add(fid)
                                existing_text_to_id[body] = fid
                                append_standard_footnote_node(fn_root, fid, body, styled=False)
                        else:
                            fid = int(number) if number.isdigit() and int(number) in existing_ids else None
                        number_to_id[number] = fid
                    fid = number_to_id[number]

                if is_styled and fid is not None:
                    styled_ids.add(fid)

                if fid is not None:
                    parts.append(("fn", fid))
                else:
                    # If no body can be resolved, keep the original token text
                    # rather than creating an empty native footnote node.
                    parts.append(("text", m.group(0)))
                pos = m.end()

            if pos < len(original):
                parts.append(("text", original[pos:]))

            idx = para.index(run)
            para.remove(run)
            offset = 0
            for kind, value in parts:
                if kind == "text":
                    text_seg = str(value)
                    if text_seg:
                        fragments = text_seg.split("\t")
                        for frag_idx, fragment in enumerate(fragments):
                            if fragment:
                                para.insert(idx + offset, make_text_run_with_rpr(deepcopy(rpr_template) if rpr_template is not None else None, fragment))
                                offset += 1
                            if frag_idx < len(fragments) - 1:
                                para.insert(idx + offset, make_tab_run(deepcopy(rpr_template) if rpr_template is not None else None))
                                offset += 1
                else:
                    para.insert(idx + offset, make_standard_footnote_reference_run(int(value), styled=is_styled))
                    offset += 1

        for fid in sorted(styled_ids):
            fn = fn_root.xpath(f'.//w:footnote[@w:id="{fid}"]', namespaces={"w": W_NS})
            if fn:
                style_footnote_node_revision(fn[0])

        # Integrity pass: every referenced footnote id must exist and carry text.
        ref_ids: set[int] = set()
        for ref in doc_root.xpath(".//w:footnoteReference", namespaces={"w": W_NS}):
            rid_raw = ref.get(w_tag("id"))
            if rid_raw is None or not str(rid_raw).lstrip("-").isdigit():
                continue
            rid = int(rid_raw)
            if rid >= 1:
                ref_ids.add(rid)

        for rid in sorted(ref_ids):
            nodes = fn_root.xpath(f'.//w:footnote[@w:id="{rid}"]', namespaces={"w": W_NS})
            fallback_text = normalize_footnote_body_text(
                existing_id_to_text.get(rid, "") or amending_footnotes.get(str(rid), "")
            )
            if not nodes:
                append_standard_footnote_node(fn_root, rid, fallback_text or f"[Missing footnote text {rid}]", styled=False)
                continue

            node = nodes[0]
            current_text = footnote_plain_text_from_node(node)
            if current_text.strip():
                continue

            text_nodes = node.xpath('.//w:t', namespaces={"w": W_NS})
            if text_nodes:
                text_nodes[-1].text = fallback_text or f"[Missing footnote text {rid}]"
            else:
                p = node.find(w_tag("p"))
                if p is None:
                    p = etree.SubElement(node, w_tag("p"))
                r = etree.SubElement(p, w_tag("r"))
                t = etree.SubElement(r, w_tag("t"))
                t.text = fallback_text or f"[Missing footnote text {rid}]"

        document_xml.write_bytes(etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        footnotes_xml.write_bytes(etree.tostring(fn_root, xml_declaration=True, encoding="UTF-8", standalone="yes"))
        rels_xml.write_bytes(ensure_footnotes_relationship(rels_xml.read_bytes()))
        content_types.write_bytes(ensure_footnotes_content_type(content_types.read_bytes()))

        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))

    return token_total


def strip_src_markers_in_document(doc: Document) -> int:
    """Remove [[SRC:...]] markers from all paragraphs, including split-run markers."""
    changed = 0
    for paragraph in iter_document_paragraphs(doc):
        runs = list(paragraph.runs)
        if not runs:
            continue
        flat: list[tuple[int, str]] = []
        for ridx, run in enumerate(runs):
            for ch in (run.text or ""):
                flat.append((ridx, ch))
        if not flat:
            continue

        text = "".join(ch for _, ch in flat)

        keep = [True] * len(flat)
        i = 0
        removed_any = False
        while i < len(text):
            if text.startswith("[[SRC:", i):
                removed_any = True
                j = i - 1
                while j >= 0 and keep[j] and text[j].isspace():
                    keep[j] = False
                    j -= 1

                k = i
                while k < len(text) and not text.startswith("]]", k):
                    keep[k] = False
                    k += 1
                if k < len(text):
                    keep[k] = False
                    if k + 1 < len(text):
                        keep[k + 1] = False
                    i = k + 2
                else:
                    while k < len(text):
                        keep[k] = False
                        k += 1
                    i = k
                continue
            i += 1

        if removed_any:
            out_by_run: list[list[str]] = [[] for _ in runs]
            for idx, (ridx, ch) in enumerate(flat):
                if keep[idx]:
                    out_by_run[ridx].append(ch)

            paragraph_changed = False
            for ridx, run in enumerate(runs):
                new_text = "".join(out_by_run[ridx])
                if (run.text or "") != new_text:
                    run.text = new_text
                    paragraph_changed = True
            if paragraph_changed:
                changed += 1

        # Secondary cleanup: remove decomposed/plain SRC residues if present.
        for run in runs:
            original = run.text or ""
            cleaned = re.sub(r"\s*\bSRC:[A-Za-z0-9_.-]+\b", "", original, flags=re.I)
            cleaned = re.sub(r"\bsrc\s+art\s+\d+[A-Za-z]?(?:\s+\d+[A-Za-z]?)?\b", "", cleaned, flags=re.I)
            cleaned = re.sub(r"\s+([,.;:])", r"\1", cleaned)
            if cleaned != original:
                run.text = cleaned
                changed += 1

    return changed


def insert_amending_title(doc: Document, insertion_index: int, title_text: str) -> None:
    if not title_text:
        return
    normalized_title = normalize_text(title_text)
    existing_title = next((p for p in doc.paragraphs if normalize_text(p.text) == normalized_title), None)

    def style_as_amended_by_paragraph(paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Cm(0)
        paragraph.paragraph_format.space_after = Cm(0)
        if not paragraph.runs:
            paragraph.add_run("as amended by")
        for run in paragraph.runs:
            if normalize_text(run.text) == "as amended by":
                style_inserted_run_no_bold_change(run)

    if existing_title is not None:
        title_idx = find_current_paragraph_index(doc, existing_title)
        if title_idx is None:
            title_idx = insertion_index

        prev_idx = title_idx - 1
        while prev_idx >= 0 and not normalize_text(doc.paragraphs[prev_idx].text):
            prev_idx -= 1
        needs_as_amended_by = prev_idx < 0 or normalize_text(doc.paragraphs[prev_idx].text) != "as amended by"

        if needs_as_amended_by:
            spacer_before = insert_plain_paragraph_before_index(doc, title_idx)
            spacer_before.paragraph_format.left_indent = Cm(0)
            spacer_before.paragraph_format.space_before = Cm(0)
            spacer_before.paragraph_format.space_after = Cm(0)
            title_idx += 1

            as_amended_by_para = insert_plain_paragraph_before_index(doc, title_idx)
            style_as_amended_by_paragraph(as_amended_by_para)
            title_idx += 1

            spacer_after = insert_plain_paragraph_before_index(doc, title_idx)
            spacer_after.paragraph_format.left_indent = Cm(0)
            spacer_after.paragraph_format.space_before = Cm(0)
            spacer_after.paragraph_format.space_after = Cm(0)
            title_idx += 1

        if prev_idx >= 0 and normalize_text(doc.paragraphs[prev_idx].text) == "as amended by":
            style_as_amended_by_paragraph(doc.paragraphs[prev_idx])

        for run in existing_title.runs:
            run.bold = True
            style_inserted_title_run(run)
        return

    spacer_before = insert_plain_paragraph_before_index(doc, insertion_index)
    spacer_before.paragraph_format.left_indent = Cm(0)
    spacer_before.paragraph_format.space_before = Cm(0)
    spacer_before.paragraph_format.space_after = Cm(0)
    insertion_index += 1

    as_amended_by_para = insert_plain_paragraph_before_index(doc, insertion_index)
    style_as_amended_by_paragraph(as_amended_by_para)
    insertion_index += 1

    spacer_after = insert_plain_paragraph_before_index(doc, insertion_index)
    spacer_after.paragraph_format.left_indent = Cm(0)
    spacer_after.paragraph_format.space_before = Cm(0)
    spacer_after.paragraph_format.space_after = Cm(0)
    insertion_index += 1

    p = insert_plain_paragraph_before_index(doc, insertion_index)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Cm(0)
    p.paragraph_format.space_after = Cm(0)
    lines = split_amending_title(title_text)
    for idx, line in enumerate(lines):
        if idx:
            p.runs[-1].add_break()
        run = p.add_run(line)
        style_inserted_title_run(run)


def insert_amending_recitals(doc: Document, insertion_index: int, recitals: list[str], style_name: str | None) -> None:
    existing = {normalize_text(p.text): p for p in doc.paragraphs if normalize_text(p.text)}
    for recital in reversed(recitals):
        raw_recital = normalize_text(recital)
        marker_match = re.match(r"^\((\d+)\)\s*(.*)$", raw_recital, flags=re.S)
        if marker_match:
            recital_marker = marker_match.group(1)
            recital_body = marker_match.group(2).strip()
            txt = f"({recital_marker})\t{recital_body}" if recital_body else f"({recital_marker})"
        else:
            recital_marker = None
            txt = format_amendment_item_text(raw_recital)
        if not txt:
            continue
        if recital_marker and not re.match(rf"^\({re.escape(recital_marker)}\)(?=\s|\t|$)", txt, flags=re.I):
            txt = f"({recital_marker})\t{txt}"
        if txt in existing:
            # On reruns, enforce insertion styling for already-present amending recitals.
            for run in existing[txt].runs:
                style_inserted_run_no_bold_change(run)
            continue
        p = insert_plain_paragraph_before_index(doc, insertion_index, style_name=style_name)
        run = p.add_run(txt)
        run.bold = False
        style_inserted_run_no_bold_change(run)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        existing[txt] = p


def enforce_amending_recital_block(doc: Document, insertion_index: int, recitals: list[str], style_name: str | None) -> None:
    for offset, recital in enumerate(recitals):
        paragraph_index = insertion_index + offset
        if paragraph_index >= len(doc.paragraphs):
            break
        raw_recital = normalize_text(recital)
        marker_match = re.match(r"^\((\d+)\)\s*(.*)$", raw_recital, flags=re.S)
        if marker_match:
            recital_marker = marker_match.group(1)
            recital_body = marker_match.group(2).strip()
            desired_text = f"({recital_marker})\t{recital_body}" if recital_body else f"({recital_marker})"
        else:
            desired_text = format_amendment_item_text(raw_recital)
        if not desired_text:
            continue

        paragraph = doc.paragraphs[paragraph_index]
        if normalize_text(paragraph.text) == normalize_text(desired_text) and paragraph.text.strip().startswith("("):
            continue

        paragraph.clear()
        run = paragraph.add_run(desired_text)
        run.bold = False
        style_inserted_run_no_bold_change(run)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Cm(1)
        paragraph.paragraph_format.first_line_indent = Cm(-1)


def insert_paragraphs_at_index(
    doc: Document,
    insertion_index: int,
    texts: list[str],
    style_name: str | None,
    rgb: tuple[int, int, int],
    bold: bool = False,
) -> None:
    if not texts:
        return

    paragraphs = doc.paragraphs
    if insertion_index < len(paragraphs):
        anchor = paragraphs[insertion_index]
        for text in texts:
            p = anchor.insert_paragraph_before(style=style_name)
            run = p.add_run(text)
            run.bold = bold
            apply_run_color(run, rgb)
    else:
        for text in texts:
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(text)
            run.bold = bold
            apply_run_color(run, rgb)


def tokenize(text: str) -> list[str]:
    return [m.group(0).lower() for m in TOKEN_RE.finditer(text)]


REPLACEMENT_TOKEN_RE = re.compile(
    r"\[\[FN:\d+\]\]|\[\[SRC:[^\]]+\]\]|\([A-Za-z0-9]+\)|[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?|\s+|[^\w\s]"
)


def tokenize_preserving_punctuation(text: str) -> list[str]:
    return [m.group(0) for m in REPLACEMENT_TOKEN_RE.finditer(text)]


def token_match_key(token: str) -> str:
    if re.fullmatch(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?", token):
        return token.lower()
    return token


def paragraph_similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def collect_prior_inserted_words(paragraph) -> set[str]:
    """Collect words that are currently visible as inserted in this paragraph.

    These are runs that carry double underline and are not struck through.
    If a later amendment deletes those words, we keep double underline on the
    new deletion run so provenance is preserved across amendment layers.
    """
    words: set[str] = set()
    for run in paragraph.runs:
        if run.font.underline != WD_UNDERLINE.DOUBLE:
            continue
        if run.font.strike:
            continue
        for tok in tokenize(run.text or ""):
            words.add(tok.lower())
    return words


def deletion_overlaps_prior_inserted_words(segment: str, prior_inserted_words: set[str]) -> bool:
    if not prior_inserted_words:
        return False
    for tok in tokenize(segment):
        if tok.lower() in prior_inserted_words:
            return True
    return False


def word_diff_stats(base_text: str, amend_text: str) -> dict[str, int]:
    a = tokenize(base_text)
    b = tokenize(amend_text)
    sm = SequenceMatcher(None, a, b)

    inserted = 0
    deleted = 0
    replaced_from = 0
    replaced_to = 0

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "insert":
            inserted += j2 - j1
        elif tag == "delete":
            deleted += i2 - i1
        elif tag == "replace":
            replaced_from += i2 - i1
            replaced_to += j2 - j1

    return {
        "inserted_words": inserted,
        "deleted_words": deleted,
        "replaced_from_words": replaced_from,
        "replaced_to_words": replaced_to,
    }


def highlighted_amendment_runs(paragraph, base_text: str, amend_text: str, rgb: tuple[int, int, int]) -> None:
    """Write amendment text with changed words highlighted in the selected color."""
    a = tokenize(base_text)
    b = tokenize(amend_text)
    sm = SequenceMatcher(None, a, b)

    paragraph.add_run("Amendment text: ")
    first = True
    for tag, _, _, j1, j2 in sm.get_opcodes():
        segment = " ".join(b[j1:j2]).strip()
        if not segment:
            continue
        if not first:
            paragraph.add_run(" ")
        run = paragraph.add_run(segment)
        if tag in {"insert", "replace"}:
            apply_run_color(run, rgb)
            run.bold = True
        first = False


def replace_paragraph_with_revision_marks(
    paragraph,
    base_text: str,
    amend_text: str,
    *,
    preserve_prior_inserted_context: bool = True,
) -> None:
    """Replace paragraph text and mark insertions/deletions in orange.

    - Insertions: orange + double underline
    - Deletions: orange + strikethrough
    """
    # Keep canonical marker + tab formatting even if earlier normalization
    # collapsed marker spacing to a plain space.
    base_text = ensure_marker_tabs_multiline(base_text)
    amend_text = ensure_marker_tabs_multiline(amend_text)

    was_inserted_paragraph = preserve_prior_inserted_context and paragraph_has_inserted_style(paragraph)
    prior_inserted_words = collect_prior_inserted_words(paragraph)

    # Annex heading special-case: preserve heading line break and only show the
    # minimal title-token delta (e.g. removing "The ").
    base_annex = split_annex_heading_and_title(base_text)
    amend_annex = split_annex_heading_and_title(amend_text)
    if base_annex and amend_annex and base_annex[0].lower() == amend_annex[0].lower():
        base_title = base_annex[1]
        amend_title = amend_annex[1]
        base_core = re.sub(r"^the\s+", "", base_title, flags=re.I)
        amend_core = re.sub(r"^the\s+", "", amend_title, flags=re.I)
        if base_core.lower() == amend_core.lower():
            paragraph.clear()
            heading_run = paragraph.add_run(f"{amend_annex[0]}\n")
            if was_inserted_paragraph:
                style_inserted_run_no_bold_change(heading_run)

            base_has_the = bool(re.match(r"^the\s+", base_title, flags=re.I))
            amend_has_the = bool(re.match(r"^the\s+", amend_title, flags=re.I))
            if base_has_the and not amend_has_the:
                removed = re.match(r"^(the\s+)", base_title, flags=re.I)
                removed_text = removed.group(1) if removed else "The "
                del_run = paragraph.add_run(removed_text)
                apply_run_color(del_run, REVISION_COLOR)
                del_run.font.strike = True
            elif amend_has_the and not base_has_the:
                added = re.match(r"^(the\s+)", amend_title, flags=re.I)
                added_text = added.group(1) if added else "The "
                ins_run = paragraph.add_run(added_text)
                style_inserted_run(ins_run)

            common_title = amend_core if amend_core else base_core
            title_run = paragraph.add_run(common_title)
            title_run.font.strike = False
            if was_inserted_paragraph:
                style_inserted_run_no_bold_change(title_run)
            return

    paragraph.clear()

    # Keep leading legal markers and their separator as structural prefix.
    # Example: "(e) for ..." vs "(e)\tfor ..." should not mark the tab itself.
    marker_re = re.compile(r"^(\([A-Za-z0-9]+\)|\d+[A-Za-z]?\.|[IVXLCDM]+\.)(\s+)(.*)$", flags=re.S)
    base_body = base_text
    amend_body = amend_text
    first = True
    last_emitted = ""
    last_emitted_word: str | None = None
    last_kind: str | None = None

    m_base = marker_re.match(base_text)
    m_amend = marker_re.match(amend_text)
    if m_base and m_amend and m_base.group(1).lower() == m_amend.group(1).lower():
        prefix = f"{m_amend.group(1)}\t"
        prefix_run = paragraph.add_run(prefix)
        if was_inserted_paragraph:
            style_inserted_run_no_bold_change(prefix_run)
        base_body = m_base.group(3)
        amend_body = m_amend.group(3)
        first = False
        last_emitted = prefix
        last_kind = "equal"

    base_tokens = tokenize_preserving_punctuation(base_body)
    amend_tokens = tokenize_preserving_punctuation(amend_body)
    base_words = tokenize(base_body)
    amend_words = tokenize(amend_body)

    # When substantive text diverges heavily, render the full sentence as
    # changed instead of preserving tiny unchanged glue tokens.
    if base_words and amend_words:
        body_similarity = SequenceMatcher(None, base_words, amend_words).ratio()
        if body_similarity < 0.34:
            if base_body:
                del_run = paragraph.add_run(base_body)
                apply_run_color(del_run, REVISION_COLOR)
                del_run.font.strike = True
                if deletion_overlaps_prior_inserted_words(base_body, prior_inserted_words):
                    del_run.font.underline = WD_UNDERLINE.DOUBLE
            if amend_body:
                ins_run = paragraph.add_run(amend_body)
                style_inserted_run(ins_run)
            return

    sm = SequenceMatcher(None, [token_match_key(token) for token in base_tokens], [token_match_key(token) for token in amend_tokens])
    suppressed_deleted_word: str | None = None
    opcodes = sm.get_opcodes()
    idx_opcode = 0
    while idx_opcode < len(opcodes):
        tag, i1, i2, j1, j2 = opcodes[idx_opcode]
        segments: list[tuple[str, str]] = []
        if tag == "equal":
            seg = "".join(amend_tokens[j1:j2])
            if seg:
                segments.append(("equal", seg))
            idx_opcode += 1
        else:
            # Collapse adjacent non-equal opcodes into a single visual
            # replacement block: one deletion section followed by one insertion
            # section. This is clearer for legal text substitutions.
            i_start, i_end = i1, i2
            j_start, j_end = j1, j2
            idx_opcode += 1
            while idx_opcode < len(opcodes):
                ntag, ni1, ni2, nj1, nj2 = opcodes[idx_opcode]
                if ntag != "equal":
                    i_end = ni2
                    j_end = nj2
                    idx_opcode += 1
                    continue
                # Keep glue-only equal spans (spaces/punctuation) inside the
                # same visual replacement block when surrounded by changes.
                eq_seg = "".join(amend_tokens[nj1:nj2])
                if (
                    re.fullmatch(r"[\s,;:.()\[\]{}'\"“”‘’/\\-]*", eq_seg or "")
                    and idx_opcode + 1 < len(opcodes)
                    and opcodes[idx_opcode + 1][0] != "equal"
                ):
                    i_end = ni2
                    j_end = nj2
                    idx_opcode += 1
                    continue
                if (
                    re.fullmatch(r"\s*(?:and|or)\s+", eq_seg or "", flags=re.I)
                    and idx_opcode + 1 < len(opcodes)
                    and opcodes[idx_opcode + 1][0] != "equal"
                ):
                    i_end = ni2
                    j_end = nj2
                    idx_opcode += 1
                    continue
                break

            del_seg = "".join(base_tokens[i_start:i_end])
            ins_seg = "".join(amend_tokens[j_start:j_end])
            if del_seg and ins_seg:
                if is_quote_only_replacement(del_seg, ins_seg):
                    segments.append(("equal", del_seg))
                    del_seg = ""
                    ins_seg = ""
                ins_seg = trim_duplicate_leading_word(del_seg, ins_seg)
            if del_seg:
                segments.append(("delete", del_seg))
            if ins_seg:
                segments.append(("insert", ins_seg))

        for kind, segment in segments:
            if not first and last_emitted and segment:
                if last_emitted[-1].isalnum() and segment[0].isalnum():
                    sep_run = paragraph.add_run(" ")
                    if last_kind == "delete" and kind == "insert":
                        apply_run_color(sep_run, REVISION_COLOR)
                        sep_run.font.strike = True
                    elif kind == "insert" or last_kind == "insert":
                        style_inserted_run(sep_run)
                    last_emitted = " "
                    last_kind = "insert" if (kind == "insert" or last_kind == "insert") else "equal"

            if kind == "equal" and last_kind == "insert" and segment:
                # Keep boundary spaces highlighted when they belong to inserted context.
                m_lead_ws = re.match(r"^\s+", segment)
                if m_lead_ws:
                    ws = m_lead_ws.group(0)
                    ws_run = paragraph.add_run(ws)
                    style_inserted_run(ws_run)
                    last_emitted = ws
                    last_kind = "insert"
                    segment = segment[len(ws):]
                    if not segment:
                        continue
            if kind == "delete":
                m_del_dup = re.match(r"^\s*([A-Za-z0-9]+)\s*$", segment)
                if m_del_dup and last_emitted_word and m_del_dup.group(1).lower() == last_emitted_word:
                    continue
            run = paragraph.add_run(segment)
            if kind == "insert":
                m_ins = re.match(r"^\s*([A-Za-z0-9]+)\b", segment)
                if (
                    suppressed_deleted_word
                    and m_ins
                    and m_ins.group(1).lower() == suppressed_deleted_word
                ):
                    suppressed_deleted_word = None
                    continue
                apply_run_color(run, REVISION_COLOR)
                run.font.strike = False
                run.font.underline = WD_UNDERLINE.DOUBLE
                run.bold = False
            elif kind == "delete":
                apply_run_color(run, REVISION_COLOR)
                run.font.strike = True
                if deletion_overlaps_prior_inserted_words(segment, prior_inserted_words):
                    run.font.underline = WD_UNDERLINE.DOUBLE
                m_del = re.match(r"^\s*([A-Za-z0-9]+)\b\s*$", segment)
                suppressed_deleted_word = m_del.group(1).lower() if m_del else None
            else:
                if was_inserted_paragraph:
                    style_inserted_run_no_bold_change(run)
                run.font.strike = False
                suppressed_deleted_word = None
            last_emitted = segment
            m_word = re.search(r"([A-Za-z0-9]+)\s*$", segment)
            if m_word:
                last_emitted_word = m_word.group(1).lower()
            last_kind = kind
            first = False

    # Cleanup: if a struck single-word run duplicates the immediately preceding
    # emitted word, drop it to avoid confusing displays like "for for".
    runs = list(paragraph.runs)
    for idx, run in enumerate(runs):
        if not run.font.strike:
            continue
        txt = run.text or ""
        m_del = re.match(r"^\s*([A-Za-z0-9]+)\s*$", txt)
        if not m_del:
            continue
        prev_word = None
        prev_idx = idx - 1
        while prev_idx >= 0:
            prev_txt = runs[prev_idx].text or ""
            m_prev = re.search(r"([A-Za-z0-9]+)\s*$", prev_txt)
            if m_prev:
                prev_word = m_prev.group(1).lower()
                break
            prev_idx -= 1
        if prev_word and prev_word == m_del.group(1).lower():
            run._element.getparent().remove(run._element)


def replace_paragraph_with_plain_text(paragraph, amend_text: str) -> None:
    paragraph.clear()
    paragraph.add_run(amend_text)


def write_analysis_docx(
    analysis_docx: Path,
    amending_html: Path,
    rgb: tuple[int, int, int],
    analysis_items: list[dict],
) -> None:
    analysis_doc = Document()
    heading = analysis_doc.add_paragraph("Amendment Change Analysis")
    if heading.runs:
        heading.runs[0].bold = True
        apply_run_color(heading.runs[0], rgb)

    meta = analysis_doc.add_paragraph(
        f"Amending source: {amending_html} | Generated: {datetime.now(timezone.utc).isoformat()}"
    )
    if meta.runs:
        apply_run_color(meta.runs[0], rgb)

    if not analysis_items:
        warning = analysis_doc.add_paragraph(
            "No actionable amending article blocks were detected in the provided HTML; "
            "no substantive paragraph amendments were applied."
        )
        if warning.runs:
            warning.runs[0].bold = True
            apply_run_color(warning.runs[0], rgb)

    for item in analysis_items:
        diff = item["diff"]
        best_score = item["best_match_score"]
        summary = analysis_doc.add_paragraph(
            f"Article {item['article_number']} | match={best_score:.3f} | +{diff['inserted_words']} "
            f"-{diff['deleted_words']} ~{diff['replaced_from_words']}->{diff['replaced_to_words']}"
        )
        if summary.runs:
            summary.runs[0].bold = True
            apply_run_color(summary.runs[0], rgb)

        analysis_doc.add_paragraph(f"Closest previous text: {item['best_match_text'][:280]}")
        highlighted = analysis_doc.add_paragraph()
        highlighted_amendment_runs(highlighted, item["best_match_text"], item["amending_text"], rgb)

    analysis_docx.parent.mkdir(parents=True, exist_ok=True)
    analysis_doc.save(str(analysis_docx))


def collect_amending_blocks(amending_html: Path) -> tuple[str, list[str], list[dict[str, str]]]:
    # Reuse Stage 1 extraction approach for title, recitals, and operative text.
    soup = load_html(amending_html)

    title_tag = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    title = normalize_text(title_tag.get_text(" ", strip=True)) if title_tag else ""

    recitals: list[str] = []
    for rec in soup.find_all(id=re.compile(r"^rct_\d+$")):
        rid = str(rec.get("id", ""))
        number = rid.replace("rct_", "")
        paras = provision_paragraphs(rec, classes=("oj-normal",))
        body_texts: list[str] = []
        for p in paras:
            txt = text_with_footnote_tokens(p)
            if txt in {number, f"({number})"}:
                continue
            txt = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", txt).strip()
            if txt:
                body_texts.append(txt)
        body = " ".join(body_texts) if body_texts else re.sub(
            rf"^\(?{re.escape(number)}\)?\s*", "", text_with_footnote_tokens(rec)
        ).strip()
        if body:
            recital_text = f"({number}) {body}"
            if not is_non_legislative_payload_text(recital_text):
                recitals.append(recital_text)

    items: list[dict[str, str]] = []
    for art in soup.find_all(id=re.compile(r"^art_[A-Za-z0-9]+$")):
        a1, a2 = article_heading_parts(art)
        heading = normalize_text(f"{a1} {a2 or ''}")
        article_number = a1.replace("Article", "").strip() or "?"
        current_target_article: str | None = None
        current_target_paragraph: str | None = None
        current_target_point: str | None = None
        current_target_annex: str | None = None
        current_target_annex_point: str | None = None
        current_inserted_article_marker: str | None = None
        current_replacement_target_point: str | None = None
        current_instruction: str | None = None
        current_parent_paragraph_marker: str | None = None
        current_parent_point_marker: str | None = None
        for p in provision_paragraphs(art):
            txt = normalize_text(text_with_footnote_tokens(p))
            if not txt:
                continue
            if is_non_legislative_payload_text(txt):
                continue
            if is_instructional_amendment_line(txt):
                current_instruction = txt
                current_replacement_target_point = None
                if re.fullmatch(r"\(\d+\)", txt):
                    # Structural numbering markers in amending acts (e.g. "(5)")
                    # start a new instruction bucket; clear target context so
                    # previous Article/paragraph anchors cannot leak into the
                    # following unrelated amendment block.
                    current_target_article = None
                    current_target_paragraph = None
                    current_target_point = None
                    current_target_annex = None
                    current_target_annex_point = None
                    current_inserted_article_marker = None
                    current_parent_paragraph_marker = None
                    current_parent_point_marker = None
                    continue
                generic_point_insertion_instruction = bool(
                    re.search(
                        r"\bthe following point(?:s)?\s+(?:is|are)\s+(?:added|inserted)\b",
                        txt,
                        flags=re.I,
                    )
                )
                preserved_parent_point = current_parent_point_marker
                # New instructions usually start a fresh point-target scope.
                current_target_point = None
                current_parent_point_marker = None
                target = extract_target_article_from_instruction(txt)
                inserted_article_marker = None
                if re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", txt, flags=re.I):
                    inserted_article_marker = extract_inserted_article_marker_from_instruction_paragraph(p)
                    if inserted_article_marker:
                        base_article, _ = split_article_number(inserted_article_marker)
                        target = target or base_article
                current_inserted_article_marker = inserted_article_marker
                if target:
                    if current_target_article != target:
                        current_target_paragraph = None
                        current_target_point = None
                    current_target_article = target
                    current_target_annex = None
                    current_target_annex_point = None
                elif re.search(r"\bAnnex\s+[IVXLC0-9]+\b", txt, flags=re.I):
                    # Prevent stale article context from leaking into Annex amendments.
                    current_target_article = None
                    current_target_paragraph = None
                    current_target_point = None
                    current_target_annex = extract_target_annex_from_instruction(txt)
                    current_target_annex_point = extract_target_annex_point_from_instruction(txt)
                target_paragraph = extract_target_paragraph_from_instruction(txt)
                if target_paragraph:
                    current_target_paragraph = target_paragraph
                    current_target_point = None
                elif (
                    current_parent_paragraph_marker
                    and re.search(r"\b(?:point|points|introductory wording|sentence|subparagraph)\b", txt, flags=re.I)
                ):
                    # Many amending lines omit the paragraph number and rely on
                    # the immediately preceding paragraph marker context.
                    current_target_paragraph = current_parent_paragraph_marker

                if re.search(r"\bis\s+deleted\b|\bare\s+deleted\b", txt, flags=re.I):
                    # Deletion instructions do not introduce follow-on payload
                    # lines. If we keep this paragraph target in context,
                    # subsequent generic instructions like "the following
                    # paragraph is inserted" can inherit a stale anchor
                    # (e.g. deleted paragraph 3a), causing misplacement.
                    deletion_points = extract_all_target_points_from_instruction(txt)
                    if deletion_points:
                        for deletion_point in deletion_points:
                            items.append(
                                {
                                    "article_number": article_number,
                                    "article_heading": heading,
                                    "text": txt,
                                    "amendment_kind": "deletion",
                                    "target_article_number": current_target_article,
                                    "target_paragraph_number": current_target_paragraph,
                                    "target_point_marker": deletion_point,
                                    "target_annex_number": current_target_annex,
                                    "target_annex_point_marker": current_target_annex_point,
                                    "inserted_article_marker": current_inserted_article_marker,
                                    "source_instruction": current_instruction,
                                }
                            )
                    else:
                        items.append(
                            {
                                "article_number": article_number,
                                "article_heading": heading,
                                "text": txt,
                                "amendment_kind": "deletion",
                                "target_article_number": current_target_article,
                                "target_paragraph_number": current_target_paragraph,
                                "target_point_marker": current_target_point,
                                "target_annex_number": current_target_annex,
                                "target_annex_point_marker": current_target_annex_point,
                                "inserted_article_marker": current_inserted_article_marker,
                                "source_instruction": current_instruction,
                            }
                        )
                    current_target_paragraph = None
                    current_target_point = None

                target_point = extract_target_point_from_instruction(txt)
                if target_point:
                    current_target_point = target_point
                    if re.search(r"\breplaced by the following\b", txt, flags=re.I):
                        current_replacement_target_point = target_point
                elif (
                    generic_point_insertion_instruction
                    and preserved_parent_point
                    and not is_roman_point_marker(preserved_parent_point)
                    and len(preserved_parent_point) >= 2
                ):
                    # Keep parent point context for nested insertions like
                    # (ca) followed by inserted (i)/(ii) under that point.
                    current_target_point = preserved_parent_point
                continue

            marker_probe = sanitize_inserted_marker_text(normalize_quote_wrapping(txt))
            point_only_marker = re.fullmatch(r"\(([a-z]+|[ivxlcdm]+)\)", marker_probe, flags=re.I)
            if point_only_marker:
                current_parent_point_marker = point_only_marker.group(1).lower()
                # For replacement instructions, the instruction's target point (e.g. point (b))
                # must remain the active target; literal (i)/(ii) rows inside the replacement
                # block are the amendment's internal structure, not the target level.
                if not (current_instruction and re.search(r"\breplaced by the following\b", current_instruction, flags=re.I)):
                    current_target_point = current_parent_point_marker
                continue

            cleaned_txt = format_amendment_item_text(txt)
            if not cleaned_txt:
                continue
            paragraph_marker = first_top_level_paragraph_marker(cleaned_txt)
            point_marker = first_top_level_point_marker(cleaned_txt)

            if (
                paragraph_marker is None
                and point_marker is None
                and current_target_point is not None
                and current_instruction
                and re.search(
                    r"\bthe following Article(?:s)?\s+(?:is|are)\s+inserted\b",
                    current_instruction,
                    flags=re.I,
                )
                and p.find_parent("table") is None
            ):
                current_target_point = None

            if paragraph_marker:
                current_parent_paragraph_marker = paragraph_marker
                current_parent_point_marker = None
                current_target_point = None
            elif point_marker:
                current_parent_point_marker = point_marker

            effective_target_paragraph = current_target_paragraph
            if effective_target_paragraph is None and paragraph_marker is None and current_target_annex is None:
                effective_target_paragraph = current_parent_paragraph_marker

            effective_target_point = current_target_point
            if current_instruction and re.search(r"\breplaced by the following\b", current_instruction, flags=re.I):
                effective_target_point = extract_target_point_from_instruction(current_instruction) or current_target_point
            if (
                effective_target_point is None
                and point_marker is None
                and current_target_annex is None
                and not (
                    current_instruction
                    and re.search(
                        r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b",
                        current_instruction,
                        flags=re.I,
                    )
                )
            ):
                effective_target_point = current_parent_point_marker

            effective_target_annex = current_target_annex
            effective_target_annex_point = current_target_annex_point
            if effective_target_annex and current_instruction and re.search(r"\breplaced by the following\b", current_instruction, flags=re.I):
                effective_target_annex_point = extract_target_annex_point_from_instruction(current_instruction) or current_target_annex_point

            items.append(
                {
                    "article_number": article_number,
                    "article_heading": heading,
                    "text": cleaned_txt,
                    "amendment_kind": classify_amendment_kind(current_instruction, cleaned_txt),
                    "target_article_number": current_target_article,
                    "target_paragraph_number": effective_target_paragraph,
                    "target_point_marker": effective_target_point,
                    "target_annex_number": effective_target_annex,
                    "target_annex_point_marker": effective_target_annex_point,
                    "inserted_article_marker": current_inserted_article_marker,
                    "source_instruction": current_instruction,
                }
            )
    if not title or not recitals or not items:
        legislation = parse_eurlex_document(amending_html)
        if not title:
            title = extract_title_from_html(amending_html) or normalize_text(legislation.title)
        if not recitals:
            recitals = [normalize_text(r.text) for r in legislation.recitals if normalize_text(r.text)]
            recitals = [r for r in recitals if not is_non_legislative_payload_text(r)]
            if not recitals:
                recitals = fallback_extract_recitals_from_html(amending_html)
        if not items:
            for article in legislation.articles:
                heading = normalize_text(article.heading)
                for block in article.blocks:
                    text = normalize_text(getattr(block, "text", ""))
                    if not text:
                        continue
                    if is_non_legislative_payload_text(text):
                        continue
                    if is_instructional_amendment_line(text):
                        continue
                    items.append(
                        {
                            "article_number": article.number,
                            "article_heading": heading,
                            "text": format_amendment_item_text(text),
                            "amendment_kind": classify_amendment_kind(None, text),
                            "target_article_number": None,
                            "target_paragraph_number": None,
                            "target_point_marker": None,
                            "source_instruction": None,
                        }
                    )
            if not items:
                items = fallback_extract_amending_items_from_html(amending_html)

    recitals = [format_amendment_item_text(recital) for recital in recitals if format_amendment_item_text(recital)]
    items = normalize_amendment_items(items)
    items = normalize_replacement_target_points(items)

    return title, recitals, items


def find_current_paragraph_index(doc: Document, paragraph) -> int | None:
    target_el = paragraph._element
    for idx, p in enumerate(doc.paragraphs):
        if p._element is target_el:
            return idx
    return None


def article_heading_number(text: str) -> str | None:
    m = re.match(r"^Article\s+(\d+[A-Za-z]?)\b", normalize_text(text), flags=re.I)
    return m.group(1) if m else None


def annex_heading_number(text: str) -> str | None:
    m = re.match(r"^ANNEX\s+([IVXLC]+)\b", normalize_text(text), flags=re.I)
    return m.group(1).upper() if m else None


def chapter_heading_number(text: str) -> str | None:
    m = re.match(r"^CHAPTER\s+([IVXLC]+)\b", normalize_text(text), flags=re.I)
    return m.group(1).upper() if m else None


def find_article_section_bounds(doc: Document, article_number: str) -> tuple[int, int] | None:
    headings: list[tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        num = article_heading_number(p.text)
        if num:
            headings.append((idx, num.lower()))

    target = article_number.lower()
    for i, (idx, num) in enumerate(headings):
        if num != target:
            continue
        end_idx = headings[i + 1][0] if i + 1 < len(headings) else len(doc.paragraphs)

        # Chapter/Annex headings can appear between article headings and must
        # bound the current article section so insertions do not spill below
        # the next chapter title.
        for probe_idx in range(idx + 1, end_idx):
            probe_text = doc.paragraphs[probe_idx].text
            if chapter_heading_number(probe_text) or annex_heading_number(probe_text):
                end_idx = probe_idx
                break
        return idx, end_idx
    return None


def find_annex_section_bounds(doc: Document, annex_number: str) -> tuple[int, int] | None:
    headings: list[tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        num = annex_heading_number(p.text)
        if num:
            headings.append((idx, num))

    target = annex_number.upper()
    for i, (idx, num) in enumerate(headings):
        if num != target:
            continue
        end_idx = headings[i + 1][0] if i + 1 < len(headings) else len(doc.paragraphs)
        return idx, end_idx
    return None


def matches_annex_point_marker(text: str, marker: str) -> bool:
    txt = normalize_text(text)
    escaped = re.escape(marker.upper())
    return bool(re.match(rf"^[‘'\"]?{escaped}\.(?=\s|$)", txt, flags=re.I))


def find_annex_point_block_bounds(
    doc: Document,
    annex_number: str,
    annex_point_marker: str,
) -> tuple[int, int] | None:
    bounds = find_annex_section_bounds(doc, annex_number)
    if bounds is None:
        return None
    section_start, section_end = bounds

    point_start: int | None = None
    for idx in range(section_start, section_end):
        if matches_annex_point_marker(doc.paragraphs[idx].text, annex_point_marker):
            point_start = idx
            break
    if point_start is None:
        return None

    point_end = section_end
    for idx in range(point_start + 1, section_end):
        if re.match(r"^[‘'\"]?[IVXLC]+\.(?=\s|$)", normalize_text(doc.paragraphs[idx].text), flags=re.I):
            point_end = idx
            break
    return point_start, point_end


def matches_top_level_paragraph_marker(text: str, marker: str) -> bool:
    txt = normalize_text(text)
    escaped = re.escape(marker)
    return bool(re.match(rf"^[‘'\"]?{escaped}\.", txt, flags=re.I))


def matches_top_level_point_marker(text: str, marker: str) -> bool:
    txt = normalize_text(text)
    escaped = re.escape(marker)
    return bool(re.match(rf"^(?:\([ivxlcdm]+\)\s+)?\({escaped}\)", txt, flags=re.I))


def first_top_level_paragraph_marker(text: str) -> str | None:
    txt = normalize_text(text)
    m = re.match(r"^[‘'\"]?(\d+[A-Za-z]?)\.", txt)
    if m:
        return m.group(1)
    return None


def paragraph_marker_sort_key(marker: str) -> tuple[int, str]:
    m = re.fullmatch(r"(\d+)([A-Za-z]*)", marker.strip())
    if not m:
        return (10**9, marker.lower())
    return (int(m.group(1)), m.group(2).lower())


def first_top_level_point_marker(text: str) -> str | None:
    txt = normalize_text(text)
    m = re.match(r"^\(([a-z]+)\)", txt, flags=re.I)
    if m:
        return m.group(1).lower()
    return None


def is_roman_point_marker(marker: str | None) -> bool:
    if not marker:
        return False
    roman_markers = {
        "i",
        "ii",
        "iii",
        "iv",
        "v",
        "vi",
        "vii",
        "viii",
        "ix",
        "x",
        "xi",
        "xii",
        "xiii",
        "xiv",
        "xv",
    }
    return marker.lower() in roman_markers


def trim_duplicate_leading_word(delete_seg: str, insert_seg: str) -> str:
    """Avoid visually confusing duplicate leading words in replace diffs.

    Example: delete='for', insert='for equity ...' -> insert='equity ...'
    """
    dm = re.match(r"^\s*([A-Za-z0-9]+)\b", delete_seg)
    im = re.match(r"^\s*([A-Za-z0-9]+)(\b\s*)", insert_seg)
    if not dm or not im:
        return insert_seg
    if dm.group(1).lower() != im.group(1).lower():
        return insert_seg
    trimmed = insert_seg[im.end():]
    return trimmed if trimmed.strip() else insert_seg


def normalize_quote_noise(text: str) -> str:
    # Strip quote glyphs/spacing to compare substantive content.
    return re.sub(r"[\s'\"“”‘’`´]+", "", text or "")


def is_quote_only_replacement(delete_seg: str, insert_seg: str) -> bool:
    if not delete_seg or not insert_seg:
        return False
    return normalize_quote_noise(delete_seg) == normalize_quote_noise(insert_seg)


def split_annex_heading_and_title(text: str) -> tuple[str, str] | None:
    txt = normalize_text(text)
    m = re.match(r"^(ANNEX\s+[IVXLC]+[A-Za-z]?)(?:\s+|\n)(.+)$", txt, flags=re.I | re.S)
    if not m:
        return None
    return m.group(1), normalize_text(m.group(2))


def inject_missing_legal_marker_for_replacement(base_text: str, amend_text: str) -> str:
    """Preserve unchanged legal markers when amend payload omits them.

    For some replacement blocks, parsed amendment rows carry target metadata
    (e.g. point (a)/(b)) but the row text excludes the leading marker. If we
    diff that directly against base text, the marker appears as a deletion even
    though numbering is unchanged.
    """
    base_marker = first_top_level_paragraph_marker(base_text)
    if base_marker and not first_top_level_paragraph_marker(amend_text):
        return f"{base_marker}.\t{amend_text}"

    base_point = first_top_level_point_marker(base_text)
    if base_point and not first_top_level_point_marker(amend_text):
        return f"({base_point})\t{amend_text}"

    return amend_text


def find_paragraph_block_end(doc: Document, section_start: int, section_end: int, paragraph_marker: str) -> int:
    """Return insertion index at end of the target top-level paragraph block."""
    marker_idx: int | None = None
    for idx in range(section_start, section_end):
        if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, paragraph_marker):
            marker_idx = idx
            break
    if marker_idx is None:
        return section_end

    for idx in range(marker_idx + 1, section_end):
        next_para = first_top_level_paragraph_marker(doc.paragraphs[idx].text)
        if next_para is not None:
            return idx
    return section_end


def find_first_subparagraph_block_end(doc: Document, paragraph_start: int, paragraph_end: int) -> int:
    """Return the end of the first subparagraph's point list within a paragraph block.

    This is a focused heuristic for instructions like "in Article X, first
    subparagraph, the following point is added", where the inserted point must
    land before the next plain-text subparagraph starts.
    """
    seen_point = False
    for idx in range(paragraph_start, paragraph_end):
        marker = first_top_level_point_marker(doc.paragraphs[idx].text)
        if marker is not None:
            seen_point = True
            continue
        if seen_point and normalize_text(doc.paragraphs[idx].text):
            return idx
    return paragraph_end


def find_insertion_index_by_paragraph_marker_order(
    doc: Document,
    section_start: int,
    section_end: int,
    new_marker: str,
) -> int:
    """Find insertion point for a top-level paragraph marker by legal order.

    Example: inserting 6a should land between existing 6 and 7 markers.
    """
    new_key = paragraph_marker_sort_key(new_marker)
    for idx in range(section_start, section_end):
        marker = first_top_level_paragraph_marker(doc.paragraphs[idx].text)
        if marker is None:
            continue
        if paragraph_marker_sort_key(marker) > new_key:
            return idx
    return section_end


def find_point_block_end_within_paragraph(
    doc: Document,
    paragraph_start: int,
    paragraph_end: int,
    point_marker: str,
) -> int:
    """Return insertion index at end of a specific top-level point block."""
    point_idx: int | None = None
    for idx in range(paragraph_start, paragraph_end):
        if matches_top_level_point_marker(doc.paragraphs[idx].text, point_marker):
            point_idx = idx
            break
    if point_idx is None:
        return paragraph_end

    for idx in range(point_idx + 1, paragraph_end):
        next_point = first_top_level_point_marker(doc.paragraphs[idx].text)
        if next_point is not None:
            return idx
    return paragraph_end


def find_target_point_block_bounds(
    doc: Document,
    *,
    target_article: str | None,
    target_paragraph: str | None,
    target_point: str | None,
) -> tuple[int, int] | None:
    if not target_article or not target_point:
        return None
    bounds = find_article_section_bounds(doc, target_article)
    if bounds is None:
        return None
    section_start, section_end = bounds

    search_start = section_start
    search_end = section_end
    if target_paragraph:
        para_start = None
        for idx in range(section_start, section_end):
            if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                para_start = idx
                break
        if para_start is not None:
            para_end = find_paragraph_block_end(doc, section_start, section_end, target_paragraph)
            search_start, search_end = para_start, para_end

    point_start = None
    for idx in range(search_start, search_end):
        if matches_top_level_point_marker(doc.paragraphs[idx].text, target_point):
            point_start = idx
            break
    if point_start is None:
        return None

    point_end = search_end
    target_is_roman = is_roman_point_marker(target_point)
    for idx in range(point_start + 1, search_end):
        next_txt = normalize_text(doc.paragraphs[idx].text)
        if first_top_level_paragraph_marker(next_txt) is not None:
            point_end = idx
            break
        next_point = first_top_level_point_marker(next_txt)
        if next_point is not None:
            # For non-roman parent points (e.g. (z)), include nested roman
            # children (i)/(ii)/... in the same logical block.
            if target_is_roman or not is_roman_point_marker(next_point):
                point_end = idx
                break
            continue

        # Guardrail: after a top-level point, a markerless paragraph that
        # starts a new sentence (capitalised, not leading conjunction) is more
        # likely paragraph-level follow-on text than part of the point block.
        if next_txt and re.match(r"^[A-Z]", next_txt) and not re.match(r"^(and|or)\b", next_txt, flags=re.I):
            point_end = idx
            break
    return point_start, point_end


def prune_deleted_point_from_multiline_paragraph(text: str, target_point: str) -> str | None:
    """Remove one point block from a multi-line paragraph while preserving trailing text.

    Some source paragraphs contain multiple logical lines (point marker, nested
    lines, and a following paragraph-level sentence) in one DOCX paragraph.
    """
    raw = strip_src_artifacts(text or "")
    if "\n" not in raw:
        return None

    lines = raw.splitlines()
    start: int | None = None
    for i, line in enumerate(lines):
        if matches_top_level_point_marker(line, target_point):
            start = i
            break
    if start is None:
        return None

    end = len(lines)
    target_is_roman = is_roman_point_marker(target_point)
    for i in range(start + 1, len(lines)):
        line = normalize_text(lines[i])
        if not line:
            continue
        if first_top_level_paragraph_marker(line) is not None:
            end = i
            break
        next_point = first_top_level_point_marker(line)
        if next_point is not None:
            if target_is_roman or not is_roman_point_marker(next_point):
                end = i
                break
            continue
        if re.match(r"^[A-Z]", line) and not re.match(r"^(and|or)\b", line, flags=re.I):
            end = i
            break

    kept = [ln for ln in (lines[:start] + lines[end:]) if normalize_text(ln)]
    if not kept:
        return ""
    new_text = "\n".join(kept)
    if normalize_text(new_text) == normalize_text(raw):
        return None
    return new_text


def find_insertion_index_by_point_marker_order_within_paragraph(
    doc: Document,
    paragraph_start: int,
    paragraph_end: int,
    new_marker: str,
) -> int:
    """Find insertion index for a new top-level point marker within a paragraph block.

    Keeps newly inserted points inside the point list (before trailing non-point
    subparagraph text of the same numbered paragraph).
    """
    if is_roman_point_marker(new_marker):
        new_key = new_marker.lower()
        last_point_idx: int | None = None

        for idx in range(paragraph_start, paragraph_end):
            marker = first_top_level_point_marker(doc.paragraphs[idx].text)
            if marker is None:
                continue
            marker_key = marker.lower()
            if marker_key > new_key:
                return idx
            last_point_idx = idx

        if last_point_idx is not None:
            return min(last_point_idx + 1, paragraph_end)
        return paragraph_end

    new_key = new_marker.lower()
    last_alpha_idx: int | None = None

    for idx in range(paragraph_start, paragraph_end):
        marker = first_top_level_point_marker(doc.paragraphs[idx].text)
        if marker is None:
            continue
        if is_roman_point_marker(marker):
            continue
        marker_key = marker.lower()
        if marker_key > new_key:
            return idx
        last_alpha_idx = idx

    if last_alpha_idx is not None:
        insert_at = last_alpha_idx + 1
        for idx in range(insert_at, paragraph_end):
            marker = first_top_level_point_marker(doc.paragraphs[idx].text)
            if marker is None:
                return idx
            if not is_roman_point_marker(marker):
                return idx
            insert_at = idx + 1
        return min(insert_at, paragraph_end)
    return paragraph_end


def find_nested_roman_insertion_index_within_parent_point(
    doc: Document,
    paragraph_start: int,
    paragraph_end: int,
    parent_point_marker: str,
) -> int:
    """Insert nested roman points after parent point and existing roman children.

    Stops before the next non-roman sibling point or before trailing non-point text.
    """
    parent_idx: int | None = None
    for idx in range(paragraph_start, paragraph_end):
        if matches_top_level_point_marker(doc.paragraphs[idx].text, parent_point_marker):
            parent_idx = idx
            break
    if parent_idx is None:
        return paragraph_end

    insert_at = parent_idx + 1
    for idx in range(parent_idx + 1, paragraph_end):
        marker = first_top_level_point_marker(doc.paragraphs[idx].text)
        if marker is None:
            return idx
        if is_roman_point_marker(marker):
            insert_at = idx + 1
            continue
        return idx
    return min(insert_at, paragraph_end)


def find_target_point_paragraph_index(
    doc: Document,
    *,
    target_article: str | None,
    target_paragraph: str | None,
    target_point: str | None,
) -> int | None:
    """Find the paragraph index for a targeted point marker within article context."""
    if not target_article or not target_point:
        return None
    bounds = find_article_section_bounds(doc, target_article)
    if bounds is None:
        return None
    section_start, section_end = bounds

    search_start = section_start
    search_end = section_end
    if target_paragraph:
        para_start = None
        for idx in range(section_start, section_end):
            if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                para_start = idx
                break
        if para_start is not None:
            para_end = find_paragraph_block_end(doc, section_start, section_end, target_paragraph)
            search_start, search_end = para_start, para_end

    matches: list[int] = []
    for idx in range(search_start, search_end):
        if matches_top_level_point_marker(doc.paragraphs[idx].text, target_point):
            matches.append(idx)

    if not matches:
        return None
    if target_paragraph:
        # Roman-looking markers such as (i)/(ii)/(iii) can exist both as
        # nested children and as top-level points within the same paragraph
        # block. Anchoring to the first textual match can therefore hijack a
        # top-level replacement instruction and mutate a newly inserted nested
        # child instead. When multiple matches exist, defer to similarity
        # scoring unless a narrower parent scope was resolved elsewhere.
        if is_roman_point_marker(target_point) and len(matches) != 1:
            return None
        return matches[0]

    # Without paragraph scope, duplicate point markers are common across the
    # article (e.g. multiple "(i)" lists). Avoid anchoring to the first match;
    # let similarity matching decide instead.
    if len(matches) != 1:
        return None
    return matches[0]


def infer_article_paragraph_from_referenced_point(
    doc: Document,
    *,
    target_article: str | None,
    amend_text: str,
) -> str | None:
    """Infer paragraph number from references like 'point (da)(iii)' in added subparagraph text."""
    if not target_article:
        return None
    bounds = find_article_section_bounds(doc, target_article)
    if bounds is None:
        return None

    referenced_points = [m.lower() for m in re.findall(r"\bpoint\s*\(([a-z]+)\)", amend_text, flags=re.I)]
    if not referenced_points:
        return None

    section_start, section_end = bounds

    def paragraph_of_point(point: str) -> str | None:
        current_paragraph: str | None = None
        for idx in range(section_start, section_end):
            para_marker = first_top_level_paragraph_marker(doc.paragraphs[idx].text)
            if para_marker is not None:
                current_paragraph = para_marker
            if matches_top_level_point_marker(doc.paragraphs[idx].text, point):
                return current_paragraph
        return None

    candidates: list[str] = []
    for point in referenced_points:
        p = paragraph_of_point(point)
        if p:
            candidates.append(p)

    if not candidates:
        return None
    # Require a single unambiguous paragraph target across referenced points.
    unique = sorted(set(candidates))
    if len(unique) == 1:
        return unique[0]
    return None


def infer_paragraph_from_instruction_point_references(
    doc: Document,
    *,
    target_article: str | None,
    source_instruction: str,
) -> str | None:
    """Infer target paragraph from points explicitly referenced in instruction text."""
    if not target_article:
        return None
    bounds = find_article_section_bounds(doc, target_article)
    if bounds is None:
        return None

    refs = [m.lower() for m in re.findall(r"\(([a-z]+)\)", source_instruction, flags=re.I)]
    if not refs:
        return None
    # Keep deterministic order and remove duplicates.
    refs = list(dict.fromkeys(refs))

    section_start, section_end = bounds
    current_paragraph: str | None = None
    point_to_paragraphs: dict[str, set[str]] = {}
    for idx in range(section_start, section_end):
        para_marker = first_top_level_paragraph_marker(doc.paragraphs[idx].text)
        if para_marker is not None:
            current_paragraph = para_marker
        point_marker = first_top_level_point_marker(doc.paragraphs[idx].text)
        if point_marker is not None and current_paragraph is not None:
            point_to_paragraphs.setdefault(point_marker.lower(), set()).add(current_paragraph)

    if not all(p in point_to_paragraphs for p in refs):
        return None

    # Resolve by intersection: all referenced points must coexist under one
    # paragraph marker. If multiple paragraphs satisfy this, keep it
    # unresolved rather than biasing to first occurrence.
    common = set(point_to_paragraphs[refs[0]])
    for ref in refs[1:]:
        common &= point_to_paragraphs[ref]
    if len(common) == 1:
        return sorted(common)[0]
    return None


def infer_paragraph_from_article_similarity(
    doc: Document,
    *,
    target_article: str | None,
    amend_text: str,
    min_score: float = 0.45,
) -> str | None:
    """Infer paragraph by best textual similarity within the target article.

    Used as a guardrail for replacement instructions that only mention an inner
    subparagraph ordinal (e.g. "in the second subparagraph...") and otherwise
    risk drifting to the wrong paragraph via point-marker references.
    """
    if not target_article:
        return None

    bounds = find_article_section_bounds(doc, target_article)
    if bounds is None:
        return None

    section_start, section_end = bounds
    best_score = 0.0
    best_paragraph: str | None = None
    ambiguous = False
    current_paragraph: str | None = None

    for idx in range(section_start, section_end):
        para_txt = strip_src_artifacts(doc.paragraphs[idx].text)
        marker = first_top_level_paragraph_marker(para_txt)
        if marker is not None:
            current_paragraph = marker

        if current_paragraph is None:
            continue

        score = paragraph_similarity(para_txt, amend_text)
        if score > best_score:
            best_score = score
            best_paragraph = current_paragraph
            ambiguous = False
        elif score == best_score and score >= min_score and current_paragraph != best_paragraph:
            ambiguous = True

    if best_paragraph and best_score >= min_score and not ambiguous:
        return best_paragraph
    return None


def extract_subparagraph_ordinal_from_instruction(text: str) -> int | None:
    txt = normalize_text(text).lower()
    m = re.search(r"\b(first|second|third|fourth|fifth)\s+subparagraph\b", txt)
    if not m:
        return None
    ordinal_map = {
        "first": 1,
        "second": 2,
        "third": 3,
        "fourth": 4,
        "fifth": 5,
    }
    return ordinal_map.get(m.group(1))


def extract_inner_paragraph_ordinal_from_instruction(text: str) -> int | None:
    """Extract ordinal for inner paragraph/subparagraph references.

    Example: "the first paragraph is replaced" or
    "the second subparagraph is replaced".
    """
    txt = normalize_text(text).lower()
    m = re.search(r"\b(first|second|third|fourth|fifth)\s+(?:sub)?paragraph\b", txt)
    if not m:
        return None
    ordinal_map = {
        "first": 1,
        "second": 2,
        "third": 3,
        "fourth": 4,
        "fifth": 5,
    }
    return ordinal_map.get(m.group(1))


def find_target_subparagraph_index(
    doc: Document,
    *,
    target_article: str | None,
    target_paragraph: str | None,
    subparagraph_ordinal: int | None,
) -> int | None:
    if not target_article or not target_paragraph or not subparagraph_ordinal or subparagraph_ordinal < 1:
        return None
    bounds = find_article_section_bounds(doc, target_article)
    if bounds is None:
        return None
    section_start, section_end = bounds

    paragraph_start: int | None = None
    for idx in range(section_start, section_end):
        if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
            paragraph_start = idx
            break
    if paragraph_start is None:
        return None

    paragraph_end = find_paragraph_block_end(doc, section_start, section_end, target_paragraph)
    candidates: list[int] = []
    for idx in range(paragraph_start, paragraph_end):
        txt = normalize_text(doc.paragraphs[idx].text)
        if not txt:
            continue
        if first_top_level_point_marker(txt) is not None:
            continue
        candidates.append(idx)

    wanted = subparagraph_ordinal - 1
    if wanted < len(candidates):
        return candidates[wanted]
    return None


def best_match_in_doc(doc_texts: list[str], amendment_text: str) -> tuple[int, float]:
    best_idx = -1
    best_score = 0.0
    for idx, candidate in enumerate(doc_texts):
        score = paragraph_similarity(candidate, amendment_text)
        if score > best_score:
            best_score = score
            best_idx = idx
    return best_idx, best_score


def split_article_number(marker: str) -> tuple[str, str]:
    m = re.fullmatch(r"(\d+)([A-Za-z]*)", (marker or "").strip())
    if not m:
        return marker, ""
    return m.group(1), m.group(2)


def infer_insert_after_article_for_unanchored_block(
    block_items: list[dict[str, str]],
    doc: Document,
) -> tuple[str | None, str | None]:
    """Infer an article anchor for a full-article insertion block.

    For blocks like "the following Article is inserted:" where parsed items lack
    explicit target article, use repeated references such as "Article 14a" in
    the block body and anchor insertion after the base article ("14").
    """
    counts: dict[str, int] = {}
    for item in block_items:
        text = normalize_text(item.get("text", ""))
        for marker in re.findall(r"\bArticle\s+(\d+[A-Za-z])\b", text, flags=re.I):
            key = marker.lower()
            counts[key] = counts.get(key, 0) + 1

    if not counts:
        return None, None

    ranked = sorted(counts.items(), key=lambda kv: kv[1], reverse=True)
    best_marker, best_count = ranked[0]
    if best_count < 2:
        return None, None

    # If the letter-suffixed article already exists, skip inference.
    if find_article_section_bounds(doc, best_marker) is not None:
        return None, None

    base_num, suffix = split_article_number(best_marker)
    if not suffix:
        return None, None
    if find_article_section_bounds(doc, base_num) is None:
        return None, None
    return base_num, best_marker


def apply_single_amending_regulation(
    base_docx: Path,
    amending_html: Path,
    color: str,
    output_docx: Path,
    analysis_json: Path,
    max_analysis_items: int = 0,
    identified_payload: dict | None = None,
    require_full_application: bool = True,
) -> Path:
    rgb = parse_color(color)
    set_revision_color(rgb)

    doc = Document(str(base_docx))
    preexisting_paragraphs = [p for p in doc.paragraphs if strip_src_artifacts(p.text)]
    preexisting_doc_texts = [strip_src_artifacts(p.text) for p in preexisting_paragraphs]

    if identified_payload is not None:
        amending_title = normalize_text(str(identified_payload.get("amending_title", "")))
        amending_recitals = [normalize_text(r) for r in identified_payload.get("recitals", []) if normalize_text(r)]
        amending_items = [it for it in identified_payload.get("items", []) if normalize_text(it.get("text", ""))]
        amending_items = backfill_targets_from_instruction(amending_items)
        entry_into_force_block = identified_payload.get("entry_into_force_block")
    else:
        amending_title, amending_recitals, amending_items = collect_amending_blocks(amending_html)
        entry_into_force_block = extract_entry_into_force_signoff_block(amending_html)
    if is_non_legislative_payload_text(amending_title):
        amending_title = ""
    amending_recitals = [format_amendment_item_text(r) for r in amending_recitals if not is_non_legislative_payload_text(r)]
    amending_items = [item for item in amending_items if not is_non_legislative_payload_text(item.get("text", ""))]
    amending_items = normalize_amendment_items(coalesce_marker_body_items(amending_items))
    amending_items = backfill_targets_from_instruction(amending_items)
    amending_items = normalize_replacement_target_points(amending_items)

    title_insert_index = min(find_title_block_end_index(doc), len(doc.paragraphs))
    insert_amending_title(doc, title_insert_index, amending_title)

    recital_style = None
    last_recital_idx = find_last_recital_index(doc)
    if last_recital_idx is not None:
        recital_style = doc.paragraphs[last_recital_idx].style.name

    adoption_idx = find_adoption_formula_index(doc)
    if adoption_idx is None:
        adoption_idx = (last_recital_idx + 1) if last_recital_idx is not None else len(doc.paragraphs)

    insert_amending_recitals(doc, adoption_idx, amending_recitals, recital_style)
    enforce_amending_recital_block(doc, adoption_idx, amending_recitals, recital_style)

    if not isinstance(entry_into_force_block, dict):
        entry_into_force_block = extract_entry_into_force_signoff_block(amending_html)

    binding_idx = find_binding_clause_index(doc)
    if binding_idx is None:
        annex_idx = next(
            (idx for idx, para in enumerate(doc.paragraphs) if annex_heading_number(para.text) is not None),
            len(doc.paragraphs),
        )
        binding_idx = annex_idx
    insert_amending_entry_into_force_block(doc, binding_idx, entry_into_force_block)

    signoff_insert_idx = next(
        (idx for idx, para in enumerate(doc.paragraphs) if annex_heading_number(para.text) is not None),
        len(doc.paragraphs),
    )
    insert_amending_final_signoff_block(doc, signoff_insert_idx, entry_into_force_block)
    insert_annex_va_from_amending_html(doc, amending_html)
    apply_annex_blocks_from_amending_html(doc, amending_html)

    analysis_items = []
    applied_count = 0
    inserted_count = 0
    inserted_paragraphs = []
    retrofit_inserted_paragraphs = []
    nested_parent_point_context: dict[tuple[str | None, str | None, str], str] = {}
    colon_parent_point_context: dict[tuple[str | None, str | None, str], str] = {}
    annex_replacement_offsets: dict[tuple[str, str, str], int] = {}
    last_resolved_point_paragraph_by_article: dict[str, str] = {}

    items_to_process = amending_items if max_analysis_items <= 0 else amending_items[:max_analysis_items]

    inferred_article_block_meta: dict[int, tuple[str, str]] = {}
    idx = 0
    while idx < len(items_to_process):
        probe = items_to_process[idx]
        probe_instruction = normalize_text(probe.get("source_instruction", ""))
        if (
            re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", probe_instruction, flags=re.I)
            and not probe.get("target_article_number")
            and not probe.get("target_annex_number")
        ):
            block_start = idx
            while idx < len(items_to_process):
                inner = items_to_process[idx]
                inner_instruction = normalize_text(inner.get("source_instruction", ""))
                if not (
                    re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", inner_instruction, flags=re.I)
                    and not inner.get("target_article_number")
                    and not inner.get("target_annex_number")
                ):
                    break
                idx += 1
            block_end = idx
            block_items = items_to_process[block_start:block_end]

            explicit_marker = None
            for block_item in block_items:
                candidate_marker = normalize_text(str(block_item.get("inserted_article_marker") or ""))
                if re.fullmatch(r"\d+[A-Za-z]", candidate_marker):
                    explicit_marker = candidate_marker.lower()
                    break

            insert_after_article: str | None = None
            new_article_marker: str | None = None
            if explicit_marker:
                base_article, suffix = split_article_number(explicit_marker)
                if suffix and find_article_section_bounds(doc, base_article) is not None:
                    insert_after_article = base_article
                    new_article_marker = explicit_marker
            if not insert_after_article or not new_article_marker:
                insert_after_article, new_article_marker = infer_insert_after_article_for_unanchored_block(block_items, doc)

            if insert_after_article and new_article_marker:
                for mark_idx in range(block_start, block_end):
                    inferred_article_block_meta[mark_idx] = (insert_after_article, new_article_marker)
            continue
        idx += 1

    inserted_inferred_article_headings: set[str] = set()

    for item_idx, item in enumerate(items_to_process):
        inferred_insert_after_article, inferred_new_article = inferred_article_block_meta.get(item_idx, (None, None))
        amend_text = item["text"]
        best_idx, best_score = best_match_in_doc(preexisting_doc_texts, amend_text)
        base_text = preexisting_doc_texts[best_idx] if best_idx >= 0 else ""
        source_instruction = normalize_text(item.get("source_instruction", ""))
        sentence_addition_instruction = bool(
            re.search(r"\bthe following sentence is added\b", source_instruction, flags=re.I)
        )
        subparagraph_addition_instruction = bool(
            re.search(
                r"\bthe following subparagraph(?:s)?\s+(?:is|are)\s+(?:added|inserted)\b",
                source_instruction,
                flags=re.I,
            )
        )
        target_article = item.get("target_article_number") or inferred_insert_after_article
        target_paragraph = item.get("target_paragraph_number")
        if (
            inferred_insert_after_article
            and not item.get("target_article_number")
            and not item.get("target_paragraph_number")
        ):
            # In inferred full-article blocks, route unanchored lines into the
            # inserted article section. Preserve parsed paragraph context when
            # available (for continuation-line indentation like Article 14a(1)).
            target_paragraph = None
        target_point = item.get("target_point_marker")
        target_annex = normalize_text(str(item.get("target_annex_number") or ""))
        if not target_annex:
            target_annex = normalize_text(extract_target_annex_from_instruction(source_instruction) or "")
        target_annex_point = normalize_text(str(item.get("target_annex_point_marker") or ""))
        if not target_annex_point:
            target_annex_point = normalize_text(extract_target_annex_point_from_instruction(source_instruction) or "")
        amendment_kind = normalize_text(str(item.get("amendment_kind") or ""))
        if (
            amendment_kind == "deletion"
            and re.search(r"\breplaced\s+by\s+the\s+following\b", source_instruction, flags=re.I)
        ):
            # Replacement language must never be treated as deletion.
            amendment_kind = "replacement"
            item["amendment_kind"] = "replacement"
        explicit_inserted_article_marker = normalize_text(str(item.get("inserted_article_marker") or ""))
        effective_insert_after_article = inferred_insert_after_article
        effective_new_article = inferred_new_article
        if (
            re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", source_instruction, flags=re.I)
            and re.fullmatch(r"\d+[A-Za-z]", explicit_inserted_article_marker)
        ):
            marker_base, marker_suffix = split_article_number(explicit_inserted_article_marker)
            if marker_suffix and find_article_section_bounds(doc, marker_base) is not None:
                effective_insert_after_article = marker_base
                effective_new_article = explicit_inserted_article_marker.lower()
        full_article_insertion_block = bool(
            re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", source_instruction, flags=re.I)
            and effective_new_article
        )
        if full_article_insertion_block:
            # Anchor the inserted body to the newly created article marker
            # (e.g. 14a), not to the base article after which it is inserted.
            # This prevents 14a content from leaking into Article 14.
            target_article = effective_new_article
            # Some extracted continuation lines in inserted-article blocks can
            # inherit a stale point marker from preceding items (e.g. after
            # (c)). For markerless body lines, clear that point context so
            # paragraph-level placement/indentation is applied.
            if (
                first_top_level_paragraph_marker(amend_text) is None
                and first_top_level_point_marker(amend_text) is None
            ):
                target_point = None
        instruction_key = (target_article, target_paragraph, source_instruction)
        insertion_instruction = bool(
            re.search(r"\bthe following\b", source_instruction, flags=re.I)
            and not re.search(r"\breplaced by the following\b", source_instruction, flags=re.I)
        )

        parent_point_for_nested = nested_parent_point_context.get(instruction_key)
        if parent_point_for_nested is None and target_point and is_roman_point_marker(target_point):
            parent_point_for_nested = colon_parent_point_context.get(instruction_key)

        if (
            target_point
            and not is_roman_point_marker(target_point)
            and insertion_instruction
        ):
            nested_parent_point_context[instruction_key] = target_point
            parent_point_for_nested = target_point
            if normalize_text(amend_text).endswith(":"):
                # A trailing colon on a non-roman point usually introduces a
                # nested list level, so keep it as fallback parent context.
                colon_parent_point_context[instruction_key] = target_point
        replacement_instruction = amendment_kind == "replacement" or bool(
            re.search(r"\breplaced by the following\b", source_instruction, flags=re.I)
        )
        article_insertion_instruction = amendment_kind == "insertion" or bool(
            re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+inserted\b", source_instruction, flags=re.I)
        )
        replacement_like_instruction = replacement_instruction and not article_insertion_instruction and not sentence_addition_instruction
        unanchored_article_insertion_instruction = bool(
            article_insertion_instruction
            and not target_article
            and not target_annex
            and not target_paragraph
            and not target_point
        )

        if (
            replacement_instruction
            and target_article
            and not target_paragraph
            and re.search(r"\b(?:first|second|third|fourth|fifth)\s+(?:sub)?paragraph\b", source_instruction, flags=re.I)
        ):
            inferred_paragraph = infer_article_paragraph_from_referenced_point(
                doc,
                target_article=target_article,
                amend_text=amend_text,
            )
            if inferred_paragraph:
                target_paragraph = inferred_paragraph
                item["target_paragraph_number"] = inferred_paragraph
            elif target_article in last_resolved_point_paragraph_by_article:
                # Keep nearby amendment context when the instruction omits the
                # paragraph number (common in chained point/subparagraph edits).
                target_paragraph = last_resolved_point_paragraph_by_article[target_article]
                item["target_paragraph_number"] = target_paragraph

        if (
            amendment_kind == "deletion"
            and target_article
            and target_point
            and not target_paragraph
        ):
            inferred_paragraph = infer_paragraph_from_instruction_point_references(
                doc,
                target_article=target_article,
                source_instruction=source_instruction,
            )
            if inferred_paragraph:
                target_paragraph = inferred_paragraph
                item["target_paragraph_number"] = inferred_paragraph
            elif (
                target_article in last_resolved_point_paragraph_by_article
                and re.search(r"\b(?:point|points)\b", source_instruction, flags=re.I)
            ):
                target_paragraph = last_resolved_point_paragraph_by_article[target_article]
                item["target_paragraph_number"] = target_paragraph

        if amendment_kind == "deletion":
            deletion_indices: list[int] = []
            if target_annex and target_annex_point:
                annex_bounds = find_annex_point_block_bounds(doc, target_annex, target_annex_point)
                if annex_bounds is not None:
                    point_start, point_end = annex_bounds
                    deletion_indices = list(range(point_start, point_end))
            else:
                if target_article and target_point:
                    point_bounds = find_target_point_block_bounds(
                        doc,
                        target_article=target_article,
                        target_paragraph=target_paragraph,
                        target_point=target_point,
                    )
                    if point_bounds is not None:
                        point_start, point_end = point_bounds
                        deletion_indices = list(range(point_start, point_end))
                elif target_article and target_paragraph:
                    bounds = find_article_section_bounds(doc, target_article)
                    if bounds is not None:
                        section_start, section_end = bounds
                        paragraph_start = None
                        for idx in range(section_start, section_end):
                            if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                                paragraph_start = idx
                                break
                        if paragraph_start is not None:
                            paragraph_end = find_paragraph_block_end(doc, section_start, section_end, target_paragraph)
                            deletion_indices = list(range(paragraph_start, paragraph_end))
                elif target_article:
                    bounds = find_article_section_bounds(doc, target_article)
                    if bounds is not None:
                        start_idx, end_idx = bounds
                        deletion_indices = list(range(start_idx, end_idx))

            deletion_indices = [idx for idx in deletion_indices if 0 <= idx < len(doc.paragraphs)]
            if deletion_indices:
                base_text = "\n".join(strip_src_artifacts(doc.paragraphs[idx].text) for idx in deletion_indices)
                best_score = 1.0 if base_text else 0.0
                diff = word_diff_stats(base_text, "")
                for idx in deletion_indices:
                    target_para = doc.paragraphs[idx]
                    if target_para._element.getparent() is not None:
                        original_text = strip_src_artifacts(target_para.text)
                        if target_point:
                            pruned = prune_deleted_point_from_multiline_paragraph(original_text, target_point)
                            if pruned is not None and normalize_text(pruned):
                                replace_paragraph_with_revision_marks(target_para, original_text, pruned)
                                continue
                        replace_paragraph_with_revision_marks(target_para, original_text, "")
                applied_mode = "replaced"
                applied_count += 1
            else:
                diff = word_diff_stats(base_text, "")
                applied_mode = "analysis_only"

            analysis_items.append(
                {
                    "article_number": item["article_number"],
                    "article_heading": item["article_heading"],
                    "amending_text": amend_text,
                    "source_instruction": item.get("source_instruction"),
                    "target_article_number": item.get("target_article_number"),
                    "target_paragraph_number": item.get("target_paragraph_number"),
                    "target_point_marker": item.get("target_point_marker"),
                    "target_annex_number": item.get("target_annex_number"),
                    "target_annex_point_marker": item.get("target_annex_point_marker"),
                    "best_match_index": best_idx,
                    "best_match_score": round(best_score, 6),
                    "best_match_text": base_text,
                    "applied_mode": applied_mode,
                    "diff": diff,
                }
            )
            continue

        inferred_target_paragraph_from_text = first_top_level_paragraph_marker(amend_text)
        if (
            replacement_instruction
            and target_article
            and not target_paragraph
            and inferred_target_paragraph_from_text
            and (
                re.search(r"\bArticle\s+\d+[A-Za-z]?\s+(?:is|are)\s+replaced\b", source_instruction, flags=re.I)
                or re.search(r"\bparagraph(?:s)?\s+\d+[a-z]?\b", source_instruction, flags=re.I)
            )
        ):
            target_paragraph = inferred_target_paragraph_from_text
            item["target_paragraph_number"] = inferred_target_paragraph_from_text

        explicit_point_replacement_instruction = bool(
            replacement_instruction
            and target_point
            and re.search(
                r"\bpoint(?:s)?\s*\([a-zivxlcdm]+\)(?:\s*(?:,|and|or|to)\s*\([a-zivxlcdm]+\))*"
                r"(?:[^.]{0,200}?)\b(?:is|are)\s+replaced\b",
                source_instruction,
                flags=re.I,
            )
        )
        whole_point_replacement_instruction = bool(
            explicit_point_replacement_instruction
            and re.search(r"\bpoint\s*\([a-z]+\)\s+is\s+replaced\s+by\s+the\s+following\b", source_instruction, flags=re.I)
            and not re.search(r"\bintroductory wording\b", source_instruction, flags=re.I)
        )
        explicit_paragraph_replacement_instruction = bool(
            replacement_instruction
            and target_article
            and target_paragraph
            and not target_point
            and re.search(
                r"\b(?:first|second|third|fourth|fifth)\s+(?:sub)?paragraph\b(?:[^.]{0,200}?)\bis\s+replaced\b",
                source_instruction,
                flags=re.I,
            )
        )
        explicit_numbered_paragraph_replacement_instruction = bool(
            replacement_instruction
            and target_article
            and target_paragraph
            and re.search(
                r"\bparagraph(?:s)?\s+\d+[a-z]?"
                r"(?:\s*(?:,|and|or|to)\s*\d+[a-z]?)*"
                r"(?:[^.]{0,200}?)\b(?:is|are)\s+replaced\b",
                source_instruction,
                flags=re.I,
            )
        )
        explicit_article_replacement_instruction = bool(
            replacement_instruction
            and target_article
            and target_paragraph
            and re.search(
                r"\bArticle\s+\d+[A-Za-z]?\s+(?:is|are)\s+replaced\b",
                source_instruction,
                flags=re.I,
            )
        )
        paragraph_replacement_instruction = bool(
            explicit_paragraph_replacement_instruction
            or explicit_numbered_paragraph_replacement_instruction
            or explicit_article_replacement_instruction
        )

        if (
            target_article
            and not target_paragraph
            and re.search(
                r"\bthe following subparagraphs?\s+(?:is|are)\s+(?:added|inserted)\b",
                source_instruction,
                flags=re.I,
            )
        ):
            inferred_paragraph = infer_article_paragraph_from_referenced_point(
                doc,
                target_article=target_article,
                amend_text=amend_text,
            )
            if inferred_paragraph:
                target_paragraph = inferred_paragraph
                item["target_paragraph_number"] = inferred_paragraph

        if (
            replacement_instruction
            and target_article
            and not target_paragraph
            and not target_point
            and re.search(r"\bpoint\s*\([a-z]+\)", amend_text, flags=re.I)
            and extract_inner_paragraph_ordinal_from_instruction(source_instruction) is not None
        ):
            # For instructions like "in the second subparagraph...", prefer
            # article-local textual alignment over point-marker inference.
            inferred_paragraph = infer_paragraph_from_article_similarity(
                doc,
                target_article=target_article,
                amend_text=amend_text,
            )
            if inferred_paragraph:
                target_paragraph = inferred_paragraph
                item["target_paragraph_number"] = inferred_paragraph

        if (
            replacement_instruction
            and target_article
            and not target_paragraph
            and not target_point
            and re.search(r"\bpoint\s*\([a-z]+\)", amend_text, flags=re.I)
        ):
            inferred_paragraph = infer_article_paragraph_from_referenced_point(
                doc,
                target_article=target_article,
                amend_text=amend_text,
            )
            if inferred_paragraph:
                target_paragraph = inferred_paragraph
                item["target_paragraph_number"] = inferred_paragraph

        if (
            target_article
            and not target_paragraph
            and re.search(r"\bpoints?\s*\([a-z]+\)", source_instruction, flags=re.I)
        ):
            inferred_paragraph = infer_paragraph_from_instruction_point_references(
                doc,
                target_article=target_article,
                source_instruction=source_instruction,
            )
            if inferred_paragraph:
                target_paragraph = inferred_paragraph
                item["target_paragraph_number"] = inferred_paragraph
        if (
            target_article
            and target_point
            and not target_paragraph
            and not replacement_instruction
            and target_article in last_resolved_point_paragraph_by_article
            and re.search(r"\b(?:point|points)\b", source_instruction, flags=re.I)
        ):
            # When a sequence of point-level amendments omits repeated
            # paragraph markers, keep using the last resolved paragraph scope
            # for that article in this run.
            target_paragraph = last_resolved_point_paragraph_by_article[target_article]
            item["target_paragraph_number"] = target_paragraph

        if (
            replacement_instruction
            and target_article
            and not target_point
            and extract_inner_paragraph_ordinal_from_instruction(source_instruction) is not None
            and not re.search(r"\bparagraph(?:s)?\s+\d+[a-z]?\b", source_instruction, flags=re.I)
            and target_article in last_resolved_point_paragraph_by_article
        ):
            # For inner-subparagraph replacements with no explicit paragraph
            # number, prefer the current article-local paragraph context from
            # immediately preceding point-level operations.
            carry_paragraph = last_resolved_point_paragraph_by_article[target_article]
            if carry_paragraph:
                target_paragraph = carry_paragraph
                item["target_paragraph_number"] = carry_paragraph

        if target_article and target_paragraph and target_point:
            last_resolved_point_paragraph_by_article[target_article] = target_paragraph
        paragraph_replacement_ordinal = extract_inner_paragraph_ordinal_from_instruction(source_instruction) or 1
        replacement_min_score = 0.60 if explicit_point_replacement_instruction else HIGH_CONFIDENCE_MATCH

        explicit_annex_point_replacement_instruction = bool(
            replacement_instruction
            and target_annex
            and target_annex_point
            and re.search(r"\bin\s+Annex\s+[IVXLC]+\b", source_instruction, flags=re.I)
            and re.search(r"\bpoint\s+[IVXLC]+\b", source_instruction, flags=re.I)
            and re.search(r"\bis\s+replaced\b", source_instruction, flags=re.I)
        )

        anchored_replace_idx = find_target_point_paragraph_index(
            doc,
            target_article=target_article,
            target_paragraph=target_paragraph,
            target_point=target_point,
        )
        if (
            anchored_replace_idx is None
            and explicit_point_replacement_instruction
            and target_article
            and target_point
        ):
            bounds = find_article_section_bounds(doc, target_article)
            if bounds is not None:
                section_start, section_end = bounds
                candidate_indices: list[int] = []
                for idx in range(section_start, section_end):
                    if matches_top_level_point_marker(doc.paragraphs[idx].text, target_point):
                        candidate_indices.append(idx)
                if candidate_indices:
                    anchored_replace_idx = max(
                        candidate_indices,
                        key=lambda idx: paragraph_similarity(strip_src_artifacts(doc.paragraphs[idx].text), amend_text),
                    )
        if anchored_replace_idx is None and explicit_paragraph_replacement_instruction:
            anchored_replace_idx = find_target_subparagraph_index(
                doc,
                target_article=target_article,
                target_paragraph=target_paragraph,
                subparagraph_ordinal=paragraph_replacement_ordinal,
            )
        if anchored_replace_idx is not None and target_article and not target_paragraph:
            bounds = find_article_section_bounds(doc, target_article)
            if bounds is not None:
                section_start, section_end = bounds
                resolved_paragraph: str | None = None
                for idx in range(section_start, min(section_end, anchored_replace_idx + 1)):
                    marker = first_top_level_paragraph_marker(doc.paragraphs[idx].text)
                    if marker is not None:
                        resolved_paragraph = marker
                if resolved_paragraph:
                    target_paragraph = resolved_paragraph
                    item["target_paragraph_number"] = resolved_paragraph
                    if target_point:
                        last_resolved_point_paragraph_by_article[target_article] = resolved_paragraph
        allow_item_replacement = bool(
            (
                explicit_point_replacement_instruction
                and (
                    anchored_replace_idx is not None
                    or (best_idx >= 0 and best_score >= replacement_min_score)
                )
            )
            or (
                paragraph_replacement_instruction
                and (
                    anchored_replace_idx is not None
                    or (best_idx >= 0 and best_score >= replacement_min_score)
                )
            )
            or explicit_annex_point_replacement_instruction
        )
        if anchored_replace_idx is not None:
            anchored_para = doc.paragraphs[anchored_replace_idx]
            anchored_text = strip_src_artifacts(anchored_para.text)
            anchored_score = paragraph_similarity(anchored_text, amend_text)
            if anchored_score >= best_score:
                base_text = anchored_text
                best_score = anchored_score
                try:
                    best_idx = preexisting_paragraphs.index(anchored_para)
                except ValueError:
                    best_idx = -1

        diff = word_diff_stats(base_text, amend_text)

        applied_mode = "analysis_only"

        if (
            ENABLE_HEURISTIC_DOC_MUTATIONS
            and unanchored_article_insertion_instruction
            and first_top_level_paragraph_marker(amend_text) is None
        ):
            # A full-article insertion without a reliable legal anchor is unsafe
            # to apply heuristically because it can spill into unrelated sections.
            # Keep non-numbered lines in analysis-only mode until an explicit
            # anchor is parsed. Numbered top-level paragraphs (e.g. "1.", "2.")
            # are allowed through so the inserted article skeleton is not dropped.
            analysis_items.append(
                {
                    "article_number": item["article_number"],
                    "article_heading": item["article_heading"],
                    "amending_text": amend_text,
                    "source_instruction": item.get("source_instruction"),
                    "target_article_number": item.get("target_article_number"),
                    "target_paragraph_number": item.get("target_paragraph_number"),
                    "target_point_marker": item.get("target_point_marker"),
                    "target_annex_number": item.get("target_annex_number"),
                    "target_annex_point_marker": item.get("target_annex_point_marker"),
                    "best_match_index": best_idx,
                    "best_match_score": round(best_score, 6),
                    "best_match_text": base_text,
                    "applied_mode": applied_mode,
                    "diff": diff,
                }
            )
            continue

        # Annex point replacements are block-anchored and applied line-by-line.
        if ENABLE_HEURISTIC_DOC_MUTATIONS and explicit_annex_point_replacement_instruction:
            annex_bounds = find_annex_point_block_bounds(doc, target_annex, target_annex_point)
            if annex_bounds is not None:
                point_start, point_end = annex_bounds
                key = (target_annex.upper(), target_annex_point.lower(), source_instruction)
                line_offset = annex_replacement_offsets.get(key, 0)

                replacement_targets: list[int] = [point_start]
                replacement_targets.extend(
                    idx
                    for idx in range(point_start + 1, point_end)
                    if first_top_level_point_marker(doc.paragraphs[idx].text) is None
                )

                if line_offset < len(replacement_targets):
                    target_idx = replacement_targets[line_offset]
                    target_para = doc.paragraphs[target_idx]
                    replacement_text = amend_text
                    if line_offset == 0 and not matches_annex_point_marker(replacement_text, target_annex_point):
                        replacement_text = f"{target_annex_point.upper()}.\t{replacement_text}"
                    base_text = strip_src_artifacts(target_para.text)
                    best_score = paragraph_similarity(base_text, replacement_text)
                    diff = word_diff_stats(base_text, replacement_text)
                    replace_paragraph_with_revision_marks(target_para, base_text, replacement_text)
                    applied_mode = "replaced"
                    applied_count += 1
                    annex_replacement_offsets[key] = line_offset + 1
                else:
                    applied_mode = "analysis_only"

            analysis_items.append(
                {
                    "article_number": item["article_number"],
                    "article_heading": item["article_heading"],
                    "amending_text": amend_text,
                    "source_instruction": item.get("source_instruction"),
                    "target_article_number": item.get("target_article_number"),
                    "target_paragraph_number": item.get("target_paragraph_number"),
                    "target_point_marker": item.get("target_point_marker"),
                    "target_annex_number": item.get("target_annex_number"),
                    "target_annex_point_marker": item.get("target_annex_point_marker"),
                    "best_match_index": best_idx,
                    "best_match_score": round(best_score, 6),
                    "best_match_text": base_text,
                    "applied_mode": applied_mode,
                    "diff": diff,
                }
            )
            continue

        # Structured paragraph replacements (e.g. "in Article 48, paragraph 2
        # is replaced by the following") should replace existing marker lines
        # against the original text where anchors exist, and only insert truly
        # new markers.
        if (
            ENABLE_HEURISTIC_DOC_MUTATIONS
            and paragraph_replacement_instruction
            and target_article
            and target_paragraph
        ):
            structured_target_idx: int | None = None
            paragraph_start: int | None = None
            paragraph_end: int | None = None
            amend_point_marker: str | None = None
            amend_paragraph_marker: str | None = first_top_level_paragraph_marker(amend_text)
            section_start: int | None = None
            section_end: int | None = None
            bounds = find_article_section_bounds(doc, target_article)
            if bounds is not None:
                section_start, section_end = bounds
                for idx in range(section_start, section_end):
                    if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                        paragraph_start = idx
                        break
                if paragraph_start is not None:
                    paragraph_end = find_paragraph_block_end(
                        doc,
                        section_start,
                        section_end,
                        target_paragraph,
                    )
                    amend_point_marker = first_top_level_point_marker(amend_text) or target_point

                    if amend_paragraph_marker and amend_paragraph_marker == target_paragraph:
                        structured_target_idx = paragraph_start
                    elif amend_point_marker:
                        for idx in range(paragraph_start + 1, paragraph_end):
                            if matches_top_level_point_marker(doc.paragraphs[idx].text, amend_point_marker):
                                structured_target_idx = idx
                                break

                    if (
                        structured_target_idx is None
                        and explicit_paragraph_replacement_instruction
                        and anchored_replace_idx is not None
                        and paragraph_start is not None
                        and paragraph_end is not None
                        and paragraph_start <= anchored_replace_idx < paragraph_end
                        and amend_paragraph_marker is None
                        and amend_point_marker is None
                    ):
                        # For explicit subparagraph replacements, reuse the
                        # resolved local anchor instead of appending as a new
                        # markerless line at paragraph end.
                        structured_target_idx = anchored_replace_idx

            # In paragraph-level replacement blocks, newly introduced points
            # (e.g. replacing Article 48(2) with expanded points (c)-(g)) may
            # have no existing anchor yet. Insert them in legal marker order
            # within the target paragraph block and treat as part of replacement.
            if (
                structured_target_idx is None
                and paragraph_start is not None
                and section_start is not None
                and section_end is not None
                and amend_paragraph_marker is not None
                and amend_paragraph_marker != target_paragraph
                and not FORCE_ANALYSIS_ON_DOUBT
            ):
                insertion_index = find_insertion_index_by_paragraph_marker_order(
                    doc,
                    section_start,
                    section_end,
                    amend_paragraph_marker,
                )
                inserted = insert_plain_paragraph_before_index(doc, insertion_index)

                display_text = amend_text
                if not re.match(rf"^{re.escape(amend_paragraph_marker)}\.(?=\s|\t|$)", display_text, flags=re.I):
                    display_text = f"{amend_paragraph_marker}.\t{display_text}"
                display_text = ensure_leading_marker_tab(display_text)
                inserted.add_run(display_text)

                layout_idx = find_previous_top_level_paragraph_index(doc, section_start, insertion_index)
                if layout_idx is not None:
                    base_para = doc.paragraphs[layout_idx]
                    if base_para._element.getparent() is not None:
                        inserted.style = base_para.style
                        inserted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        inserted.paragraph_format.left_indent = base_para.paragraph_format.left_indent
                        inserted.paragraph_format.first_line_indent = base_para.paragraph_format.first_line_indent

                inserted.paragraph_format.left_indent = Cm(1)
                inserted.paragraph_format.first_line_indent = Cm(-1)
                inserted_paragraphs.append((inserted, item.get("source_instruction")))

                base_text = ""
                best_score = 0.0
                diff = word_diff_stats(base_text, display_text)
                applied_mode = "replaced"
                applied_count += 1

                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue

            if (
                structured_target_idx is None
                and paragraph_start is not None
                and paragraph_end is not None
                and amend_point_marker is not None
                and not FORCE_ANALYSIS_ON_DOUBT
            ):
                insertion_index = find_insertion_index_by_point_marker_order_within_paragraph(
                    doc,
                    paragraph_start,
                    paragraph_end,
                    amend_point_marker,
                )
                inserted = insert_plain_paragraph_before_index(doc, insertion_index)

                display_text = amend_text
                if not re.match(rf"^\({re.escape(amend_point_marker)}\)(?=\s|\t|$)", display_text, flags=re.I):
                    display_text = f"({amend_point_marker})\t{display_text}"
                display_text = ensure_leading_marker_tab(display_text)
                inserted.add_run(display_text)

                # Keep paragraph style/layout coherent within the replaced block.
                layout_idx = find_previous_top_level_paragraph_index(doc, paragraph_start, insertion_index)
                if layout_idx is not None:
                    base_para = doc.paragraphs[layout_idx]
                    if base_para._element.getparent() is not None:
                        inserted.style = base_para.style
                        inserted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        inserted.paragraph_format.left_indent = base_para.paragraph_format.left_indent
                        inserted.paragraph_format.first_line_indent = base_para.paragraph_format.first_line_indent

                inserted.paragraph_format.left_indent = Cm(2)
                inserted.paragraph_format.first_line_indent = Cm(-1)
                inserted_paragraphs.append((inserted, item.get("source_instruction")))

                base_text = ""
                best_score = 0.0
                diff = word_diff_stats(base_text, display_text)
                applied_mode = "replaced"
                applied_count += 1

                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue

            # In point-replacement blocks, newly introduced points may not have
            # an existing anchor (e.g. adding point (ba) when replacing points
            # (a) and (b)). Insert by marker order inside the resolved
            # paragraph block and treat as replacement content.
            if (
                structured_target_idx is None
                and explicit_point_replacement_instruction
                and target_article
                and target_paragraph
                and target_point
                and not is_roman_point_marker(target_point)
                and not FORCE_ANALYSIS_ON_DOUBT
            ):
                point_bounds = find_article_section_bounds(doc, target_article)
                if point_bounds is not None:
                    section_start, section_end = point_bounds
                    paragraph_start = None
                    for idx in range(section_start, section_end):
                        if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                            paragraph_start = idx
                            break
                    if paragraph_start is not None:
                        paragraph_end = find_paragraph_block_end(
                            doc,
                            section_start,
                            section_end,
                            target_paragraph,
                        )
                        insertion_index = find_insertion_index_by_point_marker_order_within_paragraph(
                            doc,
                            paragraph_start,
                            paragraph_end,
                            target_point,
                        )
                        inserted = insert_plain_paragraph_before_index(doc, insertion_index)

                        display_text = amend_text
                        if not re.match(rf"^\({re.escape(target_point)}\)(?=\s|\t|$)", display_text, flags=re.I):
                            display_text = f"({target_point})\t{display_text}"
                        display_text = ensure_leading_marker_tab(display_text)
                        inserted.add_run(display_text)

                        inserted.paragraph_format.left_indent = Cm(2)
                        inserted.paragraph_format.first_line_indent = Cm(-1)
                        inserted_paragraphs.append((inserted, item.get("source_instruction")))

                        base_text = ""
                        best_score = 0.0
                        diff = word_diff_stats(base_text, display_text)
                        applied_mode = "replaced"
                        applied_count += 1

                        analysis_items.append(
                            {
                                "article_number": item["article_number"],
                                "article_heading": item["article_heading"],
                                "amending_text": amend_text,
                                "source_instruction": item.get("source_instruction"),
                                "target_article_number": item.get("target_article_number"),
                                "target_paragraph_number": item.get("target_paragraph_number"),
                                "target_point_marker": item.get("target_point_marker"),
                                "target_annex_number": item.get("target_annex_number"),
                                "target_annex_point_marker": item.get("target_annex_point_marker"),
                                "best_match_index": best_idx,
                                "best_match_score": round(best_score, 6),
                                "best_match_text": base_text,
                                "applied_mode": applied_mode,
                                "diff": diff,
                            }
                        )
                        continue

            if (
                structured_target_idx is None
                and paragraph_start is not None
                and paragraph_end is not None
                and paragraph_replacement_instruction
                and amend_paragraph_marker is None
                and amend_point_marker is None
                and not explicit_paragraph_replacement_instruction
                and not re.search(r"\bintroductory\s+wording\b", source_instruction, flags=re.I)
                and not FORCE_ANALYSIS_ON_DOUBT
            ):
                # Multi-line replacements often provide one numbered head line
                # followed by markerless continuation/subparagraph lines.
                # Keep those continuation lines inside the same target
                # paragraph block as replacement content.
                insertion_index = paragraph_end
                inserted = insert_plain_paragraph_before_index(doc, insertion_index)
                display_text = ensure_leading_marker_tab(amend_text)
                inserted.add_run(display_text)

                inserted.paragraph_format.left_indent = Cm(1)
                inserted.paragraph_format.first_line_indent = None
                inserted_paragraphs.append((inserted, item.get("source_instruction")))

                base_text = ""
                best_score = 0.0
                diff = word_diff_stats(base_text, display_text)
                applied_mode = "replaced"
                applied_count += 1

                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue

            if (
                structured_target_idx is None
                and paragraph_start is not None
                and paragraph_replacement_instruction
                and amend_paragraph_marker is None
                and amend_point_marker is None
                and re.search(r"\bintroductory\s+(?:wording|sentence)\b", source_instruction, flags=re.I)
            ):
                # Introductory replacements without explicit local markers must
                # still anchor to the targeted paragraph head, not to global
                # best-match text in nearby paragraphs.
                structured_target_idx = paragraph_start

            # In replacement blocks that introduce brand-new top-level
            # paragraphs (e.g. 2a, 2b, 2c), insert in legal paragraph-marker
            # order within the target article section.
            if (
                structured_target_idx is None
                and paragraph_start is None
                and amend_paragraph_marker is not None
                and section_start is not None
                and section_end is not None
                and not FORCE_ANALYSIS_ON_DOUBT
            ):
                insertion_index = find_insertion_index_by_paragraph_marker_order(
                    doc,
                    section_start,
                    section_end,
                    amend_paragraph_marker,
                )
                inserted = insert_plain_paragraph_before_index(doc, insertion_index)

                display_text = amend_text
                if not re.match(rf"^{re.escape(amend_paragraph_marker)}\.(?=\s|\t|$)", display_text, flags=re.I):
                    display_text = f"{amend_paragraph_marker}.\t{display_text}"
                display_text = ensure_leading_marker_tab(display_text)
                inserted.add_run(display_text)

                layout_idx = find_previous_top_level_paragraph_index(doc, section_start, insertion_index)
                if layout_idx is not None:
                    base_para = doc.paragraphs[layout_idx]
                    if base_para._element.getparent() is not None:
                        inserted.style = base_para.style
                        inserted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        inserted.paragraph_format.left_indent = base_para.paragraph_format.left_indent
                        inserted.paragraph_format.first_line_indent = base_para.paragraph_format.first_line_indent

                inserted.paragraph_format.left_indent = Cm(1)
                inserted.paragraph_format.first_line_indent = Cm(-1)
                inserted_paragraphs.append((inserted, item.get("source_instruction")))

                base_text = ""
                best_score = 0.0
                diff = word_diff_stats(base_text, display_text)
                applied_mode = "replaced"
                applied_count += 1

                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue

            if structured_target_idx is not None:
                target_para = doc.paragraphs[structured_target_idx]
                base_text = strip_src_artifacts(target_para.text)
                amend_text_for_replacement = inject_missing_legal_marker_for_replacement(base_text, amend_text)
                best_score = paragraph_similarity(base_text, amend_text)
                diff = word_diff_stats(base_text, amend_text_for_replacement)
                if normalize_text(base_text) == normalize_text(amend_text_for_replacement):
                    applied_mode = "already_applied"
                else:
                    replace_paragraph_with_revision_marks(
                        target_para,
                        base_text,
                        amend_text_for_replacement,
                        preserve_prior_inserted_context=not re.search(
                            r"\bintroductory\s+wording\b", source_instruction, flags=re.I
                        ),
                    )
                    applied_mode = "replaced"
                    applied_count += 1

                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue

        if (
            ENABLE_HEURISTIC_DOC_MUTATIONS
            and allow_item_replacement
            and replacement_instruction
            and (best_idx >= 0 or anchored_replace_idx is not None)
            and (
                best_score >= replacement_min_score
                or (anchored_replace_idx is not None and best_score >= 0.0)
            )
        ):
            target_para = None
            if anchored_replace_idx is not None:
                target_para = doc.paragraphs[anchored_replace_idx]
            elif best_idx < len(preexisting_paragraphs):
                target_para = preexisting_paragraphs[best_idx]
            if target_para is not None and target_para._element.getparent() is not None:
                amend_text_for_replacement = inject_missing_legal_marker_for_replacement(base_text, amend_text)
                if normalize_text(base_text) == normalize_text(amend_text_for_replacement):
                    # Already applied in a previous run; keep existing formatting as-is.
                    if not paragraph_has_inserted_style(target_para):
                        retrofit_inserted_paragraphs.append((target_para, item.get("source_instruction")))
                    applied_mode = "already_applied"
                else:
                    if whole_point_replacement_instruction and target_article and target_point:
                        point_bounds = find_target_point_block_bounds(
                            doc,
                            target_article=target_article,
                            target_paragraph=target_paragraph,
                            target_point=target_point,
                        )
                        if point_bounds is not None:
                            point_start, point_end = point_bounds
                            head_para = doc.paragraphs[point_start]
                            replace_paragraph_with_revision_marks(
                                head_para,
                                strip_src_artifacts(head_para.text),
                                amend_text_for_replacement,
                            )
                            # Keep deleted subparagraphs visible as struck-through
                            # when a whole top-level point is replaced.
                            for idx in range(point_start + 1, point_end):
                                tail_para = doc.paragraphs[idx]
                                tail_text = strip_src_artifacts(tail_para.text)
                                if tail_text:
                                    replace_paragraph_with_revision_marks(tail_para, tail_text, "")
                        else:
                            replace_paragraph_with_revision_marks(target_para, base_text, amend_text_for_replacement)
                    else:
                        replace_paragraph_with_revision_marks(
                            target_para,
                            base_text,
                            amend_text_for_replacement,
                            preserve_prior_inserted_context=not re.search(
                                r"\bintroductory\s+wording\b", source_instruction, flags=re.I
                            ),
                        )
                    applied_mode = "replaced"
                    applied_count += 1
            elif replacement_like_instruction and not (target_article or target_annex):
                applied_mode = "analysis_only"
                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue

        if (
            ENABLE_HEURISTIC_DOC_MUTATIONS
            and replacement_instruction
            and explicit_point_replacement_instruction
            and target_article
            and target_paragraph
            and target_point
            and not is_roman_point_marker(target_point)
            and anchored_replace_idx is None
            and not FORCE_ANALYSIS_ON_DOUBT
        ):
            bounds = find_article_section_bounds(doc, target_article)
            if bounds is not None:
                section_start, section_end = bounds
                paragraph_start = None
                for idx in range(section_start, section_end):
                    if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                        paragraph_start = idx
                        break
                if paragraph_start is not None:
                    paragraph_end = find_paragraph_block_end(
                        doc,
                        section_start,
                        section_end,
                        target_paragraph,
                    )
                    insertion_index = find_insertion_index_by_point_marker_order_within_paragraph(
                        doc,
                        paragraph_start,
                        paragraph_end,
                        target_point,
                    )
                    inserted = insert_plain_paragraph_before_index(doc, insertion_index)

                    display_text = amend_text
                    if not re.match(rf"^\({re.escape(target_point)}\)(?=\s|\t|$)", display_text, flags=re.I):
                        display_text = f"({target_point})\t{display_text}"
                    display_text = ensure_leading_marker_tab(display_text)
                    inserted.add_run(display_text)

                    inserted.paragraph_format.left_indent = Cm(2)
                    inserted.paragraph_format.first_line_indent = Cm(-1)
                    inserted_paragraphs.append((inserted, item.get("source_instruction")))

                    base_text = ""
                    best_score = 0.0
                    diff = word_diff_stats(base_text, display_text)
                    applied_mode = "replaced"
                    applied_count += 1

                    analysis_items.append(
                        {
                            "article_number": item["article_number"],
                            "article_heading": item["article_heading"],
                            "amending_text": amend_text,
                            "source_instruction": item.get("source_instruction"),
                            "target_article_number": item.get("target_article_number"),
                            "target_paragraph_number": item.get("target_paragraph_number"),
                            "target_point_marker": item.get("target_point_marker"),
                            "target_annex_number": item.get("target_annex_number"),
                            "target_annex_point_marker": item.get("target_annex_point_marker"),
                            "best_match_index": best_idx,
                            "best_match_score": round(best_score, 6),
                            "best_match_text": base_text,
                            "applied_mode": applied_mode,
                            "diff": diff,
                        }
                    )
                    continue
        if replacement_like_instruction:
            # Never convert unresolved replacements into insertions.
            # If replacement succeeded/already exists, keep that result.
            if applied_mode not in {"replaced", "already_applied"}:
                applied_mode = "analysis_only"
                analysis_items.append(
                    {
                        "article_number": item["article_number"],
                        "article_heading": item["article_heading"],
                        "amending_text": amend_text,
                        "source_instruction": item.get("source_instruction"),
                        "target_article_number": item.get("target_article_number"),
                        "target_paragraph_number": item.get("target_paragraph_number"),
                        "target_point_marker": item.get("target_point_marker"),
                        "target_annex_number": item.get("target_annex_number"),
                        "target_annex_point_marker": item.get("target_annex_point_marker"),
                        "best_match_index": best_idx,
                        "best_match_score": round(best_score, 6),
                        "best_match_text": base_text,
                        "applied_mode": applied_mode,
                        "diff": diff,
                    }
                )
                continue
        elif ENABLE_HEURISTIC_DOC_MUTATIONS:
            insertion_index: int | None = None
            layout_source_para = None

            if effective_insert_after_article and effective_new_article:
                inferred_bounds = find_article_section_bounds(doc, effective_insert_after_article)
                if inferred_bounds is not None:
                    _, inferred_section_end = inferred_bounds

                    if effective_new_article not in inserted_inferred_article_headings:
                        heading = insert_plain_paragraph_before_index(doc, inferred_section_end)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        heading.paragraph_format.left_indent = Cm(0)
                        heading.paragraph_format.first_line_indent = None
                        heading_run = heading.add_run(f"Article {effective_new_article}")
                        style_inserted_run_no_bold_change(heading_run)
                        inserted_inferred_article_headings.add(effective_new_article)

                    inferred_bounds = find_article_section_bounds(doc, effective_insert_after_article)
                    if effective_new_article:
                        new_article_bounds = find_article_section_bounds(doc, effective_new_article)
                        if new_article_bounds is not None:
                            _, new_article_end = new_article_bounds
                            insertion_index = new_article_end

            if sentence_addition_instruction and target_article and target_paragraph and not target_point:
                target_subparagraph_idx = find_target_subparagraph_index(
                    doc,
                    target_article=target_article,
                    target_paragraph=target_paragraph,
                    subparagraph_ordinal=extract_subparagraph_ordinal_from_instruction(source_instruction),
                )
                if target_subparagraph_idx is not None:
                    target_para = doc.paragraphs[target_subparagraph_idx]
                    sentence_text = amend_text.strip()
                    if sentence_text:
                        if target_para.text and not target_para.text.endswith((" ", "\t")):
                            prefix = " "
                        else:
                            prefix = ""
                        run = target_para.add_run(prefix + sentence_text)
                        style_inserted_run(run)
                        applied_mode = "inserted"
                        inserted_count += 1
                        analysis_items.append(
                            {
                                "article_number": item["article_number"],
                                "article_heading": item["article_heading"],
                                "amending_text": amend_text,
                                "source_instruction": item.get("source_instruction"),
                                "target_article_number": item.get("target_article_number"),
                                "target_paragraph_number": item.get("target_paragraph_number"),
                                "target_point_marker": item.get("target_point_marker"),
                                "best_match_index": best_idx,
                                "best_match_score": round(best_score, 6),
                                "best_match_text": base_text,
                                "applied_mode": applied_mode,
                                "diff": diff,
                            }
                        )
                        continue

            if target_article:
                bounds = find_article_section_bounds(doc, target_article)
                if bounds is not None:
                    section_start, section_end = bounds
                    paragraph_start: int | None = None
                    if target_paragraph:
                        for idx in range(section_start, section_end):
                            if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                                paragraph_start = idx
                                break
                    if target_paragraph:
                        paragraph_block_end = find_paragraph_block_end(
                            doc,
                            section_start,
                            section_end,
                            target_paragraph,
                        )
                        if (
                            target_point
                            and not is_roman_point_marker(target_point)
                            and extract_subparagraph_ordinal_from_instruction(source_instruction) == 1
                        ):
                            cutoff_start = paragraph_start if paragraph_start is not None else section_start
                            paragraph_block_end = min(
                                paragraph_block_end,
                                find_first_subparagraph_block_end(doc, cutoff_start, paragraph_block_end),
                            )
                        if paragraph_block_end != section_end:
                            insertion_index = paragraph_block_end

                            if target_point:
                                if (
                                    is_roman_point_marker(target_point)
                                    and parent_point_for_nested
                                    and not is_roman_point_marker(parent_point_for_nested)
                                ):
                                    insertion_index = find_nested_roman_insertion_index_within_parent_point(
                                        doc,
                                        paragraph_start,
                                        paragraph_block_end,
                                        parent_point_for_nested,
                                    )
                                else:
                                    insertion_index = find_point_block_end_within_paragraph(
                                        doc,
                                        paragraph_start,
                                        paragraph_block_end,
                                        target_point,
                                    )
                                    if insertion_index == paragraph_block_end:
                                        insertion_index = find_insertion_index_by_point_marker_order_within_paragraph(
                                            doc,
                                            paragraph_start,
                                            paragraph_block_end,
                                            target_point,
                                        )
                        # For added/replacement blocks that introduce new
                        # top-level numbered paragraphs under a target
                        # paragraph anchor, use article-wide marker ordering.
                        # Otherwise repeated insertions reuse the same anchor
                        # and can reverse order (e.g. 8,7,6,5,4,3).
                        amend_paragraph_marker = first_top_level_paragraph_marker(amend_text)
                        if amend_paragraph_marker and target_point is None:
                            insertion_index = find_insertion_index_by_paragraph_marker_order(
                                doc,
                                section_start,
                                section_end,
                                amend_paragraph_marker,
                            )
                    elif target_point:
                        # Point-level insertions without an explicit paragraph anchor
                        # still need article-order placement. Falling back to the
                        # nearest textual match can misplace late-alphabet points
                        # like da/db or replacement points such as r/z.
                        insertion_index = find_insertion_index_by_point_marker_order_within_paragraph(
                            doc,
                            section_start,
                            section_end,
                            target_point,
                        )

                    if insertion_index is None:
                        new_marker = first_top_level_paragraph_marker(amend_text)
                        if new_marker:
                            insertion_index = find_insertion_index_by_paragraph_marker_order(
                                doc,
                                section_start,
                                section_end,
                                new_marker,
                            )

                    if 0 <= best_idx < len(preexisting_paragraphs):
                        best_para = preexisting_paragraphs[best_idx]
                        best_cur_idx = find_current_paragraph_index(doc, best_para)
                        if (
                            insertion_index is None
                            and best_cur_idx is not None
                            and section_start <= best_cur_idx < section_end
                        ):
                            insertion_index = best_cur_idx + 1
                    if insertion_index is None:
                        insertion_index = section_end

                    if target_paragraph and target_point is None:
                        layout_idx = find_previous_top_level_paragraph_index(doc, section_start, insertion_index)
                        if layout_idx is not None:
                            layout_source_para = doc.paragraphs[layout_idx]

            if insertion_index is None and 0 <= best_idx < len(preexisting_paragraphs):
                best_para = preexisting_paragraphs[best_idx]
                best_cur_idx = find_current_paragraph_index(doc, best_para)
                if best_cur_idx is not None:
                    insertion_index = best_cur_idx + 1

            if insertion_index is None:
                insertion_index = len(doc.paragraphs)

            inserted = insert_plain_paragraph_before_index(doc, insertion_index)
            display_text = amend_text
            point_marker = normalize_text(str(item.get("target_point_marker") or ""))
            display_point_marker = first_top_level_point_marker(display_text)
            suppress_point_prefix = bool(
                full_article_insertion_block
                and point_marker
                and display_point_marker is None
                and re.match(r"^[A-Z]", display_text)
            )
            if suppress_point_prefix:
                point_marker = ""
            if point_marker and not re.match(rf"^\({re.escape(point_marker)}\)(?=\s|\t|$)", display_text, flags=re.I):
                display_text = f"({point_marker})\t{display_text}"
                display_point_marker = point_marker
            display_text = ensure_leading_marker_tab(display_text)
            inserted.add_run(display_text)

            # First inherit formatting from closest matched paragraph, then mark as inserted.
            base_para = layout_source_para
            if base_para is None and 0 <= best_idx < len(preexisting_paragraphs):
                base_para = preexisting_paragraphs[best_idx]
            if base_para is not None:
                if base_para._element.getparent() is not None:
                    inserted.style = base_para.style
                    inserted.alignment = base_para.alignment
                    inserted.paragraph_format.left_indent = base_para.paragraph_format.left_indent
                    inserted.paragraph_format.first_line_indent = base_para.paragraph_format.first_line_indent
                    inserted.paragraph_format.right_indent = base_para.paragraph_format.right_indent
                    inserted.paragraph_format.space_before = base_para.paragraph_format.space_before
                    inserted.paragraph_format.space_after = base_para.paragraph_format.space_after
                    inserted.paragraph_format.line_spacing = base_para.paragraph_format.line_spacing

                    # Keep insertions visually in body-text flow.
                    inserted.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Apply standard legal indentation for the detected level so top-level
            # numbered insertions and nested point insertions are formatted consistently.
            detected_paragraph_marker = first_top_level_paragraph_marker(display_text)
            detected_point_marker = first_top_level_point_marker(display_text)
            if (
                full_article_insertion_block
                and detected_paragraph_marker is None
                and detected_point_marker is None
            ):
                # In full inserted Articles, markerless lines that belong to a
                # numbered paragraph (e.g. Article 14a(1) continuation
                # subparagraphs) should follow paragraph-body indentation.
                if target_paragraph and target_point is None:
                    inserted.paragraph_format.left_indent = Cm(1)
                    inserted.paragraph_format.first_line_indent = None
                else:
                    # Standalone body lines in inserted Articles remain flush-left.
                    inserted.paragraph_format.left_indent = Cm(0)
                    inserted.paragraph_format.first_line_indent = None
            elif detected_paragraph_marker is not None:
                inserted.paragraph_format.left_indent = Cm(1)
                inserted.paragraph_format.first_line_indent = Cm(-1)
            elif detected_point_marker is not None:
                fallback_nested_parent_is_non_roman = False
                if (
                    display_point_marker
                    and is_roman_point_marker(display_point_marker)
                    and target_article
                    and target_paragraph
                    and insertion_index is not None
                ):
                    bounds = find_article_section_bounds(doc, target_article)
                    if bounds is not None:
                        section_start, section_end = bounds
                        paragraph_start_idx = None
                        for idx in range(section_start, section_end):
                            if matches_top_level_paragraph_marker(doc.paragraphs[idx].text, target_paragraph):
                                paragraph_start_idx = idx
                                break
                        if paragraph_start_idx is not None:
                            last_point_before_insert: str | None = None
                            seen_non_roman_parent_point = False
                            for idx in range(paragraph_start_idx, max(paragraph_start_idx, insertion_index)):
                                marker = first_top_level_point_marker(doc.paragraphs[idx].text)
                                if marker is not None:
                                    last_point_before_insert = marker
                                    if not is_roman_point_marker(marker):
                                        seen_non_roman_parent_point = True
                            if seen_non_roman_parent_point or (
                                last_point_before_insert and not is_roman_point_marker(last_point_before_insert)
                            ):
                                fallback_nested_parent_is_non_roman = True

                if (
                    display_point_marker
                    and is_roman_point_marker(display_point_marker)
                    and (
                        (
                            parent_point_for_nested
                            and not is_roman_point_marker(parent_point_for_nested)
                        )
                        or fallback_nested_parent_is_non_roman
                    )
                ):
                    # Nested point levels like (ca) -> (i)/(ii).
                    inserted.paragraph_format.left_indent = Cm(3)
                    inserted.paragraph_format.first_line_indent = Cm(-1)
                else:
                    if display_point_marker and not is_roman_point_marker(display_point_marker):
                        template_para = None
                        if insertion_index is not None:
                            for back_idx in range(min(insertion_index - 1, len(doc.paragraphs) - 1), -1, -1):
                                marker = first_top_level_point_marker(doc.paragraphs[back_idx].text)
                                if marker and not is_roman_point_marker(marker):
                                    template_para = doc.paragraphs[back_idx]
                                    break
                        if template_para is not None:
                            inserted.paragraph_format.left_indent = template_para.paragraph_format.left_indent
                            inserted.paragraph_format.first_line_indent = template_para.paragraph_format.first_line_indent
                        else:
                            inserted.paragraph_format.left_indent = Cm(2)
                            inserted.paragraph_format.first_line_indent = Cm(-1)
                    else:
                        inserted.paragraph_format.left_indent = Cm(2)
                        inserted.paragraph_format.first_line_indent = Cm(-1)
            elif subparagraph_addition_instruction and target_paragraph and target_point is None:
                inserted.paragraph_format.left_indent = Cm(1)
                inserted.paragraph_format.first_line_indent = None
            elif (
                target_paragraph
                and target_point is None
                and detected_paragraph_marker is None
                and detected_point_marker is None
                and re.search(
                    r"\bthe following paragraph(?:s)?\s+(?:is|are)\s+inserted\b",
                    source_instruction,
                    flags=re.I,
                )
            ):
                # Continuation lines inside an inserted paragraph block (e.g.
                # Article 7(12a) second/third subparagraphs) should be flush.
                inserted.paragraph_format.left_indent = Cm(1)
                inserted.paragraph_format.first_line_indent = None
            if not inserted.runs:
                display_text = ensure_leading_marker_tab(display_text)
                inserted.add_run(display_text)
            inserted_paragraphs.append((inserted, item.get("source_instruction")))
            inserted_count += 1
            applied_mode = "inserted"

        analysis_items.append(
            {
                "article_number": item["article_number"],
                "article_heading": item["article_heading"],
                "amending_text": amend_text,
                "source_instruction": item.get("source_instruction"),
                "target_article_number": item.get("target_article_number"),
                "target_paragraph_number": item.get("target_paragraph_number"),
                "target_point_marker": item.get("target_point_marker"),
                "target_annex_number": item.get("target_annex_number"),
                "target_annex_point_marker": item.get("target_annex_point_marker"),
                "best_match_index": best_idx,
                "best_match_score": round(best_score, 6),
                "best_match_text": base_text,
                "applied_mode": applied_mode,
                "diff": diff,
            }
        )

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    enforce_inserted_paragraph_styles(inserted_paragraphs)
    enforce_inserted_paragraph_styles(retrofit_inserted_paragraphs)
    inserted_ids = {id(p._element) for p, _ in inserted_paragraphs}
    retrofit_ids = {id(p._element) for p, _ in retrofit_inserted_paragraphs}
    retrofit_styles_for_matching_amendment_paragraphs(
        doc,
        amending_items,
        allowed_paragraph_ids=inserted_ids | retrofit_ids,
    )
    assert_inserted_annexes_present(doc, amending_items)
    assert_inserted_annex_order(doc, amending_items)

    strip_src_markers_in_document(doc)
    doc.save(str(output_docx))
    amending_footnotes = extract_footnotes_map(load_html(amending_html))
    patch_docx_with_native_footnotes_merge(output_docx, amending_footnotes)
    strip_src_markers_from_docx(output_docx)
    normalize_footnote_ids_by_appearance(output_docx)

    analysis_docx = default_analysis_docx_path(output_docx)
    write_analysis_docx(analysis_docx, amending_html, rgb, analysis_items)

    analysis_report = {
        "created_at_utc": datetime.now(timezone.utc).isoformat(),
        "base_docx": str(base_docx),
        "amending_html": str(amending_html),
        "amendment_number": next_amendment_number_from_base(base_docx),
        "output_docx": str(output_docx),
        "analysis_docx": str(analysis_docx),
        "color": color,
        "amending_title": amending_title,
        "recitals_inserted": len(amending_recitals),
        "items_detected": len(amending_items),
        "items_analyzed": len(analysis_items),
        "items_applied_by_replacement": applied_count,
        "items_inserted_as_new": inserted_count,
        "analysis": analysis_items,
        "note": (
            "Two-phase amendment application: identified relevant provisions are applied. "
            "Replacements are attempted first; unresolved replacement instructions are left as analysis_only "
            "(never inserted) to avoid creating duplicated legal provisions."
            " In strict doubt mode, unresolved low-confidence items are intentionally deferred as analysis_only."
        ),
    }
    operation_proof = build_operation_proof_report(analysis_items)
    analysis_report["operation_proof"] = operation_proof

    precondition_failures = operation_proof.get("summary", {}).get("precondition_failures", 0)
    if require_full_application and precondition_failures:
        raise RuntimeError(
            "Operation precondition check failed: one or more amendment operations are missing required targets. "
            f"See analysis report: {analysis_json}"
        )

    one_go_check = summarize_one_go_application_check(analysis_items, doc)
    analysis_report["one_go_application_check"] = one_go_check

    analysis_json.parent.mkdir(parents=True, exist_ok=True)
    analysis_json.write_text(json.dumps(analysis_report, indent=2), encoding="utf-8")

    if require_full_application and not one_go_check["passed"]:
        raise RuntimeError(
            "One-go application check failed: not all detected substantive amendment items were applied in this run. "
            f"See analysis report: {analysis_json}"
        )

    return output_docx


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Stage 3 draft: apply one amending regulation layer to an existing DOCX "
            "with visible title/recitals and word-level change analysis."
        )
    )
    parser.add_argument("base_docx", help="Path to current consolidated DOCX (Stage 1 or previous Stage 3 output)")
    parser.add_argument(
        "arg2",
        help=(
            "Either output DOCX path (new syntax) or amending HTML path (legacy syntax)."
        ),
    )
    parser.add_argument(
        "arg3",
        help=(
            "Either amending HTML path (new syntax) or color (legacy syntax)."
        ),
    )
    parser.add_argument(
        "arg4",
        nargs="?",
        help="Color when using new syntax with explicit output DOCX path.",
    )
    parser.add_argument(
        "--analysis-json",
        help="Path for JSON analysis report (default: outputs/qa/<output_stem>_amendment_analysis.json)",
    )
    parser.add_argument(
        "--identified-amendments-json",
        help="Path for identified/reviewable amendment provisions JSON.",
    )
    parser.add_argument(
        "--identify-only",
        action="store_true",
        help="Phase 1 only: identify relevant amendment provisions and exit.",
    )
    parser.add_argument(
        "--apply-from-json",
        help="Phase 2: apply amendments from a previously identified JSON file.",
    )
    parser.add_argument(
        "--target-regulation",
        help="Override target regulation id filter (e.g. 2017/1129).",
    )
    parser.add_argument(
        "--max-analysis-items",
        type=int,
        default=0,
        help="Maximum number of amending blocks to analyze/apply (0 = all detected blocks)",
    )
    parser.add_argument(
        "--allow-partial-application",
        action="store_true",
        help=(
            "Allow Stage 3 to complete even when some detected amendment items remain analysis-only. "
            "Default behavior enforces a one-go full-application check."
        ),
    )
    args = parser.parse_args(argv)

    base_docx = Path(args.base_docx).expanduser().resolve()

    # New syntax (preferred):
    #   stage3.py <base_docx> <output_docx> <amending_html> <color>
    # Legacy syntax (still supported):
    #   stage3.py <base_docx> <amending_html> <color>
    if args.arg4 is not None:
        output_docx = Path(args.arg2).expanduser().resolve()
        amending_html = Path(args.arg3).expanduser().resolve()
        color = args.arg4
    else:
        amending_html = Path(args.arg2).expanduser().resolve()
        color = args.arg3
        output_docx = default_output_path(base_docx)

    resolved_base_docx, resolution_notes = resolve_base_docx(base_docx)
    if resolved_base_docx is None:
        print(f"Base DOCX not found: {base_docx}")
        print("No fallback base could be resolved. Provide an existing Stage 1/Stage 3 DOCX path.")
        return 2
    for note in resolution_notes:
        print(note)
    base_docx = resolved_base_docx
    if not amending_html.exists():
        print(f"Amending HTML not found: {amending_html}")
        return 2

    analysis_json = (
        Path(args.analysis_json).expanduser().resolve()
        if args.analysis_json
        else default_analysis_path(output_docx)
    )

    identified_json = (
        Path(args.identified_amendments_json).expanduser().resolve()
        if args.identified_amendments_json
        else DEFAULT_ANALYSIS_DIR / f"{output_docx.stem}_identified_amendments.json"
    )

    identified_payload: dict | None = None
    if args.apply_from_json:
        src = Path(args.apply_from_json).expanduser().resolve()
        if not src.exists():
            print(f"Identified amendments JSON not found: {src}")
            return 2
        identified_payload = json.loads(src.read_text(encoding="utf-8"))
        if isinstance(identified_payload, dict) and isinstance(identified_payload.get("items"), list):
            identified_payload["items"] = backfill_targets_from_instruction(identified_payload.get("items", []))
            identified_payload["items"] = normalize_replacement_target_points(identified_payload.get("items", []))
            identified_payload["items"] = normalize_amendment_items(identified_payload.get("items", []))
            if not isinstance(identified_payload.get("entry_into_force_block"), dict):
                identified_payload["entry_into_force_block"] = extract_entry_into_force_signoff_block(amending_html)
    else:
        identified_payload = identify_relevant_amendments(
            base_docx=base_docx,
            amending_html=amending_html,
            target_regulation=args.target_regulation,
        )

    identified_json.parent.mkdir(parents=True, exist_ok=True)
    identified_json.write_text(json.dumps(identified_payload, indent=2), encoding="utf-8")

    if args.identify_only:
        print(f"Identified amendments saved: {identified_json}")
        print(
            f"Target regulation: {identified_payload.get('target_regulation')} | "
            f"detected={identified_payload.get('detected_items')} | "
            f"relevant={identified_payload.get('relevant_items')}"
        )
        print("Review this file, then run Stage 3 again with --apply-from-json to apply exactly those items.")
        return 0

    # Do not pre-delete output. When base_docx and output_docx are the same
    # (e.g. stage3_1 -> stage3_1 overwrite), deleting first would remove input.

    try:
        written = apply_single_amending_regulation(
            base_docx=base_docx,
            amending_html=amending_html,
            color=color,
            output_docx=output_docx,
            analysis_json=analysis_json,
            max_analysis_items=args.max_analysis_items,
            identified_payload=identified_payload,
            require_full_application=not args.allow_partial_application,
        )
    except Exception as exc:
        print(f"Stage 3 draft failed: {exc}")
        return 1

    print(f"Stage 3 draft output saved: {written}")
    print(f"Identified amendments saved: {identified_json}")
    print(f"Analysis report saved: {analysis_json}")
    print("Use this output DOCX as base_docx for the next amending regulation run with a new color.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
