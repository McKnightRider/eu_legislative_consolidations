"""
Stage 1 Word writer: structured legislative model -> .docx.

The styles created here are intentionally simple and stable.  The aim is a canonical
Word baseline suitable for later amendment application and QA.  Later stages can add
more sophisticated legislative numbering and colour-coded amendment mark-up.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE

from stage1_conversion_helpers import LegislativeDocument, ParagraphBlock, TableBlock


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _set_font(style, name: str = "Arial", size: int = 11, bold: bool = False, italic: bool = False):
    """Apply a basic font definition to a Word style."""
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic


def _set_paragraph_format(style, *, space_before=0, space_after=6, line_spacing=1.0,
                          left_indent_cm=0, first_line_indent_cm=None,
                          alignment=None):
    """Apply basic paragraph formatting to a Word style."""
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.left_indent = Cm(left_indent_cm)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if alignment is not None:
        pf.alignment = alignment


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Return an existing style or create it if it does not already exist."""
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def configure_legislative_styles(document: Document) -> None:
    """Create the Word styles used by the Stage 1 conversion."""
    normal = document.styles["Normal"]
    _set_font(normal, "Arial", 11)
    _set_paragraph_format(normal, space_after=6, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    title = ensure_style(document, "Legislation Title")
    _set_font(title, "Arial", 12, bold=True)
    _set_paragraph_format(title, space_before=0, space_after=12, line_spacing=1.0,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    recital = ensure_style(document, "Legislation Recital")
    _set_font(recital, "Arial", 11)
    _set_paragraph_format(recital, space_before=0, space_after=6, line_spacing=1.0,
                          left_indent_cm=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    article_heading = ensure_style(document, "Legislation Article Heading")
    _set_font(article_heading, "Arial", 11, bold=True)
    _set_paragraph_format(article_heading, space_before=12, space_after=6, line_spacing=1.0,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    article_para = ensure_style(document, "Legislation Article Paragraph")
    _set_font(article_para, "Arial", 11)
    _set_paragraph_format(article_para, space_before=0, space_after=6, line_spacing=1.0,
                          left_indent_cm=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    annex_heading = ensure_style(document, "Legislation Annex Heading")
    _set_font(annex_heading, "Arial", 11, bold=True)
    _set_paragraph_format(annex_heading, space_before=12, space_after=6, line_spacing=1.0,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    warning = ensure_style(document, "Conversion Warning")
    _set_font(warning, "Arial", 10, italic=True)
    _set_paragraph_format(warning, space_before=3, space_after=3, line_spacing=1.0,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)


def set_document_margins(document: Document, margin_cm: float = 2.54) -> None:
    """Set margins for all sections."""
    for section in document.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def add_core_properties(document: Document, model: LegislativeDocument) -> None:
    """Populate basic document core properties."""
    props = document.core_properties
    props.title = model.title[:255]
    props.subject = "Stage 1 legislative source conversion"
    props.comments = "Generated from XML/HTML source by Stage 1 conversion script. Requires Stage 2 QA before use as a master consolidation text."


# ---------------------------------------------------------------------------
# Block rendering
# ---------------------------------------------------------------------------

def style_for_role(role: str) -> str:
    """Map parser roles to Word style names."""
    return {
        "title": "Legislation Title",
        "recital": "Legislation Recital",
        "article_heading": "Legislation Article Heading",
        "article_paragraph": "Legislation Article Paragraph",
        "annex_heading": "Legislation Annex Heading",
        "paragraph": "Normal",
    }.get(role, "Normal")


def add_paragraph_block(document: Document, block: ParagraphBlock) -> None:
    """Add one ParagraphBlock to the Word document."""
    para = document.add_paragraph(style=style_for_role(block.role))
    run = para.add_run(block.text)
    if block.role == "title":
        run.bold = True
    if block.source_id:
        # Store the source id as invisible-ish text in a comment-like convention?
        # python-docx does not expose Word comments directly.  We therefore do not
        # add this to the visible document.  The information remains in the JSON/QA
        # layer if that is generated in a later stage.
        pass


def _set_cell_shading(cell, fill: str = "F2F2F2") -> None:
    """Apply light shading to a Word table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_block(document: Document, block: TableBlock) -> None:
    """Add a simplified table to the Word document."""
    if not block.rows:
        return
    max_cols = max(len(row) for row in block.rows)
    table = document.add_table(rows=len(block.rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row in enumerate(block.rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            text = row[c_idx].text if c_idx < len(row) else ""
            cell.text = text
            for para in cell.paragraphs:
                para.style = document.styles["Normal"]
            if r_idx == 0:
                _set_cell_shading(cell)


def add_warnings_section(document: Document, warnings: Iterable[str]) -> None:
    """Add parser warnings to the end of the document for transparent QA."""
    warnings = list(warnings)
    if not warnings:
        return
    document.add_page_break()
    heading = document.add_paragraph("Stage 1 conversion warnings", style="Legislation Article Heading")
    for warning in warnings:
        document.add_paragraph(f"• {warning}", style="Conversion Warning")


# ---------------------------------------------------------------------------
# Main writer
# ---------------------------------------------------------------------------

def write_legislative_docx(model: LegislativeDocument, output_path: str | Path) -> Path:
    """Write the structured LegislativeDocument model to a .docx file."""
    output_path = Path(output_path)
    document = Document()
    configure_legislative_styles(document)
    set_document_margins(document)
    add_core_properties(document, model)

    # Visible title at the start of the document.
    document.add_paragraph(model.title, style="Legislation Title")

    for block in model.blocks:
        if isinstance(block, ParagraphBlock):
            add_paragraph_block(document, block)
        elif isinstance(block, TableBlock):
            add_table_block(document, block)

    add_warnings_section(document, model.warnings)
    document.save(output_path)
    return output_path
