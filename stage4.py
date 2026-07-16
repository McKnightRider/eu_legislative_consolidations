"""Stage 4 QA checks for Stage 3 consolidated outputs.

This stage validates two things:
1) Stage 3 insertions/replacements were applied and visibly revision-formatted.
2) The effective consolidated text (computed by ignoring struck-through runs,
   while keeping the DOCX unchanged) matches official EU consolidated sources.

Typical usage:
    python stage4.py outputs/stage3_1.docx --html consolidated/stage3_1.html --pdf consolidated/stage3_1.pdf

If --html is omitted, Stage 4 tries:
    consolidated/<docx_stem>.html

If --pdf is omitted, Stage 4 tries:
    consolidated/<docx_stem>.pdf

Outputs:
- JSON QA report under outputs/qa by default.
"""

from __future__ import annotations

import argparse
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from difflib import SequenceMatcher
import json
from pathlib import Path
import re
from typing import Optional

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_UNDERLINE


DEFAULT_QA_DIR = Path("outputs") / "qa"
DEFAULT_CONSOLIDATED_DIR = Path("consolidated")
WORD_RE = re.compile(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?")
ROMAN_RE = re.compile(r"^(?=[ivxlcdm]+$)m{0,4}(cm|cd|d?c{0,3})(xc|xl|l?x{0,3})(ix|iv|v?i{0,3})$", re.IGNORECASE)

# High-noise structural/editorial tokens that often differ across HTML vs DOCX export paths.
PRACTICAL_EXCLUDE_TOKENS: set[str] = {
    "article",
    "annex",
    "chapter",
    "section",
    "subsection",
    "point",
    "text",
    "oj",
    "en",
    "eu",
    "eec",
    "ec",
    "eur",
}

COLOR_MAP: dict[str, tuple[int, int, int]] = {
    "red": (192, 0, 0),
    "blue": (0, 102, 204),
    "green": (0, 128, 0),
    "orange": (230, 120, 0),
    "purple": (128, 0, 128),
    "teal": (0, 128, 128),
    "brown": (128, 64, 0),
    "black": (0, 0, 0),
}


@dataclass
class CoverageResult:
    total_source_tokens: int
    total_docx_tokens: int
    matched_tokens: int
    missing_tokens_total: int
    missing_unique_tokens: int
    extra_tokens_total: int
    extra_unique_tokens: int
    coverage_ratio: float
    top_missing: list[tuple[str, int]]
    top_extra: list[tuple[str, int]]


@dataclass
class PracticalCoverageResult:
    total_source_tokens: int
    total_docx_tokens: int
    matched_tokens: int
    missing_tokens_total: int
    missing_unique_tokens: int
    extra_tokens_total: int
    extra_unique_tokens: int
    coverage_ratio: float
    top_missing: list[tuple[str, int]]
    top_extra: list[tuple[str, int]]
    excluded_source_tokens: int
    excluded_docx_tokens: int


@dataclass
class LevenshteinResult:
    docx_chars: int
    source_chars: int
    distance: int
    normalized_similarity: float
    method: str
    is_estimate: bool


@dataclass
class TokenLevenshteinResult:
    docx_tokens: int
    source_tokens: int
    distance: int
    normalized_similarity: float
    method: str
    is_estimate: bool


def clean_text(text: str | None) -> str:
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_text(text: str | None) -> str:
    return clean_text(text).lower()


def tokenize(text: str) -> list[str]:
    return [m.group(0).lower() for m in WORD_RE.finditer(text)]


def is_practical_noise_token(token: str) -> bool:
    if not token:
        return True
    if token in PRACTICAL_EXCLUDE_TOKENS:
        return True
    if token.startswith("m") and token[1:].isdigit():
        return True
    if token.isdigit():
        return True
    if len(token) == 1 and token.isalpha():
        return True
    if ROMAN_RE.fullmatch(token):
        return True
    return False


def filter_practical_tokens(tokens: list[str]) -> tuple[list[str], int]:
    kept: list[str] = []
    excluded = 0
    for token in tokens:
        if is_practical_noise_token(token):
            excluded += 1
        else:
            kept.append(token)
    return kept, excluded


def rgb_to_hex(rgb: tuple[int, int, int]) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def parse_color_value(color: str | None) -> tuple[int, int, int]:
    if not color:
        return COLOR_MAP["orange"]
    value = color.strip().lower()
    if value in COLOR_MAP:
        return COLOR_MAP[value]
    if value.startswith("#") and len(value) == 7:
        return (int(value[1:3], 16), int(value[3:5], 16), int(value[5:7], 16))
    if re.fullmatch(r"[0-9a-fA-F]{6}", value):
        return (int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))
    raise ValueError(f"Unsupported color format: {color}")


def run_color_hex(run) -> Optional[str]:
    color = run.font.color.rgb
    if color is None:
        return None
    return str(color).upper()


def iter_document_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def collect_docx_text(doc: Document, *, exclude_strike: bool) -> str:
    parts: list[str] = []
    for p in iter_document_paragraphs(doc):
        chunks: list[str] = []
        for run in p.runs:
            if exclude_strike and run.font.strike:
                continue
            txt = run.text or ""
            if txt:
                chunks.append(txt)
        line = clean_text("".join(chunks))
        if line:
            parts.append(line)
    return clean_text("\n".join(parts))


def paragraph_text(p, *, exclude_strike: bool) -> str:
    if not exclude_strike:
        return clean_text(p.text)

    chunks: list[str] = []
    for run in p.runs:
        if run.font.strike:
            continue
        txt = run.text or ""
        if txt:
            chunks.append(txt)
    return clean_text("".join(chunks))


def extract_html_text(path: str | Path) -> str:
    html = Path(path).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    for noisy in soup(["script", "style", "noscript"]):
        noisy.decompose()
    return clean_text(soup.get_text(" ", strip=True))


def extract_pdf_text(path: str | Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF comparison requires pypdf. Install it with: pip install pypdf") from exc

    reader = PdfReader(str(path))
    page_text: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        page_text.append(extracted)
    return clean_text(" ".join(page_text))


def compare_token_coverage(docx_text: str, source_text: str, top_n: int) -> CoverageResult:
    docx_tokens = tokenize(docx_text)
    source_tokens = tokenize(source_text)

    docx_counter = Counter(docx_tokens)
    source_counter = Counter(source_tokens)

    missing_counter = source_counter - docx_counter
    extra_counter = docx_counter - source_counter

    matched_tokens = sum(min(source_counter[t], docx_counter[t]) for t in source_counter)
    total_source = len(source_tokens)
    coverage = (matched_tokens / total_source) if total_source else 1.0

    return CoverageResult(
        total_source_tokens=total_source,
        total_docx_tokens=len(docx_tokens),
        matched_tokens=matched_tokens,
        missing_tokens_total=sum(missing_counter.values()),
        missing_unique_tokens=len(missing_counter),
        extra_tokens_total=sum(extra_counter.values()),
        extra_unique_tokens=len(extra_counter),
        coverage_ratio=coverage,
        top_missing=missing_counter.most_common(top_n),
        top_extra=extra_counter.most_common(top_n),
    )


def compare_practical_token_coverage(docx_text: str, source_text: str, top_n: int) -> PracticalCoverageResult:
    docx_tokens_raw = tokenize(docx_text)
    source_tokens_raw = tokenize(source_text)

    docx_tokens, excluded_docx = filter_practical_tokens(docx_tokens_raw)
    source_tokens, excluded_source = filter_practical_tokens(source_tokens_raw)

    docx_counter = Counter(docx_tokens)
    source_counter = Counter(source_tokens)

    missing_counter = source_counter - docx_counter
    extra_counter = docx_counter - source_counter

    matched_tokens = sum(min(source_counter[t], docx_counter[t]) for t in source_counter)
    total_source = len(source_tokens)
    coverage = (matched_tokens / total_source) if total_source else 1.0

    return PracticalCoverageResult(
        total_source_tokens=total_source,
        total_docx_tokens=len(docx_tokens),
        matched_tokens=matched_tokens,
        missing_tokens_total=sum(missing_counter.values()),
        missing_unique_tokens=len(missing_counter),
        extra_tokens_total=sum(extra_counter.values()),
        extra_unique_tokens=len(extra_counter),
        coverage_ratio=coverage,
        top_missing=missing_counter.most_common(top_n),
        top_extra=extra_counter.most_common(top_n),
        excluded_source_tokens=excluded_source,
        excluded_docx_tokens=excluded_docx,
    )


def levenshtein_distance(a: str, b: str) -> int:
    if a == b:
        return 0
    if not a:
        return len(b)
    if not b:
        return len(a)

    if len(a) > len(b):
        a, b = b, a

    previous = list(range(len(a) + 1))
    for i, ch_b in enumerate(b, start=1):
        current = [i]
        for j, ch_a in enumerate(a, start=1):
            cost = 0 if ch_a == ch_b else 1
            current.append(
                min(
                    previous[j] + 1,
                    current[j - 1] + 1,
                    previous[j - 1] + cost,
                )
            )
        previous = current
    return previous[-1]


def levenshtein_distance_sequence(a: list[str], b: list[str]) -> int:
    if a == b:
        return 0
    if not a:
        return len(b)
    if not b:
        return len(a)

    if len(a) > len(b):
        a, b = b, a

    previous = list(range(len(a) + 1))
    for i, token_b in enumerate(b, start=1):
        current = [i]
        for j, token_a in enumerate(a, start=1):
            cost = 0 if token_a == token_b else 1
            current.append(
                min(
                    previous[j] + 1,
                    current[j - 1] + 1,
                    previous[j - 1] + cost,
                )
            )
        previous = current
    return previous[-1]


def compare_levenshtein(docx_text: str, source_text: str, *, max_exact_chars: int) -> LevenshteinResult:
    a = normalize_text(docx_text)
    b = normalize_text(source_text)
    denom = max(len(a), len(b))

    # Exact dynamic-programming Levenshtein is O(n*m); switch to a fast estimate
    # for very large consolidated texts to keep Stage 4 responsive.
    if denom <= max_exact_chars:
        dist = levenshtein_distance(a, b)
        similarity = 1.0 - (dist / denom) if denom else 1.0
        method = "exact_dynamic_programming"
        is_estimate = False
    else:
        ratio = SequenceMatcher(None, a, b).ratio()
        similarity = ratio
        dist = int(round((1.0 - ratio) * denom)) if denom else 0
        method = "estimated_from_sequence_matcher"
        is_estimate = True

    return LevenshteinResult(
        docx_chars=len(a),
        source_chars=len(b),
        distance=dist,
        normalized_similarity=similarity,
        method=method,
        is_estimate=is_estimate,
    )


def compare_token_levenshtein(docx_text: str, source_text: str, *, max_exact_tokens: int) -> TokenLevenshteinResult:
    a = tokenize(normalize_text(docx_text))
    b = tokenize(normalize_text(source_text))
    denom = max(len(a), len(b))

    # Exact token DP is O(n*m); fall back to estimate for very large token lists.
    if denom <= max_exact_tokens:
        dist = levenshtein_distance_sequence(a, b)
        similarity = 1.0 - (dist / denom) if denom else 1.0
        method = "exact_dynamic_programming"
        is_estimate = False
    else:
        ratio = SequenceMatcher(None, a, b).ratio()
        similarity = ratio
        dist = int(round((1.0 - ratio) * denom)) if denom else 0
        method = "estimated_from_sequence_matcher"
        is_estimate = True

    return TokenLevenshteinResult(
        docx_tokens=len(a),
        source_tokens=len(b),
        distance=dist,
        normalized_similarity=similarity,
        method=method,
        is_estimate=is_estimate,
    )


def expected_analysis_path(docx_path: Path) -> Path:
    return DEFAULT_QA_DIR / f"{docx_path.stem}_amendment_analysis.json"


def expected_identified_path(docx_path: Path) -> Path:
    return DEFAULT_QA_DIR / f"{docx_path.stem}_identified_amendments.json"


def default_report_path(docx_path: Path) -> Path:
    return DEFAULT_QA_DIR / f"{docx_path.stem}_stage4_qa_report.json"


def load_stage3_analysis(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_stage3_identified(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def article_heading_number(text: str) -> str | None:
    m = re.match(r"^Article\s+(\d+[A-Za-z]?)\b", clean_text(text), flags=re.I)
    return m.group(1).lower() if m else None


def is_hard_article_boundary_heading(text: str) -> bool:
    txt = clean_text(text)
    return bool(re.match(r"^(?:CHAPTER|ANNEX)\b", txt, flags=re.I))


def top_level_paragraph_marker(text: str) -> str | None:
    m = re.match(r"^[‘'\"]?(\d+[A-Za-z]?)\.(?=\s|\t|$)", clean_text(text))
    return m.group(1).lower() if m else None


def paragraph_marker_sort_key(marker: str) -> tuple[int, str]:
    m = re.fullmatch(r"(\d+)([A-Za-z]*)", (marker or "").strip())
    if not m:
        return (10**9, marker.lower())
    return (int(m.group(1)), m.group(2).lower())


def find_article_section_bounds_from_paragraphs(doc: Document, article_number: str) -> tuple[int, int] | None:
    headings: list[tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        num = article_heading_number(p.text)
        if num:
            headings.append((idx, num))

    target = article_number.lower()
    for i, (idx, num) in enumerate(headings):
        if num != target:
            continue
        end_idx = headings[i + 1][0] if i + 1 < len(headings) else len(doc.paragraphs)
        for probe_idx in range(idx + 1, end_idx):
            if is_hard_article_boundary_heading(doc.paragraphs[probe_idx].text):
                end_idx = probe_idx
                break
        return idx, end_idx
    return None


def find_exact_article_heading_indices(doc: Document, article_number: str) -> list[int]:
    target = article_number.lower()
    hits: list[int] = []
    for idx, p in enumerate(doc.paragraphs):
        txt = clean_text(p.text)
        if re.fullmatch(rf"Article\s+{re.escape(target)}", txt, flags=re.I):
            hits.append(idx)
    return hits


def check_article_paragraph_order(doc: Document) -> dict:
    headings: list[tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        num = article_heading_number(p.text)
        if num:
            headings.append((idx, num))

    inversions: list[dict] = []
    sections_checked = 0
    markers_checked = 0

    for h_idx, (start, article_num) in enumerate(headings):
        end = headings[h_idx + 1][0] if h_idx + 1 < len(headings) else len(doc.paragraphs)
        for probe_idx in range(start + 1, end):
            if is_hard_article_boundary_heading(doc.paragraphs[probe_idx].text):
                end = probe_idx
                break
        section_markers: list[tuple[int, str]] = []
        for idx in range(start + 1, end):
            marker = top_level_paragraph_marker(doc.paragraphs[idx].text)
            if marker:
                section_markers.append((idx, marker))

        if not section_markers:
            continue

        sections_checked += 1
        markers_checked += len(section_markers)

        prev_idx, prev_marker = section_markers[0]
        prev_key = paragraph_marker_sort_key(prev_marker)
        for cur_idx, cur_marker in section_markers[1:]:
            cur_key = paragraph_marker_sort_key(cur_marker)
            if cur_key < prev_key:
                inversions.append(
                    {
                        "article_number": article_num,
                        "previous_marker": prev_marker,
                        "previous_paragraph_index": prev_idx,
                        "current_marker": cur_marker,
                        "current_paragraph_index": cur_idx,
                    }
                )
            prev_idx, prev_marker, prev_key = cur_idx, cur_marker, cur_key

    return {
        "article_sections_checked": sections_checked,
        "paragraph_markers_checked": markers_checked,
        "marker_order_inversions": len(inversions),
        "marker_order_inversion_examples": inversions[:25],
    }


def check_inserted_article_block_placement(doc: Document, identified: dict | None) -> dict:
    if not identified:
        return {
            "identified_json_found": False,
            "inserted_article_items_checked": 0,
            "missing_inserted_article_heading": 0,
            "missing_from_inserted_article_section": 0,
            "found_before_inserted_article_heading": 0,
            "missing_heading_examples": [],
            "missing_section_examples": [],
            "found_before_heading_examples": [],
        }

    items = identified.get("items", [])
    inserted_article_items_without_marker = [
        i for i in items
        if re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", clean_text(i.get("source_instruction", "")), flags=re.I)
        and not clean_text(i.get("inserted_article_marker", ""))
        and clean_text(i.get("text", ""))
    ]
    inserted_article_items = [
        i for i in items
        if clean_text(i.get("inserted_article_marker", ""))
        and re.search(r"\bthe following Article(?:s)?\s+(?:is|are)\s+(?:inserted|added)\b", clean_text(i.get("source_instruction", "")), flags=re.I)
        and clean_text(i.get("text", ""))
    ]
    inserted_article_markers = sorted(
        {
            clean_text(i.get("inserted_article_marker", "")).lower()
            for i in inserted_article_items
            if clean_text(i.get("inserted_article_marker", ""))
        }
    )

    missing_heading: list[dict] = []
    missing_exact_heading: list[dict] = []
    duplicate_exact_heading: list[dict] = []
    missing_in_section: list[dict] = []
    found_before_heading: list[dict] = []
    found_in_other_sections: list[dict] = []

    article_headings: list[tuple[int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        num = article_heading_number(p.text)
        if num:
            article_headings.append((idx, num))

    for marker in inserted_article_markers:
        exact_hits = find_exact_article_heading_indices(doc, marker)
        if len(exact_hits) == 0:
            missing_exact_heading.append({"inserted_article_marker": marker})
        elif len(exact_hits) > 1:
            duplicate_exact_heading.append(
                {
                    "inserted_article_marker": marker,
                    "exact_heading_occurrences": len(exact_hits),
                }
            )

    for item in inserted_article_items:
        marker = clean_text(item.get("inserted_article_marker", "")).lower()
        section = find_article_section_bounds_from_paragraphs(doc, marker)
        if section is None:
            missing_heading.append(
                {
                    "inserted_article_marker": marker,
                    "source_instruction": clean_text(item.get("source_instruction", "")),
                }
            )
            continue

        section_start, section_end = section
        section_text = clean_text(
            "\n".join(paragraph_text(p, exclude_strike=True) for p in doc.paragraphs[section_start:section_end])
        )

        # Restrict pre-heading leakage scan to the expected base-article window.
        # This avoids matching repeated/legal boilerplate elsewhere in the document.
        before_start = 0
        target_article = clean_text(item.get("target_article_number", "")).lower()
        target_section = find_article_section_bounds_from_paragraphs(doc, target_article) if target_article else None
        if target_section is not None:
            target_start, _ = target_section
            if target_start < section_start:
                before_start = target_start
        before_text = clean_text(
            "\n".join(paragraph_text(p, exclude_strike=True) for p in doc.paragraphs[before_start:section_start])
        )

        probe_tokens = probe_tokens_for_matching(clean_text(item.get("text", "")), limit=24)
        # Ultra-short probes (e.g. single token "smes") produce noisy
        # cross-section false positives; require a minimally distinctive probe.
        if len(probe_tokens) < 6:
            continue

        section_found, section_match_len = has_probe_match_with_backoff(
            tokenize(normalize_text(section_text)),
            probe_tokens,
            fallback_lengths=[18, 12, 8],
        )
        before_found, before_match_len = has_probe_match_with_backoff(
            tokenize(normalize_text(before_text)),
            probe_tokens,
            fallback_lengths=[],
        )

        if not section_found:
            missing_in_section.append(
                {
                    "inserted_article_marker": marker,
                    "probe": " ".join(probe_tokens[:section_match_len or len(probe_tokens)]),
                }
            )
        if before_found:
            found_before_heading.append(
                {
                    "inserted_article_marker": marker,
                    "target_article_number": target_article or None,
                    "probe": " ".join(probe_tokens[:before_match_len or len(probe_tokens)]),
                }
            )

        # Detect cross-section leakage: inserted-article probe text should not
        # also appear in other article sections.
        other_hits: list[str] = []
        for h_idx, (_, heading_num) in enumerate(article_headings):
            if heading_num == marker:
                continue
            sec_start = article_headings[h_idx][0]
            sec_end = article_headings[h_idx + 1][0] if h_idx + 1 < len(article_headings) else len(doc.paragraphs)
            sec_text = clean_text(
                "\n".join(paragraph_text(p, exclude_strike=True) for p in doc.paragraphs[sec_start:sec_end])
            )
            hit, _ = has_probe_match_with_backoff(
                tokenize(normalize_text(sec_text)),
                probe_tokens,
                fallback_lengths=[],
            )
            if hit:
                other_hits.append(heading_num)

        unique_other_hits = sorted(set(other_hits))
        # A single overlap can be legitimate template language in adjacent
        # inserted/replaced articles. Flag only broader leakage patterns.
        if len(unique_other_hits) >= 2:
            found_in_other_sections.append(
                {
                    "inserted_article_marker": marker,
                    "probe": " ".join(probe_tokens),
                    "other_article_sections": unique_other_hits,
                }
            )

    return {
        "identified_json_found": True,
        "inserted_article_items_without_marker": len(inserted_article_items_without_marker),
        "inserted_article_markers_checked": len(inserted_article_markers),
        "missing_exact_inserted_article_heading": len(missing_exact_heading),
        "duplicate_exact_inserted_article_heading": len(duplicate_exact_heading),
        "inserted_article_items_checked": len(inserted_article_items),
        "missing_inserted_article_heading": len(missing_heading),
        "missing_from_inserted_article_section": len(missing_in_section),
        "found_before_inserted_article_heading": len(found_before_heading),
        "found_in_other_article_sections": len(found_in_other_sections),
        "without_marker_examples": [
            {
                "article_number": i.get("article_number"),
                "target_article_number": i.get("target_article_number"),
                "target_paragraph_number": i.get("target_paragraph_number"),
                "source_instruction": clean_text(i.get("source_instruction", "")),
            }
            for i in inserted_article_items_without_marker[:25]
        ],
        "missing_exact_heading_examples": missing_exact_heading[:25],
        "duplicate_exact_heading_examples": duplicate_exact_heading[:25],
        "missing_heading_examples": missing_heading[:25],
        "missing_section_examples": missing_in_section[:25],
        "found_before_heading_examples": found_before_heading[:25],
        "found_in_other_sections_examples": found_in_other_sections[:25],
    }


def check_revision_formatting(doc: Document, expected_color_hex: str) -> dict:
    inserted_total = 0
    inserted_expected = 0
    strike_total = 0
    strike_expected = 0

    for p in iter_document_paragraphs(doc):
        for run in p.runs:
            text = clean_text(run.text)
            if not text:
                continue
            color_hex = run_color_hex(run)
            is_inserted = run.font.underline == WD_UNDERLINE.DOUBLE
            is_deleted = bool(run.font.strike)

            if is_inserted:
                inserted_total += 1
                if color_hex == expected_color_hex:
                    inserted_expected += 1

            if is_deleted:
                strike_total += 1
                if color_hex == expected_color_hex:
                    strike_expected += 1

    return {
        "color_check_mode": "layer-aware-by-expected-color",
        "expected_revision_color_hex": expected_color_hex,
        "all_revision_runs": {
            "inserted_runs_total": inserted_total,
            "deleted_runs_total": strike_total,
        },
        "current_layer_runs": {
            "inserted_runs_total": inserted_expected,
            "deleted_runs_total": strike_expected,
        },
        "legacy_or_other_layer_runs": {
            "inserted_runs_total": max(0, inserted_total - inserted_expected),
            "deleted_runs_total": max(0, strike_total - strike_expected),
        },
    }


def top_level_tokens(text: str, limit: int = 30) -> list[str]:
    return tokenize(text)[:limit]


def strip_footnote_tokens(text: str) -> str:
    return re.sub(r"\[\[FN:\d+\]\]", " ", text)


def probe_tokens_for_matching(text: str, limit: int) -> list[str]:
    cleaned = clean_text(strip_footnote_tokens(text))
    return tokenize(cleaned)[:limit]


def contains_contiguous_token_sequence(haystack: list[str], needle: list[str]) -> bool:
    if not needle:
        return False
    n = len(needle)
    if n > len(haystack):
        return False
    first = needle[0]
    max_start = len(haystack) - n
    for i in range(max_start + 1):
        if haystack[i] != first:
            continue
        if haystack[i:i + n] == needle:
            return True
    return False


def has_probe_match_with_backoff(
    haystack_tokens: list[str],
    probe_tokens: list[str],
    *,
    fallback_lengths: list[int],
) -> tuple[bool, int]:
    if not probe_tokens:
        return False, 0

    lengths: list[int] = []
    seen: set[int] = set()
    ordered = [len(probe_tokens)] + fallback_lengths
    for length in ordered:
        trimmed = min(length, len(probe_tokens))
        if trimmed <= 0 or trimmed in seen:
            continue
        seen.add(trimmed)
        lengths.append(trimmed)

    for length in lengths:
        candidate = probe_tokens[:length]
        if contains_contiguous_token_sequence(haystack_tokens, candidate):
            return True, length

    return False, lengths[-1] if lengths else 0


def check_stage3_application(doc: Document, analysis: dict, identified: dict | None = None) -> dict:
    analysis_items = analysis.get("analysis", [])
    docx_visible = normalize_text(collect_docx_text(doc, exclude_strike=False))
    docx_visible_tokens = tokenize(docx_visible)

    inserted_items = [i for i in analysis_items if i.get("applied_mode") == "inserted"]
    replaced_items = [i for i in analysis_items if i.get("applied_mode") == "replaced"]
    replacement_instruction_items = [
        i for i in analysis_items
        if re.search(r"\breplaced by the following\b", clean_text(i.get("source_instruction", "")), flags=re.I)
    ]

    missing_insertions: list[dict] = []
    for item in inserted_items:
        amend_text = clean_text(item.get("amending_text", ""))
        if not amend_text:
            continue
        probe_tokens = probe_tokens_for_matching(amend_text, limit=24)
        if not probe_tokens:
            continue
        found, matched_length = has_probe_match_with_backoff(
            docx_visible_tokens,
            probe_tokens,
            fallback_lengths=[18, 12, 8],
        )
        probe = " ".join(probe_tokens[:matched_length or len(probe_tokens)])
        if not found:
            missing_insertions.append(
                {
                    "article_number": item.get("article_number"),
                    "target_article_number": item.get("target_article_number"),
                    "probe": probe,
                }
            )

    replaced_hits = 0
    for item in replaced_items:
        amend_text = clean_text(item.get("amending_text", ""))
        if not amend_text:
            continue
        probe_tokens = probe_tokens_for_matching(amend_text, limit=16)
        if not probe_tokens:
            continue
        found, _ = has_probe_match_with_backoff(
            docx_visible_tokens,
            probe_tokens,
            fallback_lengths=[12, 8, 6],
        )
        if found:
            replaced_hits += 1

    replacement_instruction_not_applied: list[dict] = []
    replacement_groups: dict[tuple[str, str, str, str, str, str], list[dict]] = {}
    for item in replacement_instruction_items:
        key = (
            clean_text(item.get("source_instruction", "")),
            clean_text(item.get("target_article_number", "")),
            clean_text(item.get("target_paragraph_number", "")),
            clean_text(item.get("target_point_marker", "")),
            clean_text(item.get("target_annex_number", "")),
            clean_text(item.get("target_annex_point_marker", "")),
        )
        replacement_groups.setdefault(key, []).append(item)

    for group_items in replacement_groups.values():
        has_applied = any(
            clean_text(row.get("applied_mode", "")) in {"replaced", "already_applied"}
            for row in group_items
        )
        if has_applied:
            continue

        # Some replacement rows remain marked analysis_only even though the
        # replacement text is visibly present in the DOCX (for example where
        # Stage 3 applied via a neighboring structural path). Do not count
        # such groups as unresolved QA failures.
        text_present = False
        for row in group_items:
            amend_text = clean_text(row.get("amending_text", ""))
            if not amend_text:
                continue
            probe_tokens = probe_tokens_for_matching(amend_text, limit=16)
            if not probe_tokens:
                continue
            found, _ = has_probe_match_with_backoff(
                docx_visible_tokens,
                probe_tokens,
                fallback_lengths=[12, 8, 6],
            )
            if found:
                text_present = True
                break
        if text_present:
            continue

        sample = group_items[0]
        mode = clean_text(sample.get("applied_mode", ""))
        source_instruction = clean_text(sample.get("source_instruction", ""))
        target_point_marker = clean_text(sample.get("target_point_marker", "")).lower()

        # Nested roman continuation rows in replacement payloads are often
        # represented as analysis-only leaf rows while the enclosing
        # replacement line is applied.
        if mode == "analysis_only" and target_point_marker and ROMAN_RE.fullmatch(target_point_marker):
            continue

        # Introductory wording/sentence replacement instructions frequently
        # produce continuation fragments that are not one-to-one paragraph
        # replacements; avoid flagging these fragments as hard failures.
        if mode == "analysis_only" and re.search(r"\bintroductory\s+(?:wording|sentence)\b", source_instruction, flags=re.I):
            continue

        replacement_instruction_not_applied.append(
            {
                "article_number": sample.get("article_number"),
                "target_article_number": sample.get("target_article_number"),
                "target_paragraph_number": sample.get("target_paragraph_number"),
                "target_point_marker": sample.get("target_point_marker"),
                "target_annex_number": sample.get("target_annex_number"),
                "target_annex_point_marker": sample.get("target_annex_point_marker"),
                "applied_mode": mode or None,
                "source_instruction": source_instruction,
                "group_rows": len(group_items),
            }
        )

    inserted_article_placement = check_inserted_article_block_placement(doc, identified)
    paragraph_order = check_article_paragraph_order(doc)

    return {
        "items_detected": analysis.get("items_detected"),
        "items_analyzed": analysis.get("items_analyzed"),
        "items_inserted_as_new": analysis.get("items_inserted_as_new"),
        "items_applied_by_replacement": analysis.get("items_applied_by_replacement"),
        "inserted_items_checked": len(inserted_items),
        "inserted_items_missing_from_output": len(missing_insertions),
        "missing_insertions_examples": missing_insertions[:25],
        "replaced_items_checked": len(replaced_items),
        "replaced_items_with_new_text_detected": replaced_hits,
        "replacement_instruction_items_checked": len(replacement_instruction_items),
        "replacement_instruction_items_not_properly_applied": len(replacement_instruction_not_applied),
        "replacement_instruction_not_applied_examples": replacement_instruction_not_applied[:25],
        "inserted_article_placement": inserted_article_placement,
        "article_paragraph_order": paragraph_order,
    }


def resolve_optional_source(docx_path: Path, explicit_path: Optional[str], consolidated_dir: Path, suffix: str) -> Optional[Path]:
    if explicit_path:
        p = Path(explicit_path).expanduser().resolve()
        return p if p.exists() else None

    candidate = (consolidated_dir / f"{docx_path.stem}.{suffix}").expanduser().resolve()
    return candidate if candidate.exists() else None


def resolve_eu_html(docx_path: Path, explicit_html: Optional[str], consolidated_dir: Path) -> Optional[Path]:
    return resolve_optional_source(docx_path, explicit_html, consolidated_dir, "html")


def resolve_eu_pdf(docx_path: Path, explicit_pdf: Optional[str], consolidated_dir: Path) -> Optional[Path]:
    return resolve_optional_source(docx_path, explicit_pdf, consolidated_dir, "pdf")


def build_source_comparison(
    *,
    source_type: str,
    source_path: Path,
    effective_docx_text: str,
    top_n: int,
    min_coverage: float,
    practical_min_coverage: float,
    levenshtein_max_exact_chars: int,
    token_levenshtein_max_exact_tokens: int,
) -> dict:
    if source_type == "html":
        source_text = extract_html_text(source_path)
    elif source_type == "pdf":
        source_text = extract_pdf_text(source_path)
    else:
        raise ValueError(f"Unsupported source_type: {source_type}")

    coverage = compare_token_coverage(effective_docx_text, source_text, top_n)
    practical_coverage = compare_practical_token_coverage(effective_docx_text, source_text, top_n)
    levenshtein = compare_levenshtein(
        effective_docx_text,
        source_text,
        max_exact_chars=levenshtein_max_exact_chars,
    )
    token_levenshtein = compare_token_levenshtein(
        effective_docx_text,
        source_text,
        max_exact_tokens=token_levenshtein_max_exact_tokens,
    )
    strict_passed = coverage.coverage_ratio >= min_coverage
    practical_passed = practical_coverage.coverage_ratio >= practical_min_coverage
    return {
        "source_type": source_type,
        "source_path": str(source_path),
        "strict_coverage": asdict(coverage),
        "strict_min_coverage_threshold": min_coverage,
        "strict_coverage_passed": strict_passed,
        "practical_coverage": asdict(practical_coverage),
        "practical_min_coverage_threshold": practical_min_coverage,
        "practical_coverage_passed": practical_passed,
        "levenshtein": asdict(levenshtein),
        "token_levenshtein": asdict(token_levenshtein),
        "coverage_passed": strict_passed or practical_passed,
    }


def missing_source_issue(source_type: str, explicit: Optional[str]) -> str:
    if explicit:
        return f"Provided {source_type.upper()} file was not found: {explicit}"
    return f"EU consolidated {source_type.upper()} was not provided/found; {source_type.upper()} equivalence was not checked."


def source_fail_issue(source_type: str) -> str:
    return f"Strike-filtered output DOCX does not meet minimum coverage vs EU consolidated {source_type.upper()}."


def source_strict_note(source_type: str) -> str:
    return (
        f"Strict token coverage is below threshold for EU consolidated {source_type.upper()}, "
        "but practical coverage passed (likely structural/editorial token drift)."
    )


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description=(
            "Stage 4 QA: validate Stage 3 amendment formatting and compare effective "
            "consolidated text against official EU consolidated HTML/PDF."
        )
    )
    parser.add_argument("docx", help="Path to Stage 3 output DOCX")
    parser.add_argument(
        "--analysis-json",
        help="Path to Stage 3 analysis JSON (default: outputs/qa/<docx_stem>_amendment_analysis.json)",
    )
    parser.add_argument(
        "--identified-json",
        help="Path to Stage 3 identified amendments JSON (default: outputs/qa/<docx_stem>_identified_amendments.json)",
    )
    parser.add_argument(
        "--html",
        help="Path to official EU consolidated HTML for final-text comparison",
    )
    parser.add_argument(
        "--eu-html",
        dest="html",
        help=argparse.SUPPRESS,
    )
    parser.add_argument(
        "--pdf",
        help="Path to official EU consolidated PDF for final-text comparison",
    )
    parser.add_argument(
        "--consolidated-dir",
        default=str(DEFAULT_CONSOLIDATED_DIR),
        help=(
            "Directory containing official consolidated EU HTML files "
            "(default: consolidated/)"
        ),
    )
    parser.add_argument(
        "--report",
        help="Output JSON report path (default: outputs/qa/<docx_stem>_stage4_qa_report.json)",
    )
    parser.add_argument("--top", type=int, default=50, help="Top missing/extra tokens to include")
    parser.add_argument(
        "--min-coverage",
        type=float,
        default=0.995,
        help="Minimum required token coverage per source for strike-filtered DOCX vs EU consolidated HTML/PDF",
    )
    parser.add_argument(
        "--practical-min-coverage",
        type=float,
        default=0.985,
        help=(
            "Minimum practical token coverage per source (filters high-noise structural/editorial tokens). "
            "A source passes if strict OR practical coverage passes."
        ),
    )
    parser.add_argument(
        "--practical-min-coverage-html",
        type=float,
        default=None,
        help=(
            "Optional HTML-specific practical coverage threshold. If not set, "
            "--practical-min-coverage is used."
        ),
    )
    parser.add_argument(
        "--practical-min-coverage-pdf",
        type=float,
        default=None,
        help=(
            "Optional PDF-specific practical coverage threshold. If not set, "
            "--practical-min-coverage is used."
        ),
    )
    parser.add_argument(
        "--levenshtein-max-exact-chars",
        type=int,
        default=12000,
        help=(
            "Maximum character length for exact Levenshtein; above this, Stage 4 uses "
            "a fast SequenceMatcher-based distance estimate."
        ),
    )
    parser.add_argument(
        "--token-levenshtein-max-exact-tokens",
        type=int,
        default=3000,
        help=(
            "Maximum token length for exact token-level Levenshtein; above this, Stage 4 "
            "uses a SequenceMatcher-based estimate."
        ),
    )
    parser.add_argument(
        "--fail-on-issues",
        action="store_true",
        help="Exit non-zero when Stage 4 checks report issues",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    docx_path = Path(args.docx).expanduser().resolve()
    if not docx_path.exists():
        print(f"DOCX file not found: {docx_path}")
        return 2

    analysis_path = Path(args.analysis_json).expanduser().resolve() if args.analysis_json else expected_analysis_path(docx_path)
    analysis = None
    if analysis_path.exists():
        analysis = load_stage3_analysis(analysis_path)

    identified_path = Path(args.identified_json).expanduser().resolve() if args.identified_json else expected_identified_path(docx_path)
    identified = None
    if identified_path.exists():
        identified = load_stage3_identified(identified_path)

    report_path = Path(args.report).expanduser().resolve() if args.report else default_report_path(docx_path)
    report_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(docx_path))

    expected_rgb = parse_color_value(analysis.get("color") if analysis else "orange")
    expected_hex = rgb_to_hex(expected_rgb)
    formatting = check_revision_formatting(doc, expected_hex)

    application = None
    if analysis:
        application = check_stage3_application(doc, analysis, identified)

    consolidated_dir = Path(args.consolidated_dir).expanduser().resolve()
    eu_html = resolve_eu_html(docx_path, args.html, consolidated_dir)
    eu_pdf = resolve_eu_pdf(docx_path, args.pdf, consolidated_dir)

    effective_docx_text = collect_docx_text(doc, exclude_strike=True)
    source_comparisons: list[dict] = []
    source_compare_errors: list[str] = []

    practical_min_html = (
        args.practical_min_coverage_html
        if args.practical_min_coverage_html is not None
        else args.practical_min_coverage
    )
    practical_min_pdf = (
        args.practical_min_coverage_pdf
        if args.practical_min_coverage_pdf is not None
        else args.practical_min_coverage
    )

    if eu_html is not None:
        source_comparisons.append(
            build_source_comparison(
                source_type="html",
                source_path=eu_html,
                effective_docx_text=effective_docx_text,
                top_n=args.top,
                min_coverage=args.min_coverage,
                practical_min_coverage=practical_min_html,
                levenshtein_max_exact_chars=args.levenshtein_max_exact_chars,
                token_levenshtein_max_exact_tokens=args.token_levenshtein_max_exact_tokens,
            )
        )
    elif args.html is not None:
        source_compare_errors.append(missing_source_issue("html", args.html))

    if eu_pdf is not None:
        try:
            source_comparisons.append(
                build_source_comparison(
                    source_type="pdf",
                    source_path=eu_pdf,
                    effective_docx_text=effective_docx_text,
                    top_n=args.top,
                    min_coverage=args.min_coverage,
                    practical_min_coverage=practical_min_pdf,
                    levenshtein_max_exact_chars=args.levenshtein_max_exact_chars,
                    token_levenshtein_max_exact_tokens=args.token_levenshtein_max_exact_tokens,
                )
            )
        except RuntimeError as exc:
            source_compare_errors.append(str(exc))
    elif args.pdf is not None:
        source_compare_errors.append(missing_source_issue("pdf", args.pdf))

    final_compare = {
        "docx_text_mode": "all_text_except_strikethrough_runs",
        "comparisons": source_comparisons,
    }

    issues: list[str] = []
    notes: list[str] = []
    current_inserted = formatting["current_layer_runs"]["inserted_runs_total"]
    current_deleted = formatting["current_layer_runs"]["deleted_runs_total"]
    if current_inserted == 0:
        issues.append("No inserted (double-underline) revision runs found for current amendment layer color.")

    # Deletions are expected primarily when replacements were applied in this layer.
    replacements_applied = int((analysis or {}).get("items_applied_by_replacement", 0) or 0)
    if replacements_applied > 0 and current_deleted == 0:
        issues.append("No deleted (strikethrough) revision runs found for current amendment layer color despite replacement items.")

    if application and application["inserted_items_missing_from_output"] > 0:
        issues.append("One or more inserted amendment items were not found in output text.")
    if application and application["replacement_instruction_items_not_properly_applied"] > 0:
        issues.append("One or more replacement instructions were not applied as replacement/already_applied.")
    if application and application.get("inserted_article_placement", {}).get("missing_inserted_article_heading", 0) > 0:
        issues.append("One or more inserted-article headings were not found in output DOCX.")
    if application and application.get("inserted_article_placement", {}).get("inserted_article_items_without_marker", 0) > 0:
        issues.append("One or more inserted-article instruction items are missing inserted article markers (high risk of misplacement).")
    if application and application.get("inserted_article_placement", {}).get("missing_exact_inserted_article_heading", 0) > 0:
        issues.append("One or more inserted articles do not have an exact heading match in output DOCX (expected 'Article Xy').")
    if application and application.get("inserted_article_placement", {}).get("duplicate_exact_inserted_article_heading", 0) > 0:
        issues.append("One or more inserted articles have duplicate exact heading matches in output DOCX.")
    if application and application.get("inserted_article_placement", {}).get("missing_from_inserted_article_section", 0) > 0:
        issues.append("One or more inserted-article amendment lines were not found under their inserted article heading.")
    if application and application.get("inserted_article_placement", {}).get("found_before_inserted_article_heading", 0) > 0:
        issues.append("One or more inserted-article amendment lines appear before the inserted article heading (possible placement leak).")
    if application and application.get("inserted_article_placement", {}).get("found_in_other_article_sections", 0) > 0:
        issues.append("One or more inserted-article amendment lines also appear in non-target article sections (possible cross-section leakage).")
    if application and application.get("article_paragraph_order", {}).get("marker_order_inversions", 0) > 0:
        issues.append("One or more article sections contain out-of-order top-level paragraph markers (possible paragraph placement error).")

    issues.extend(source_compare_errors)
    if not source_comparisons:
        issues.append(
            "No EU consolidated sources were compared. Provide --html and/or --pdf, "
            "or place matching files in consolidated/<docx_stem>.html/.pdf."
        )
    for comparison in source_comparisons:
        if not comparison["coverage_passed"]:
            issues.append(source_fail_issue(comparison["source_type"]))
        elif not comparison["strict_coverage_passed"]:
            notes.append(source_strict_note(comparison["source_type"]))

    report = {
        "created_at_utc": datetime.now(timezone.utc).isoformat(),
        "docx_path": str(docx_path),
        "analysis_json_path": str(analysis_path) if analysis_path.exists() else None,
        "analysis_json_found": analysis is not None,
        "identified_json_path": str(identified_path) if identified_path.exists() else None,
        "identified_json_found": identified is not None,
        "checks": {
            "revision_formatting": formatting,
            "stage3_application": application,
            "final_consolidated_match": final_compare,
        },
        "notes": notes,
        "issues": issues,
        "passed": len(issues) == 0,
    }

    report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(f"Stage 4 QA report written: {report_path}")

    for comparison in source_comparisons:
        strict_cov = comparison["strict_coverage"]["coverage_ratio"]
        practical_cov = comparison["practical_coverage"]["coverage_ratio"]
        lev_sim = comparison["levenshtein"]["normalized_similarity"]
        lev_method = comparison["levenshtein"]["method"]
        tok_lev_sim = comparison["token_levenshtein"]["normalized_similarity"]
        tok_lev_method = comparison["token_levenshtein"]["method"]
        st = comparison["source_type"]
        print(
            f"[{st}] strict={strict_cov:.4f} practical={practical_cov:.4f} "
            f"levenshtein_similarity={lev_sim:.4f} ({lev_method}) "
            f"token_levenshtein_similarity={tok_lev_sim:.4f} ({tok_lev_method})"
        )

    if report["passed"]:
        print("Stage 4 checks passed.")
    else:
        print(f"Stage 4 checks found {len(issues)} issue(s).")
        for issue in issues:
            print(f" - {issue}")

    if notes:
        print(f"Stage 4 notes ({len(notes)}):")
        for note in notes:
            print(f" - {note}")

    if args.fail_on_issues and not report["passed"]:
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
