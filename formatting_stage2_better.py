"""
formatting_stage2_better.py

Stage 2 formatting engine for EUR-Lex / Official Journal XHTML.

This version is a corrective rewrite of the prior Stage 2 script. It deliberately
uses fewer fragile ancestry assumptions and avoids broad inheritance of bold/italic
formatting that caused the previous output to appear almost wholly bold/italic.

Key fixes compared with formatting_stage2_corrected.py:
    1. Recitals: reconstruct from all leaf p.oj-normal text inside each rct_* block,
       ignoring the standalone number cell, so recital body text is retained.
    2. Adoption formula: direct narrow search for "HAVE ADOPTED THIS REGULATION:".
    3. Chapter headings: use cpt_* ID for "CHAPTER X" and scoped heading text only
       outside nested article containers.
    4. Article headings: construct "Article N" from art_* ID and use scoped
       oj-sti-art / title text as the second line; never concatenate number + title.
    5. Inline formatting: preserve explicit <b>/<strong>/<i>/<em>/<sup> formatting,
       but do not inherit arbitrary EUR-Lex CSS class styling into whole paragraphs.
    6. Footnotes: create native Word footnotes using python-docx + direct OOXML.
    7. Annex VI: generate a two-column Word table from the source XHTML table.

Dependencies:
    beautifulsoup4
    lxml
    python-docx

Run:
    python formatting_stage2_better.py L_2017168EN.01001201.xml.html Prospectus_Regulation_Stage2.docx
"""

from __future__ import annotations

from pathlib import Path
import argparse
import re
import shutil
import tempfile
import zipfile

from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


# =============================================================================
# Constants
# =============================================================================

NBSP = "\u00A0"
FN_PLACEHOLDER_RE = re.compile(r"\[\[FN:(\d+)\]\]")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


# =============================================================================
# Basic helpers
# =============================================================================

def clean_text(text: str | None) -> str:
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    return " ".join(text.split()).strip()


def load_html(path: str | Path) -> BeautifulSoup:
    return BeautifulSoup(Path(path).read_text(encoding="utf-8", errors="replace"), "lxml")


def tag_classes(tag: Tag | None) -> list[str]:
    if not isinstance(tag, Tag):
        return []
    classes = tag.get("class", [])
    if isinstance(classes, str):
        return classes.split()
    return list(classes)


def has_class(tag: Tag | None, class_name: str) -> bool:
    return class_name in tag_classes(tag)


def has_ancestor_id_prefix(tag: Tag, prefix: str) -> bool:
    parent = tag.parent
    while isinstance(parent, Tag):
        pid = parent.get("id", "")
        if pid.startswith(prefix):
            return True
        parent = parent.parent
    return False


def is_inside_other_structural_container(tag: Tag, current_id: str) -> bool:
    """Return True if tag is inside another cit/rct/cpt/art/anx container.

    This is deliberately simpler and safer than nearest-ancestor matching. It lets us
    avoid Article text while collecting Chapter headings, without losing real content
    nested in tables within the current container.
    """
    parent = tag.parent
    while isinstance(parent, Tag):
        pid = parent.get("id", "")
        if pid and pid != current_id and re.match(r"^(cit|rct|cpt|art|anx)_", pid):
            return True
        if pid == current_id:
            return False
        parent = parent.parent
    return False


def apply_non_breaking_spaces(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"(?<=\d) (?=\d{3}\b)", NBSP, text)
    text = re.sub(r"\b(EUR)\s+(?=\d)", r"\1" + NBSP, text)
    text = re.sub(r"€\s+(?=\d)", "€" + NBSP, text)
    text = re.sub(r"(?<=\d)\s+%", NBSP + "%", text)
    return text


def set_run_font(run, size: int = 11, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, *, before=6, after=6, line_spacing=1.16) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def new_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


# =============================================================================
# Footnote extraction and placeholders
# =============================================================================

def element_text_without_classes(tag: Tag, classes_to_remove: set[str]) -> str:
    tmp = BeautifulSoup(str(tag), "lxml")
    for removable in tmp.find_all(lambda t: isinstance(t, Tag) and set(tag_classes(t)).intersection(classes_to_remove)):
        removable.decompose()
    return clean_text(tmp.get_text(" ", strip=True))


def extract_footnotes_map(soup: BeautifulSoup) -> dict[str, str]:
    """Extract actual footnotes from p.oj-note only.

    The EUR-Lex source also contains hr.oj-note as a separator; that is ignored.
    """
    notes: dict[str, str] = {}
    for fallback_idx, note in enumerate(soup.select("p.oj-note"), start=1):
        marker_tag = note.select_one(".oj-note-tag")
        if marker_tag:
            marker_text = clean_text(marker_tag.get_text(" ", strip=True))
            m = re.search(r"\(?\s*(\d+)\s*\)?", marker_text)
            number = m.group(1) if m else str(fallback_idx)
            body = element_text_without_classes(note, {"oj-note-tag"})
        else:
            full = clean_text(note.get_text(" ", strip=True))
            m = re.match(r"^\(?\s*(\d+)\s*\)?\s*(.*)$", full)
            if m:
                number, body = m.groups()
            else:
                number, body = str(fallback_idx), full
        notes[number] = body.strip()
    return notes


def footnote_number_from_reference(node: Tag) -> str | None:
    if not isinstance(node, Tag):
        return None
    classes = tag_classes(node)
    if node.name == "sup" or "oj-super" in classes:
        text = clean_text(node.get_text(" ", strip=True))
        m = re.search(r"\(?\s*(\d+)\s*\)?", text)
        if m:
            return m.group(1)
    return None


def add_footnote_placeholder(paragraph, number: str) -> None:
    run = paragraph.add_run(f"[[FN:{number}]]")
    set_run_font(run, size=11)
    run.font.superscript = True


# =============================================================================
# Inline text handling
# =============================================================================

def add_inline_runs(paragraph, element: Tag, *, default_size=11) -> None:
    """Preserve explicit bold/italic/superscript only.

    Important: the earlier code allowed class-based formatting to cascade through
    large EUR-Lex containers, which made nearly the whole document bold/italic. This
    function only uses genuine inline tags (<b>, <strong>, <i>, <em>, <sup>) and
    footnote reference classes.
    """

    def walk(node, bold=False, italic=False, superscript=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(apply_non_breaking_spaces(text))
                set_run_font(run, size=default_size, bold=bold, italic=italic)
                run.font.superscript = superscript
            return
        if not isinstance(node, Tag):
            return

        fn = footnote_number_from_reference(node)
        if fn:
            add_footnote_placeholder(paragraph, fn)
            return

        next_bold = bold or node.name in {"b", "strong"}
        next_italic = italic or node.name in {"i", "em"}
        next_sup = superscript or node.name == "sup"
        for child in node.children:
            walk(child, next_bold, next_italic, next_sup)

    for child in element.children:
        walk(child)


# =============================================================================
# Word paragraph builders
# =============================================================================

def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_para_spacing(p)
    run = p.add_run(clean_text(text))
    set_run_font(run, bold=True)
    return p


def add_centered_heading(doc: Document, first_line: str, second_line: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_para_spacing(p)
    r1 = p.add_run(clean_text(first_line))
    set_run_font(r1, bold=True)
    if second_line:
        r1.add_break(WD_BREAK.LINE)
        r2 = p.add_run(clean_text(second_line))
        set_run_font(r2, bold=True)
    return p


def add_justified_text(doc: Document, text: str, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    set_para_spacing(p)
    run = p.add_run(apply_non_breaking_spaces(clean_text(text)))
    set_run_font(run)
    return p


def add_justified_element(doc: Document, element: Tag, *, left_cm=0, first_line_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    set_para_spacing(p)
    add_inline_runs(p, element)
    return p


# =============================================================================
# Extraction helpers
# =============================================================================

def p_texts_in_container(container: Tag, *, include_classes=("oj-normal", "oj-enumeration-spacing")) -> list[Tag]:
    """Return p elements with useful text, excluding nested structural containers."""
    cid = container.get("id", "")
    out: list[Tag] = []
    for p in container.find_all("p"):
        if is_inside_other_structural_container(p, cid):
            continue
        classes = tag_classes(p)
        if not any(cls in classes for cls in include_classes):
            continue
        if "oj-note" in classes:
            continue
        if clean_text(p.get_text(" ", strip=True)):
            out.append(p)
    return out


def non_number_recital_paragraphs(recital: Tag, number: str) -> list[Tag]:
    """Find recital body paragraphs, excluding standalone number paragraphs/cells."""
    paragraphs = p_texts_in_container(recital, include_classes=("oj-normal",))
    body: list[Tag] = []
    num_patterns = {
        number,
        f"({number})",
    }
    for p in paragraphs:
        txt = clean_text(p.get_text(" ", strip=True))
        if txt in num_patterns:
            continue
        # Sometimes number and body appear in the same p; keep it and strip later.
        body.append(p)
    return body


# =============================================================================
# Main document sections
# =============================================================================

def add_citations(doc: Document, soup: BeautifulSoup) -> None:
    for cit in soup.select('div.eli-subdivision[id^="cit_"]'):
        paragraphs = p_texts_in_container(cit, include_classes=("oj-normal",))
        if paragraphs:
            for p in paragraphs:
                add_justified_element(doc, p, left_cm=0)
        else:
            add_justified_text(doc, cit.get_text(" ", strip=True), left_cm=0)


def add_recitals(doc: Document, soup: BeautifulSoup) -> None:
    for rec in soup.select('div.eli-subdivision[id^="rct_"]'):
        rid = rec.get("id", "")
        number = rid.replace("rct_", "") if rid.startswith("rct_") else ""
        body_ps = non_number_recital_paragraphs(rec, number)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        set_para_spacing(p)
        rn = p.add_run(f"({number})\t" if number else "")
        set_run_font(rn)

        if body_ps:
            first = True
            for bp in body_ps:
                # Strip leading duplicate number if present inside first body paragraph.
                if first:
                    txt = clean_text(bp.get_text(" ", strip=True))
                    txt = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", txt).strip() if number else txt
                    if txt and txt != clean_text(bp.get_text(" ", strip=True)):
                        run = p.add_run(apply_non_breaking_spaces(txt))
                        set_run_font(run)
                    else:
                        add_inline_runs(p, bp)
                    first = False
                else:
                    # Very rare for recitals, but preserve additional paragraphs with a space.
                    p.add_run(" ")
                    add_inline_runs(p, bp)
        else:
            txt = clean_text(rec.get_text(" ", strip=True))
            txt = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", txt).strip() if number else txt
            run = p.add_run(apply_non_breaking_spaces(txt))
            set_run_font(run)


def add_adoption_formula(doc: Document, soup: BeautifulSoup) -> None:
    formula_re = re.compile(r"^HAVE\s+ADOPTED\s+THIS\s+REGULATION\s*:$", re.I)
    for tag in soup.find_all("p"):
        text = clean_text(tag.get_text(" ", strip=True))
        if formula_re.match(text):
            add_justified_element(doc, tag, left_cm=0)
            return
    # fallback if source has extra spacing/text in a div
    for tag in soup.find_all(["p", "div"]):
        text = clean_text(tag.get_text(" ", strip=True))
        if "HAVE ADOPTED THIS REGULATION" in text and len(text) <= 120:
            add_justified_element(doc, tag, left_cm=0)
            return


def roman_from_cpt_id(cpt_id: str) -> str:
    return cpt_id.replace("cpt_", "")


def chapter_heading(chapter: Tag) -> tuple[str, str | None]:
    cid = chapter.get("id", "")
    first = f"CHAPTER {roman_from_cpt_id(cid)}" if cid.startswith("cpt_") else "CHAPTER"

    # Find the first plausible chapter title not nested inside an article and not equal to CHAPTER X.
    for selector in [".oj-ti-section-1", ".oj-ti-section-2", ".eli-title", ".oj-doc-ti"]:
        for node in chapter.select(selector):
            if has_ancestor_id_prefix(node, "art_"):
                continue
            txt = clean_text(node.get_text(" ", strip=True))
            if not txt:
                continue
            if re.match(r"^CHAPTER\s+", txt, re.I):
                continue
            if re.match(r"^Article\s+\d+", txt, re.I):
                continue
            if len(txt) > 120:
                continue
            return first, txt
    return first, None


def article_heading(article: Tag) -> tuple[str, str | None]:
    art_id = article.get("id", "")
    num = art_id.replace("art_", "") if art_id.startswith("art_") else ""
    first = f"Article {num}" if num else "Article"

    # Prefer the official sub-title. Do NOT use oj-ti-art as it can concatenate number+title.
    for selector in [".oj-sti-art"]:
        node = article.select_one(selector)
        if node:
            txt = clean_text(node.get_text(" ", strip=True))
            if txt and not re.match(r"^Article\s+\d+", txt, re.I):
                return first, txt

    # Fallback: explicit title ID where it is not merely Article N.
    title_node = article.find(id=f"{art_id}.tit_1")
    if isinstance(title_node, Tag):
        txt = clean_text(title_node.get_text(" ", strip=True))
        # Some title nodes concatenate "Article 1Subject matter"; split it.
        m = re.match(r"^Article\s+\d+\s*(.+)$", txt, flags=re.I)
        if m and m.group(1).strip():
            return first, m.group(1).strip()
        if txt and not re.match(r"^Article\s+\d+$", txt, re.I):
            return first, txt
    return first, None


def paragraph_level(text: str) -> int:
    txt = text.strip()
    if re.match(r"^\d+\.\s+", txt):
        return 1
    if re.match(r"^\([a-z]\)\s+", txt):
        return 2
    if re.match(r"^\([ivxlcdm]+\)\s+", txt, flags=re.I):
        return 3
    if txt in {"—", "–"} or txt.startswith("—"):
        return 3
    return 0


def add_article_paragraph(doc: Document, ptag: Tag) -> None:
    text = clean_text(ptag.get_text(" ", strip=True))
    if not text:
        return
    level = paragraph_level(text)
    if level == 1:
        left, first = 1, -1
    elif level == 2:
        left, first = 2, -1
    elif level == 3:
        left, first = 3, -1
    else:
        left, first = 1, None
    add_justified_element(doc, ptag, left_cm=left, first_line_cm=first)


def article_paragraphs(article: Tag) -> list[Tag]:
    return p_texts_in_container(article, include_classes=("oj-normal", "oj-enumeration-spacing"))


def add_single_article(doc: Document, article: Tag) -> None:
    first, second = article_heading(article)
    add_centered_heading(doc, first, second)
    for p in article_paragraphs(article):
        # Exclude article heading/subheading paragraphs just in case.
        classes = tag_classes(p)
        if any(cls in classes for cls in ["oj-ti-art", "oj-sti-art", "eli-title", "oj-ti-section-1", "oj-ti-section-2"]):
            continue
        add_article_paragraph(doc, p)


def add_operatives(doc: Document, soup: BeautifulSoup) -> None:
    processed: set[str] = set()
    for cpt in soup.select('div.eli-subdivision[id^="cpt_"]'):
        first, second = chapter_heading(cpt)
        add_centered_heading(doc, first, second)
        for art in cpt.select('div.eli-subdivision[id^="art_"]'):
            aid = art.get("id", "")
            if aid in processed:
                continue
            add_single_article(doc, art)
            processed.add(aid)

    # Fallback for any articles outside chapters.
    for art in soup.select('div.eli-subdivision[id^="art_"]'):
        aid = art.get("id", "")
        if aid in processed:
            continue
        add_single_article(doc, art)
        processed.add(aid)


# =============================================================================
# Annexes and Annex VI table
# =============================================================================

def set_cell_font(cell, bold=False, size=11) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)


def cell_text_from_html(cell: Tag) -> str:
    parts = []
    children = cell.find_all(["p", "div"], recursive=False)
    if children:
        for child in children:
            txt = clean_text(child.get_text(" ", strip=True))
            if txt:
                parts.append(apply_non_breaking_spaces(txt))
    else:
        txt = clean_text(cell.get_text(" ", strip=True))
        if txt:
            parts.append(apply_non_breaking_spaces(txt))
    return "\n".join(parts)


def find_largest_html_table(annex: Tag):
    tables = annex.select("table.oj-table") or annex.find_all("table")
    if not tables:
        return None
    return max(tables, key=lambda t: len(t.find_all("tr")) * 10 + len(t.find_all(["td", "th"])))


def build_annex_vi_table(doc: Document, annex: Tag) -> None:
    html_table = find_largest_html_table(annex)
    if html_table is None:
        add_justified_text(doc, annex.get_text(" ", strip=True))
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for tr in html_table.find_all("tr"):
        cells = tr.find_all(["th", "td"], recursive=False)
        if not cells:
            continue
        extracted = [cell_text_from_html(c) for c in cells]
        extracted = [e for e in extracted if e]
        if not extracted:
            continue
        left = extracted[0]
        right = "\n".join(extracted[1:]) if len(extracted) > 1 else ""
        row = table.add_row()
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].text = left
        row.cells[1].text = right
        is_header = any(c.name == "th" or "oj-tbl-hdr" in tag_classes(c) for c in cells)
        set_cell_font(row.cells[0], bold=is_header)
        set_cell_font(row.cells[1], bold=is_header)


def add_annexes(doc: Document, soup: BeautifulSoup) -> None:
    for annex in soup.select('div.eli-container[id^="anx_"]'):
        heading = annex.select_one("p.oj-doc-ti")
        heading_text = clean_text(heading.get_text(" ", strip=True)) if heading else annex.get("id", "Annex")
        add_centered_heading(doc, heading_text)
        if annex.get("id") == "anx_VI":
            build_annex_vi_table(doc, annex)
        else:
            # Preserve each visible paragraph in order rather than flattening the whole annex.
            paras = p_texts_in_container(annex, include_classes=("oj-normal", "oj-enumeration-spacing", "oj-doc-ti", "oj-ti-section-1", "oj-ti-section-2"))
            emitted = False
            for p in paras:
                txt = clean_text(p.get_text(" ", strip=True))
                if not txt or txt == heading_text:
                    continue
                classes = tag_classes(p)
                if any(c in classes for c in ["oj-doc-ti", "oj-ti-section-1", "oj-ti-section-2"]):
                    add_centered_heading(doc, txt)
                else:
                    add_justified_element(doc, p, left_cm=0)
                emitted = True
            if not emitted:
                body = clean_text(annex.get_text(" ", strip=True))
                if body.startswith(heading_text):
                    body = body[len(heading_text):].strip()
                add_justified_text(doc, body)


# =============================================================================
# OOXML native footnote patcher
# =============================================================================

def w_tag(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def rel_tag(local: str) -> str:
    return f"{{{REL_NS}}}{local}"


def ct_tag(local: str) -> str:
    return f"{{{CT_NS}}}{local}"


def make_text_run(text: str) -> etree._Element:
    r = etree.Element(w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_footnote_reference_run(fid: int) -> etree._Element:
    r = etree.Element(w_tag("r"))
    rpr = etree.SubElement(r, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    ref = etree.SubElement(r, w_tag("footnoteReference"))
    ref.set(w_tag("id"), str(fid))
    return r


def replace_placeholders_in_document_xml(document_xml: bytes, footnote_texts: dict[str, str]) -> tuple[bytes, dict[int, str]]:
    root = etree.fromstring(document_xml, etree.XMLParser(remove_blank_text=False))
    next_id = 1
    number_to_id: dict[str, int] = {}
    id_to_text: dict[int, str] = {}
    for t in list(root.xpath(".//w:t", namespaces={"w": W_NS})):
        if not t.text or "[[FN:" not in t.text:
            continue
        run = t.getparent()
        paragraph = run.getparent() if run is not None else None
        if run is None or paragraph is None:
            continue
        original = t.text
        parts: list[tuple[str, str | int]] = []
        pos = 0
        for m in FN_PLACEHOLDER_RE.finditer(original):
            if m.start() > pos:
                parts.append(("text", original[pos:m.start()]))
            number = m.group(1)
            if number not in number_to_id:
                number_to_id[number] = next_id
                id_to_text[next_id] = footnote_texts.get(number, "")
                next_id += 1
            parts.append(("fn", number_to_id[number]))
            pos = m.end()
        if pos < len(original):
            parts.append(("text", original[pos:]))
        idx = paragraph.index(run)
        paragraph.remove(run)
        for offset, (kind, value) in enumerate(parts):
            new_run = make_text_run(str(value)) if kind == "text" else make_footnote_reference_run(int(value))
            paragraph.insert(idx + offset, new_run)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"), id_to_text


def create_footnotes_xml(id_to_text: dict[int, str]) -> bytes:
    root = etree.Element(w_tag("footnotes"), nsmap={"w": W_NS})
    for fid, ftype, marker in [(-1, "separator", "separator"), (0, "continuationSeparator", "continuationSeparator")]:
        fn = etree.SubElement(root, w_tag("footnote"))
        fn.set(w_tag("id"), str(fid))
        fn.set(w_tag("type"), ftype)
        p = etree.SubElement(fn, w_tag("p"))
        r = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r, w_tag(marker))
    for fid, text in sorted(id_to_text.items()):
        fn = etree.SubElement(root, w_tag("footnote"))
        fn.set(w_tag("id"), str(fid))
        p = etree.SubElement(fn, w_tag("p"))
        ppr = etree.SubElement(p, w_tag("pPr"))
        etree.SubElement(ppr, w_tag("pStyle")).set(w_tag("val"), "FootnoteText")
        ind = etree.SubElement(ppr, w_tag("ind"))
        ind.set(w_tag("left"), "567")
        ind.set(w_tag("hanging"), "567")
        spacing = etree.SubElement(ppr, w_tag("spacing"))
        spacing.set(w_tag("before"), "0")
        spacing.set(w_tag("after"), "0")
        spacing.set(w_tag("line"), "240")
        spacing.set(w_tag("lineRule"), "auto")
        r_ref = etree.SubElement(p, w_tag("r"))
        rpr = etree.SubElement(r_ref, w_tag("rPr"))
        etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
        etree.SubElement(r_ref, w_tag("footnoteRef"))
        r_tab = etree.SubElement(p, w_tag("r"))
        etree.SubElement(r_tab, w_tag("tab"))
        r_text = etree.SubElement(p, w_tag("r"))
        rpr_text = etree.SubElement(r_text, w_tag("rPr"))
        rfonts = etree.SubElement(rpr_text, w_tag("rFonts"))
        rfonts.set(w_tag("ascii"), "Arial")
        rfonts.set(w_tag("hAnsi"), "Arial")
        sz = etree.SubElement(rpr_text, w_tag("sz"))
        sz.set(w_tag("val"), "18")  # 9 pt
        t = etree.SubElement(r_text, w_tag("t"))
        t.text = apply_non_breaking_spaces(text)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_relationship(rels_xml: bytes) -> bytes:
    root = etree.fromstring(rels_xml)
    existing = root.xpath(
        "./rel:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']",
        namespaces={"rel": REL_NS},
    )
    if existing:
        return rels_xml
    used = {rel.get("Id") for rel in root.findall(rel_tag("Relationship"))}
    rid = "rIdFootnotes"
    i = 1
    while rid in used:
        i += 1
        rid = f"rIdFootnotes{i}"
    rel = etree.SubElement(root, rel_tag("Relationship"))
    rel.set("Id", rid)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    rel.set("Target", "footnotes.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = etree.fromstring(content_types_xml)
    existing = root.xpath("./ct:Override[@PartName='/word/footnotes.xml']", namespaces={"ct": CT_NS})
    if existing:
        return content_types_xml
    override = etree.SubElement(root, ct_tag("Override"))
    override.set("PartName", "/word/footnotes.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_with_native_footnotes(docx_path: str | Path, footnote_texts: dict[str, str]) -> None:
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)
        document_xml = tmp / "word" / "document.xml"
        rels_xml = tmp / "word" / "_rels" / "document.xml.rels"
        content_types = tmp / "[Content_Types].xml"
        footnotes_xml = tmp / "word" / "footnotes.xml"
        new_doc_xml, id_to_text = replace_placeholders_in_document_xml(document_xml.read_bytes(), footnote_texts)
        document_xml.write_bytes(new_doc_xml)
        if id_to_text:
            footnotes_xml.write_bytes(create_footnotes_xml(id_to_text))
            rels_xml.write_bytes(ensure_footnotes_relationship(rels_xml.read_bytes()))
            content_types.write_bytes(ensure_footnotes_content_type(content_types.read_bytes()))
        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


# =============================================================================
# Build document
# =============================================================================

def build_document(html_path: str | Path, output_path: str | Path) -> Path:
    soup = load_html(html_path)
    footnotes = extract_footnotes_map(soup)
    doc = new_document()

    title = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if title:
        add_title(doc, title.get_text(" ", strip=True))

    add_citations(doc, soup)
    add_recitals(doc, soup)
    add_adoption_formula(doc, soup)
    add_operatives(doc, soup)
    add_annexes(doc, soup)

    output_path = Path(output_path)
    doc.save(output_path)
    patch_docx_with_native_footnotes(output_path, footnotes)
    return output_path


# =============================================================================
# CLI
# =============================================================================

def main() -> int:
    parser = argparse.ArgumentParser(description="Stage 2 EUR-Lex XHTML to formatted Word, with native OOXML footnotes.")
    parser.add_argument("source", help="Path to EUR-Lex XHTML/HTML input file")
    parser.add_argument("output", help="Path to output .docx file")
    args = parser.parse_args()
    output = build_document(args.source, args.output)
    print(f"Stage 2 document saved: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
