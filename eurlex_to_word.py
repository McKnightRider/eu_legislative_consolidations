from pathlib import Path
import re

from bs4 import BeautifulSoup
from docx import Document


def clean(text):
    return " ".join(text.split())


def load_soup(path):
    html = Path(path).read_text(
        encoding="utf-8",
        errors="replace"
    )
    return BeautifulSoup(html, "lxml")


def add_heading(doc, text, level=1):
    doc.add_heading(clean(text), level=level)


def extract_title(soup):

    title = soup.select_one(
        "div.eli-main-title"
    )

    if title:
        return clean(title.get_text(" "))

    return "EUR-Lex Document"


def extract_citations(doc, soup):

    citations = soup.select(
        'div.eli-subdivision[id^="cit_"]'
    )

    if not citations:
        return

    add_heading(doc, "Citations", 1)

    for citation in citations:
        doc.add_paragraph(
            clean(citation.get_text(" "))
        )


def extract_recitals(doc, soup):

    recitals = soup.select(
        'div.eli-subdivision[id^="rct_"]'
    )

    if not recitals:
        return

    add_heading(doc, "Recitals", 1)

    for recital in recitals:

        recital_id = recital["id"]

        number = recital_id.replace(
            "rct_",
            ""
        )

        doc.add_paragraph(
            clean(recital.get_text(" "))
        )


def extract_articles(doc, soup):

    articles = soup.select(
        'div.eli-subdivision[id^="art_"]'
    )

    add_heading(doc, "Articles", 1)

    for article in articles:

        article_id = article["id"]

        title = soup.find(
            id=f"{article_id}.tit_1"
        )

        heading = (
            clean(title.get_text(" "))
            if title
            else article_id
        )

        add_heading(
            doc,
            heading,
            level=2
        )

        body = clean(
            article.get_text(" ")
        )

        doc.add_paragraph(body)


def extract_annexes(doc, soup):

    annexes = soup.select(
        'div.eli-container[id^="anx_"]'
    )

    if not annexes:
        return

    add_heading(doc, "Annexes", 1)

    for annex in annexes:

        heading_tag = annex.select_one(
            "p.oj-doc-ti"
        )

        heading = (
            clean(heading_tag.get_text(" "))
            if heading_tag
            else annex["id"]
        )

        add_heading(
            doc,
            heading,
            level=2
        )

        body = clean(
            annex.get_text(" ")
        )

        doc.add_paragraph(body)


def extract_footnotes(doc, soup):

    notes = soup.select(
        "p.oj-note"
    )

    if not notes:
        return

    add_heading(
        doc,
        "Footnotes",
        level=1
    )

    for note in notes:
        doc.add_paragraph(
            clean(note.get_text(" "))
        )


def convert_to_word(
    source_file,
    output_file
):

    soup = load_soup(source_file)

    doc = Document()

    title = extract_title(soup)

    doc.add_heading(
        title,
        level=0
    )

    extract_citations(doc, soup)

    extract_recitals(doc, soup)

    extract_articles(doc, soup)

    extract_annexes(doc, soup)

    extract_footnotes(doc, soup)

    doc.save(output_file)

    print(
        f"Saved: {output_file}"
    )


if __name__ == "__main__":

    convert_to_word(
        "../html_to_word/1260706_EUPR_Original.html",
        "Prospectus_Regulation_Structured.docx"
    )
