
"""
formatting_stage2_v6.py

EUR-Lex / Official Journal XHTML -> formatted Word document.

This version updates formatting_stage2_v2.py for the further adjustments requested:
    - main title line breaks around "of 14 June 2017" and before "(Text with EEA relevance)";
    - CHAPTER number lines not bold;
    - Article number lines italic, not bold;
    - improved contextual indentation for numbered/lettered/roman paragraphs;
    - Article 2 top-level definitions start at the margin;
    - continuation paragraphs after numbered paragraphs are indented by 1 cm;
    - .tit_ metadata nodes excluded from content extraction;
    - Annex headings start on a new page;
    - Annex A./B./etc. markers are merged with following text using a tab.

Dependencies:
    beautifulsoup4
    lxml
    python-docx

Run:
    python formatting_stage2_v6.py L_2017168EN.01001201.xml.html Prospectus_Regulation_Stage2.docx
"""

from __future__ import annotations

from pathlib import Path
import argparse
import re
import shutil
import tempfile
import zipfile

from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
from docx.oxml import OxmlElement

NBSP = "\u00A0"
FN_TOKEN_RE = re.compile(r"\[\[FN:(\d+)\]\]")
STRUCTURAL_ID_RE = re.compile(r"^(cit_\d+|rct_\d+|cpt_[IVXLCDM]+|art_\d+[A-Za-z]?|anx_[IVXLCDM]+)$")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def clean_text(text: str | None) -> str:
    if not text:
        return ""
    return " ".join(text.replace("\xa0", " ").split()).strip()


def load_html(path: str | Path) -> BeautifulSoup:
    return BeautifulSoup(Path(path).read_text(encoding="utf-8", errors="replace"), "lxml")


def tag_classes(tag: Tag | None) -> list[str]:
    if not isinstance(tag, Tag):
        return []
    classes = tag.get("class", [])
    if isinstance(classes, str):
        return classes.split()
    return list(classes)


def has_class(tag: Tag | None, class_name: str) -> bool:
    return class_name in tag_classes(tag)


def apply_nbsp_rules(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"(?<=\d) (?=\d{3}\b)", NBSP, text)
    text = re.sub(r"\b(EUR)\s+(?=\d)", r"\1" + NBSP, text)
    text = re.sub(r"€\s+(?=\d)", "€" + NBSP, text)
    text = re.sub(r"(?<=\d)\s+%", NBSP + "%", text)
    return text


def set_run_font(run, *, size: int = 11, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, *, before=6, after=6, line_spacing=1.16) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def new_document() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def nearest_structural_id(tag: Tag | None) -> str | None:
    cur = tag
    while isinstance(cur, Tag):
        sid = str(cur.get("id", ""))
        if sid and STRUCTURAL_ID_RE.match(sid):
            return sid
        cur = cur.parent
    return None


def is_within_structural_container(tag: Tag, container_id: str) -> bool:
    return nearest_structural_id(tag) == container_id


def has_ancestor_with_id_prefix(tag: Tag, prefix: str) -> bool:
    cur = tag.parent
    while isinstance(cur, Tag):
        if str(cur.get("id", "")).startswith(prefix):
            return True
        cur = cur.parent
    return False


# =============================================================================
# Text extraction / footnote tokenisation
# =============================================================================

def footnote_number_from_reference(node: Tag) -> str | None:
    if not isinstance(node, Tag):
        return None
    if node.name == "sup" or "oj-super" in tag_classes(node):
        txt = clean_text(node.get_text(" ", strip=True))
        m = re.search(r"\(?\s*(\d+)\s*\)?", txt)
        if m:
            return m.group(1)
    return None


def text_with_footnote_tokens(element: Tag) -> str:
    parts: list[str] = []

    def walk(node):
        if isinstance(node, NavigableString):
            txt = clean_text(str(node))
            if txt:
                parts.append(txt)
            return
        if not isinstance(node, Tag):
            return
        if "oj-note" in tag_classes(node):
            return
        fn = footnote_number_from_reference(node)
        if fn:
            parts.append(f"[[FN:{fn}]]")
            return
        for child in node.children:
            walk(child)

    walk(element)
    text = clean_text(" ".join(parts))
    text = re.sub(r"\s*\(\s*\[\[FN:(\d+)\]\]\s*\)", r"[[FN:\1]]", text)
    text = re.sub(r"\s*\(\s*(\[\[FN:\d+\]\])\s*\)", r"\1", text)
    text = re.sub(r"\s+(\[\[FN:\d+\]\])", r"\1", text)
    text = re.sub(r"(\[\[FN:\d+\]\])\s+([,.;:])", r"\1\2", text)
    return apply_nbsp_rules(text)


def add_text_with_footnotes(paragraph, text: str, *, size=11, bold=False, italic=False) -> None:
    if not text:
        return
    pos = 0
    for m in FN_TOKEN_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run_font(r, size=size, bold=bold, italic=italic)
        r = paragraph.add_run(f"[[FN:{m.group(1)}]]")
        set_run_font(r, size=11)
        r.font.superscript = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, size=size, bold=bold, italic=italic)


# =============================================================================
# Word paragraph helpers
# =============================================================================

def split_main_title(title: str) -> list[str]:
    """Insert title line breaks requested by Simon.

    Intended result:
        REGULATION ... COUNCIL
        of 14 June 2017
        on the prospectus ... Directive 2003/71/EC
        (Text with EEA relevance)
    """
    title = clean_text(title)
    title = re.sub(r"\s+(of\s+14\s+June\s+2017)\s+", r"\n\1\n", title, flags=re.I)
    title = re.sub(r"\s+(\(Text\s+with\s+EEA\s+relevance\))", r"\n\1", title, flags=re.I)
    return [line.strip() for line in title.split("\n") if line.strip()]


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_para_spacing(p)
    lines = split_main_title(text)
    for idx, line in enumerate(lines):
        if idx:
            p.runs[-1].add_break(WD_BREAK.LINE)
        r = p.add_run(line)
        set_run_font(r, bold=True)
    return p


def add_heading_lines(doc: Document, first_line: str, second_line: str | None = None, *, first_bold=True, first_italic=False, second_bold=True, second_italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0)
    set_para_spacing(p)
    r1 = p.add_run(clean_text(first_line))
    set_run_font(r1, bold=first_bold, italic=first_italic)
    if second_line:
        r1.add_break(WD_BREAK.LINE)
        r2 = p.add_run(clean_text(second_line))
        set_run_font(r2, bold=second_bold, italic=second_italic)
    return p


def add_chapter_heading(doc: Document, first_line: str, second_line: str | None = None):
    # CHAPTER # not bold; title remains bold for readability.
    return add_heading_lines(doc, first_line, second_line, first_bold=False, first_italic=False, second_bold=True, second_italic=False)


def add_article_heading(doc: Document, first_line: str, second_line: str | None = None):
    # Article # italic, not bold; article title remains bold.
    return add_heading_lines(doc, first_line, second_line, first_bold=False, first_italic=True, second_bold=True, second_italic=False)


def add_center_heading(doc: Document, first_line: str, second_line: str | None = None):
    # General heading, used mainly for annex headings/subheadings.
    return add_heading_lines(doc, first_line, second_line, first_bold=True, first_italic=False, second_bold=True, second_italic=False)


def split_annex_heading(text: str) -> tuple[str, str | None]:
    text = clean_text(text)
    match = re.match(
        r"^(ANNEX(?:E)?S?)\s*([IVXLCDM]+)(?:\s*[\.\-–—:]+\s*|\s+)(.+)$",
        text,
        flags=re.I
    )
    if match:
        first_line = f"{match.group(1).upper()} {match.group(2)}"
        second_line = match.group(3).strip()
        return first_line, second_line
    match = re.match(r"^(ANNEX(?:E)?S?)\s*([IVXLCDM]+)\s*$", text, flags=re.I)
    if match:
        return f"{match.group(1).upper()} {match.group(2)}", None
    return text, None


def add_para(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left_cm=0, first_line_cm=None, size=11, bold=False, italic=False, before=6, after=6, line_spacing=1.16, src: str | None = None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.left_indent = Cm(left_cm)
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    set_para_spacing(p, before=before, after=after, line_spacing=line_spacing)
    add_text_with_footnotes(p, text, size=size, bold=bold, italic=italic)
    if src:
        _append_anchor(p, src)
    return p


# =============================================================================
# Footnote bodies
# =============================================================================

def element_text_without_classes(tag: Tag, classes_to_remove: set[str]) -> str:
    tmp = BeautifulSoup(str(tag), "lxml")
    for removable in tmp.find_all(lambda t: isinstance(t, Tag) and set(tag_classes(t)).intersection(classes_to_remove)):
        removable.decompose()
    return clean_text(tmp.get_text(" ", strip=True))


def extract_footnotes_map(soup: BeautifulSoup) -> dict[str, str]:
    def normalise_footnote_body(text: str) -> str:
        # EUR-Lex note text can contain spacing artefacts like "p. 1 ." or
        # "p. 64 ).". Remove spaces after a number when punctuation follows.
        text = re.sub(r"(?<=\d)\s+([\)\]\.,;:!?])", r"\1", text)
        # Also normalise bracketed OJ references: "( OJ ..." -> "(OJ ...".
        text = re.sub(r"\(\s+(?=OJ\b)", "(", text)
        return text

    notes: dict[str, str] = {}
    for fallback, note in enumerate(soup.select("p.oj-note"), start=1):
        marker = note.select_one(".oj-note-tag")
        if marker:
            marker_text = clean_text(marker.get_text(" ", strip=True))
            m = re.search(r"\(?\s*(\d+)\s*\)?", marker_text)
            number = m.group(1) if m else str(fallback)
            body = element_text_without_classes(note, {"oj-note-tag"})
        else:
            whole = clean_text(note.get_text(" ", strip=True))
            m = re.match(r"^\(?\s*(\d+)\s*\)?\s*(.*)$", whole)
            if m:
                number, body = m.groups()
            else:
                number, body = str(fallback), whole
        body = apply_nbsp_rules(body.strip())
        notes[number] = normalise_footnote_body(body)
    return notes


# =============================================================================
# Paragraph extraction and indentation
# =============================================================================

def provision_paragraphs(container: Tag, classes: tuple[str, ...] = ("oj-normal", "oj-enumeration-spacing")) -> list[Tag]:
    cid = str(container.get("id", ""))
    out: list[Tag] = []
    for p in container.find_all("p"):
        pid = str(p.get("id", ""))
        if ".tit_" in pid:
            continue
        if cid and not is_within_structural_container(p, cid):
            continue
        p_classes = tag_classes(p)
        if "oj-note" in p_classes:
            continue
        if any(c in p_classes for c in classes):
            if text_with_footnote_tokens(p):
                out.append(p)
    return out


def marker_type(text: str) -> str | None:
    t = clean_text(text)
    if re.match(r"^\d+\.", t):
        return "digit"
    # prefer single-letter markers as 'letter' to avoid misclassifying (c)/(d)
    if re.match(r"^\([a-z]\)", t):
        return "letter"
    if re.match(r"^\([ivxlcdm]+\)", t, flags=re.I):
        return "roman"
    if t.startswith("—") or t.startswith("–"):
        return "dash"
    if re.match(r"^[A-Z]\.", t):
        return "annex-letter"
    if re.match(r"^[IVXLCDM]+\.", t, flags=re.I):
        return "annex-roman"
    return None


def is_standalone_marker(text: str) -> bool:
    return bool(re.match(r"^(\([ivxlcdm]+\)|\([a-z]\)|[A-Z]\.|[IVXLCDM]+\.|—|–)$", clean_text(text), flags=re.I))


def normalise_marker_spacing(text: str, *, annex_mode=False) -> str:
    if annex_mode:
        # Annex A./B. markers should be followed by a tab, not a return.
        text = re.sub(r"^([A-Z]\.|[IVXLCDM]+\.)\s+", r"\1\t", text, flags=re.I)
    text = re.sub(r"^(\d+\.)\s+", r"\1\t", text)
    text = re.sub(r"^(\([a-z]\))\s+", r"\1\t", text)
    text = re.sub(r"^(\([ivxlcdm]+\))\s+", r"\1\t", text, flags=re.I)
    return text


def count_table_ancestors(tag: Tag) -> int:
    """Count <table> ancestors between this tag and its nearest structural container.

    EUR-Lex encodes list nesting with nested <table> elements:
      0 tables → top-level (digit markers, chapeau / continuation text)
      1 table  → letter-level items (a), (b), ..., (i), (j)
      2 tables → roman sub-items (i), (ii), (iii), ...
    """
    count = 0
    cur = tag.parent
    while isinstance(cur, Tag):
        if cur.name == "table":
            count += 1
        sid = str(cur.get("id", ""))
        if sid and STRUCTURAL_ID_RE.match(sid):
            break
        cur = cur.parent
    return count


def _append_anchor(para_p, src: str) -> None:
    """Append a Word-hidden anchor run to an existing paragraph for traceability."""
    r = para_p.add_run(f" [[SRC:{src}]]")
    set_run_font(r, size=1)
    try:
        r.font.hidden = True
    except Exception:
        try:
            r.font.color.rgb = RGBColor(255, 255, 255)
        except Exception:
            pass


def emit_numbered_paragraph_sequence(doc: Document, paragraphs: list[Tag], *, annex_mode: bool = False) -> None:
    """Emit legal-style paragraphs using HTML table-nesting depth for indentation.

    The EUR-Lex HTML encodes list nesting via nested <table> elements.
    count_table_ancestors() returns 0, 1, 2 ... which maps directly to left_cm
    for hanging-indent marker paragraphs:
      tbl_depth=0 → left=0cm, hanging=-1cm  (digit markers)
      tbl_depth=1 → left=1cm, hanging=-1cm  (letter items)
      tbl_depth=2 → left=2cm, hanging=-1cm  (roman sub-items)

    No marker-type stack or disambiguation is needed; the HTML structure is
    authoritative and avoids (c)/(d)/(i) misclassification as roman numerals.
    """
    i = 0
    while i < len(paragraphs):
        p_tag = paragraphs[i]
        txt = text_with_footnote_tokens(p_tag)
        if not txt:
            i += 1
            continue

        base_id = nearest_structural_id(p_tag) or "unknown"
        src_marker = f"{base_id}.{i}"
        tbl_depth = count_table_ancestors(p_tag)
        left_cm = tbl_depth  # 0, 1, 2, ... maps directly to user indent rules

        if is_standalone_marker(txt) and i + 1 < len(paragraphs):
            # Combine standalone marker + following body into one hanging paragraph
            nxt_txt = text_with_footnote_tokens(paragraphs[i + 1])
            src_body = f"{base_id}.{i + 1}"
            joined = normalise_marker_spacing(f"{txt}\t{nxt_txt}", annex_mode=annex_mode)
            p = add_para(doc, joined, left_cm=left_cm, first_line_cm=-1, src=src_marker)
            _append_anchor(p, src_body)
            i += 2
        elif marker_type(txt):
            # Marker already has inline body text (e.g. "1.\tThis Regulation...")
            combined = normalise_marker_spacing(txt, annex_mode=annex_mode)
            add_para(doc, combined, left_cm=left_cm, first_line_cm=-1, src=src_marker)
            i += 1
        else:
            # Continuation / body-only text sits at the same left_cm as the
            # enclosing marker level (tbl_depth).  tbl_depth=0 → flush left (0cm);
            # tbl_depth=1 → 1cm (aligned with letter marker); etc.
            add_para(doc, normalise_marker_spacing(txt, annex_mode=annex_mode), left_cm=tbl_depth, src=src_marker)
            i += 1


# =============================================================================
# Main sections
# =============================================================================

def add_enacting_entities(doc: Document, soup: BeautifulSoup) -> None:
    """Emit the opening formula (e.g. 'THE EUROPEAN PARLIAMENT AND THE COUNCIL…')
    that appears directly before the citations in the preamble block."""
    pbl = soup.find(id="pbl_1")
    if not pbl:
        return
    for child in pbl.children:
        if not isinstance(child, Tag):
            continue
        # Stop as soon as we reach the first citation container
        if re.match(r"^cit_", str(child.get("id", ""))):
            break
        if "oj-normal" in tag_classes(child):
            txt = text_with_footnote_tokens(child)
            if txt:
                add_para(doc, txt, left_cm=0)


def add_citations(doc: Document, soup: BeautifulSoup) -> None:
    for cit in soup.find_all(id=re.compile(r"^cit_\d+$")):
        paras = provision_paragraphs(cit, classes=("oj-normal",))
        if paras:
            for p in paras:
                add_para(doc, text_with_footnote_tokens(p), left_cm=0)
        else:
            add_para(doc, text_with_footnote_tokens(cit), left_cm=0)


def add_pre_recital_bridge(doc: Document, soup: BeautifulSoup) -> None:
    """Emit plain pre-recital lines (e.g. 'Whereas:') that sit between
    the last citation container and the first recital container."""
    first_recital = soup.find(id=re.compile(r"^rct_\d+$"))
    if not first_recital:
        return

    collected: list[str] = []
    for sib in first_recital.previous_siblings:
        if not isinstance(sib, Tag):
            continue

        sid = str(sib.get("id", ""))
        if re.match(r"^cit_\d+$", sid):
            break

        # Skip other structural containers; we only want free-standing bridge text.
        if sid and STRUCTURAL_ID_RE.match(sid):
            continue

        paras = provision_paragraphs(sib, classes=("oj-normal",))
        if paras:
            texts = [text_with_footnote_tokens(p) for p in paras]
        elif "oj-normal" in tag_classes(sib) or sib.name in {"p", "div"}:
            texts = [text_with_footnote_tokens(sib)]
        else:
            texts = []

        for txt in texts:
            txt = clean_text(txt)
            if txt:
                collected.append(txt)

    for txt in reversed(collected):
        add_para(doc, txt, left_cm=0)


def add_recitals(doc: Document, soup: BeautifulSoup) -> None:
    for rec in soup.find_all(id=re.compile(r"^rct_\d+$")):
        rid = str(rec.get("id", ""))
        number = rid.replace("rct_", "")
        paras = provision_paragraphs(rec, classes=("oj-normal",))
        body_texts: list[str] = []
        for p in paras:
            txt = text_with_footnote_tokens(p)
            if txt in {number, f"({number})"}:
                continue
            txt = re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", txt).strip()
            if txt:
                body_texts.append(txt)
        body = " ".join(body_texts) if body_texts else re.sub(rf"^\(?{re.escape(number)}\)?\s*", "", text_with_footnote_tokens(rec)).strip()
        add_para(doc, f"({number})\t{body}", left_cm=1, first_line_cm=-1)


def add_adoption_formula(doc: Document, soup: BeautifulSoup) -> None:
    for p in soup.find_all("p"):
        txt = text_with_footnote_tokens(p)
        if re.match(r"^HAVE\s+ADOPTED\s+THIS\s+REGULATION\s*:$", txt, flags=re.I):
            add_para(doc, txt, left_cm=0)
            return
    for tag in soup.find_all(["p", "div"]):
        txt = text_with_footnote_tokens(tag)
        if "HAVE ADOPTED THIS REGULATION" in txt and len(txt) <= 120:
            add_para(doc, txt, left_cm=0)
            return


def chapter_heading_parts(chapter: Tag) -> tuple[str, str | None]:
    cid = str(chapter.get("id", ""))
    roman = cid.replace("cpt_", "") if cid.startswith("cpt_") else ""
    first = f"CHAPTER {roman}" if roman else "CHAPTER"
    candidates: list[str] = []
    for node in chapter.find_all(id=re.compile(rf"^{re.escape(cid)}\.tit_\d+$")):
        txt = clean_text(node.get_text(" ", strip=True))
        if txt and not re.match(r"^CHAPTER\s+", txt, re.I):
            candidates.append(txt)
    if not candidates:
        for selector in [".oj-ti-section-1", ".oj-ti-section-2", ".eli-title", ".oj-doc-ti"]:
            for node in chapter.select(selector):
                if has_ancestor_with_id_prefix(node, "art_"):
                    continue
                txt = clean_text(node.get_text(" ", strip=True))
                if not txt or ".tit_" in txt:
                    continue
                if re.match(r"^CHAPTER\s+", txt, re.I) or re.match(r"^Article\s+\d+", txt, re.I):
                    continue
                if len(txt) <= 120 and txt not in candidates:
                    candidates.append(txt)
    return first, candidates[0] if candidates else None


def article_heading_parts(article: Tag) -> tuple[str, str | None]:
    aid = str(article.get("id", ""))
    num = aid.replace("art_", "") if aid.startswith("art_") else ""
    first = f"Article {num}" if num else "Article"
    title = None
    for node in article.select(".oj-sti-art"):
        if is_within_structural_container(node, aid):
            cand = clean_text(node.get_text(" ", strip=True))
            if cand and not re.match(r"^Article\s+\d+", cand, re.I):
                title = cand
                break
    if not title:
        for node in article.find_all(id=re.compile(rf"^{re.escape(aid)}\.tit_\d+$")):
            cand = clean_text(node.get_text(" ", strip=True))
            if not cand:
                continue
            m = re.match(r"^Article\s+\d+\s*(.+)$", cand, flags=re.I)
            if m and m.group(1).strip():
                title = m.group(1).strip()
                break
            if not re.match(r"^Article\s+\d+$", cand, flags=re.I):
                title = cand
                break
    return first, title


def add_operatives(doc: Document, soup: BeautifulSoup) -> None:
    processed: set[str] = set()
    for cpt in soup.find_all(id=re.compile(r"^cpt_[IVXLCDM]+$")):
        first, second = chapter_heading_parts(cpt)
        add_chapter_heading(doc, first, second)
        for art in cpt.find_all(id=re.compile(r"^art_\d+[A-Za-z]?$")):
            aid = str(art.get("id", ""))
            if aid in processed:
                continue
            a1, a2 = article_heading_parts(art)
            add_article_heading(doc, a1, a2)
            emit_numbered_paragraph_sequence(doc, provision_paragraphs(art))
            processed.add(aid)

    for art in soup.find_all(id=re.compile(r"^art_\d+[A-Za-z]?$")):
        aid = str(art.get("id", ""))
        if aid in processed:
            continue
        a1, a2 = article_heading_parts(art)
        add_article_heading(doc, a1, a2)
        emit_numbered_paragraph_sequence(doc, provision_paragraphs(art))
        processed.add(aid)


# =============================================================================
# Annexes
# =============================================================================

def set_cell_font(cell, *, bold=False, size=11) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            set_run_font(r, size=size, bold=bold)


def cell_text_from_html(cell: Tag) -> str:
    children = cell.find_all(["p", "div"], recursive=False)
    parts = [text_with_footnote_tokens(child) for child in children] if children else [text_with_footnote_tokens(cell)]
    return "\n".join([p for p in parts if p])


def find_largest_html_table(annex: Tag):
    tables = annex.select("table.oj-table") or annex.find_all("table")
    if not tables:
        return None
    return max(tables, key=lambda t: len(t.find_all("tr")) * 10 + len(t.find_all(["td", "th"])))


def emit_annex_structure(
    doc: Document,
    annex: Tag
) -> None:

    children = list(
        annex.find_all(
            recursive=False
        )
    )

    for node in children:

        text = clean_text(
            node.get_text(
                " ",
                strip=True
            )
        )

        if not text:
            continue

        #
        # Skip ANNEX heading itself
        #
        if (
            node.name == "p"
            and "oj-doc-ti"
            in tag_classes(node)
        ):
            continue

        #
        # Main Annex headings and their explanatory text.
        #
        # Example:
        #   I. Summary
        #   II. Identity of directors...
        #   The purpose is to identify...
        #
        if (
            node.name == "div"
            and "oj-enumeration-spacing"
            in tag_classes(node)
        ):

            heading_parts: list[str] = []
            body_parts: list[str] = []
            table_nodes: list[Tag] = []

            for child in node.children:
                if isinstance(child, NavigableString):
                    txt = clean_text(str(child))
                    if txt:
                        heading_parts.append(txt)
                elif isinstance(child, Tag):
                    if child.name == "table":
                        table_nodes.append(child)
                    elif "oj-normal" in tag_classes(child):
                        body_text = text_with_footnote_tokens(child)
                        if body_text:
                            body_parts.append(body_text)
                    else:
                        child_text = text_with_footnote_tokens(child)
                        if child_text:
                            heading_parts.append(child_text)

            heading_text = " ".join(part for part in heading_parts if part)
            if heading_text:
                # normalise annex markers (I., A.) to use a tab after the marker
                heading_text = normalise_marker_spacing(heading_text, annex_mode=True)
                add_para(
                    doc,
                    heading_text,
                    left_cm=1,
                    first_line_cm=-1
                )

            for body_text in body_parts:
                if body_text:
                    add_para(
                        doc,
                        body_text,
                        left_cm=1,
                        first_line_cm=0
                    )

            for table_node in table_nodes:
                for row in table_node.find_all("tr"):
                    cells = row.find_all(["td", "th"])
                    if not cells:
                        continue
                    cell_texts = [cell_text_from_html(cell) for cell in cells]
                    cell_texts = [t for t in cell_texts if t]
                    if not cell_texts:
                        continue
                    row_text = "\t".join(cell_texts)
                    # normalise markers inside annex table rows (A., I.)
                    row_text = normalise_marker_spacing(row_text, annex_mode=True)
                    add_para(
                        doc,
                        row_text,
                        left_cm=2,
                        first_line_cm=-1
                    )

            continue

        #
        # A. Offer statistics
        # B. Method and expected timetable
        #
        if node.name == "table":

            for row in node.find_all("tr"):

                cells = row.find_all(["td", "th"])

                if not cells:
                    continue

                cell_texts = [cell_text_from_html(cell) for cell in cells]
                cell_texts = [t for t in cell_texts if t]
                if not cell_texts:
                    continue

                row_text = "\t".join(cell_texts)

                add_para(
                    doc,
                    row_text,
                    left_cm=2,
                    first_line_cm=-1
                )

            continue

        #
        # Explanatory text
        #
        if (
            node.name == "p"
            and "oj-normal"
            in tag_classes(node)
        ):

            add_para(
                doc,
                text,
                left_cm=0
            )


def configure_annex_vi_borders(table):

    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement("w:tblBorders")

    for edge in ("left", "right", "bottom", "insideH"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)

    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    borders.append(top)

    inside_v = OxmlElement("w:insideV")
    inside_v.set(qn("w:val"), "single")
    borders.append(inside_v)

    tblPr.append(borders)


def set_header_row_border(row):

    for cell in row.cells:

        tcPr = cell._tc.get_or_add_tcPr()

        borders = OxmlElement("w:tcBorders")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")

        borders.append(bottom)

        tcPr.append(borders)



def build_annex_vi_table(doc: Document, annex: Tag) -> None:
    heading_lines = [
        clean_text(p.get_text(" ", strip=True))
        for p in annex.select("p.oj-ti-tbl")
        if clean_text(p.get_text(" ", strip=True))
    ]
    if heading_lines:
        first_line = heading_lines[0]
        second_line = heading_lines[1] if len(heading_lines) > 1 else None
        add_heading_lines(
            doc,
            first_line,
            second_line,
            first_bold=False,
            first_italic=False,
            second_bold=False,
            second_italic=False,
        )

    html_table = find_largest_html_table(annex)
    if html_table is None:
        add_para(doc, text_with_footnote_tokens(annex), left_cm=0)
        return
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    configure_annex_vi_borders(table)
    row_idx = 0
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for tr in html_table.find_all("tr"):
        cells = tr.find_all(["th", "td"], recursive=False)
        if not cells:
            continue
        extracted = [cell_text_from_html(c) for c in cells]
        extracted = [e for e in extracted if e]
        if not extracted:
            continue
        row = table.add_row()
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row.cells[0].text = extracted[0]
        row.cells[1].text = "\n".join(extracted[1:]) if len(extracted) > 1 else ""
        is_header = any(c.name == "th" or "oj-tbl-hdr" in tag_classes(c) for c in cells)
        set_cell_font(row.cells[0], bold=is_header)
        set_cell_font(row.cells[1], bold=is_header)
        if row_idx == 0:
            set_header_row_border(row)
        row_idx += 1


def add_annexes(doc: Document, soup: BeautifulSoup) -> None:

    for annex in soup.find_all(id=re.compile(r"^anx_[IVXLCDM]+$")):

        headings = annex.select("p.oj-doc-ti")

        if headings:
            heading_texts = [clean_text(h.get_text(" ", strip=True)) for h in headings[:2]]
            if len(heading_texts) == 1:
                heading_texts.append(None)
            first_line, second_line = split_annex_heading(heading_texts[0])
            if heading_texts[1]:
                second_line = heading_texts[1]
        else:
            first_line, second_line = str(annex.get("id", "Annex")), None

        first_line = first_line.upper()
        hp = add_heading_lines(
            doc,
            first_line,
            second_line,
            first_bold=False,
            first_italic=True,
            second_bold=True,
            second_italic=False,
        )

        hp.paragraph_format.page_break_before = True

        if annex.get("id") == "anx_VI":
            build_annex_vi_table(
                doc,
                annex
            )
            continue

        emit_annex_structure(
            doc,
            annex
        )


# =============================================================================
# OOXML footnote patcher
# =============================================================================

def w_tag(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def rel_tag(local: str) -> str:
    return f"{{{REL_NS}}}{local}"


def ct_tag(local: str) -> str:
    return f"{{{CT_NS}}}{local}"


def make_text_run(text: str) -> etree._Element:
    r = etree.Element(w_tag("r"))
    t = etree.SubElement(r, w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def make_footnote_reference_run(fid: int) -> etree._Element:
    r = etree.Element(w_tag("r"))
    rpr = etree.SubElement(r, w_tag("rPr"))
    etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
    rfonts = etree.SubElement(rpr, w_tag("rFonts"))
    rfonts.set(w_tag("ascii"), "Arial")
    rfonts.set(w_tag("hAnsi"), "Arial")
    sz = etree.SubElement(rpr, w_tag("sz")); sz.set(w_tag("val"), "22")
    szcs = etree.SubElement(rpr, w_tag("szCs")); szcs.set(w_tag("val"), "22")
    vert = etree.SubElement(rpr, w_tag("vertAlign")); vert.set(w_tag("val"), "superscript")
    ref = etree.SubElement(r, w_tag("footnoteReference")); ref.set(w_tag("id"), str(fid))
    return r


def replace_placeholders_in_document_xml(document_xml: bytes, footnote_texts: dict[str, str]) -> tuple[bytes, dict[int, str]]:
    root = etree.fromstring(document_xml, etree.XMLParser(remove_blank_text=False))
    next_id = 1
    number_to_id: dict[str, int] = {}
    id_to_text: dict[int, str] = {}
    for t in list(root.xpath(".//w:t", namespaces={"w": W_NS})):
        if not t.text or "[[FN:" not in t.text:
            continue
        run = t.getparent(); para = run.getparent() if run is not None else None
        if run is None or para is None:
            continue
        original = t.text
        parts: list[tuple[str, str | int]] = []
        pos = 0
        for m in FN_TOKEN_RE.finditer(original):
            if m.start() > pos:
                parts.append(("text", original[pos:m.start()]))
            num = m.group(1)
            if num not in number_to_id:
                number_to_id[num] = next_id
                id_to_text[next_id] = footnote_texts.get(num, "")
                next_id += 1
            parts.append(("fn", number_to_id[num]))
            pos = m.end()
        if pos < len(original):
            parts.append(("text", original[pos:]))
        idx = para.index(run)
        para.remove(run)
        for offset, (kind, value) in enumerate(parts):
            para.insert(idx + offset, make_text_run(str(value)) if kind == "text" else make_footnote_reference_run(int(value)))
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes"), id_to_text


def create_footnotes_xml(id_to_text: dict[int, str]) -> bytes:
    root = etree.Element(w_tag("footnotes"), nsmap={"w": W_NS})
    for fid, ftype, marker in [(-1, "separator", "separator"), (0, "continuationSeparator", "continuationSeparator")]:
        fn = etree.SubElement(root, w_tag("footnote")); fn.set(w_tag("id"), str(fid)); fn.set(w_tag("type"), ftype)
        p = etree.SubElement(fn, w_tag("p")); r = etree.SubElement(p, w_tag("r")); etree.SubElement(r, w_tag(marker))
    for fid, text in sorted(id_to_text.items()):
        fn = etree.SubElement(root, w_tag("footnote")); fn.set(w_tag("id"), str(fid))
        p = etree.SubElement(fn, w_tag("p"))
        ppr = etree.SubElement(p, w_tag("pPr"))
        # OOXML schema order: tabs → spacing → ind → jc
        tabs = etree.SubElement(ppr, w_tag("tabs"))
        tab_stop = etree.SubElement(tabs, w_tag("tab"))
        tab_stop.set(w_tag("val"), "left"); tab_stop.set(w_tag("pos"), "284")
        spacing = etree.SubElement(ppr, w_tag("spacing")); spacing.set(w_tag("before"), "0"); spacing.set(w_tag("after"), "0"); spacing.set(w_tag("line"), "240"); spacing.set(w_tag("lineRule"), "auto")
        ind = etree.SubElement(ppr, w_tag("ind")); ind.set(w_tag("left"), "284"); ind.set(w_tag("hanging"), "284")
        jc = etree.SubElement(ppr, w_tag("jc")); jc.set(w_tag("val"), "both")
        r_ref = etree.SubElement(p, w_tag("r")); rpr = etree.SubElement(r_ref, w_tag("rPr"))
        etree.SubElement(rpr, w_tag("rStyle")).set(w_tag("val"), "FootnoteReference")
        rfonts_ref = etree.SubElement(rpr, w_tag("rFonts")); rfonts_ref.set(w_tag("ascii"), "Arial"); rfonts_ref.set(w_tag("hAnsi"), "Arial")
        sz_ref = etree.SubElement(rpr, w_tag("sz")); sz_ref.set(w_tag("val"), "18")
        szcs_ref = etree.SubElement(rpr, w_tag("szCs")); szcs_ref.set(w_tag("val"), "18")
        vert_ref = etree.SubElement(rpr, w_tag("vertAlign")); vert_ref.set(w_tag("val"), "superscript")
        etree.SubElement(r_ref, w_tag("footnoteRef"))
        r_tab = etree.SubElement(p, w_tag("r")); etree.SubElement(r_tab, w_tag("tab"))
        r_text = etree.SubElement(p, w_tag("r")); rpr_text = etree.SubElement(r_text, w_tag("rPr"))
        rfonts = etree.SubElement(rpr_text, w_tag("rFonts")); rfonts.set(w_tag("ascii"), "Arial"); rfonts.set(w_tag("hAnsi"), "Arial")
        sz = etree.SubElement(rpr_text, w_tag("sz")); sz.set(w_tag("val"), "18")
        szcs = etree.SubElement(rpr_text, w_tag("szCs")); szcs.set(w_tag("val"), "18")
        t = etree.SubElement(r_text, w_tag("t")); t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_relationship(rels_xml: bytes) -> bytes:
    root = etree.fromstring(rels_xml)
    existing = root.xpath("./rel:Relationship[@Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes']", namespaces={"rel": REL_NS})
    if existing:
        return rels_xml
    used = {rel.get("Id") for rel in root.findall(rel_tag("Relationship"))}
    rid = "rIdFootnotes"; i = 1
    while rid in used:
        i += 1; rid = f"rIdFootnotes{i}"
    rel = etree.SubElement(root, rel_tag("Relationship"))
    rel.set("Id", rid); rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"); rel.set("Target", "footnotes.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = etree.fromstring(content_types_xml)
    existing = root.xpath("./ct:Override[@PartName='/word/footnotes.xml']", namespaces={"ct": CT_NS})
    if not existing:
        override = etree.SubElement(root, ct_tag("Override"))
        override.set("PartName", "/word/footnotes.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_with_native_footnotes(docx_path: str | Path, footnotes: dict[str, str]) -> None:
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        with zipfile.ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp)
        document_xml = tmp / "word" / "document.xml"
        rels_xml = tmp / "word" / "_rels" / "document.xml.rels"
        content_types = tmp / "[Content_Types].xml"
        footnotes_xml = tmp / "word" / "footnotes.xml"
        new_doc_xml, id_to_text = replace_placeholders_in_document_xml(document_xml.read_bytes(), footnotes)
        document_xml.write_bytes(new_doc_xml)
        if id_to_text:
            footnotes_xml.write_bytes(create_footnotes_xml(id_to_text))
            rels_xml.write_bytes(ensure_footnotes_relationship(rels_xml.read_bytes()))
            content_types.write_bytes(ensure_footnotes_content_type(content_types.read_bytes()))
        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        shutil.move(str(tmp_docx), str(docx_path))


def build_document(html_path: str | Path, output_path: str | Path) -> Path:
    soup = load_html(html_path)
    footnotes = extract_footnotes_map(soup)
    doc = new_document()
    title = soup.select_one("div.eli-main-title") or soup.select_one(".eli-main-title")
    if title:
        add_title(doc, clean_text(title.get_text(" ", strip=True)))
    add_enacting_entities(doc, soup)
    add_citations(doc, soup)
    add_pre_recital_bridge(doc, soup)
    add_recitals(doc, soup)
    add_adoption_formula(doc, soup)
    add_operatives(doc, soup)
    add_annexes(doc, soup)
    output_path = Path(output_path)
    doc.save(output_path)
    patch_docx_with_native_footnotes(output_path, footnotes)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="EUR-Lex XHTML to formatted Word with native footnotes.")
    parser.add_argument("source", help="Path to EUR-Lex XHTML/HTML input file")
    parser.add_argument("output", help="Path to output .docx file")
    args = parser.parse_args()
    output = build_document(args.source, args.output)
    print(f"Stage 2 document saved: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
